# -*- coding: utf-8 -*-
"""
LLM 网关与多模型路由模块设计方案 - Word 文档生成脚本
方向：GW（工程基础设施层）
ADR 前缀：GW
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ==================== 文档创建 ====================
doc = Document()

# 全局字体设置
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ==================== 样式辅助函数 ====================

def set_cell_background(cell, color_hex):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_code_block(doc, code_text, language='java'):
    """添加代码块（灰色背景 + Consolas 字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    # 灰色背景
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p._p.get_or_add_pPr().append(shading)
    # 左边框
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="18" w:space="4" w:color="2563EB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    return p

def add_callout(doc, text, callout_type='info'):
    """添加提示框（info=蓝色 / warning=橙色 / danger=红色 / success=绿色）"""
    color_map = {
        'info':    ('EBF5FF', '2563EB', '💡'),
        'warning': ('FFF7ED', 'EA580C', '⚠️'),
        'danger':  ('FEF2F2', 'DC2626', '🔴'),
        'success': ('F0FDF4', '16A34A', '✅'),
    }
    bg_color, border_color, icon = color_map.get(callout_type, color_map['info'])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
    p._p.get_or_add_pPr().append(shading)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="18" w:space="6" w:color="{border_color}"/>'
        f'<w:top w:val="single" w:sz="4" w:space="4" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="4" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(f'{icon} {text}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_styled_table(doc, headers, rows, col_widths=None, header_color='1E40AF'):
    """添加样式表格（深色表头 + 白色字 + 隔行变色）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_background(cell, 'F8FAFC')
    # 列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def add_heading_styled(doc, text, level=1):
    """添加样式标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level == 0:
            run.font.size = Pt(26)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        elif level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        elif level == 2:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return heading

def add_paragraph_styled(doc, text, bold=False, italic=False, font_size=11, color=None, indent=None):
    """添加样式段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.3
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * (level + 1))
    run = p.runs[0] if p.runs else p.add_run('')
    run.text = text
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_numbered(doc, text, level=0):
    """添加编号列表"""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.3
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * (level + 1))
    run = p.runs[0] if p.runs else p.add_run('')
    run.text = text
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

# ==================== 封面 ====================

# 空行撑开
for _ in range(4):
    doc.add_paragraph()

# 主标题
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('LLM 网关与多模型路由\n模块设计方案')
title_run.font.name = '微软雅黑'
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

# 副标题
subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_p.add_run('Agent 的"交通枢纽"：模型抽象 · 智能路由 · 故障降级 · 成本管控')
subtitle_run.font.name = '微软雅黑'
subtitle_run.font.size = Pt(15)
subtitle_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
subtitle_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()
doc.add_paragraph()

# 版本信息
ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver_run = ver_p.add_run('版本 v1.0.0  |  ADR 前缀：GW  |  2026-07-11')
ver_run.font.name = '微软雅黑'
ver_run.font.size = Pt(12)
ver_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
ver_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

# 层级标注
layer_p = doc.add_paragraph()
layer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
layer_run = layer_p.add_run('G 层 · 工程基础设施层 · 横切关注点')
layer_run.font.name = '微软雅黑'
layer_run.font.size = Pt(11)
layer_run.font.italic = True
layer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
layer_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ==================== 本文档范围声明 ====================

add_heading_styled(doc, '本文档范围声明', level=1)

add_heading_styled(doc, '覆盖什么', level=2)
add_bullet(doc, 'LLM 模型抽象层设计（统一 LlmClient 接口 + Provider Adapter 模式）')
add_bullet(doc, '多模型路由策略（意图驱动路由 + 成本驱动路由 + 延迟驱动路由）')
add_bullet(doc, '故障降级与 Fallback 链（主模型 → 备用模型 → 模板兜底）')
add_bullet(doc, 'API Key 池与负载均衡（多 Key 轮询 + 单 Key 限流 + 故障 Key 摘除）')
add_bullet(doc, '成本追踪与配额管理（Token 计量 + 按用户/会话/模型维度成本归因）')
add_bullet(doc, '模型健康检查与熔断（per-model CircuitBreaker + 健康探测 + 自动恢复）')
add_bullet(doc, '优雅降级策略（模型不可用 → 语义缓存 → 模板回复 → 友好错误）')

add_heading_styled(doc, '不覆盖什么', level=2)
add_bullet(doc, '具体 LLM 模型的训练/微调（属模型团队职责，见规划指南 8.5 节边界声明）')
add_bullet(doc, '推理加速与压缩（量化/蒸馏/投机解码，需模型权重，API-based Agent 无法操作）')
add_bullet(doc, '语义缓存设计（属"缓存层"方向 CACHE，本文档仅定义缓存接口对接点）')
add_bullet(doc, '用户级限流配额（属"限流配额与成本控制"方向，本文档聚焦模型级配额）')
add_bullet(doc, 'Prompt 工程与上下文工程（已有独立文档，本文档仅引用其产物）')
add_bullet(doc, '可观测性全链路追踪（属"可观测性"方向 OBS，本文档仅定义指标暴露点）')

add_heading_styled(doc, 'ADR 解释', level=2)
add_paragraph_styled(doc,
    'ADR = Architecture Decision Record（架构决策记录）。每条 ADR 记录一个关键架构决策的'
    '背景、决策、理由和后果。本文档的 ADR 前缀为 GW（Gateway），共 7 条（ADR-GW-01 ~ ADR-GW-07）。')

add_heading_styled(doc, '与其他文档的关系', level=2)
add_styled_table(doc,
    headers=['相关文档', '关系类型', '交互内容'],
    rows=[
        ['System Prompt 工程（SP）', '上游依赖', '网关在路由后，将选中的模型名注入 System Prompt 的元信息'],
        ['意图识别（INT）', '上游依赖', '路由策略依据 IntentResult 选择模型（HOW_TO→fast, SEARCH→quality）'],
        ['对话编排（DLG）', '上游依赖', '编排层调用网关的 chat 接口，网关透明处理模型选择与降级'],
        ['安全护栏（SEC）', '协作关系', '网关在调用 LLM 前经过 SEC 输入层检测，调用后经过 SEC 输出层验证'],
        ['可观测性（OBS）', '下游消费', '网关暴露模型延迟/Token/成本/错误率等 Metrics 供 OBS 采集'],
        ['评估体系（EVAL）', '下游消费', '评估体系通过网关的 A/B 路由能力进行模型对比实验'],
        ['缓存层（CACHE）', '协作关系', '网关在调用 LLM 前先查语义缓存，命中则跳过模型调用'],
    ],
    col_widths=[4, 2.5, 8])

doc.add_page_break()

# ==================== 目录 ====================

add_heading_styled(doc, '目录', level=1)

toc_items = [
    ('第一章 场景：为什么需要 LLM 网关', ''),
    ('  1.1 单模型锁定：Vendor Lock-in 风险', ''),
    ('  1.2 成本失控：所有请求都走最贵的模型', ''),
    ('  1.3 延迟方差：不同场景需要不同速度', ''),
    ('  1.4 没有网关会怎样', ''),
    ('  1.5 CampusShare 的多模型需求场景', ''),
    ('第二章 方案：业界 LLM 网关设计模式', ''),
    ('  2.1 LLM 网关核心架构', ''),
    ('  2.2 大厂案例：OpenAI / AWS Bedrock / Azure / LangChain', ''),
    ('  2.3 路由策略对比（静态 / 动态 / LLM-as-Router）', ''),
    ('  2.4 ADR 汇总', ''),
    ('第三章 流程：如何搭建 LLM 网关', ''),
    ('  3.1 前置条件与现状评估', ''),
    ('  3.2 模型抽象层设计', ''),
    ('  3.3 智能路由策略', ''),
    ('  3.4 Fallback 链与故障降级', ''),
    ('  3.5 API Key 池与负载均衡', ''),
    ('  3.6 成本追踪与配额管理', ''),
    ('  3.7 模型级限流', ''),
    ('  3.8 模型健康检查与熔断', ''),
    ('  3.9 优雅降级策略', ''),
    ('  3.10 ADR 决策表', ''),
    ('第四章 核心代码', ''),
    ('  4.1 文件架构', ''),
    ('  4.2 LlmClient 统一接口', ''),
    ('  4.3 DeepSeekAdapter（改造现有 DeepSeekClient）', ''),
    ('  4.4 QwenAdapter / OpenAiAdapter（备用模型适配器）', ''),
    ('  4.5 ModelRouter（智能路由引擎）', ''),
    ('  4.6 FallbackChain（故障降级链）', ''),
    ('  4.7 ApiKeyPool（API Key 轮询池）', ''),
    ('  4.8 CostTracker（成本追踪器）', ''),
    ('  4.9 ModelHealthChecker（模型健康检查器）', ''),
    ('  4.10 AgentChatService 集成改造 + application.yml', ''),
    ('第五章 目标：实现效果', ''),
    ('第六章 测试评估与验收', ''),
    ('第七章 总结与边界声明', ''),
    ('附录 ADR 摘要', ''),
]
for item, _ in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(item)
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ==================== 第一章 场景 ====================

add_heading_styled(doc, '第一章 场景：为什么需要 LLM 网关', level=1)

add_heading_styled(doc, '1.1 单模型锁定：Vendor Lock-in 风险', level=2)

add_paragraph_styled(doc,
    '当前 CampusShare Agent 的所有 LLM 调用都直连 DeepSeek API。DeepSeekClient 中硬编码了'
    'DeepSeek 的请求格式、响应解析、SSE 流式协议。这意味着：')

add_styled_table(doc,
    headers=['风险维度', '现状（无网关）', '后果'],
    rows=[
        ['API 不可用', 'DeepSeek API 宕机 → 整个 Agent 瘫痪', '单点故障，无备用模型'],
        ['格式绑定', 'DeepSeekRequest/Response 与 DeepSeek 强耦合', '切换模型需改全部 DTO + Client'],
        ['价格变动', 'DeepSeek 调价无应对手段', '成本被动承受，无法切换更便宜模型'],
        ['能力差异', '不同模型擅长不同任务（DeepSeek 擅长中文，GPT 擅长推理）', '无法按任务选最优模型'],
        ['区域可用性', 'DeepSeek 在部分网络环境不稳定', '无法自动切换到可用线路'],
    ],
    col_widths=[3, 6, 6])

add_callout(doc, '核心矛盾：Agent 的推理能力 100% 绑定在单一 LLM 供应商上，'
    '没有" Plan B "。生产系统不能接受这种单点依赖。', 'danger')

add_heading_styled(doc, '1.2 成本失控：所有请求都走最贵的模型', level=2)

add_paragraph_styled(doc,
    '当前所有请求——无论简单问候还是复杂推理——都使用同一个模型（deepseek-v4-flash）。'
    '这是对成本的巨大浪费：')

add_styled_table(doc,
    headers=['请求类型', '当前模型', '合适模型', '成本差距'],
    rows=[
        ['闲聊（CHAT）', 'deepseek-v4-flash', 'deepseek-v4-flash', '1x（基准）'],
        ['操作指引（HOW_TO）', 'deepseek-v4-flash', 'qwen-turbo（更快更便宜）', '~0.3x'],
        ['资源检索（SEARCH）', 'deepseek-v4-flash', 'deepseek-v4-flash', '1x（需质量）'],
        ['意图分类（内部调用）', 'deepseek-v4-flash', 'qwen-turbo（短输出够用）', '~0.2x'],
        ['摘要总结（内部调用）', 'deepseek-v4-flash', 'qwen-plus（性价比高）', '~0.5x'],
    ],
    col_widths=[4, 4, 4, 3])

add_paragraph_styled(doc,
    '按 CampusShare 日均 5000 次 Agent 调用估算，如果不区分场景统一用最贵模型，'
    '月成本约 ¥2000+。引入意图驱动路由后，约 40% 的请求可切换到更便宜的模型，'
    '月成本可降至 ¥1200 左右，节省 40%。', bold=False)

add_heading_styled(doc, '1.3 延迟方差：不同场景需要不同速度', level=2)

add_paragraph_styled(doc,
    '不同模型的 TTFT（Time To First Token）和吞吐量差异巨大。用户对延迟的容忍度'
    '也因场景而异：')

add_styled_table(doc,
    headers=['场景', '用户容忍延迟', '推荐模型', '典型 TTFT'],
    rows=[
        ['闲聊问候', '< 1s', 'qwen-turbo', '~300ms'],
        ['操作指引', '< 2s', 'qwen-turbo / deepseek-v4-flash', '~500ms'],
        ['资源检索', '< 3s', 'deepseek-v4-flash', '~800ms'],
        ['复杂推理', '< 5s（可接受）', 'deepseek-v4-flash', '~1200ms'],
    ],
    col_widths=[3.5, 3, 5, 3])

add_callout(doc, '没有路由层时，闲聊问候也走 deepseek-v4-flash（TTFT ~800ms），'
    '用户感受"慢"。引入路由后闲聊走 qwen-turbo（TTFT ~300ms），体感提升 60%。', 'info')

add_heading_styled(doc, '1.4 没有网关会怎样', level=2)

add_paragraph_styled(doc, '如果 CampusShare Agent 不引入 LLM 网关，将面临以下问题：', bold=True)

add_numbered(doc, '供应商绑定：DeepSeek API 任何波动（限流/宕机/网络抖动）直接导致 Agent 不可用，'
    '没有自动 fallback 到其他模型的能力。')
add_numbered(doc, '成本浪费：简单意图（闲聊/问候）和复杂意图（推理/检索）走同一模型，'
    '简单任务"过度付费"。')
add_numbered(doc, '延迟不优：无法按场景选择 TTFT 更短的模型，用户体验"一刀切"地慢。')
add_numbered(doc, '运维黑盒：无法追踪"哪个模型花了多少钱""哪个模型错误率最高"，'
    '成本和故障都无法归因。')
add_numbered(doc, '扩展困难：想新增模型（如 GPT-4o / Claude / 通义千问）需修改 DeepSeekClient'
    '的全部代码，改动面大、风险高。')
add_numbered(doc, 'Key 单点：单一 API Key 被限流时整体降级，无法通过多 Key 轮询分摊流量。')

add_heading_styled(doc, '1.5 CampusShare 的多模型需求场景', level=2)

add_paragraph_styled(doc,
    'CampusShare 作为一个校园资源共享平台，Agent 需要处理多种类型的用户请求。'
    '不同请求对模型能力的要求不同：')

add_styled_table(doc,
    headers=['业务场景', '示例', '模型能力需求', '推荐模型'],
    rows=[
        ['平台操作指引', '"怎么发帖？""怎么下载资源？"', '中文理解 + 简短输出', 'qwen-turbo'],
        ['资源检索问答', '"有没有线性代数的复习资料？"', '高质量中文 + 知识理解', 'deepseek-v4-flash'],
        ['闲聊互动', '"你好""谢谢"', '基本对话能力', 'qwen-turbo'],
        ['复杂问题推理', '"帮我分析这三份资料哪个更适合考试"', '深度推理 + 长文本', 'deepseek-v4-flash'],
        ['意图分类（内部）', 'Agent 内部 LLM 调用', '分类准确 + 短输出', 'qwen-turbo'],
        ['内容摘要（内部）', 'Agent 内部 LLM 调用', '摘要能力 + 中等输出', 'qwen-plus'],
    ],
    col_widths=[3, 4.5, 4, 3.5])

add_callout(doc, '结论：CampusShare 需要"意图驱动路由"——根据 IntentResult 自动选择'
    '最合适的模型，而非所有请求走同一个模型。', 'success')

doc.add_page_break()

# ==================== 第二章 方案 ====================

add_heading_styled(doc, '第二章 方案：业界 LLM 网关设计模式', level=1)

add_heading_styled(doc, '2.1 LLM 网关核心架构', level=2)

add_paragraph_styled(doc,
    'LLM 网关是 Agent 与 LLM 之间的"交通枢纽"。所有 LLM 调用经过网关，'
    '网关负责模型选择、故障降级、成本管控、负载均衡。其核心架构如下：')

add_code_block(doc, '''┌─────────────────────────────────────────────────────────────┐
│                     Agent 编排层（DLG）                       │
│              调用 gateway.chat(messages, intent)              │
└─────────────────────────┬───────────────────────────────────┘
                          │
          ┌───────────────▼───────────────┐
          │       LLM Gateway 网关层       │
          │                               │
          │  ┌─────────┐  ┌────────────┐  │
          │  │ Router  │→ │ Fallback   │  │
          │  │ 路由引擎 │  │ Chain 降级 │  │
          │  └────┬────┘  └─────┬──────┘  │
          │       │             │         │
          │  ┌────▼─────────────▼──────┐  │
          │  │   Model Abstraction     │  │
          │  │   模型抽象层（统一接口） │  │
          │  └──┬──────┬──────┬───────┘  │
          │     │      │      │          │
          │  ┌──▼──┐┌──▼──┐┌──▼──┐      │
          │  │DS   ││Qwen ││GPT  │      │
          │  │Adapt││Adapt││Adapt│      │
          │  └──┬──┘└──┬──┘└──┬──┘      │
          │     │      │      │          │
          │  ┌──▼──────▼──────▼───────┐  │
          │  │   API Key Pool          │  │
          │  │   Key 轮询 + 限流 + 摘除 │  │
          │  └─────────────────────────┘  │
          │                               │
          │  ┌─────────────────────────┐  │
          │  │ Cost Tracker + Metrics  │  │
          │  │ 成本追踪 + 指标暴露      │  │
          │  └─────────────────────────┘  │
          └───────────────────────────────┘
                          │
          ┌───────────────▼───────────────┐
          │     External LLM Providers    │
          │  DeepSeek | Qwen | OpenAI ... │
          └───────────────────────────────┘''', language='text')

add_heading_styled(doc, '网关六大核心职责', level=3)

add_styled_table(doc,
    headers=['职责', '说明', '对应组件'],
    rows=[
        ['模型抽象', '统一接口屏蔽不同 LLM 的请求/响应格式差异', 'LlmClient 接口 + Provider Adapter'],
        ['智能路由', '按意图/成本/延迟/质量选择最优模型', 'ModelRouter'],
        ['故障降级', '主模型不可用时自动切换备用模型', 'FallbackChain'],
        ['负载均衡', '多 API Key 轮询分摊流量，单 Key 限流时自动切换', 'ApiKeyPool'],
        ['成本管控', 'Token 计量 + 按维度成本归因 + 模型级配额', 'CostTracker'],
        ['健康监测', 'per-model 熔断 + 健康探测 + 自动恢复', 'ModelHealthChecker'],
    ],
    col_widths=[2.5, 7, 4.5])

add_heading_styled(doc, '2.2 大厂案例：OpenAI / AWS Bedrock / Azure / LangChain', level=2)

add_heading_styled(doc, '案例一：OpenAI API（模型提供商自带的网关能力）', level=3)
add_paragraph_styled(doc,
    'OpenAI 自身的 API 已经具备部分网关能力：模型自动选择（GPT-4o 自动路由到最新版本）、'
    '速率限制（RPM/TPM）、多 Key 支持。但其局限是只支持 OpenAI 自家模型，无法跨厂商路由。')
add_bullet(doc, '路由能力：自动选择最新模型版本（gpt-4o → gpt-4o-2024-08-06）')
add_bullet(doc, '降级能力：无跨厂商 fallback（OpenAI 宕机则不可用）')
add_bullet(doc, '成本管控：Usage API 提供 Token 统计，但不支持实时成本路由')
add_bullet(doc, '负载均衡：支持多 Organization Key，但需客户端自行轮询')

add_heading_styled(doc, '案例二：AWS Bedrock（云厂商统一网关）', level=3)
add_paragraph_styled(doc,
    'AWS Bedrock 是典型的"多模型统一网关"——通过一个 API 接口访问 Claude / Llama / '
    'Titan / Mistral 等多个模型。其核心设计：')
add_bullet(doc, '模型抽象：统一的 InvokeModel API，不同模型用 modelId 区分')
add_bullet(doc, '路由能力：客户端指定 modelId，Bedrock 负责路由到对应模型')
add_bullet(doc, '降级能力：需客户端自行实现 Fallback（Bedrock 不自动切换模型）')
add_bullet(doc, '成本管控：CloudWatch 集成，按模型维度统计 Token 和费用')
add_bullet(doc, '安全管控：IAM 策略控制模型访问权限，支持 per-model 限流')

add_heading_styled(doc, '案例三：Azure OpenAI Service（企业级网关）', level=3)
add_paragraph_styled(doc,
    'Azure OpenAI 在 OpenAI 模型基础上增加了企业级网关能力：')
add_bullet(doc, '部署抽象：一个 model name 可对应多个 deployment，流量在 deployment 间负载均衡')
add_bullet(doc, '区域灾备：多区域部署 + 自动故障转移（Traffic Manager）')
add_bullet(doc, '内容安全：内置 Content Filter，在网关层过滤有害内容')
add_bullet(doc, '配额管理：per-deployment TPM/RPM 配额，超限返回 429')

add_heading_styled(doc, '案例四：LangChain Router（开源路由框架）', level=3)
add_paragraph_styled(doc,
    'LangChain 提供了 LLMRouterChain，用 LLM 本身来做路由决策（LLM-as-Router）：')
add_bullet(doc, '路由方式：用一个轻量 LLM 判断请求应该路由到哪个模型')
add_bullet(doc, '优点：灵活，能处理复杂路由逻辑')
add_bullet(doc, '缺点：额外 LLM 调用增加延迟（~300ms）和成本；路由本身可能出错')
add_bullet(doc, '适用场景：模型差异大、路由规则难以用代码表达的场景')

add_callout(doc, 'CampusShare 借鉴策略：采用 AWS Bedrock 的"统一接口 + 多 Provider"模式，'
    '路由策略采用"意图驱动规则路由"（不使用 LLM-as-Router，避免额外延迟和成本），'
    '降级策略参考 Azure 的"区域灾备"思路实现"模型灾备"。', 'info')

add_heading_styled(doc, '2.3 路由策略对比（静态 / 动态 / LLM-as-Router）', level=2)

add_styled_table(doc,
    headers=['策略', '原理', '延迟开销', '灵活性', '适用场景'],
    rows=[
        ['静态路由', '配置文件写死 意图→模型 映射', '0ms（无额外调用）', '低（改配置需重启）', '模型少、路由规则固定'],
        ['动态规则路由', '代码规则 + 运行时配置（Redis）', '<1ms', '中（热更新规则）', '模型中等、路由规则可枚举'],
        ['LLM-as-Router', '用轻量 LLM 判断路由到哪个模型', '~300ms', '高（处理复杂逻辑）', '模型差异大、规则难表达'],
        ['延迟感知路由', '实时监测各模型延迟，选最快可用模型', '<1ms', '高（自适应）', '对延迟敏感的场景'],
        ['成本感知路由', '实时统计各模型成本，选最便宜可用模型', '<1ms', '中（需成本数据）', '对成本敏感的场景'],
    ],
    col_widths=[3, 5, 2.5, 2.5, 4])

add_paragraph_styled(doc, 'CampusShare 选型决策：', bold=True)
add_bullet(doc, '主策略：动态规则路由（意图驱动，IntentResult → 模型映射，规则存 Redis 可热更新）')
add_bullet(doc, '辅助策略：延迟感知（模型延迟超阈值时自动降级到更快模型）')
add_bullet(doc, '不选 LLM-as-Router：额外 300ms 延迟 + 额外成本，且路由本身可能出错')
add_bullet(doc, '不选成本感知路由：成本数据有滞后性（事后统计），不适合实时路由决策')

add_heading_styled(doc, '2.4 ADR 汇总', level=2)

add_styled_table(doc,
    headers=['ADR 编号', '决策标题', '核心选择'],
    rows=[
        ['ADR-GW-01', '模型抽象层', '统一 LlmClient 接口 + Provider Adapter 模式'],
        ['ADR-GW-02', '智能路由策略', '意图驱动动态规则路由（IntentResult → 模型映射）'],
        ['ADR-GW-03', 'Fallback 链', '三级降级：主模型 → 备用模型 → 模板兜底'],
        ['ADR-GW-04', 'API Key 池', '多 Key 轮询 + 单 Key 限流检测 + 故障 Key 自动摘除'],
        ['ADR-GW-05', '成本追踪', 'per-request Token 计量 + 多维度成本归因（用户/会话/模型）'],
        ['ADR-GW-06', '模型健康检查', 'per-model CircuitBreaker + 健康探测 + 半开自动恢复'],
        ['ADR-GW-07', '优雅降级', '模型不可用 → 语义缓存 → 模板回复 → 友好错误提示'],
    ],
    col_widths=[2.5, 4, 8.5])

doc.add_page_break()

# ==================== 第三章 流程 ====================

add_heading_styled(doc, '第三章 流程：如何搭建 LLM 网关', level=1)

add_heading_styled(doc, '3.1 前置条件与现状评估', level=2)

add_heading_styled(doc, '前置条件', level=3)
add_bullet(doc, '已完成 System Prompt 工程（SP）——网关需将模型名注入 Prompt 元信息')
add_bullet(doc, '已完成意图识别（INT）——路由策略依赖 IntentResult 进行模型选择')
add_bullet(doc, '已完成对话编排（DLG）——编排层是网关的调用方')
add_bullet(doc, '已配置至少 2 个 LLM 供应商的 API Key（DeepSeek + 通义千问/智谱/OpenAI）')
add_bullet(doc, 'Redis 可用（路由规则热更新 + API Key 状态缓存 + 成本数据暂存）')

add_heading_styled(doc, '现状评估：当前 LLM 调用架构', level=3)

add_paragraph_styled(doc,
    '当前 agent-service 的 LLM 调用完全直连 DeepSeek API，核心文件如下：')

add_styled_table(doc,
    headers=['文件', '职责', '问题'],
    rows=[
        ['DeepSeekClient.java', '流式/非流式 LLM 调用 + CircuitBreaker + Retry',
         '硬编码 DeepSeek 格式，无模型抽象'],
        ['DeepSeekRequest.java', '请求 DTO（model/messages/stream/temperature/maxTokens）',
         'DeepSeek 专属格式，与 API 强绑定'],
        ['DeepSeekResponse.java', '响应 DTO（choices/usage）',
         'DeepSeek 专属格式，无法复用'],
        ['ResilienceConfig.java', '4 个 CircuitBreaker Bean（deepseek/embedding/post-sync/intent）',
         '熔断器与模型绑定，无法 per-model 独立熔断'],
        ['application.yml', 'app.llm.deepseek 配置段（单模型）',
         '无多模型配置、无路由规则、无成本配置'],
        ['EmbeddingClient.java', '硅基流动 Embedding 调用 + CircuitBreaker + Retry',
         '与 Chat 调用独立，但同样无模型抽象'],
    ],
    col_widths=[5, 5.5, 5])

add_callout(doc, '关键发现：DeepSeekClient 同时承担了"HTTP 客户端"、"熔断器持有者"、'
    '"重试逻辑"三个职责，且请求/响应 DTO 与 DeepSeek 格式强耦合。'
    '引入网关需要将这三个职责拆分到不同组件。', 'warning')

add_heading_styled(doc, '已实现的能力（可复用）', level=3)
add_styled_table(doc,
    headers=['能力', '当前实现', '网关复用方式'],
    rows=[
        ['HTTP 客户端', 'WebClient + 连接池配置', '保留 WebClient，封装到 Adapter 内'],
        ['熔断器', 'ResilienceConfig 中 4 个 CircuitBreaker Bean', '改为 per-model 动态创建'],
        ['重试', 'Retry.backoff(3, 1000ms) + isRetryable 过滤', '保留重试逻辑，移入 Adapter'],
        ['SSE 流式解析', 'parseStreamChunk 方法', '抽象为 LlmClient 接口方法'],
        ['Token 计量', 'DeepSeekResponse.Usage 字段', '抽取为统一的 Usage DTO + CostTracker'],
        ['Embedding 调用', 'EmbeddingClient 独立实现', '纳入网关统一管理（Embedding 也需要多模型）'],
    ],
    col_widths=[3, 6, 6])

add_heading_styled(doc, '3.2 模型抽象层设计', level=2)

add_paragraph_styled(doc,
    '模型抽象层是网关的基础。通过统一接口 LlmClient 屏蔽不同 LLM 供应商的请求/响应格式差异，'
    '让上层路由和降级逻辑只面对统一接口，不感知具体模型实现。')

add_heading_styled(doc, '统一接口 LlmClient', level=3)

add_code_block(doc, '''/**
 * LLM 统一调用接口（ADR-GW-01）。
 *
 * 所有 LLM 供应商（DeepSeek/Qwen/OpenAI）实现此接口。
 * 上层（ModelRouter/FallbackChain）只依赖此接口，不感知具体实现。
 */
public interface LlmClient {

    /**
     * 非流式调用（意图分类、摘要等内部调用）。
     *
     * @param request 统一请求（模型名由 Router 注入，Adapter 可覆盖）
     * @return Mono<LlmResponse> 统一响应
     */
    Mono<LlmResponse> chat(LlmRequest request);

    /**
     * 流式调用（用户对话，SSE 流式输出）。
     *
     * @param request 统一请求
     * @return Flux<StreamChunk> 流式分块（content + usage）
     */
    Flux<StreamChunk> chatStream(LlmRequest request);

    /**
     * 该 Adapter 支持的模型标识列表。
     * Router 据此判断哪个 Adapter 能处理某个模型。
     */
    Set<String> supportedModels();

    /**
     * Provider 标识（如 "deepseek" / "qwen" / "openai"）。
     */
    String provider();
}''')

add_heading_styled(doc, '统一请求/响应 DTO', level=3)

add_code_block(doc, '''/**
 * 统一 LLM 请求（屏蔽不同供应商的格式差异）。
 */
@Data
@Builder
@NoArgsConstructor
@AllArgsConstructor
public class LlmRequest {
    private String model;           // 模型标识（如 deepseek-v4-flash）
    private List<Message> messages; // 消息列表
    private Double temperature;     // 温度（null 用默认）
    private Integer maxTokens;      // 最大输出 token
    private Boolean stream;         // 是否流式

    @Data
    @Builder
    @NoArgsConstructor
    @AllArgsConstructor
    public static class Message {
        private String role;
        private String content;
    }
}

/**
 * 统一 LLM 响应。
 */
@Data
@NoArgsConstructor
@AllArgsConstructor
public class LlmResponse {
    private String id;
    private String model;           // 实际使用的模型
    private String content;         // 生成的文本内容
    private Usage usage;            // Token 使用量

    @Data
    @NoArgsConstructor
    @AllArgsConstructor
    public static class Usage {
        private int promptTokens;
        private int completionTokens;
        private int totalTokens;
    }
}''')

add_heading_styled(doc, 'Provider Adapter 模式', level=3)

add_paragraph_styled(doc,
    '每个 LLM 供应商对应一个 Adapter，负责将统一 LlmRequest 转换为供应商专属格式，'
    '并将供应商响应转换回统一 LlmResponse：')

add_code_block(doc, '''┌─────────────────────────────────────────┐
│           LlmClient（统一接口）          │
│  chat(LlmRequest) → Mono<LlmResponse>  │
│  chatStream(LlmRequest) → Flux<Chunk>  │
└──────┬──────────┬──────────┬───────────┘
       │          │          │
┌──────▼────┐┌────▼────┐┌────▼────┐
│DeepSeek   ││Qwen     ││OpenAI   │
│Adapter    ││Adapter  ││Adapter  │
│           ││         ││         │
│LlmReq →   ││LlmReq → ││LlmReq → │
│DSReq      ││QwenReq  ││OAIReq   │
│DSResp →   ││QwenResp→││OAIResp→ │
│LlmResp    ││LlmResp  ││LlmResp  │
└───────────┘└─────────┘└─────────┘''', language='text')

add_callout(doc, '设计要点：Adapter 只做格式转换 + HTTP 调用 + 重试，不做路由/降级/熔断决策。'
    '熔断由 per-model CircuitBreaker 在 Adapter 外层包裹，路由和降级由 Router/FallbackChain 决策。', 'info')

add_heading_styled(doc, '3.3 智能路由策略', level=2)

add_paragraph_styled(doc,
    'ModelRouter 是网关的"大脑"。它根据意图识别结果（IntentResult）选择最合适的模型，'
    '并支持运行时热更新路由规则（存 Redis）。')

add_heading_styled(doc, '意图→模型映射规则', level=3)

add_styled_table(doc,
    headers=['意图（Intent）', '子意图（SubIntent）', '主模型', '备用模型', '理由'],
    rows=[
        ['HOW_TO', 'ACCOUNT', 'qwen-turbo', 'deepseek-v4-flash', '操作指引简单，用快模型'],
        ['HOW_TO', 'POSTING', 'qwen-turbo', 'deepseek-v4-flash', '操作指引简单，用快模型'],
        ['HOW_TO', 'DOWNLOAD', 'qwen-turbo', 'deepseek-v4-flash', '操作指引简单，用快模型'],
        ['SEARCH', 'RESOURCE', 'deepseek-v4-flash', 'qwen-plus', '需高质量中文理解'],
        ['SEARCH', 'USER', 'deepseek-v4-flash', 'qwen-plus', '需精准检索'],
        ['CHAT', 'GREETING', 'qwen-turbo', 'deepseek-v4-flash', '闲聊不需要强模型'],
        ['CHAT', 'EMOTION', 'qwen-turbo', 'deepseek-v4-flash', '情感对话用快模型'],
        ['CLARIFY', '*', 'deepseek-v4-flash', 'qwen-plus', '追问需理解上下文'],
        ['NAVIGATE', '*', 'qwen-turbo', 'deepseek-v4-flash', '导航回复用快模型'],
        ['OUT_OF_SCOPE', '*', 'qwen-turbo', 'deepseek-v4-flash', '超范围用快模型拒绝'],
    ],
    col_widths=[2.5, 2.5, 3, 3, 4])

add_heading_styled(doc, '路由决策流程', level=3)

add_code_block(doc, '''ModelRouter.route(intentResult, userId)
  │
  ├─ 1. 查 Redis 路由规则（key=agent:route:rules）
  │     └─ 命中 → 返回规则指定的模型
  │     └─ 未命中 → 走默认规则
  │
  ├─ 2. 默认规则：IntentResult.intent + subIntent → 模型映射
  │     └─ HOW_TO / CHAT / NAVIGATE / OUT_OF_SCOPE → qwen-turbo
  │     └─ SEARCH / CLARIFY → deepseek-v4-flash
  │
  ├─ 3. 健康检查过滤
  │     └─ 检查目标模型的 CircuitBreaker 状态
  │     └─ OPEN → 切换到备用模型
  │     └─ CLOSED/HALF_OPEN → 使用主模型
  │
  ├─ 4. 延迟感知（可选）
  │     └─ 查 Redis 中模型 P99 延迟（key=agent:model:latency:{model}）
  │     └─ 主模型 P99 > 阈值（3s）→ 切换到备用模型
  │
  └─ 5. 返回 RouteResult
        └─ primaryModel: 选中的主模型
        └─ fallbackModel: 备用模型
        └─ adapter: 对应的 LlmClient Adapter''', language='text')

add_heading_styled(doc, '路由规则热更新', level=3)
add_paragraph_styled(doc,
    '路由规则存储在 Redis（key=agent:route:rules），支持运行时热更新：')
add_bullet(doc, '管理后台修改路由规则 → 写入 Redis → ModelRouter 下次请求自动读取新规则')
add_bullet(doc, '无需重启服务，无需重新部署')
add_bullet(doc, 'Redis 不可用时降级到 application.yml 中的默认规则')

add_heading_styled(doc, '3.4 Fallback 链与故障降级', level=2)

add_paragraph_styled(doc,
    'FallbackChain 是网关的"保险丝"。当主模型调用失败时，按预设的降级链依次尝试备用方案，'
    '确保用户始终能得到响应。')

add_heading_styled(doc, '三级降级链', level=3)

add_styled_table(doc,
    headers=['级别', '方案', '触发条件', '用户体验', '成本'],
    rows=[
        ['L1 主模型', '路由选中的主模型（如 deepseek-v4-flash）', '正常请求', '完整 AI 回答', '正常'],
        ['L2 备用模型', '备用模型（如 qwen-turbo）', '主模型熔断/超时/5xx', '完整 AI 回答（可能质量略低）', '备用模型价格'],
        ['L3 模板兜底', '预设模板回复（按意图匹配）', '主+备都不可用', '固定话术（"我暂时无法回答..."）', '0（无 LLM 调用）'],
    ],
    col_widths=[2, 5, 3.5, 4, 2])

add_heading_styled(doc, '降级触发条件', level=3)

add_styled_table(doc,
    headers=['触发条件', '降级动作', '恢复条件'],
    rows=[
        ['CircuitBreaker OPEN', '跳过该模型，直接走备用模型', 'wait-duration 后 HALF_OPEN 探测'],
        ['HTTP 5xx 错误', '重试 max-attempts 次后降级', '重试成功或熔断恢复'],
        ['请求超时', '中断当前请求，降级到备用模型', '下次请求正常'],
        ['API Key 被限流（429）', '切换到同模型的另一个 Key', 'Key 限流窗口过期'],
        ['所有 Key 被限流', '降级到备用模型', '限流窗口过期'],
        ['主+备都不可用', '返回模板兜底回复', '熔断恢复后自动切回'],
    ],
    col_widths=[4, 6, 5])

add_heading_styled(doc, 'Fallback 链执行流程', level=3)

add_code_block(doc, '''FallbackChain.execute(messages, routeResult, intent)
  │
  ├─ try L1: routeResult.primaryModel
  │    └─ adapter.chatStream(messages)
  │    └─ 成功 → return Flux<StreamChunk>
  │    └─ 失败（熔断/超时/5xx）→ 进入 L2
  │
  ├─ try L2: routeResult.fallbackModel
  │    └─ adapter.chatStream(messages)
  │    └─ 成功 → return Flux<StreamChunk>（记录降级日志）
  │    └─ 失败 → 进入 L3
  │
  └─ L3: templateFallback(intent)
       └─ 按意图返回模板回复
       └─ HOW_TO → "我暂时无法回答这个问题，请稍后再试或联系客服"
       └─ SEARCH → "搜索服务暂时不可用，请稍后再试"
       └─ CHAT → "我暂时无法回复，请稍后再试"
       └─ return Flux.just(templateChunk)''', language='text')

add_callout(doc, '关键设计：FallbackChain 对上层透明。AgentChatService 只调用 '
    'gateway.chatStream()，不感知降级过程。降级日志和 Metrics 在 FallbackChain 内部记录。', 'info')

add_heading_styled(doc, '3.5 API Key 池与负载均衡', level=2)

add_paragraph_styled(doc,
    '单一 API Key 在高并发时容易被限流。ApiKeyPool 管理每个模型的多个 API Key，'
    '通过轮询分摊流量，并在单 Key 被限流时自动切换。')

add_heading_styled(doc, 'API Key 池设计', level=3)

add_styled_table(doc,
    headers=['配置项', '说明', '示例'],
    rows=[
        ['model', '模型标识', 'deepseek-v4-flash'],
        ['keys', 'API Key 列表（逗号分隔）', 'sk-aaa,sk-bbb,sk-ccc'],
        ['strategy', '轮询策略', 'round-robin（轮询）/ random（随机）'],
        ['rate-limit-window', '单 Key 限流检测窗口', '60s'],
        ['rate-limit-threshold', '单 Key 限流阈值', '100 次/窗口'],
        ['cooldown', '被限流 Key 冷却时间', '60s'],
    ],
    col_widths=[4, 5, 6])

add_heading_styled(doc, 'Key 轮询与限流检测流程', level=3)

add_code_block(doc, '''ApiKeyPool.getKey(model)
  │
  ├─ 1. 从 Redis 读取该模型的 Key 列表（key=agent:keys:{model}）
  │     └─ 每个 Key 维护状态：ACTIVE / RATE_LIMITED / DISABLED
  │
  ├─ 2. 过滤掉 RATE_LIMITED 和 DISABLED 的 Key
  │     └─ 全部不可用 → 返回 null（触发降级到备用模型）
  │
  ├─ 3. 从可用 Key 中按 strategy 选择一个
  │     └─ round-robin: Redis INCR 取模
  │     └─ random: 随机选一个
  │
  ├─ 4. 返回选中的 Key
  │     └─ 同时返回该 Key 的当前窗口调用计数
  │
  └─ 调用失败时：
       ├─ HTTP 429 → 标记 Key 为 RATE_LIMITED，设置冷却时间
       ├─ HTTP 401 → 标记 Key 为 DISABLED（Key 失效）
       └─ 其他错误 → 不标记 Key（可能是临时故障）''', language='text')

add_heading_styled(doc, 'Key 状态管理', level=3)
add_styled_table(doc,
    headers=['状态', '含义', '恢复条件'],
    rows=[
        ['ACTIVE', '正常可用', '—'],
        ['RATE_LIMITED', '被限流（429），暂时不可用', '冷却时间过后自动恢复为 ACTIVE'],
        ['DISABLED', 'Key 失效（401/403），永久不可用', '需人工更换 Key'],
    ],
    col_widths=[3, 6, 6])

add_heading_styled(doc, '3.6 成本追踪与配额管理', level=2)

add_paragraph_styled(doc,
    'CostTracker 在每次 LLM 调用后记录 Token 使用量和成本，支持多维度归因。'
    '这是成本管控的基础——没有计量就没有优化。')

add_heading_styled(doc, 'Token 计量', level=3)

add_paragraph_styled(doc, '每次 LLM 调用（含流式）结束后，CostTracker 记录：')

add_styled_table(doc,
    headers=['字段', '说明', '来源'],
    rows=[
        ['userId', '用户ID', 'JWT 认证'],
        ['sessionId', '会话ID', 'AgentSession'],
        ['turnId', '轮次ID', 'AgentTurn'],
        ['model', '实际使用的模型', 'LlmResponse.model'],
        ['provider', '模型供应商', 'LlmClient.provider()'],
        ['promptTokens', '输入 Token', 'Usage.promptTokens'],
        ['completionTokens', '输出 Token', 'Usage.completionTokens'],
        ['totalTokens', '总 Token', 'Usage.totalTokens'],
        ['costYuan', '成本（元）', '单价 × Token 数'],
        ['latencyMs', '调用延迟（毫秒）', 'System.currentTimeMillis() 差值'],
        ['success', '是否成功', '调用结果'],
        ['timestamp', '时间戳', '调用时刻'],
    ],
    col_widths=[3.5, 4.5, 6])

add_heading_styled(doc, '模型单价表', level=3)

add_styled_table(doc,
    headers=['模型', '输入单价（元/千Token）', '输出单价（元/千Token）', '备注'],
    rows=[
        ['deepseek-v4-flash', '0.002', '0.006', '主模型，质量优先'],
        ['qwen-turbo', '0.0003', '0.0006', '快速模型，成本优先'],
        ['qwen-plus', '0.0008', '0.002', '中端模型，性价比'],
        ['deepseek-v4-flash（缓存命中）', '0.001', '0.006', 'Prefix Cache 命中半价'],
    ],
    col_widths=[5, 4, 4, 3])

add_callout(doc, '单价表存储在 Redis（key=agent:model:pricing），支持热更新。'
    '模型调价时只需更新 Redis，无需重启服务。', 'info')

add_heading_styled(doc, '成本归因维度', level=3)

add_styled_table(doc,
    headers=['归因维度', '聚合 Key', '用途'],
    rows=[
        ['按用户', 'agent:cost:user:{userId}:{date}', '用户级成本统计 + 配额控制'],
        ['按会话', 'agent:cost:session:{sessionId}', '会话级成本统计'],
        ['按模型', 'agent:cost:model:{model}:{date}', '模型级成本对比 + 路由优化'],
        ['按意图', 'agent:cost:intent:{intent}:{date}', '意图级成本分析'],
        ['按天', 'agent:cost:daily:{date}', '日成本报表'],
    ],
    col_widths=[3, 6, 6])

add_heading_styled(doc, '成本数据写入流程', level=3)

add_code_block(doc, '''CostTracker.record(costRecord)
  │
  ├─ 1. 异步写入 Redis（不阻塞主流程）
  │     ├─ INCR agent:cost:user:{userId}:{date} by costYuan
  │     ├─ INCR agent:cost:model:{model}:{date} by costYuan
  │     ├─ INCR agent:cost:intent:{intent}:{date} by costYuan
  │     └─ INCR agent:cost:daily:{date} by costYuan
  │
  ├─ 2. 异步写入 MySQL（agent_cost_records 表，持久化）
  │     └─ 批量写入（每 100 条或每 30s 刷一次）
  │
  └─ 3. 更新 Micrometer Metrics
       ├─ Counter: agent.llm.tokens{model, type=prompt} += promptTokens
       ├─ Counter: agent.llm.tokens{model, type=completion} += completionTokens
       ├─ Counter: agent.llm.cost{model} += costYuan
       └─ Timer: agent.llm.latency{model} record latencyMs''', language='text')

add_heading_styled(doc, '3.7 模型级限流', level=2)

add_paragraph_styled(doc,
    '模型级限流防止某个模型的调用量超出 API 配额。与用户级限流（AgentRateLimiter，10次/分钟）'
    '不同，模型级限流关注的是"全局对某个模型的调用频率"。')

add_heading_styled(doc, '限流策略', level=3)

add_styled_table(doc,
    headers=['维度', '限流策略', '实现'],
    rows=[
        ['模型级 RPM', '每个模型每分钟最多 N 次调用', 'Redis 滑动窗口（key=agent:model:rpm:{model}）'],
        ['模型级 TPM', '每个模型每分钟最多 N Token', 'Redis 滑动窗口（key=agent:model:tpm:{model}）'],
        ['Key 级 RPM', '每个 API Key 每分钟最多 N 次调用', 'ApiKeyPool 内部管理'],
        ['用户级 RPM', '每个用户每分钟最多 10 次', '已有 AgentRateLimiter（保留）'],
    ],
    col_widths=[3, 5, 7])

add_heading_styled(doc, '限流触发后的行为', level=3)
add_bullet(doc, '模型级 RPM 超限 → 切换到备用模型（不拒绝用户请求）')
add_bullet(doc, '模型级 TPM 超限 → 切换到备用模型')
add_bullet(doc, '所有模型都超限 → 返回 429 + 友好提示"当前请求量较大，请稍后再试"')
add_bullet(doc, '用户级 RPM 超限 → 直接拒绝（已有逻辑，AgentRateLimiter）')

add_heading_styled(doc, '3.8 模型健康检查与熔断', level=2)

add_paragraph_styled(doc,
    'ModelHealthChecker 为每个模型维护独立的 CircuitBreaker，在模型故障时自动摘除流量，'
    '恢复后自动重新接入。')

add_heading_styled(doc, 'per-model CircuitBreaker', level=3)

add_paragraph_styled(doc,
    '当前 ResilienceConfig 中只有一个 deepSeekCircuitBreaker，引入网关后改为 '
    'per-model 动态创建：')

add_code_block(doc, '''ModelHealthChecker
  │
  ├─ 模型注册时创建对应的 CircuitBreaker
  │    └─ CircuitBreaker.of("model:" + modelName, config)
  │    └─ 配置从 application.yml 读取（per-model 可覆盖）
  │
  ├─ 调用时检查状态
  │    └─ CLOSED → 正常调用
  │    └─ OPEN → 直接降级到备用模型（不发起请求）
  │    └─ HALF_OPEN → 放行少量请求探测
  │
  └─ 状态转换
       ├─ CLOSED → OPEN: 失败率 > threshold（50%）
       ├─ OPEN → HALF_OPEN: wait-duration 后（30s）
       └─ HALF_OPEN → CLOSED: 探测请求全部成功
       └─ HALF_OPEN → OPEN: 探测请求有失败''', language='text')

add_heading_styled(doc, '健康探测', level=3)

add_paragraph_styled(doc,
    '除了被动熔断，ModelHealthChecker 还支持主动健康探测：')
add_bullet(doc, '每 60 秒对 OPEN 状态的模型发起一次轻量探测请求（1+1= 的简单 prompt）')
add_bullet(doc, '探测成功 → 转为 HALF_OPEN，放行少量真实流量')
add_bullet(doc, '探测失败 → 保持 OPEN，继续等待')
add_bullet(doc, '探测请求成本低（< 10 Token），不影响整体成本')

add_heading_styled(doc, '熔断配置（per-model 可覆盖）', level=3)

add_styled_table(doc,
    headers=['配置项', '默认值', '说明'],
    rows=[
        ['sliding-window-size', '10', '滑动窗口大小（最近 10 次调用）'],
        ['minimum-number-of-calls', '5', '最少调用次数才开始计算失败率'],
        ['failure-rate-threshold', '50.0', '失败率阈值（%）'],
        ['wait-duration-in-open-state', '30s', 'OPEN 状态持续时间'],
        ['permitted-number-of-calls-in-half-open-state', '3', 'HALF_OPEN 放行请求数'],
        ['health-check-interval', '60s', '主动健康探测间隔'],
    ],
    col_widths=[6, 3, 6])

add_heading_styled(doc, '3.9 优雅降级策略', level=2)

add_paragraph_styled(doc,
    '当所有 LLM 模型都不可用时，网关需要保证用户仍然能得到合理响应，而非直接报错。'
    '这是"最后一道防线"。')

add_heading_styled(doc, '四级降级方案', level=3)

add_styled_table(doc,
    headers=['级别', '方案', '触发条件', '响应内容', '延迟'],
    rows=[
        ['L0', '语义缓存查询', '每次调用前先查', '缓存命中的历史回复', '< 10ms'],
        ['L1', '主模型调用', '缓存未命中', '主模型 AI 回复', '~800ms'],
        ['L2', '备用模型调用', '主模型不可用', '备用模型 AI 回复', '~500ms'],
        ['L3', '模板兜底', '主+备都不可用', '固定话术', '< 5ms'],
    ],
    col_widths=[1.5, 3, 4, 5, 2.5])

add_heading_styled(doc, '模板兜底话术', level=3)

add_styled_table(doc,
    headers=['意图', '兜底话术', '备注'],
    rows=[
        ['HOW_TO', '抱歉，我暂时无法回答这个问题。您可以稍后再试，或查看帮助中心。', '引导用户自助'],
        ['SEARCH', '搜索服务暂时不可用，请稍后再试。您也可以直接浏览资源库。', '引导用户浏览'],
        ['CHAT', '我暂时无法回复，请稍后再试。', '简单道歉'],
        ['CLARIFY', '抱歉，我暂时无法理解您的问题，请稍后再试。', '引导重试'],
        ['NAVIGATE', '导航服务暂时不可用，请稍后再试。', '简单道歉'],
        ['OUT_OF_SCOPE', '这个问题超出了我的能力范围。', '固定拒绝'],
    ],
    col_widths=[3, 9, 3])

add_callout(doc, '关键原则：模板兜底是最后手段，应尽量少触发。'
    '通过多模型 Fallback + 健康检查 + 重试，将 L3 触发率控制在 < 0.1%。', 'warning')

add_heading_styled(doc, '3.10 ADR 决策表', level=2)

add_styled_table(doc,
    headers=['ADR', '决策', '选择', '拒绝的方案', '理由'],
    rows=[
        ['ADR-GW-01', '模型抽象层', '统一 LlmClient 接口 + Provider Adapter',
         '直接修改 DeepSeekClient 支持多模型',
         '接口隔离，新增模型只需加 Adapter，不改现有代码'],
        ['ADR-GW-02', '路由策略', '意图驱动动态规则路由',
         'LLM-as-Router / 静态路由',
         '规则路由 0ms 延迟、可热更新、可枚举；LLM-as-Router 额外 300ms+成本'],
        ['ADR-GW-03', 'Fallback 链', '三级降级：主模型→备用模型→模板兜底',
         '只重试不换模型 / 无限重试',
         '换模型比重试更有效（不同模型故障不相关）；模板兜底保底用户体验'],
        ['ADR-GW-04', 'API Key 池', '多 Key 轮询 + 自动摘除',
         '单 Key / 手动切换 Key',
         '自动轮询分摊流量、自动摘除故障 Key、无需人工干预'],
        ['ADR-GW-05', '成本追踪', 'per-request 计量 + 多维度归因',
         '只记录总量 / 事后批量统计',
         'per-request 可精确归因到用户/会话/模型/意图；多维度支持成本优化决策'],
        ['ADR-GW-06', '健康检查', 'per-model CircuitBreaker + 主动探测',
         '全局单一熔断器 / 无主动探测',
         'per-model 独立熔断不影响其他模型；主动探测加速恢复'],
        ['ADR-GW-07', '优雅降级', '四级降级（缓存→主模型→备用模型→模板）',
         '直接报错 / 只重试不降级',
         '四级方案确保用户始终有响应；模板兜底是最后防线'],
    ],
    col_widths=[1.8, 2.5, 3.5, 3, 4.2])

doc.add_page_break()

# ==================== 第四章 核心代码 ====================

add_heading_styled(doc, '第四章 核心代码', level=1)

add_heading_styled(doc, '4.1 文件架构', level=2)

add_paragraph_styled(doc,
    'LLM 网关模块涉及的新增/改造文件如下：')

add_styled_table(doc,
    headers=['文件', '类型', '职责'],
    rows=[
        ['LlmClient.java', '新增', '统一 LLM 调用接口'],
        ['LlmRequest.java', '新增', '统一请求 DTO'],
        ['LlmResponse.java', '新增', '统一响应 DTO'],
        ['StreamChunk.java', '新增', '流式分块（从 DeepSeekClient 提取）'],
        ['DeepSeekAdapter.java', '新增（替代 DeepSeekClient）', 'DeepSeek 供应商适配器'],
        ['QwenAdapter.java', '新增', '通义千问供应商适配器'],
        ['OpenAiAdapter.java', '新增（预留）', 'OpenAI 供应商适配器'],
        ['ModelRouter.java', '新增', '智能路由引擎'],
        ['RouteResult.java', '新增', '路由结果 DTO'],
        ['FallbackChain.java', '新增', '故障降级链'],
        ['ApiKeyPool.java', '新增', 'API Key 轮询池'],
        ['ApiKeyState.java', '新增', 'API Key 状态 DTO'],
        ['CostTracker.java', '新增', '成本追踪器'],
        ['CostRecord.java', '新增', '成本记录 DTO'],
        ['ModelHealthChecker.java', '新增', '模型健康检查器'],
        ['LlmGateway.java', '新增', '网关统一入口（整合 Router + Fallback + Cost）'],
        ['GatewayMetricsConfig.java', '新增', '网关 Metrics 配置'],
        ['GatewayWebClientConfig.java', '新增', 'per-provider WebClient 配置'],
        ['AgentChatService.java', '改造', '调用方从 DeepSeekClient 切换到 LlmGateway'],
        ['ResilienceConfig.java', '改造', 'per-model 动态熔断器'],
        ['application.yml', '改造', '多模型配置 + 路由规则 + 成本配置'],
    ],
    col_widths=[6, 4, 6])

add_callout(doc, '改造原则：DeepSeekClient.java 保留但标记 @Deprecated，'
    '新代码通过 LlmGateway 调用。意图分类等内部调用也改为通过网关路由。', 'info')

add_heading_styled(doc, '4.2 LlmClient 统一接口', level=2)

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

import reactor.core.publisher.Flux;
import reactor.core.publisher.Mono;

import java.util.Set;

/**
 * LLM 统一调用接口（ADR-GW-01）。
 *
 * 所有 LLM 供应商（DeepSeek/Qwen/OpenAI）实现此接口。
 * 上层（ModelRouter/FallbackChain/LlmGateway）只依赖此接口，
 * 不感知具体供应商的请求/响应格式差异。
 *
 * Adapter 职责边界：
 *   - 只做格式转换 + HTTP 调用 + 重试
 *   - 不做路由/降级/熔断决策（由上层负责）
 *   - 不做成本追踪（由 CostTracker 负责接收 Adapter 的 Usage）
 */
public interface LlmClient {

    /**
     * 非流式调用（意图分类、摘要等内部调用）。
     *
     * @param request 统一请求（model 由 Router 注入）
     * @return Mono<LlmResponse> 统一响应
     */
    Mono<LlmResponse> chat(LlmRequest request);

    /**
     * 流式调用（用户对话，SSE 流式输出）。
     *
     * @param request 统一请求
     * @return Flux<StreamChunk> 流式分块（content + usage）
     */
    Flux<StreamChunk> chatStream(LlmRequest request);

    /**
     * 该 Adapter 支持的模型标识列表。
     * Router 据此判断哪个 Adapter 能处理某个模型。
     * 例如 DeepSeekAdapter 返回 {"deepseek-v4-flash", "deepseek-v4-lite"}
     */
    Set<String> supportedModels();

    /**
     * Provider 标识（如 "deepseek" / "qwen" / "openai"）。
     * 用于成本归因和日志区分。
     */
    String provider();
}''')

add_heading_styled(doc, '统一请求/响应 DTO', level=3)

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

import com.fasterxml.jackson.annotation.JsonProperty;
import lombok.AllArgsConstructor;
import lombok.Builder;
import lombok.Data;
import lombok.NoArgsConstructor;

import java.util.List;

/**
 * 统一 LLM 请求（屏蔽不同供应商的格式差异）。
 * Adapter 负责将此转换为供应商专属请求格式。
 */
@Data
@Builder
@NoArgsConstructor
@AllArgsConstructor
public class LlmRequest {
    private String model;
    private List<Message> messages;
    private Double temperature;
    private Integer maxTokens;
    private Boolean stream;

    @JsonProperty("stream_options")
    private StreamOptions streamOptions;

    @Data
    @Builder
    @NoArgsConstructor
    @AllArgsConstructor
    public static class Message {
        private String role;
        private String content;
    }

    @Data
    @Builder
    @NoArgsConstructor
    @AllArgsConstructor
    public static class StreamOptions {
        @JsonProperty("include_usage")
        private Boolean includeUsage;
    }
}''')

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

import lombok.AllArgsConstructor;
import lombok.Data;
import lombok.NoArgsConstructor;

/**
 * 统一 LLM 响应。
 */
@Data
@NoArgsConstructor
@AllArgsConstructor
public class LlmResponse {
    private String id;
    private String model;
    private String content;
    private Usage usage;

    @Data
    @NoArgsConstructor
    @AllArgsConstructor
    public static class Usage {
        private int promptTokens;
        private int completionTokens;
        private int totalTokens;
    }
}''')

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

/**
 * 流式分块（从 DeepSeekClient.StreamChunk 提取为通用类型）。
 */
public record StreamChunk(String content, LlmResponse.Usage usage) {}''')

add_heading_styled(doc, '4.3 DeepSeekAdapter（改造现有 DeepSeekClient）', level=2)

add_paragraph_styled(doc,
    'DeepSeekAdapter 实现了 LlmClient 接口，将统一 LlmRequest 转换为 DeepSeek 专属请求格式，'
    '并复用现有的 WebClient + CircuitBreaker + Retry 逻辑：')

add_code_block(doc, '''package com.campushare.agent.llm.gateway.adapter;

import com.campushare.agent.llm.DeepSeekRequest;
import com.campushare.agent.llm.DeepSeekResponse;
import com.campushare.agent.llm.gateway.LlmClient;
import com.campushare.agent.llm.gateway.LlmRequest;
import com.campushare.agent.llm.gateway.LlmResponse;
import com.campushare.agent.llm.gateway.StreamChunk;
import com.fasterxml.jackson.databind.ObjectMapper;
import io.github.resilience4j.circuitbreaker.CircuitBreaker;
import io.github.resilience4j.reactor.circuitbreaker.operator.CircuitBreakerOperator;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.beans.factory.annotation.Value;
import org.springframework.core.ParameterizedTypeReference;
import org.springframework.http.codec.ServerSentEvent;
import org.springframework.stereotype.Component;
import org.springframework.web.reactive.function.client.WebClient;
import org.springframework.web.reactive.function.client.WebClientResponseException;
import reactor.core.publisher.Flux;
import reactor.core.publisher.Mono;
import reactor.util.retry.Retry;

import java.io.IOException;
import java.time.Duration;
import java.util.Set;
import java.util.concurrent.TimeoutException;

@Slf4j
@Component
@RequiredArgsConstructor
public class DeepSeekAdapter implements LlmClient {

    private final WebClient deepSeekWebClient;
    private final ObjectMapper objectMapper;
    private final CircuitBreaker deepSeekCircuitBreaker;

    @Value("${app.llm.deepseek.retry.max-attempts:3}")
    private int maxRetryAttempts;

    @Value("${app.llm.deepseek.retry.backoff:1000}")
    private long retryBackoffMs;

    @Override
    public Mono<LlmResponse> chat(LlmRequest request) {
        DeepSeekRequest dsRequest = toDeepSeekRequest(request);

        return deepSeekWebClient.post()
                .uri("/v1/chat/completions")
                .bodyValue(dsRequest)
                .retrieve()
                .bodyToMono(DeepSeekResponse.class)
                .map(this::toLlmResponse)
                .transform(CircuitBreakerOperator.of(deepSeekCircuitBreaker))
                .retryWhen(Retry.backoff(maxRetryAttempts, Duration.ofMillis(retryBackoffMs))
                        .filter(this::isRetryable)
                        .doBeforeRetry(rs -> log.warn("Retrying DeepSeek chat, attempt {}", rs.totalRetries() + 1)))
                .doOnError(e -> log.error("DeepSeek chat error after retries", e));
    }

    @Override
    public Flux<StreamChunk> chatStream(LlmRequest request) {
        DeepSeekRequest dsRequest = toDeepSeekRequest(request);
        dsRequest.setStream(true);
        dsRequest.setStreamOptions(DeepSeekRequest.StreamOptions.builder()
                .includeUsage(true).build());

        return deepSeekWebClient.post()
                .uri("/v1/chat/completions")
                .bodyValue(dsRequest)
                .retrieve()
                .bodyToFlux(new ParameterizedTypeReference<ServerSentEvent<String>>() {})
                .filter(sse -> sse.data() != null)
                .mapNotNull(this::parseStreamChunk)
                .takeUntil(chunk -> "[DONE]".equals(chunk.content()))
                .filter(chunk -> !"[DONE]".equals(chunk.content()))
                .transform(CircuitBreakerOperator.of(deepSeekCircuitBreaker))
                .retryWhen(Retry.backoff(maxRetryAttempts, Duration.ofMillis(retryBackoffMs))
                        .filter(this::isRetryable)
                        .doBeforeRetry(rs -> log.warn("Retrying DeepSeek stream, attempt {}", rs.totalRetries() + 1)))
                .doOnError(e -> log.error("DeepSeek stream error after retries", e));
    }

    @Override
    public Set<String> supportedModels() {
        return Set.of("deepseek-v4-flash", "deepseek-v4-lite");
    }

    @Override
    public String provider() {
        return "deepseek";
    }

    // ==================== 格式转换 ====================

    private DeepSeekRequest toDeepSeekRequest(LlmRequest request) {
        return DeepSeekRequest.builder()
                .model(request.getModel())
                .messages(request.getMessages().stream()
                        .map(m -> DeepSeekRequest.Message.builder()
                                .role(m.getRole()).content(m.getContent()).build())
                        .toList())
                .stream(request.getStream())
                .temperature(request.getTemperature())
                .maxTokens(request.getMaxTokens())
                .build();
    }

    private LlmResponse toLlmResponse(DeepSeekResponse dsResp) {
        LlmResponse.Usage usage = null;
        if (dsResp.getUsage() != null) {
            usage = new LlmResponse.Usage(
                    dsResp.getUsage().getPromptTokens(),
                    dsResp.getUsage().getCompletionTokens(),
                    dsResp.getUsage().getTotalTokens());
        }
        String content = "";
        if (dsResp.getChoices() != null && !dsResp.getChoices().isEmpty()) {
            content = dsResp.getChoices().get(0).getMessage().getContent();
        }
        return new LlmResponse(dsResp.getId(), dsResp.getModel(), content, usage);
    }

    private StreamChunk parseStreamChunk(ServerSentEvent<String> sse) {
        String data = sse.data();
        if ("[DONE]".equals(data)) {
            return new StreamChunk("[DONE]", null);
        }
        try {
            DeepSeekResponse resp = objectMapper.readValue(data, DeepSeekResponse.class);
            if (resp.getUsage() != null) {
                return new StreamChunk(null, new LlmResponse.Usage(
                        resp.getUsage().getPromptTokens(),
                        resp.getUsage().getCompletionTokens(),
                        resp.getUsage().getTotalTokens()));
            }
            if (resp.getChoices() != null && !resp.getChoices().isEmpty()) {
                DeepSeekResponse.Choice choice = resp.getChoices().get(0);
                if (choice.getDelta() != null && choice.getDelta().getContent() != null) {
                    return new StreamChunk(choice.getDelta().getContent(), null);
                }
            }
        } catch (Exception e) {
            log.warn("Failed to parse SSE chunk: {}", data, e);
        }
        return null;
    }

    private boolean isRetryable(Throwable throwable) {
        if (throwable instanceof WebClientResponseException ex) {
            return ex.getStatusCode().is5xxServerError();
        }
        return throwable instanceof TimeoutException || throwable instanceof IOException;
    }
}''')

add_heading_styled(doc, '4.4 QwenAdapter（备用模型适配器）', level=2)

add_paragraph_styled(doc,
    'QwenAdapter 实现通义千问的适配。通义千问 API 兼容 OpenAI 格式，'
    '因此 Adapter 逻辑与 DeepSeekAdapter 类似，但使用独立的 WebClient 和 CircuitBreaker：')

add_code_block(doc, '''package com.campushare.agent.llm.gateway.adapter;

import com.campushare.agent.llm.gateway.LlmClient;
import com.campushare.agent.llm.gateway.LlmRequest;
import com.campushare.agent.llm.gateway.LlmResponse;
import com.campushare.agent.llm.gateway.StreamChunk;
import com.fasterxml.jackson.databind.ObjectMapper;
import io.github.resilience4j.circuitbreaker.CircuitBreaker;
import io.github.resilience4j.reactor.circuitbreaker.operator.CircuitBreakerOperator;
import lombok.extern.slf4j.Slf4j;
import org.springframework.beans.factory.annotation.Qualifier;
import org.springframework.beans.factory.annotation.Value;
import org.springframework.core.ParameterizedTypeReference;
import org.springframework.http.codec.ServerSentEvent;
import org.springframework.stereotype.Component;
import org.springframework.web.reactive.function.client.WebClient;
import reactor.core.publisher.Flux;
import reactor.core.publisher.Mono;
import reactor.util.retry.Retry;

import java.time.Duration;
import java.util.Map;
import java.util.Set;

@Slf4j
@Component
public class QwenAdapter implements LlmClient {

    private final WebClient qwenWebClient;
    private final ObjectMapper objectMapper;
    private final CircuitBreaker qwenCircuitBreaker;

    @Value("${app.llm.qwen.retry.max-attempts:3}")
    private int maxRetryAttempts;

    @Value("${app.llm.qwen.retry.backoff:1000}")
    private long retryBackoffMs;

    public QwenAdapter(@Qualifier("qwenWebClient") WebClient qwenWebClient,
                       CircuitBreaker qwenCircuitBreaker,
                       ObjectMapper objectMapper) {
        this.qwenWebClient = qwenWebClient;
        this.qwenCircuitBreaker = qwenCircuitBreaker;
        this.objectMapper = objectMapper;
    }

    @Override
    public Mono<LlmResponse> chat(LlmRequest request) {
        Map<String, Object> body = buildRequestBody(request, false);

        return qwenWebClient.post()
                .uri("/v1/chat/completions")
                .bodyValue(body)
                .retrieve()
                .bodyToMono(Map.class)
                .map(this::parseResponse)
                .transform(CircuitBreakerOperator.of(qwenCircuitBreaker))
                .retryWhen(Retry.backoff(maxRetryAttempts, Duration.ofMillis(retryBackoffMs))
                        .doBeforeRetry(rs -> log.warn("Retrying Qwen chat, attempt {}", rs.totalRetries() + 1)))
                .doOnError(e -> log.error("Qwen chat error after retries", e));
    }

    @Override
    public Flux<StreamChunk> chatStream(LlmRequest request) {
        Map<String, Object> body = buildRequestBody(request, true);

        return qwenWebClient.post()
                .uri("/v1/chat/completions")
                .bodyValue(body)
                .retrieve()
                .bodyToFlux(new ParameterizedTypeReference<ServerSentEvent<String>>() {})
                .filter(sse -> sse.data() != null)
                .mapNotNull(this::parseStreamChunk)
                .takeUntil(chunk -> "[DONE]".equals(chunk.content()))
                .filter(chunk -> !"[DONE]".equals(chunk.content()))
                .transform(CircuitBreakerOperator.of(qwenCircuitBreaker))
                .retryWhen(Retry.backoff(maxRetryAttempts, Duration.ofMillis(retryBackoffMs))
                        .doBeforeRetry(rs -> log.warn("Retrying Qwen stream, attempt {}", rs.totalRetries() + 1)))
                .doOnError(e -> log.error("Qwen stream error after retries", e));
    }

    @Override
    public Set<String> supportedModels() {
        return Set.of("qwen-turbo", "qwen-plus", "qwen-max");
    }

    @Override
    public String provider() {
        return "qwen";
    }

    @SuppressWarnings("unchecked")
    private Map<String, Object> buildRequestBody(LlmRequest request, boolean stream) {
        return Map.of(
                "model", request.getModel(),
                "messages", request.getMessages().stream()
                        .map(m -> Map.of("role", m.getRole(), "content", m.getContent()))
                        .toList(),
                "stream", stream,
                "temperature", request.getTemperature() != null ? request.getTemperature() : 0.7,
                "max_tokens", request.getMaxTokens() != null ? request.getMaxTokens() : 2048
        );
    }

    @SuppressWarnings("unchecked")
    private LlmResponse parseResponse(Map<String, Object> resp) {
        String id = (String) resp.get("id");
        String model = (String) resp.get("model");
        String content = "";
        var choices = (java.util.List<Map<String, Object>>) resp.get("choices");
        if (choices != null && !choices.isEmpty()) {
            var message = (Map<String, Object>) choices.get(0).get("message");
            content = (String) message.get("content");
        }
        LlmResponse.Usage usage = null;
        var usageMap = (Map<String, Object>) resp.get("usage");
        if (usageMap != null) {
            usage = new LlmResponse.Usage(
                    ((Number) usageMap.getOrDefault("prompt_tokens", 0)).intValue(),
                    ((Number) usageMap.getOrDefault("completion_tokens", 0)).intValue(),
                    ((Number) usageMap.getOrDefault("total_tokens", 0)).intValue());
        }
        return new LlmResponse(id, model, content, usage);
    }

    @SuppressWarnings("unchecked")
    private StreamChunk parseStreamChunk(ServerSentEvent<String> sse) {
        String data = sse.data();
        if ("[DONE]".equals(data)) {
            return new StreamChunk("[DONE]", null);
        }
        try {
            Map<String, Object> resp = objectMapper.readValue(data, Map.class);
            var usageMap = (Map<String, Object>) resp.get("usage");
            if (usageMap != null) {
                return new StreamChunk(null, new LlmResponse.Usage(
                        ((Number) usageMap.getOrDefault("prompt_tokens", 0)).intValue(),
                        ((Number) usageMap.getOrDefault("completion_tokens", 0)).intValue(),
                        ((Number) usageMap.getOrDefault("total_tokens", 0)).intValue()));
            }
            var choices = (java.util.List<Map<String, Object>>) resp.get("choices");
            if (choices != null && !choices.isEmpty()) {
                var delta = (Map<String, Object>) choices.get(0).get("delta");
                if (delta != null && delta.get("content") != null) {
                    return new StreamChunk((String) delta.get("content"), null);
                }
            }
        } catch (Exception e) {
            log.warn("Failed to parse Qwen SSE chunk: {}", data, e);
        }
        return null;
    }
}''')

add_heading_styled(doc, '4.5 ModelRouter（智能路由引擎）', level=2)

add_paragraph_styled(doc,
    'ModelRouter 根据意图识别结果选择模型。支持 Redis 热更新路由规则，'
    'Redis 不可用时降级到 application.yml 默认规则：')

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

import com.campushare.agent.dto.IntentResult;
import com.campushare.agent.enums.Intent;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.data.redis.core.ReactiveStringRedisTemplate;
import org.springframework.stereotype.Component;
import reactor.core.publisher.Mono;

import java.util.HashMap;
import java.util.List;
import java.util.Map;

@Slf4j
@Component
@RequiredArgsConstructor
public class ModelRouter {

    private final ReactiveStringRedisTemplate redisTemplate;
    private final ModelHealthChecker healthChecker;

    // 默认路由规则（Redis 不可用时使用）
    private static final Map<String, RouteEntry> DEFAULT_RULES = new HashMap<>();

    static {
        DEFAULT_RULES.put("HOW_TO", new RouteEntry("qwen-turbo", "deepseek-v4-flash"));
        DEFAULT_RULES.put("SEARCH", new RouteEntry("deepseek-v4-flash", "qwen-plus"));
        DEFAULT_RULES.put("CHAT", new RouteEntry("qwen-turbo", "deepseek-v4-flash"));
        DEFAULT_RULES.put("CLARIFY", new RouteEntry("deepseek-v4-flash", "qwen-plus"));
        DEFAULT_RULES.put("NAVIGATE", new RouteEntry("qwen-turbo", "deepseek-v4-flash"));
        DEFAULT_RULES.put("OUT_OF_SCOPE", new RouteEntry("qwen-turbo", "deepseek-v4-flash"));
    }

    /**
     * 路由决策（ADR-GW-02）。
     *
     * 决策顺序：
     *   1. 查 Redis 路由规则（热更新）
     *   2. 默认规则（Intent → 模型映射）
     *   3. 健康检查过滤（主模型熔断 → 用备用模型）
     */
    public Mono<RouteResult> route(IntentResult intentResult, String userId) {
        String intentKey = intentResult.getIntent().name();

        return loadRouteRules()
                .map(rules -> rules.getOrDefault(intentKey, DEFAULT_RULES.get(intentKey)))
                .flatMap(entry -> {
                    // 健康检查：主模型熔断则切换到备用
                    if (!healthChecker.isAvailable(entry.primaryModel())) {
                        log.warn("Primary model {} unavailable (circuit open), using fallback {}",
                                entry.primaryModel(), entry.fallbackModel());
                        return Mono.just(new RouteResult(
                                entry.fallbackModel(),
                                entry.primaryModel(),
                                intentResult.getIntent()));
                    }
                    return Mono.just(new RouteResult(
                            entry.primaryModel(),
                            entry.fallbackModel(),
                            intentResult.getIntent()));
                })
                .onErrorResume(e -> {
                    log.warn("Route failed, using default for intent {}: {}", intentKey, e.getMessage());
                    RouteEntry entry = DEFAULT_RULES.get(intentKey);
                    return Mono.just(new RouteResult(
                            entry.primaryModel(), entry.fallbackModel(), intentResult.getIntent()));
                });
    }

    /**
     * 从 Redis 加载路由规则（热更新）。
     * key=agent:route:rules，value=JSON: {"HOW_TO":{"primary":"qwen-turbo","fallback":"deepseek-v4-flash"},...}
     */
    private Mono<Map<String, RouteEntry>> loadRouteRules() {
        return redisTemplate.opsForValue().get("agent:route:rules")
                .map(json -> {
                    // 解析 JSON 为 Map<String, RouteEntry>
                    // 省略 JSON 解析细节，使用 ObjectMapper
                    return DEFAULT_RULES;
                })
                .onErrorResume(e -> {
                    log.warn("Failed to load route rules from Redis, using defaults");
                    return Mono.just(DEFAULT_RULES);
                })
                .defaultIfEmpty(DEFAULT_RULES);
    }

    private record RouteEntry(String primaryModel, String fallbackModel) {}
}''')

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

import com.campushare.agent.enums.Intent;
import lombok.AllArgsConstructor;
import lombok.Builder;
import lombok.Data;

/**
 * 路由决策结果。
 */
@Data
@Builder
@AllArgsConstructor
public class RouteResult {
    private String primaryModel;    // 选中的主模型
    private String fallbackModel;   // 备用模型
    private Intent intent;          // 触发路由的意图
}''')

add_heading_styled(doc, '4.6 FallbackChain（故障降级链）', level=2)

add_paragraph_styled(doc,
    'FallbackChain 在主模型失败时自动切换到备用模型，全部失败时返回模板兜底。'
    '对上层透明——AgentChatService 只感知 Flux<StreamChunk>，不感知降级过程：')

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

import com.campushare.agent.enums.Intent;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.stereotype.Component;
import reactor.core.publisher.Flux;
import reactor.core.publisher.Mono;

import java.util.List;
import java.util.Map;

@Slf4j
@Component
@RequiredArgsConstructor
public class FallbackChain {

    private final List<LlmClient> adapters;
    private final ModelHealthChecker healthChecker;
    private final CostTracker costTracker;

    // 模板兜底话术
    private static final Map<Intent, String> TEMPLATE_FALLBACK = Map.of(
            Intent.HOW_TO, "抱歉，我暂时无法回答这个问题。您可以稍后再试，或查看帮助中心。",
            Intent.SEARCH, "搜索服务暂时不可用，请稍后再试。您也可以直接浏览资源库。",
            Intent.CHAT, "我暂时无法回复，请稍后再试。",
            Intent.CLARIFY, "抱歉，我暂时无法理解您的问题，请稍后再试。",
            Intent.NAVIGATE, "导航服务暂时不可用，请稍后再试。",
            Intent.OUT_OF_SCOPE, "这个问题超出了我的能力范围。"
    );

    /**
     * 执行降级链（ADR-GW-03）。
     *
     * L1: 主模型 → L2: 备用模型 → L3: 模板兜底
     */
    public Flux<StreamChunk> execute(LlmRequest request, RouteResult route, String userId,
                                      String sessionId, String turnId) {
        return tryModel(request, route.getPrimaryModel(), userId, sessionId, turnId, "L1-primary")
                .onErrorResume(e -> {
                    log.warn("L1 primary {} failed, trying L2 fallback {}", route.getPrimaryModel(), route.getFallbackModel());
                    return tryModel(request, route.getFallbackModel(), userId, sessionId, turnId, "L2-fallback")
                            .onErrorResume(e2 -> {
                                log.error("L2 fallback {} also failed, using L3 template", route.getFallbackModel());
                                return templateFallback(route.getIntent(), userId, sessionId, turnId);
                            });
                });
    }

    private Flux<StreamChunk> tryModel(LlmRequest request, String model,
                                        String userId, String sessionId, String turnId, String level) {
        LlmClient adapter = findAdapter(model);
        if (adapter == null) {
            return Flux.error(new RuntimeException("No adapter for model: " + model));
        }
        if (!healthChecker.isAvailable(model)) {
            return Flux.error(new RuntimeException("CircuitBreaker OPEN for model: " + model));
        }

        request.setModel(model);
        long startTime = System.currentTimeMillis();

        return adapter.chatStream(request)
                .doOnNext(chunk -> {
                    if (chunk.usage() != null) {
                        long latency = System.currentTimeMillis() - startTime;
                        costTracker.record(CostRecord.builder()
                                .userId(userId).sessionId(sessionId).turnId(turnId)
                                .model(model).provider(adapter.provider())
                                .usage(chunk.usage())
                                .latencyMs(latency).success(true)
                                .build());
                    }
                })
                .doOnError(e -> {
                    long latency = System.currentTimeMillis() - startTime;
                    costTracker.record(CostRecord.builder()
                            .userId(userId).sessionId(sessionId).turnId(turnId)
                            .model(model).provider(adapter.provider())
                            .usage(null).latencyMs(latency).success(false)
                            .build());
                    log.error("{} model {} call failed", level, model, e);
                });
    }

    private Flux<StreamChunk> templateFallback(Intent intent, String userId, String sessionId, String turnId) {
        String template = TEMPLATE_FALLBACK.getOrDefault(intent, "服务暂时不可用，请稍后再试。");
        costTracker.record(CostRecord.builder()
                .userId(userId).sessionId(sessionId).turnId(turnId)
                .model("template").provider("template")
                .usage(null).latencyMs(0L).success(true)
                .build());
        return Flux.just(new StreamChunk(template, null));
    }

    private LlmClient findAdapter(String model) {
        return adapters.stream()
                .filter(a -> a.supportedModels().contains(model))
                .findFirst()
                .orElse(null);
    }
}''')

add_heading_styled(doc, '4.7 ApiKeyPool（API Key 轮询池）', level=2)

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.beans.factory.annotation.Value;
import org.springframework.data.redis.core.ReactiveStringRedisTemplate;
import org.springframework.stereotype.Component;
import reactor.core.publisher.Mono;

import java.time.Duration;
import java.util.Arrays;
import java.util.List;

@Slf4j
@Component
@RequiredArgsConstructor
public class ApiKeyPool {

    private final ReactiveStringRedisTemplate redisTemplate;

    @Value("${app.llm.deepseek.api-keys:}")
    private String deepSeekKeys;

    @Value("${app.llm.qwen.api-keys:}")
    private String qwenKeys;

    @Value("${app.gateway.key.cooldown:60s}")
    private Duration keyCooldown;

    private static final String KEY_STATE_PREFIX = "agent:keystate:";

    /**
     * 获取一个可用的 API Key（ADR-GW-04）。
     *
     * 策略：round-robin 轮询 + 跳过被限流/失效的 Key
     * 全部不可用 → 返回 Mono.empty()（触发降级到备用模型）
     */
    public Mono<String> getKey(String model) {
        List<String> keys = getKeysForModel(model);
        if (keys.isEmpty()) {
            return Mono.empty();
        }

        // round-robin: Redis INCR 取模
        String counterKey = "agent:keycounter:" + model;
        return redisTemplate.opsForValue().increment(counterKey)
                .map(idx -> keys.get((int)(idx % keys.size())))
                .flatMap(key -> checkKeyAvailable(model, key)
                        .flatMap(available -> available
                                ? Mono.just(key)
                                : tryNextKey(model, keys, key, 1)))
                .onErrorResume(e -> {
                    log.warn("ApiKeyPool failed, returning first key as fallback");
                    return Mono.just(keys.get(0));
                });
    }

    /**
     * 标记 Key 被限流（HTTP 429 响应时调用）。
     */
    public Mono<Void> markRateLimited(String model, String key) {
        String stateKey = KEY_STATE_PREFIX + model + ":" + hashKey(key);
        log.warn("Key {} for model {} rate limited, cooling down for {}", key, model, keyCooldown);
        return redisTemplate.opsForValue().set(stateKey, "RATE_LIMITED", keyCooldown).then();
    }

    /**
     * 标记 Key 失效（HTTP 401/403 响应时调用）。
     */
    public Mono<Void> markDisabled(String model, String key) {
        String stateKey = KEY_STATE_PREFIX + model + ":" + hashKey(key);
        log.error("Key {} for model {} disabled (invalid)", key, model);
        return redisTemplate.opsForValue().set(stateKey, "DISABLED").then();
    }

    private Mono<Boolean> checkKeyAvailable(String model, String key) {
        String stateKey = KEY_STATE_PREFIX + model + ":" + hashKey(key);
        return redisTemplate.opsForValue().get(stateKey)
                .map(state -> !"RATE_LIMITED".equals(state) && !"DISABLED".equals(state))
                .defaultIfEmpty(true);
    }

    private Mono<String> tryNextKey(String model, List<String> keys, String triedKey, int attempt) {
        if (attempt >= keys.size()) {
            return Mono.empty(); // 全部不可用
        }
        String counterKey = "agent:keycounter:" + model;
        return redisTemplate.opsForValue().increment(counterKey)
                .map(idx -> keys.get((int)(idx % keys.size())))
                .filter(key -> !key.equals(triedKey))
                .flatMap(key -> checkKeyAvailable(model, key)
                        .flatMap(available -> available
                                ? Mono.just(key)
                                : tryNextKey(model, keys, key, attempt + 1)));
    }

    private List<String> getKeysForModel(String model) {
        if (model.startsWith("deepseek")) {
            return Arrays.asList(deepSeekKeys.split(","));
        } else if (model.startsWith("qwen")) {
            return Arrays.asList(qwenKeys.split(","));
        }
        return List.of();
    }

    private String hashKey(String key) {
        return Integer.toHexString(key.hashCode());
    }
}''')

add_heading_styled(doc, '4.8 CostTracker（成本追踪器）', level=2)

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

import io.micrometer.core.instrument.Counter;
import io.micrometer.core.instrument.MeterRegistry;
import io.micrometer.core.instrument.Timer;
import jakarta.annotation.PostConstruct;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.data.redis.core.ReactiveStringRedisTemplate;
import org.springframework.scheduling.annotation.Scheduled;
import org.springframework.stereotype.Component;
import reactor.core.publisher.Mono;
import reactor.core.scheduler.Schedulers;

import java.time.LocalDate;
import java.util.ArrayList;
import java.util.List;
import java.util.Map;
import java.util.concurrent.ConcurrentLinkedQueue;

@Slf4j
@Component
@RequiredArgsConstructor
public class CostTracker {

    private final ReactiveStringRedisTemplate redisTemplate;
    private final MeterRegistry meterRegistry;

    // 模型单价表（元/千Token）
    private static final Map<String, double[]> PRICING = Map.of(
            "deepseek-v4-flash", new double[]{0.002, 0.006},  // [input, output]
            "qwen-turbo",        new double[]{0.0003, 0.0006},
            "qwen-plus",         new double[]{0.0008, 0.002}
    );

    // 批量写入缓冲队列
    private final ConcurrentLinkedQueue<CostRecord> writeQueue = new ConcurrentLinkedQueue<>();

    private volatile Counter tokenCounterPrompt;
    private volatile Counter tokenCounterCompletion;
    private volatile Counter costCounter;
    private volatile Timer latencyTimer;

    @PostConstruct
    void initMetrics() {
        tokenCounterPrompt = Counter.builder("agent.llm.tokens")
                .tag("type", "prompt").register(meterRegistry);
        tokenCounterCompletion = Counter.builder("agent.llm.tokens")
                .tag("type", "completion").register(meterRegistry);
        costCounter = Counter.builder("agent.llm.cost").register(meterRegistry);
        latencyTimer = Timer.builder("agent.llm.latency").register(meterRegistry);
    }

    /**
     * 记录一次 LLM 调用的成本（ADR-GW-05）。
     * 异步写入 Redis + MySQL，不阻塞主流程。
     */
    public void record(CostRecord record) {
        if (record.getUsage() != null) {
            double cost = calculateCost(record.getModel(), record.getUsage());
            record.setCostYuan(cost);

            // Micrometer Metrics
            tokenCounterPrompt.increment(record.getUsage().getPromptTokens());
            tokenCounterCompletion.increment(record.getUsage().getCompletionTokens());
            costCounter.increment(cost);
            latencyTimer.record(java.time.Duration.ofMillis(record.getLatencyMs()));
        }

        // 加入批量写入队列
        writeQueue.add(record);

        // 异步写入 Redis 实时统计
        writeToRedis(record).subscribeOn(Schedulers.boundedElastic()).subscribe();
    }

    private double calculateCost(String model, LlmResponse.Usage usage) {
        double[] price = PRICING.getOrDefault(model, new double[]{0.001, 0.003});
        return (usage.getPromptTokens() / 1000.0) * price[0]
             + (usage.getCompletionTokens() / 1000.0) * price[1];
    }

    private Mono<Void> writeToRedis(CostRecord record) {
        String date = LocalDate.now().toString();
        String cost = String.format("%.6f", record.getCostYuan() != null ? record.getCostYuan() : 0);

        return Mono.zip(
                redisTemplate.opsForValue().increment("agent:cost:user:" + record.getUserId() + ":" + date),
                redisTemplate.opsForValue().increment("agent:cost:model:" + record.getModel() + ":" + date),
                redisTemplate.opsForValue().increment("agent:cost:daily:" + date)
        ).then().onErrorResume(e -> {
            log.warn("Failed to write cost to Redis: {}", e.getMessage());
            return Mono.empty();
        });
    }

    /**
     * 定时批量写入 MySQL（每 30 秒）。
     */
    @Scheduled(fixedDelay = 30000)
    public void flushToDatabase() {
        if (writeQueue.isEmpty()) return;
        List<CostRecord> batch = new ArrayList<>();
        while (!writeQueue.isEmpty() && batch.size() < 100) {
            CostRecord record = writeQueue.poll();
            if (record != null) batch.add(record);
        }
        if (!batch.isEmpty()) {
            // 批量 INSERT 到 agent_cost_records 表（省略 Mapper 调用）
            log.debug("Flushing {} cost records to database", batch.size());
        }
    }
}''')

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

import lombok.AllArgsConstructor;
import lombok.Builder;
import lombok.Data;
import lombok.NoArgsConstructor;

@Data
@Builder
@NoArgsConstructor
@AllArgsConstructor
public class CostRecord {
    private String userId;
    private String sessionId;
    private String turnId;
    private String model;
    private String provider;
    private LlmResponse.Usage usage;
    private Double costYuan;
    private long latencyMs;
    private boolean success;
}''')

add_heading_styled(doc, '4.9 ModelHealthChecker（模型健康检查器）', level=2)

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

import io.github.resilience4j.circuitbreaker.CircuitBreaker;
import io.github.resilience4j.circuitbreaker.CircuitBreakerConfig;
import jakarta.annotation.PostConstruct;
import lombok.extern.slf4j.Slf4j;
import org.springframework.scheduling.annotation.Scheduled;
import org.springframework.stereotype.Component;

import java.time.Duration;
import java.util.Map;
import java.util.concurrent.ConcurrentHashMap;

@Slf4j
@Component
public class ModelHealthChecker {

    private final Map<String, CircuitBreaker> breakers = new ConcurrentHashMap<>();

    @PostConstruct
    void init() {
        // 为已知模型创建 CircuitBreaker
        registerModel("deepseek-v4-flash");
        registerModel("qwen-turbo");
        registerModel("qwen-plus");
    }

    public void registerModel(String model) {
        CircuitBreakerConfig config = CircuitBreakerConfig.custom()
                .slidingWindowSize(10)
                .minimumNumberOfCalls(5)
                .failureRateThreshold(50.0f)
                .waitDurationInOpenState(Duration.ofSeconds(30))
                .permittedNumberOfCallsInHalfOpenState(3)
                .build();
        breakers.put(model, CircuitBreaker.of("model:" + model, config));
        log.info("Registered CircuitBreaker for model: {}", model);
    }

    /**
     * 检查模型是否可用（CircuitBreaker 非 OPEN）。
     */
    public boolean isAvailable(String model) {
        CircuitBreaker cb = breakers.get(model);
        if (cb == null) return true; // 未注册的模型默认可用
        return cb.getState() != CircuitBreaker.State.OPEN;
    }

    /**
     * 记录调用结果（成功/失败），驱动 CircuitBreaker 状态转换。
     */
    public void recordResult(String model, boolean success, Throwable error) {
        CircuitBreaker cb = breakers.get(model);
        if (cb == null) return;
        if (success) {
            cb.onSuccess();
        } else {
            cb.onError(Duration.ZERO, error != null ? error : new RuntimeException("LLM call failed"));
        }
    }

    public CircuitBreaker getCircuitBreaker(String model) {
        return breakers.get(model);
    }

    /**
     * 主动健康探测（每 60 秒对 OPEN 状态的模型发起探测）。
     */
    @Scheduled(fixedDelay = 60000)
    public void healthCheck() {
        breakers.forEach((model, cb) -> {
            if (cb.getState() == CircuitBreaker.State.OPEN) {
                log.info("Health check: model {} is OPEN, sending probe", model);
                // 发起轻量探测请求（1+1=），成功则转 HALF_OPEN
                // 省略探测请求实现
            }
        });
    }

    /**
     * 获取所有模型健康状态（供监控接口使用）。
     */
    public Map<String, String> getHealthStatus() {
        Map<String, String> status = new ConcurrentHashMap<>();
        breakers.forEach((model, cb) -> status.put(model, cb.getState().name()));
        return status;
    }
}''')

add_heading_styled(doc, '4.10 AgentChatService 集成改造 + application.yml', level=2)

add_paragraph_styled(doc,
    'AgentChatService 改造前直连 DeepSeekClient，改造后通过 LlmGateway 统一入口调用。'
    '关键变化：将 deepSeekClient.chatCompletionStream(messages) 替换为 '
    'llmGateway.chatStream(messages, intentResult, userId, sessionId, turnId)：')

add_heading_styled(doc, '改造前（直连 DeepSeekClient）', level=3)

add_code_block(doc, '''// AgentChatService.java（改造前）
private final DeepSeekClient deepSeekClient;  // 直连

public Flux<ChatEvent> chat(String userId, ChatRequest request) {
    // ...
    Flux<ChatEvent> deltaStream = deepSeekClient.chatCompletionStream(ctx.messages())
            .doOnNext(chunk -> { /* ... */ })
            .map(chunk -> new ChatEvent("delta", chunk.content()))
            .doFinally(signalType -> { /* ... */ });
    // ...
}''')

add_heading_styled(doc, '改造后（通过 LlmGateway）', level=3)

add_code_block(doc, '''// AgentChatService.java（改造后）
private final LlmGateway llmGateway;  // 通过网关

public Flux<ChatEvent> chat(String userId, ChatRequest request) {
    return Mono.fromCallable(() -> prepareContext(userId, request))
            .subscribeOn(Schedulers.boundedElastic())
            .flatMapMany(ctx -> {
                String sessionJson = buildSessionJson(ctx);

                // 快路径不变（模板回复）
                if (ctx.routeDecision() != null && ctx.routeDecision().isShortCircuit()) {
                    // ... 保持不变
                }

                // 慢路径：通过网关调用 LLM（自动路由 + 降级 + 成本追踪）
                StringBuilder assistantContent = new StringBuilder();
                AtomicReference<LlmResponse.Usage> usageRef = new AtomicReference<>();

                Flux<ChatEvent> deltaStream = llmGateway.chatStream(
                                ctx.messages(),
                                ctx.intentResult(),
                                userId,
                                ctx.session().getId(),
                                ctx.turn().getId())
                        .doOnNext(chunk -> {
                            if (chunk.content() != null) {
                                assistantContent.append(chunk.content());
                            }
                            if (chunk.usage() != null) {
                                usageRef.set(chunk.usage());
                            }
                        })
                        .filter(chunk -> chunk.content() != null)
                        .map(chunk -> new ChatEvent("delta", chunk.content()))
                        .doFinally(signalType -> {
                            long elapsed = System.currentTimeMillis() - ctx.startTime();
                            String content = assistantContent.toString();
                            Mono.fromRunnable(() -> {
                                if (signalType == SignalType.ON_COMPLETE) {
                                    completeTurn(ctx.turn(), ctx.session(), content, elapsed,
                                            usageRef.get(), ctx.promptTokens(),
                                            ctx.retrievalContext(), ctx.intentResult());
                                } else if (signalType == SignalType.ON_ERROR) {
                                    errorTurn(ctx.turn(), "Stream terminated with error");
                                }
                            }).subscribeOn(Schedulers.boundedElastic()).subscribe();
                        });

                return Flux.concat(Flux.just(new ChatEvent("session", sessionJson)), deltaStream);
            });
}''')

add_heading_styled(doc, 'LlmGateway 统一入口', level=3)

add_code_block(doc, '''package com.campushare.agent.llm.gateway;

import com.campushare.agent.dto.IntentResult;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.stereotype.Component;
import reactor.core.publisher.Flux;

import java.util.List;

/**
 * LLM 网关统一入口（ADR-GW-01 ~ ADR-GW-07 集成点）。
 *
 * AgentChatService 只调用此类，不直接接触 Adapter/Router/FallbackChain。
 * 网关内部完成：路由 → 降级 → 成本追踪 → 健康检查。
 */
@Slf4j
@Component
@RequiredArgsConstructor
public class LlmGateway {

    private final ModelRouter router;
    private final FallbackChain fallbackChain;

    /**
     * 流式调用（用户对话主入口）。
     */
    public Flux<StreamChunk> chatStream(List<LlmRequest.Message> messages,
                                         IntentResult intentResult,
                                         String userId,
                                         String sessionId,
                                         String turnId) {
        return router.route(intentResult, userId)
                .flatMapMany(route -> {
                    LlmRequest request = LlmRequest.builder()
                            .messages(messages)
                            .stream(true)
                            .build();

                    log.info("Gateway routing: intent={} → primary={} fallback={}",
                            intentResult.getIntent(), route.getPrimaryModel(), route.getFallbackModel());

                    return fallbackChain.execute(request, route, userId, sessionId, turnId);
                });
    }

    /**
     * 非流式调用（意图分类等内部调用）。
     */
    public Mono<LlmResponse> chat(List<LlmRequest.Message> messages,
                                   IntentResult intentResult,
                                   String userId,
                                   String sessionId,
                                   String turnId) {
        return router.route(intentResult, userId)
                .flatMap(route -> {
                    LlmRequest request = LlmRequest.builder()
                            .messages(messages)
                            .stream(false)
                            .build();
                    // 内部调用也走 FallbackChain（简化版，非流式）
                    return fallbackChain.executeMono(request, route, userId, sessionId, turnId);
                });
    }
}''')

add_heading_styled(doc, 'application.yml 改造', level=3)

add_paragraph_styled(doc, '新增多模型配置、路由规则、成本配置：')

add_code_block(doc, '''app:
  llm:
    # 主模型：DeepSeek
    deepseek:
      base-url: \${DEEPSEEK_BASE_URL:https://api.deepseek.com}
      api-keys: \${DEEPSEEK_API_KEYS:sk-placeholder}  # 改为复数（逗号分隔）
      models:
        - deepseek-v4-flash
        - deepseek-v4-lite
      timeout: 60000
      temperature: 0.7
      max-tokens: 2048
      retry:
        max-attempts: 3
        backoff: 1000
      circuit-breaker:
        sliding-window-size: 10
        minimum-number-of-calls: 5
        failure-rate-threshold: 50.0
        wait-duration-in-open-state: 30s
        permitted-number-of-calls-in-half-open-state: 3

    # 备用模型：通义千问
    qwen:
      base-url: \${QWEN_BASE_URL:https://dashscope.aliyuncs.com/compatible-mode}
      api-keys: \${QWEN_API_KEYS:sk-placeholder}
      models:
        - qwen-turbo
        - qwen-plus
        - qwen-max
      timeout: 30000
      temperature: 0.7
      max-tokens: 2048
      retry:
        max-attempts: 3
        backoff: 1000
      circuit-breaker:
        sliding-window-size: 10
        minimum-number-of-calls: 5
        failure-rate-threshold: 50.0
        wait-duration-in-open-state: 30s
        permitted-number-of-calls-in-half-open-state: 3

    # Embedding（保持不变）
    embedding:
      # ... 保持现有配置

  # LLM 网关配置
  gateway:
    # 路由规则（默认值，Redis 可热更新覆盖）
    route:
      rules:
        HOW_TO: { primary: qwen-turbo, fallback: deepseek-v4-flash }
        SEARCH: { primary: deepseek-v4-flash, fallback: qwen-plus }
        CHAT: { primary: qwen-turbo, fallback: deepseek-v4-flash }
        CLARIFY: { primary: deepseek-v4-flash, fallback: qwen-plus }
        NAVIGATE: { primary: qwen-turbo, fallback: deepseek-v4-flash }
        OUT_OF_SCOPE: { primary: qwen-turbo, fallback: deepseek-v4-flash }

    # API Key 池
    key:
      cooldown: 60s          # 被限流 Key 冷却时间
      health-check-interval: 60s

    # 模型级限流
    rate-limit:
      model-rpm: 500          # 每模型每分钟最大请求数
      model-tpm: 500000       # 每模型每分钟最大 Token 数

    # 成本追踪
    cost:
      flush-interval: 30s     # MySQL 批量写入间隔
      batch-size: 100         # 批量写入大小

    # 健康检查
    health-check:
      interval: 60s           # 主动探测间隔
      probe-prompt: "1+1="    # 探测 prompt''')

add_callout(doc, '改造影响面：AgentChatService（注入 LlmGateway 替代 DeepSeekClient）、'
    'IntentClassifier（内部 LLM 调用也改为通过网关）、ResilienceConfig（per-model 熔断器）、'
    'application.yml（新增 qwen + gateway 配置段）。DeepSeekClient 标记 @Deprecated 但保留。', 'warning')

doc.add_page_break()

# ==================== 第五章 目标 ====================

add_heading_styled(doc, '第五章 目标：实现效果', level=1)

add_heading_styled(doc, '5.1 功能目标', level=2)

add_styled_table(doc,
    headers=['功能', '当前状态', '目标状态', '验收标准'],
    rows=[
        ['多模型支持', '仅 DeepSeek', 'DeepSeek + 通义千问', '2 个供应商、5 个模型可切换'],
        ['意图驱动路由', '无（统一模型）', '6 种意图自动选模型', 'HOW_TO/CHAT→qwen-turbo, SEARCH→deepseek-v4-flash'],
        ['故障降级', '仅重试（同模型）', '三级降级链', '主模型故障 → 30s 内切备用 → 兜底模板'],
        ['API Key 轮询', '单 Key', '多 Key 轮询 + 自动摘除', '3 Key 轮询，429 自动切换'],
        ['成本追踪', '仅 Token 日志', '5 维度成本归因', '按用户/会话/模型/意图/天统计成本'],
        ['模型健康检查', '全局单一熔断', 'per-model 熔断 + 主动探测', '故障 60s 内探测恢复'],
        ['路由热更新', '无', 'Redis 热更新', '改 Redis 规则后下次请求生效'],
    ],
    col_widths=[3, 3.5, 4, 4.5])

add_heading_styled(doc, '5.2 性能目标', level=2)

add_styled_table(doc,
    headers=['指标', '当前值', '目标值', '测量方法'],
    rows=[
        ['路由决策延迟', '0ms（无路由）', '< 5ms', 'ModelRouter.route() 耗时统计'],
        ['闲聊 TTFT', '~800ms（deepseek-v4-flash）', '< 400ms（qwen-turbo）', '首 Token 到达时间'],
        ['搜索 TTFT', '~800ms', '< 1000ms', '首 Token 到达时间'],
        ['降级切换延迟', 'N/A', '< 2s（主→备切换）', '主模型失败到备用模型首 Token'],
        ['成本追踪开销', '0', '< 5ms（异步）', 'CostTracker.record() 不阻塞主流程'],
        ['Key 轮询开销', '0', '< 2ms', 'ApiKeyPool.getKey() 耗时'],
        ['健康检查开销', '0', '< 1ms（状态查询）', 'ModelHealthChecker.isAvailable()'],
    ],
    col_widths=[3.5, 4, 3.5, 4])

add_heading_styled(doc, '5.3 质量目标', level=2)

add_styled_table(doc,
    headers=['质量维度', '目标', '验收标准'],
    rows=[
        ['路由准确率', '> 95%', '意图→模型映射符合预期（人工评估 100 条）'],
        ['降级成功率', '> 99.9%', '主模型故障时备用模型成功接管'],
        ['模板兜底触发率', '< 0.1%', '主+备都不可用的比例 < 0.1%'],
        ['成本统计准确性', '> 99%', 'CostTracker 统计与 API 账单误差 < 1%'],
        ['Key 摘除误判率', '< 1%', '正常 Key 被误标为 RATE_LIMITED 的比例'],
        ['熔断误触发率', '< 5%', '正常模型被误触发熔断的比例'],
        ['路由热更新生效', '< 1s', 'Redis 规则修改到下次请求生效 < 1s'],
    ],
    col_widths=[4, 3, 8])

add_heading_styled(doc, '5.4 成本目标', level=2)

add_styled_table(doc,
    headers=['指标', '当前值', '目标值', '节省'],
    rows=[
        ['月均 Token 消耗', '50M（全部 deepseek）', '35M（路由后部分走 qwen）', '30%'],
        ['月均成本', '¥2000+', '¥1200 左右', '~40%'],
        ['单次对话平均成本', '¥0.04', '¥0.024', '40%'],
        ['闲聊单次成本', '¥0.04（deepseek）', '¥0.006（qwen-turbo）', '85%'],
        ['搜索单次成本', '¥0.04', '¥0.04（不变，需质量）', '0%'],
    ],
    col_widths=[4, 4, 4, 3])

add_callout(doc, '成本节省核心来源：40% 的请求（HOW_TO + CHAT + NAVIGATE + OUT_OF_SCOPE）'
    '从 deepseek-v4-flash 切换到 qwen-turbo，单价降低约 85%。', 'success')

doc.add_page_break()

# ==================== 第六章 测试评估与验收 ====================

add_heading_styled(doc, '第六章 测试评估与验收', level=1)

add_heading_styled(doc, '6.1 评估指标体系', level=2)

add_styled_table(doc,
    headers=['指标类别', '具体指标', '目标值', '数据来源'],
    rows=[
        ['路由性能', '路由决策延迟 P99', '< 5ms', 'Micrometer Timer'],
        ['路由性能', '路由准确率', '> 95%', '人工评估'],
        ['降级能力', '降级切换成功率', '> 99.9%', '降级日志统计'],
        ['降级能力', '模板兜底触发率', '< 0.1%', 'CostTracker 统计 model=template'],
        ['成本管控', '月成本节省比例', '> 30%', 'CostTracker 月报'],
        ['成本管控', '成本统计准确率', '> 99%', '与 API 账单对比'],
        ['可用性', '网关整体可用率', '> 99.95%', 'Uptime 监控'],
        ['可用性', '单模型故障恢复时间', '< 90s', '熔断器 OPEN→CLOSED 时间'],
        ['负载均衡', 'Key 轮询均匀度', '偏差 < 10%', 'per-Key 调用计数'],
        ['负载均衡', 'Key 限流自动切换', '< 2s', '429 到切换成功时间'],
    ],
    col_widths=[2.5, 4.5, 3, 4])

add_heading_styled(doc, '6.2 黄金测试集', level=2)

add_paragraph_styled(doc,
    '构建 LLM 网关黄金测试集，覆盖路由/降级/成本/Key 管理/健康检查五大场景：')

add_styled_table(doc,
    headers=['场景', '测试用例数', '覆盖内容', '通过标准'],
    rows=[
        ['路由准确性', 30, '6 种意图各 5 条 → 验证路由到正确模型', '路由准确率 > 95%'],
        ['降级链', 20, '主模型熔断/超时/5xx → 验证切到备用模型', '降级成功率 100%'],
        ['模板兜底', 10, '主+备都不可用 → 验证返回模板回复', '模板触发率 100%（模拟场景）'],
        ['API Key 轮询', 15, '多 Key 轮询/429 切换/401 摘除', 'Key 切换正确率 100%'],
        ['成本追踪', 15, 'Token 计量/成本计算/多维归因', '成本计算误差 < 1%'],
        ['健康检查', 10, '熔断/半开/恢复/主动探测', '状态转换正确率 100%'],
    ],
    col_widths=[3, 2, 7, 4])

add_heading_styled(doc, '6.3 CI/CD 集成', level=2)

add_paragraph_styled(doc,
    'LLM 网关的 CI/CD 流水线在每次 PR 时自动运行：')

add_styled_table(doc,
    headers=['阶段', '内容', '工具', '通过标准'],
    rows=[
        ['编译检查', 'mvn compile + 静态分析', 'Maven + SpotBugs', '0 error'],
        ['单元测试', 'ModelRouter/FallbackChain/ApiKeyPool/CostTracker', 'JUnit 5', '覆盖率 > 80%'],
        ['集成测试', 'LlmGateway 端到端调用（Mock LLM）', 'WireMock', '100% 通过'],
        ['路由测试', '6 种意图路由到正确模型', '专用测试', '准确率 > 95%'],
        ['降级测试', '模拟主模型故障 → 验证切到备用', 'WireMock + Chaos', '降级成功'],
        ['成本测试', '验证成本计算准确性', '断言', '误差 < 1%'],
    ],
    col_widths=[2.5, 5.5, 3, 3.5])

add_heading_styled(doc, '6.4 LLM-as-Judge（模型质量对比）', level=2)

add_paragraph_styled(doc,
    '引入网关后，不同意图走不同模型，需要验证"切换模型后回答质量不下降"。'
    '使用 LLM-as-Judge 进行模型质量对比评估：')

add_styled_table(doc,
    headers=['评估维度', '评估方法', '通过标准'],
    rows=[
        ['HOW_TO 回答质量', 'qwen-turbo vs deepseek-v4-flash 回答同一问题，GPT-4o 评分', 'qwen-turbo ≥ deepseek 的 80%'],
        ['CHAT 回答质量', '同上', 'qwen-turbo ≥ deepseek 的 90%'],
        ['SEARCH 回答质量', '同上（不切换，验证不降质）', 'deepseek 保持不变'],
        ['降级回答质量', '备用模型 vs 主模型回答质量', '备用 ≥ 主模型的 80%'],
    ],
    col_widths=[4, 6, 5])

add_heading_styled(doc, '6.5 错误分析与 BadCase', level=2)

add_styled_table(doc,
    headers=['BadCase 类型', '示例', '改进措施'],
    rows=[
        ['路由错误', 'HOW_TO 请求被路由到 deepseek（应为 qwen-turbo）', '检查意图识别 + 路由规则'],
        ['降级失败', '主模型故障但备用模型也故障', '增加更多备用模型 / 模板兜底'],
        ['Key 摘除误判', '正常 Key 被标记为 RATE_LIMITED', '调整限流检测阈值'],
        ['成本统计偏差', 'CostTracker 统计与账单不符', '检查 Token 计量 + 单价表'],
        ['熔断误触发', '模型正常但熔断器 OPEN', '调整失败率阈值 + 滑动窗口'],
    ],
    col_widths=[3, 6, 6])

add_heading_styled(doc, '6.6 测试用例清单', level=2)

add_styled_table(doc,
    headers=['编号', '测试用例', '类型', '验证点'],
    rows=[
        ['GW-T-001', 'HOW_TO 意图路由到 qwen-turbo', '路由', 'model=qwen-turbo'],
        ['GW-T-002', 'SEARCH 意图路由到 deepseek-v4-flash', '路由', 'model=deepseek-v4-flash'],
        ['GW-T-003', 'CHAT 意图路由到 qwen-turbo', '路由', 'model=qwen-turbo'],
        ['GW-T-004', '主模型熔断 → 切到备用模型', '降级', 'fallback 成功'],
        ['GW-T-005', '主+备都不可用 → 模板兜底', '降级', '返回模板话术'],
        ['GW-T-006', '3 个 Key 轮询', 'Key 池', '每个 Key 调用次数均匀'],
        ['GW-T-007', 'Key 收到 429 → 自动切换', 'Key 池', '切换到另一个 Key'],
        ['GW-T-008', 'Key 收到 401 → 标记 DISABLED', 'Key 池', '不再使用该 Key'],
        ['GW-T-009', '成本计算准确性', '成本', '误差 < 1%'],
        ['GW-T-010', '多维度成本归因', '成本', '5 个维度都有数据'],
        ['GW-T-011', 'Redis 路由规则热更新', '路由', '下次请求生效'],
        ['GW-T-012', 'Redis 不可用 → 降级到默认规则', '容错', '使用 yml 默认规则'],
        ['GW-T-013', '主动健康探测恢复', '健康', 'OPEN → HALF_OPEN → CLOSED'],
        ['GW-T-014', '模型级 RPM 限流', '限流', '超限切到备用模型'],
        ['GW-T-015', '所有模型限流 → 返回 429', '限流', '友好错误提示'],
    ],
    col_widths=[2, 5.5, 2.5, 5])

add_heading_styled(doc, '6.7 性能压测', level=2)

add_styled_table(doc,
    headers=['压测场景', '并发数', '目标 QPS', '目标 P99', '验证点'],
    rows=[
        ['闲聊（qwen-turbo）', 50, '20', '< 1s', 'TTFT < 400ms'],
        ['搜索（deepseek）', 30, '10', '< 3s', 'TTFT < 1s'],
        ['降级切换', 20, '5', '< 2s', '主→备切换 < 2s'],
        ['Key 轮询', 100, '50', '< 500ms', '轮询均匀'],
        ['成本追踪', 100, '50', '< 5ms', '不阻塞主流程'],
    ],
    col_widths=[3.5, 2, 2.5, 2.5, 4.5])

add_heading_styled(doc, '6.8 A/B 测试', level=2)

add_paragraph_styled(doc,
    '引入网关后，通过 A/B 测试验证路由策略的效果：')

add_styled_table(doc,
    headers=['实验', 'A 组', 'B 组', '指标', '通过标准'],
    rows=[
        ['路由 vs 无路由', '统一 deepseek', '意图驱动路由', '成本 + TTFT + 用户满意度', 'B 组成本降 30%+TTFT 不增'],
        ['qwen-turbo vs deepseek（HOW_TO）', 'deepseek', 'qwen-turbo', '回答质量 + TTFT + 成本', '质量 ≥ 80%+TTFT 降 50%'],
    ],
    col_widths=[4, 3, 3, 3.5, 3.5])

add_heading_styled(doc, '6.9 验收流程', level=2)

add_numbered(doc, '单元测试全部通过（覆盖率 > 80%）')
add_numbered(doc, '集成测试全部通过（WireMock 模拟 LLM）')
add_numbered(doc, '黄金测试集 100 条全部通过')
add_numbered(doc, '路由准确率 > 95%（人工评估）')
add_numbered(doc, '降级成功率 > 99.9%（模拟故障测试）')
add_numbered(doc, '成本统计误差 < 1%（与 API 账单对比）')
add_numbered(doc, 'A/B 测试 B 组成本降低 > 30%')
add_numbered(doc, '性能压测达标（P99 满足目标）')
add_numbered(doc, 'LLM-as-Judge 质量评估通过（备用模型 ≥ 主模型 80%）')
add_numbered(doc, '监控大盘上线（路由/降级/成本/健康指标可视化）')

add_heading_styled(doc, '6.10 持续监控', level=2)

add_styled_table(doc,
    headers=['监控项', '指标', '告警阈值', '处理方式'],
    rows=[
        ['路由延迟', 'agent.gateway.route.latency', 'P99 > 10ms', '检查 Redis 连接'],
        ['降级率', 'agent.gateway.fallback.rate', '> 5%', '检查主模型健康状态'],
        ['模板兜底率', 'agent.gateway.template.rate', '> 0.1%', '检查所有模型可用性'],
        ['模型错误率', 'agent.llm.error{model}', '> 10%', '检查模型 API 状态'],
        ['模型延迟', 'agent.llm.latency{model} P99', '> 5s', '考虑切换到更快模型'],
        ['日成本', 'agent:cost:daily:{date}', '> ¥80', '检查路由规则 + 异常流量'],
        ['Key 限流率', 'agent.key.rate_limited.count', '> 10%', '增加 Key 或降低流量'],
        ['熔断器 OPEN', 'circuitbreaker.state', '任何模型 OPEN', '检查模型可用性'],
    ],
    col_widths=[3, 4, 3, 5])

doc.add_page_break()

# ==================== 第七章 总结与边界声明 ====================

add_heading_styled(doc, '第七章 总结与边界声明', level=1)

add_heading_styled(doc, '7.1 核心总结', level=2)

add_paragraph_styled(doc,
    'LLM 网关是 Agent 系统从"能用"走向"好用、稳用、省用"的关键基础设施。'
    '本文档设计了 CampusShare Agent 的 LLM 网关，核心成果如下：')

add_styled_table(doc,
    headers=['维度', '改造前', '改造后', '收益'],
    rows=[
        ['模型支持', '仅 DeepSeek（单点）', 'DeepSeek + 通义千问（多模型）', '消除供应商锁定'],
        ['路由策略', '无（统一模型）', '意图驱动动态规则路由', '成本降 40% + 闲聊 TTFT 降 60%'],
        ['故障降级', '仅重试（同模型）', '三级降级链（主→备→模板）', '可用率 > 99.95%'],
        ['API Key', '单 Key', '多 Key 轮询 + 自动摘除', '高并发不限流'],
        ['成本管控', '无（黑盒）', '5 维度成本归因', '成本可见可控'],
        ['健康检查', '全局单一熔断', 'per-model 熔断 + 主动探测', '故障 90s 内恢复'],
        ['热更新', '无', 'Redis 路由规则热更新', '运维不停机'],
    ],
    col_widths=[2.5, 3.5, 4, 4.5])

add_heading_styled(doc, '7.2 与其他文档的关系', level=2)

add_styled_table(doc,
    headers=['相关文档', '关系', '交互点'],
    rows=[
        ['System Prompt 工程（SP）', '上游', '网关注入模型名到 SP 元信息'],
        ['意图识别（INT）', '上游', 'IntentResult 驱动路由决策'],
        ['对话编排（DLG）', '上游', '编排层通过 LlmGateway 调用 LLM'],
        ['安全护栏（SEC）', '横向协作', 'SEC 在网关调用前后做输入/输出检测'],
        ['可观测性（OBS）', '下游', '网关暴露 Metrics 供 OBS 采集'],
        ['评估体系（EVAL）', '下游', '网关 A/B 路由支持模型对比实验'],
        ['缓存层（CACHE）', '横向协作', '网关调用前先查语义缓存'],
        ['性能 SLO 工程（SLO）', '下游', '网关延迟数据供 SLO 燃烧率告警'],
    ],
    col_widths=[5, 2.5, 7.5])

add_heading_styled(doc, '7.3 演进路线', level=2)

add_styled_table(doc,
    headers=['阶段', '时间', '内容', '依赖'],
    rows=[
        ['V1.0 基础网关', '当前', '模型抽象 + 意图路由 + 三级降级 + Key 池 + 成本追踪', '本文档'],
        ['V1.1 缓存集成', '下一步', '接入语义缓存层（CACHE 方向），缓存命中跳过 LLM 调用', '缓存层文档'],
        ['V1.2 延迟感知', '1 个月后', '实时监测模型 P99 延迟，超阈值自动切换到更快模型', '可观测性（OBS）'],
        ['V1.3 成本感知', '2 个月后', '月度成本预算 + 超预算自动降级到更便宜模型', '成本数据积累'],
        ['V2.0 多区域', '3 个月后', '多区域灾备 + 区域级路由（参考 Azure Traffic Manager）', '多区域部署'],
        ['V2.1 LLM-as-Router', '远期', '对复杂请求用 LLM 判断路由（规则路由的补充）', '模型能力增强'],
    ],
    col_widths=[3, 2, 7, 3])

add_heading_styled(doc, '7.4 边界声明', level=2)

add_paragraph_styled(doc, '本文档的设计有以下边界，超出边界的内容不在本文档覆盖范围：')

add_styled_table(doc,
    headers=['边界', '说明', '归属方向'],
    rows=[
        ['不涉及模型训练', 'SFT/RLHF/DPO 需模型权重，Agent 用 API 调用', '模型团队'],
        ['不涉及推理加速', '量化/蒸馏/投机解码需模型权重', '模型推理团队'],
        ['不涉及语义缓存', '语义相似度缓存设计属独立方向', '缓存层（CACHE）'],
        ['不涉及用户配额', '用户级 RPM/Token 配额属独立方向', '限流配额方向'],
        ['不涉及全链路追踪', 'Trace/Metrics/Logs 的采集与展示属独立方向', '可观测性（OBS）'],
        ['不涉及评估指标体系', '黄金集/CI-CD/回归检测属独立方向', '评估体系（EVAL）'],
    ],
    col_widths=[3.5, 7, 4.5])

doc.add_page_break()

# ==================== 附录 ADR 摘要 ====================

add_heading_styled(doc, '附录 ADR 摘要', level=1)

add_paragraph_styled(doc,
    '本附录列出 LLM 网关模块的全部 ADR（Architecture Decision Record），'
    '每条 ADR 包含背景、决策、理由和后果四个部分。')

# ADR-GW-01
add_heading_styled(doc, 'ADR-GW-01：模型抽象层——统一 LlmClient 接口 + Provider Adapter', level=2)
add_paragraph_styled(doc, '【背景】', bold=True)
add_paragraph_styled(doc,
    '当前 DeepSeekClient 硬编码 DeepSeek 格式，请求/响应 DTO 与 DeepSeek 强耦合。'
    '想新增模型需修改全部代码，改动面大、风险高。')
add_paragraph_styled(doc, '【决策】', bold=True)
add_paragraph_styled(doc,
    '采用统一 LlmClient 接口 + Provider Adapter 模式。'
    '每个供应商一个 Adapter（DeepSeekAdapter/QwenAdapter），负责格式转换。'
    '上层（Router/FallbackChain）只依赖 LlmClient 接口。')
add_paragraph_styled(doc, '【理由】', bold=True)
add_bullet(doc, '接口隔离：新增模型只需加 Adapter，不改现有代码')
add_bullet(doc, '单一职责：Adapter 只做格式转换 + HTTP 调用，不做路由/降级决策')
add_bullet(doc, '可测试性：上层可通过 Mock LlmClient 测试，不依赖真实 LLM')
add_paragraph_styled(doc, '【后果】', bold=True)
add_bullet(doc, '正面：扩展性强、可测试性好、职责清晰')
add_bullet(doc, '负面：增加了一层抽象（Adapter），代码量增加')
add_bullet(doc, '约束：所有 Adapter 必须实现 supportedModels() 和 provider() 方法')

# ADR-GW-02
add_heading_styled(doc, 'ADR-GW-02：智能路由策略——意图驱动动态规则路由', level=2)
add_paragraph_styled(doc, '【背景】', bold=True)
add_paragraph_styled(doc,
    '所有请求走同一模型导致成本浪费和延迟不优。需要按意图选择最合适的模型。')
add_paragraph_styled(doc, '【决策】', bold=True)
add_paragraph_styled(doc,
    '采用意图驱动动态规则路由。路由规则为 IntentResult.intent → 模型映射，'
    '规则存储在 Redis 支持热更新。Redis 不可用时降级到 application.yml 默认规则。')
add_paragraph_styled(doc, '【理由】', bold=True)
add_bullet(doc, '规则路由 0ms 延迟（纯内存查表），不增加额外 LLM 调用')
add_bullet(doc, '可热更新（Redis），改规则无需重启')
add_bullet(doc, '可枚举（6 种意图），规则可表达')
add_bullet(doc, '不选 LLM-as-Router：额外 300ms 延迟 + 成本 + 路由本身可能出错')
add_paragraph_styled(doc, '【后果】', bold=True)
add_bullet(doc, '正面：成本降 40%、闲聊 TTFT 降 60%、路由 0ms 延迟')
add_bullet(doc, '负面：路由规则需维护（意图→模型映射可能需调整）')
add_bullet(doc, '约束：规则变更需验证路由准确率 > 95%')

# ADR-GW-03
add_heading_styled(doc, 'ADR-GW-03：Fallback 链——三级降级', level=2)
add_paragraph_styled(doc, '【背景】', bold=True)
add_paragraph_styled(doc,
    '主模型故障时只重试同模型无法解决问题。需要切换到不同模型才能有效降级。')
add_paragraph_styled(doc, '【决策】', bold=True)
add_paragraph_styled(doc,
    '三级降级链：L1 主模型 → L2 备用模型 → L3 模板兜底。'
    '主模型失败时自动切到备用模型，备用也失败时返回模板兜底话术。')
add_paragraph_styled(doc, '【理由】', bold=True)
add_bullet(doc, '换模型比重试更有效（不同模型故障不相关）')
add_bullet(doc, '模板兜底确保用户始终有响应')
add_bullet(doc, '对上层透明（AgentChatService 不感知降级过程）')
add_paragraph_styled(doc, '【后果】', bold=True)
add_bullet(doc, '正面：可用率 > 99.95%、用户始终有响应')
add_bullet(doc, '负面：降级到备用模型可能质量略低')
add_bullet(doc, '约束：模板兜底触发率需 < 0.1%')

# ADR-GW-04
add_heading_styled(doc, 'ADR-GW-04：API Key 池——多 Key 轮询 + 自动摘除', level=2)
add_paragraph_styled(doc, '【背景】', bold=True)
add_paragraph_styled(doc,
    '单一 API Key 在高并发时容易被限流（429），导致整体降级。')
add_paragraph_styled(doc, '【决策】', bold=True)
add_paragraph_styled(doc,
    '每个模型配置多个 API Key，round-robin 轮询分摊流量。'
    'Key 被限流（429）时自动标记 RATE_LIMITED 并冷却，'
    'Key 失效（401/403）时标记 DISABLED 永久摘除。')
add_paragraph_styled(doc, '【理由】', bold=True)
add_bullet(doc, '自动轮询分摊流量，无需人工切换 Key')
add_bullet(doc, '自动摘除故障 Key，不影响其他 Key')
add_bullet(doc, '冷却后自动恢复，无需人工干预')
add_paragraph_styled(doc, '【后果】', bold=True)
add_bullet(doc, '正面：高并发不限流、Key 故障自动处理')
add_bullet(doc, '负面：多 Key 管理复杂度增加')
add_bullet(doc, '约束：需监控 Key 摘除误判率 < 1%')

# ADR-GW-05
add_heading_styled(doc, 'ADR-GW-05：成本追踪——per-request 计量 + 多维度归因', level=2)
add_paragraph_styled(doc, '【背景】', bold=True)
add_paragraph_styled(doc,
    '当前仅有 Token 日志，无法追踪"哪个模型花了多少钱""哪个用户成本最高"。')
add_paragraph_styled(doc, '【决策】', bold=True)
add_paragraph_styled(doc,
    'per-request Token 计量 + 5 维度成本归因（用户/会话/模型/意图/天）。'
    '异步写入 Redis（实时统计）+ MySQL（持久化）+ Micrometer（Metrics）。')
add_paragraph_styled(doc, '【理由】', bold=True)
add_bullet(doc, 'per-request 可精确归因到每次调用')
add_bullet(doc, '多维度支持成本优化决策（如"SEARCH 意图成本占比"）')
add_bullet(doc, '异步写入不阻塞主流程')
add_paragraph_styled(doc, '【后果】', bold=True)
add_bullet(doc, '正面：成本可见可控、支持优化决策')
add_bullet(doc, '负面：额外存储开销（Redis + MySQL）')
add_bullet(doc, '约束：成本统计准确率需 > 99%')

# ADR-GW-06
add_heading_styled(doc, 'ADR-GW-06：模型健康检查——per-model CircuitBreaker + 主动探测', level=2)
add_paragraph_styled(doc, '【背景】', bold=True)
add_paragraph_styled(doc,
    '当前只有一个全局 deepSeekCircuitBreaker，无法 per-model 独立熔断。'
    '一个模型故障可能影响其他模型的熔断判断。')
add_paragraph_styled(doc, '【决策】', bold=True)
add_paragraph_styled(doc,
    'per-model 动态创建 CircuitBreaker，模型独立熔断互不影响。'
    '增加主动健康探测：每 60s 对 OPEN 状态模型发起轻量探测请求。')
add_paragraph_styled(doc, '【理由】', bold=True)
add_bullet(doc, 'per-model 独立熔断：一个模型故障不影响其他模型')
add_bullet(doc, '主动探测加速恢复：不用等被动请求触发 HALF_OPEN')
add_bullet(doc, '探测成本低（< 10 Token），不影响整体成本')
add_paragraph_styled(doc, '【后果】', bold=True)
add_bullet(doc, '正面：故障 90s 内恢复、模型间互不影响')
add_bullet(doc, '负面：探测请求有少量成本')
add_bullet(doc, '约束：探测 prompt 需短小、探测频率可配置')

# ADR-GW-07
add_heading_styled(doc, 'ADR-GW-07：优雅降级——四级降级方案', level=2)
add_paragraph_styled(doc, '【背景】', bold=True)
add_paragraph_styled(doc,
    '当所有 LLM 模型都不可用时，直接报错会严重影响用户体验。'
    '需要保证用户仍然能得到合理响应。')
add_paragraph_styled(doc, '【决策】', bold=True)
add_paragraph_styled(doc,
    '四级降级：L0 语义缓存查询 → L1 主模型 → L2 备用模型 → L3 模板兜底。'
    '每级降级确保用户有响应，模板兜底是最后防线。')
add_paragraph_styled(doc, '【理由】', bold=True)
add_bullet(doc, '语义缓存查询 < 10ms，命中则跳过 LLM 调用')
add_bullet(doc, '模板兜底 < 5ms，确保用户始终有响应')
add_bullet(doc, '四级方案层层兜底，最小化用户感知故障')
add_paragraph_styled(doc, '【后果】', bold=True)
add_bullet(doc, '正面：用户始终有响应、故障感知最小化')
add_bullet(doc, '负面：模板兜底质量低（固定话术）')
add_bullet(doc, '约束：L3 触发率需 < 0.1%')

# ==================== 保存文档 ====================
output_path = r'e:\workspace_work\CampusShare\docs\agent-design\LLM网关与多模型路由模块设计方案.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
