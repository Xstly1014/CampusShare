# -*- coding: utf-8 -*-
"""
生成《CampusShare Agent 上下文工程模块设计方案》Word 文档
专注主题：上下文装载 / Token 预算 / 压缩策略（不涉及 System Prompt 内容/对话编排流程/RAG 算法）
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 样式辅助函数
# ============================================================

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border)
    pPr.append(pBdr)
    lines = code_text.split('\n')
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if i < len(lines) - 1:
            run.add_break()
    return p


def add_callout(doc, text, color='FFF3CD', border_color='FFC107'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), border_color)
        pBdr.append(border)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def style_table_header(row, bg_color='2563EB'):
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_styled_table(doc, headers, rows, col_widths=None, header_bg='2563EB'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    style_table_header(table.rows[0], header_bg)
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_background(row_cells[c_idx], 'F8F9FA')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            run.font.size = Pt(20)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            run.font.size = Pt(13)
    return h


def add_paragraph_styled(doc, text, bold=False, size=10.5, color=None, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


# ============================================================
# 文档生成开始
# ============================================================

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(10.5)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ==================== 封面 ====================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CampusShare Agent\n上下文工程模块设计方案')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('管理 Agent 的"工作记忆"：分层装载 · Token 预算 · 三级压缩')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

for _ in range(8):
    doc.add_paragraph()

info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('文档版本', 'v1.0'),
    ('文档日期', '2026-07-05'),
    ('文档状态', '设计中'),
    ('适用范围', 'campushare-agent 服务 / 上下文工程模块'),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    for p in info_table.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)
    for p in info_table.rows[i].cells[1].paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)
    set_cell_background(info_table.rows[i].cells[0], 'F0F4FF')

doc.add_page_break()

# ==================== 本文档范围声明 ====================
add_heading_styled(doc, '本文档范围声明', level=1)

add_callout(doc,
    '本文档专注讨论上下文工程这一个细小方向。上下文工程是 Agent 的"工作记忆管理"——'
    '它决定每次调用 LLM 时喂什么、喂多少、太长了怎么压缩。这是搭建 Agent 时数据层的核心工作，'
    '承接 System Prompt（人格）、RAG（知识）、意图识别（路由）的输出，'
    '为对话编排（流程）提供组装好的上下文。',
    color='E8F4FD', border_color='2196F3')

add_paragraph_styled(doc, '本文档覆盖：', bold=True)
add_bullet(doc, 'L0-L5 六层上下文分层装载（System Prompt / 用户画像 / 工具 schema / 检索结果 / 对话历史 / 当前输入）')
add_bullet(doc, 'Token 预算分配（按意图动态分配 8K Token）')
add_bullet(doc, '三级渐进压缩策略（Rolling Summary + Slot Freezing + Pin Message）')
add_bullet(doc, '上下文优先级与淘汰策略（当 Token 超预算时按优先级裁剪）')
add_bullet(doc, '上下文工程专属的评估指标、压缩质量测试、A/B 测试、漂移检测')

add_paragraph_styled(doc, '本文档不覆盖（避免主题混乱）：', bold=True)
add_bullet(doc, 'System Prompt 怎么写（角色/边界/格式/护栏）→ 属于《System Prompt 工程》文档')
add_bullet(doc, '对话状态机 / SSE 流式 / 中断恢复 → 属于《对话编排》文档')
add_bullet(doc, 'RAG 检索算法 / 向量索引 / RRF 融合 → 属于《RAG 检索增强生成》文档')
add_bullet(doc, '意图识别算法 / 规则短路 / LLM 分类器 → 属于《意图识别》文档')
add_bullet(doc, 'ReAct / CoT / Plan-and-Execute 范式对比 → 属于《Agent 范式选型》文档')

add_callout(doc,
    '关于 ADR：本文档会引用 ADR 编号（如 ADR-CTX-01）。ADR = Architecture Decision Record（架构决策记录），'
    '是业界记录重要架构决策的实践，每条 ADR 包含「上下文 / 决策 / 后果」三段式。'
    '本文档末尾附录列出所有引用的 ADR 摘要，编号以 CTX 前缀表示 Context Engineering 专用。',
    color='FFF3CD', border_color='FFC107')

doc.add_page_break()

# ==================== 目录 ====================
add_heading_styled(doc, '目录', level=1)

toc_items = [
    '一、场景：为什么上下文工程是 Agent 的"工作记忆"',
    '    1.1 业务背景：LLM 的上下文窗口限制',
    '    1.2 没有上下文工程会怎样：四大问题',
    '    1.3 上下文工程 vs System Prompt 工程',
    '    1.4 CampusShare 的具体场景与挑战',
    '二、方案：业界上下文工程设计模式',
    '    2.1 上下文工程的三个核心问题',
    '    2.2 大厂案例研究',
    '    2.3 上下文压缩技术演进',
    '    2.4 CampusShare 选型决策',
    '三、流程：如何搭建上下文工程',
    '    3.1 前置条件',
    '    3.2 L0-L5 六层上下文装载',
    '    3.3 Token 预算分配（按意图动态分配）',
    '    3.4 三级渐进压缩策略',
    '    3.5 上下文优先级与淘汰策略',
    '    3.6 上下文工程的工程实现',
    '四、核心代码',
    '    4.1 文件架构',
    '    4.2 ContextLayer 枚举（六层定义）',
    '    4.3 ContextAssembler（上下文装配器）',
    '    4.4 TokenBudgetAllocator（Token 预算分配）',
    '    4.5 ContextCompressor（三级压缩）',
    '    4.6 RollingSummaryService（滚动摘要）',
    '    4.7 数据库 Schema',
    '五、目标：实现效果',
    '    5.1 上下文装载目标',
    '    5.2 Token 预算目标',
    '    5.3 压缩质量目标',
    '    5.4 性能与成本目标',
    '六、测试评估与验收',
    '    6.1 评估指标体系',
    '    6.2 黄金测试集构建',
    '    6.3 评估流水线与 CI/CD',
    '    6.4 LLM-as-Judge 评估',
    '    6.5 测试用例设计',
    '    6.6 压缩质量专项测试',
    '    6.7 A/B 测试设计',
    '    6.8 验收流程与准入准出',
    '    6.9 持续监控与漂移检测',
    '七、总结与边界声明',
    '    7.1 核心总结',
    '    7.2 本文档与其他文档的关系',
    '    7.3 演进路线',
    '附录：ADR 摘要',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    if not item.startswith('    '):
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ==================== 一、场景 ====================
add_heading_styled(doc, '一、场景：为什么上下文工程是 Agent 的"工作记忆"', level=1)

add_heading_styled(doc, '1.1 业务背景：LLM 的上下文窗口限制', level=2)

add_paragraph_styled(doc,
    'LLM 有一个硬性限制：**上下文窗口**。DeepSeek-V3 的窗口是 128K Token，'
    '听起来很大，但实际可用远没有那么多——输入和输出共享这 128K，'
    '而我们要在输入里塞 System Prompt、用户画像、工具 schema、检索结果、对话历史、当前输入。')

add_paragraph_styled(doc,
    '更关键的是：**上下文越长，成本越高、延迟越大、质量越差**。'
    '业界有个反直觉的发现——"Lost in the Middle"：LLM 对上下文中间的信息记忆力差，'
    '即使没超窗口，把 100 轮历史全塞进去，LLM 也会"忘记"中间的关键信息。'
    '所以上下文工程不是"塞得越多越好"，而是"塞得越精准越好"。', bold=True)

add_styled_table(doc,
    ['上下文长度', '输入成本', '延迟（TTFT）', '质量（准确率）', '适用场景'],
    [
        ['2K Token', '1x', '~0.8s', '95%', '单轮对话'],
        ['8K Token', '4x', '~1.2s', '92%', '多轮对话（推荐）'],
        ['32K Token', '16x', '~2.5s', '85%', '长文档分析'],
        ['128K Token', '64x', '~5s', '78%', '超长上下文（慎用）'],
    ],
    col_widths=[3, 2.5, 3, 3, 4])

add_callout(doc,
    '关键认知：上下文工程的本质不是"塞满窗口"，而是"在有限预算内最大化信息密度"。'
    'Claude 团队有句名言："Context is the working memory of an Agent." '
    '人的工作记忆只有 7±2 个块，LLM 的工作记忆也不是无限的——'
    '上下文工程就是 LLM 的"认知负荷管理"。')

add_heading_styled(doc, '1.2 没有上下文工程会怎样：四大问题', level=2)

add_paragraph_styled(doc,
    '如果直接把"System Prompt + 全量历史 + 当前输入"拼起来喂给 LLM，不做任何管理，会出现：',
    bold=True)

add_styled_table(doc,
    ['问题', '表现', '后果', '量化影响'],
    [
        ['Token 超限', '多轮对话后历史累积超过窗口',
         'LLM 报错或自动截断', '第 20 轮必崩'],
        ['成本爆炸', '每轮都带全量历史，Token 线性增长',
         'API 费用超预算', '成本翻 5-10 倍'],
        ['Lost in Middle', '关键信息被淹没在长历史中',
         'LLM "忘记"早期承诺', '准确率下降 15%+'],
        ['信息冗余', '寒暄、重复、废话都进上下文',
         '有效信息占比低', '信噪比 < 30%'],
    ],
    col_widths=[2.5, 4.5, 4, 3.5])

add_heading_styled(doc, '1.3 上下文工程 vs System Prompt 工程', level=2)

add_paragraph_styled(doc,
    '上下文工程和 System Prompt 工程常被混为一谈，但它们是不同层次的工作：')

add_styled_table(doc,
    ['维度', 'System Prompt 工程', '上下文工程'],
    [
        ['关注什么', 'Prompt 内容（角色/边界/格式/护栏）',
         'Prompt 怎么组装（分层/预算/压缩）'],
        ['变更频率', '低（按版本发布）',
         '高（每轮对话都重新组装）'],
        ['核心问题', '"LLM 是谁"',
         '"每次喂什么给 LLM"'],
        ['输入', '业务需求 + 合规要求',
         'System Prompt + RAG 结果 + 对话历史 + 意图'],
        ['输出', '一段固定的 System Prompt 文本',
         '一个组装好的 messages 数组'],
        ['失败模式', '身份混乱/能力越界',
         'Token 超限/信息丢失/成本爆炸'],
        ['依赖关系', '独立',
         '依赖 System Prompt + RAG + 意图识别的输出'],
    ],
    col_widths=[3, 6, 6])

add_paragraph_styled(doc,
    '本文档假设 System Prompt、RAG、意图识别的输出已就绪，'
    '专注讨论"如何把这些输出组装成 LLM 的输入"。', bold=True)

add_heading_styled(doc, '1.4 CampusShare 的具体场景与挑战', level=2)

add_paragraph_styled(doc,
    'CampusShare「小享」的上下文工程面临三个特殊挑战：')

add_styled_table(doc,
    ['挑战', '具体场景', '对上下文工程的要求'],
    [
        ['多轮对话长', '用户可能连续问 20+ 轮（找资源时反复筛选）',
         '必须压缩历史，否则第 10 轮就超预算'],
        ['RAG 结果大', '检索类问题返回 Top-5 文档，每个 500-1000 Token',
         'RAG 占预算 30-40%，必须与历史平衡'],
        ['意图差异大', 'HOW_TO 需要完整历史，CHAT 只需要近 3 轮',
         '按意图动态分配 Token 预算'],
        ['成本敏感', '校园平台预算有限，不能像大厂那样无脑用 32K',
         '严控在 8K Token 以内'],
    ],
    col_widths=[3, 5, 7])

doc.add_page_break()

# ==================== 二、方案 ====================
add_heading_styled(doc, '二、方案：业界上下文工程设计模式', level=1)

add_heading_styled(doc, '2.1 上下文工程的三个核心问题', level=2)

add_paragraph_styled(doc,
    '上下文工程要回答三个核心问题，对应"装什么、装多少、超了怎么办"：')

add_styled_table(doc,
    ['问题', '通俗说法', '工程方案', '本文档章节'],
    [
        ['装什么', '哪些信息该进上下文？',
         'L0-L5 六层分层装载', '3.2'],
        ['装多少', '每层该分多少 Token？',
         '按意图动态分配 8K 预算', '3.3'],
        ['超了怎么办', '历史太长怎么压缩？',
         '三级渐进压缩', '3.4'],
    ],
    col_widths=[3, 4, 5, 3])

add_callout(doc,
    '三个问题的顺序很重要：先确定"装什么"（分层），再确定"装多少"（预算），'
    '最后处理"超了怎么办"（压缩）。跳过分层直接做压缩，会压缩错对象——'
    '把该留的 RAG 结果压缩了，却保留了无用的寒暄。')

add_heading_styled(doc, '2.2 大厂案例研究', level=2)

add_heading_styled(doc, '2.2.1 Anthropic Claude', level=3)
add_bullet(doc, '窗口：200K Token（业界最大）')
add_bullet(doc, '策略：明确区分 System / Tool / RAG / History / User 五层；用 XML 标签 <context> 分隔')
add_bullet(doc, '压缩：不主动压缩，依赖长窗口；但研究表明超过 100K 后质量下降')
add_bullet(doc, '可借鉴：XML 标签分层（让 LLM 明确每段的性质）')

add_heading_styled(doc, '2.2.2 OpenAI GPT-4o', level=3)
add_bullet(doc, '窗口：128K Token')
add_bullet(doc, '策略：自动截断旧消息（保最近 N 轮）；推荐用 System Prompt 携带固定信息')
add_bullet(doc, '压缩：不提供摘要，靠用户自行管理；推出了 "Assistants API" 自动管理历史')
add_bullet(doc, '可借鉴：自动截断策略简单可靠；Assistants API 的 thread 概念')

add_heading_styled(doc, '2.2.3 Google Gemini', level=3)
add_bullet(doc, '窗口：1M Token（最大）')
add_bullet(doc, '策略：用 "context caching" 缓存前缀；长上下文不压缩')
add_bullet(doc, '可借鉴：前缀缓存思路（与 DeepSeek Prefix Cache 一致）')

add_heading_styled(doc, '2.2.4 阿里通义千问', level=3)
add_bullet(doc, '窗口：32K-128K')
add_bullet(doc, '策略：对话历史分级存储（热/温/冷）；热历史原文，温历史摘要，冷历史归档')
add_bullet(doc, '可借鉴：分级存储思路（对应本文档的三级压缩）')

add_paragraph_styled(doc, '案例对比总结：', bold=True)

add_styled_table(doc,
    ['厂商', '窗口', '分层策略', '压缩策略', 'CampusShare 可借鉴度'],
    [
        ['Claude', '200K', 'XML 标签五层', '不压缩（靠长窗口）', '★★★★★（分层）'],
        ['GPT-4o', '128K', '自动截断', '不压缩（靠 Assistants）', '★★★（截断太粗暴）'],
        ['Gemini', '1M', '前缀缓存', '不压缩', '★★★★（缓存思路）'],
        ['通义', '32K+', '热温冷分级', '摘要 + 归档', '★★★★★（分级压缩）'],
    ],
    col_widths=[2.5, 2, 3, 3, 4])

add_paragraph_styled(doc,
    'CampusShare 选型：以通义的"热温冷分级压缩"为骨架，融合 Claude 的"XML 标签分层"，'
    '在 8K Token 预算内实现高信噪比的上下文管理。', bold=True)

add_heading_styled(doc, '2.3 上下文压缩技术演进', level=2)

add_styled_table(doc,
    ['阶段', '技术', '原理', '优缺点'],
    [
        ['1.0', '滑动窗口截断', '保留最近 N 轮，丢弃更早的',
         '简单但信息丢失严重'],
        ['2.0', '全量摘要', '把整个历史压缩成一段摘要',
         '省 Token 但丢失细节，且摘要本身耗 Token'],
        ['3.0', '分级压缩', '热原文 / 温摘要 / 冷归档',
         '平衡信息保留与成本，但实现复杂'],
        ['4.0（本方案）', '三级渐进 + 槽位冻结', '摘要 + 槽位 + Pin Message',
         '保留关键事实，渐进压缩，但工程量大'],
    ],
    col_widths=[2, 3, 4, 5])

add_heading_styled(doc, '2.4 CampusShare 选型决策', level=2)

add_paragraph_styled(doc, '基于上述调研，CampusShare 上下文工程的选型决策如下（ADR 摘要）：')

add_styled_table(doc,
    ['ADR 编号', '决策', '理由'],
    [
        ['ADR-CTX-01', '采用 L0-L5 六层分层装载',
         '明确每层职责，避免信息混杂'],
        ['ADR-CTX-02', 'Token 预算固定 8K（不动态扩容）',
         '平衡成本与质量；超 8K 后质量下降明显'],
        ['ADR-CTX-03', '按意图动态分配预算',
         'HOW_TO 多分历史，SEARCH 多分 RAG，CHAT 少分历史'],
        ['ADR-CTX-04', '三级渐进压缩（Summary + Slot + Pin）',
         '保留关键事实，渐进压缩而非一刀切'],
        ['ADR-CTX-05', '用 XML 标签分层（<context> <history>）',
         '让 LLM 明确每段性质，防隐式注入'],
        ['ADR-CTX-06', '压缩触发阈值 6K（8K 的 75%）',
         '提前压缩，避免到 8K 才被动截断'],
        ['ADR-CTX-07', '压缩质量用 LLM-as-Judge 评估',
         '摘要不能丢关键信息，必须可量化'],
    ],
    col_widths=[3, 6, 6])

doc.add_page_break()

# ==================== 三、流程 ====================
add_heading_styled(doc, '三、流程：如何搭建上下文工程', level=1)

add_heading_styled(doc, '3.1 前置条件', level=2)

add_styled_table(doc,
    ['前置条件', '具体要求', '来源文档'],
    [
        ['System Prompt 已定义', '六要素结构已就绪，L1 平台级固定',
         '《System Prompt 工程》'],
        ['RAG 检索结果可用', '能返回 Top-K 文档，每篇带元数据',
         '《RAG 检索增强生成》'],
        ['意图识别已上线', '能输出 HOW_TO / SEARCH / CHAT 标签',
         '《意图识别》'],
        ['Token 计数工具', 'JTokkit 本地估算，误差 < 5%',
         '基础设施'],
        ['对话历史存储', 'agent_turns 表记录每轮消息',
         '数据库'],
    ],
    col_widths=[4, 7, 4])

add_heading_styled(doc, '3.2 L0-L5 六层上下文装载', level=2)

add_paragraph_styled(doc,
    '上下文分六层装载，每层职责清晰，用 XML 标签分隔。'
    '这个分层是上下文工程的骨架——所有信息都要归入某一层，不允许"游离信息"。')

add_styled_table(doc,
    ['层', '名称', '内容', 'Token 预算', '变更频率', '缓存策略'],
    [
        ['L0', 'System Prompt', '角色/边界/格式/护栏',
         '~1500', '低（版本化）', 'Prefix Cache'],
        ['L1', '用户画像', '用户 ID / 学校 / 偏好',
         '~200', '中（按用户）', 'Redis 缓存'],
        ['L2', '工具 Schema', '可用工具的 JSON Schema',
         '~300', '低（按版本）', 'Prefix Cache'],
        ['L3', 'RAG 检索结果', 'Top-K 文档 + 引用编号',
         '~2000', '高（每轮）', '不缓存'],
        ['L4', '对话历史', '历史消息（含压缩摘要）',
         '~3500', '高（每轮）', '不缓存'],
        ['L5', '当前输入', '用户本轮 query',
         '~500', '高（每轮）', '不缓存'],
    ],
    col_widths=[1, 2.5, 4, 2, 2.5, 2.5])

add_callout(doc,
    'L0-L2 是"固定层"（同一用户同一会话内基本不变），可以命中 Prefix Cache；'
    'L3-L5 是"变动层"（每轮都变），无法缓存。'
    '固定层在前，变动层在后——这样 LLM 能"先认识环境，再看资料，最后回答"。')

add_code_block(doc, '''# 上下文六层结构（XML 标签分隔）

<system>
{L0 System Prompt}
</system>

<user_profile>
{L1 用户画像：用户ID=12345, 学校=清华, 偏好=计算机}
</user_profile>

<tools>
{L2 工具 Schema：search_resources, get_post_detail}
</tools>

<context>
{L3 RAG 检索结果：
[1] 创作者认证申请指南：获赞≥10000且发帖≥50...
[2] 平台用户角色说明：普通用户/创作者/管理员...}
</context>

<history>
{L4 对话历史（含压缩摘要）：
[摘要] 用户之前询问了发帖方法，已解答。
[最近3轮原文]
User: 怎么发帖？
Assistant: ...
User: 发帖需要什么条件？
Assistant: ...}
</history>

<current_input>
{L5 用户本轮输入：怎么申请创作者？}
</current_input>''')

add_heading_styled(doc, '3.3 Token 预算分配（按意图动态分配）', level=2)

add_paragraph_styled(doc,
    '8K Token 总预算按意图动态分配给六层。'
    '不同意图对"历史 vs RAG"的需求不同——'
    'HOW_TO 需要完整历史（理解上下文），SEARCH 需要大 RAG（看更多资料），CHAT 不需要太多历史。')

add_styled_table(doc,
    ['层', 'HOW_TO 预算', 'SEARCH 预算', 'CHAT 预算', '说明'],
    [
        ['L0 System', '1500', '1500', '1500', '固定不变'],
        ['L1 画像', '200', '200', '200', '固定不变'],
        ['L2 工具', '300', '300', '0', 'CHAT 无工具'],
        ['L3 RAG', '1500', '3000', '0', 'SEARCH 多分'],
        ['L4 历史', '3500', '2000', '5000', 'HOW_TO/CHAT 多分'],
        ['L5 当前', '500', '500', '500', '用户输入'],
        ['L6 预留输出', '500', '500', '800', '给 LLM 留输出空间'],
        ['合计', '8000', '8000', '8000', '总预算 8K'],
    ],
    col_widths=[2.5, 2.5, 2.5, 2.5, 4])

add_callout(doc,
    '关键设计：必须给 LLM 预留输出 Token（L6）。'
    '如果 8K 全被输入占满，LLM 没有空间生成回复。'
    '一般预留 500-1000 Token 给输出，实际输入预算 = 8K - 预留输出。')

add_heading_styled(doc, '3.4 三级渐进压缩策略', level=2)

add_paragraph_styled(doc,
    '当 L4 对话历史超过预算时，触发三级渐进压缩。'
    '"渐进"意味着不是一刀切，而是按信息重要性逐级压缩——'
    '先压缩最不重要的，保留最重要的。')

add_styled_table(doc,
    ['级别', '触发条件', '压缩方式', '保留什么', '丢失什么'],
    [
        ['一级：Rolling Summary',
         '历史 > 2000 Token',
         '把最早 N 轮压缩成一段摘要',
         '关键事实（用户问什么/答什么）',
         '具体措辞/语气'],
        ['二级：Slot Freezing',
         '历史 > 3500 Token',
         '提取关键槽位（学校/分类/资源类型）冻结',
         '结构化关键信息',
         '非结构化细节'],
        ['三级：Pin Message',
         '历史 > 5000 Token',
         '用户/Agent 标记的关键消息永久保留',
         '被 pin 的消息原文',
         '未 pin 的旧消息'],
    ],
    col_widths=[3, 3, 3.5, 3, 3])

add_code_block(doc, '''# 三级压缩示例

## 原始历史（10 轮，4500 Token）
T1: User: 找清华的操作系统卷子
T2: Assistant: 找到以下资源 [1][2]...
T3: User: 要 2023 年的
T4: Assistant: 筛选后找到 [1] 2023 年的...
T5: User: 有答案吗？
T6: Assistant: [1] 含答案...
T7: User: 怎么下载？
T8: Assistant: 点击下载按钮...
T9: User: 下载要登录吗？
T10: Assistant: 需要...

## 一级压缩后（T1-T6 压缩成摘要，200 Token）
[摘要 T1-T6] 用户寻找清华操作系统 2023 年期末卷子（含答案），
  已找到资源 [1]，用户询问了是否含答案。

T7-T10 原文保留（1500 Token）

## 二级压缩后（提取槽位）
[摘要 T1-T6] 用户寻找清华操作系统 2023 年期末卷子...
[槽位] 学校=清华, 科目=操作系统, 年份=2023, 资源=[1]

T7-T10 原文保留

## 三级压缩后（Pin T2 因为含资源列表）
[摘要 T1-T6] ...
[槽位] 学校=清华, 科目=操作系统, 年份=2023, 资源=[1]
[Pin] T2: Assistant: 找到以下资源 [1][2]...

T7-T10 原文保留''')

add_callout(doc,
    '三级压缩不是"三选一"，而是"叠加"——一级压缩后仍超预算，触发二级；'
    '二级后仍超，触发三级。这样能在保留最多信息的前提下逐步压缩。')

add_heading_styled(doc, '3.5 上下文优先级与淘汰策略', level=2)

add_paragraph_styled(doc,
    '当所有压缩都做完后仍超预算，启动"淘汰策略"——按优先级裁剪。'
    '优先级决定"谁先被丢"。')

add_styled_table(doc,
    ['优先级', '层', '淘汰顺序', '理由'],
    [
        ['P0（绝不丢）', 'L0 System / L5 当前输入',
         '永不淘汰', '丢了就身份混乱/无法回答'],
        ['P1（尽量留）', 'Pin Message / 槽位',
         '最后淘汰', '用户标记的关键信息'],
        ['P2（可压缩）', 'L3 RAG / L4 近 3 轮历史',
         '先减 RAG 数量', 'RAG 可减 Top-5 到 Top-3'],
        ['P3（可淘汰）', 'L4 早期历史 / 寒暄',
         '最先淘汰', '信息密度低'],
        ['P4（可省略）', 'L1 画像 / L2 工具',
         'CHAT 时可省 L2', '不影响主流程'],
    ],
    col_widths=[3, 4, 4, 4])

add_heading_styled(doc, '3.6 上下文工程的工程实现', level=2)

add_paragraph_styled(doc, '完整的上下文工程流程（每轮对话执行）：')

add_styled_table(doc,
    ['步骤', '动作', '输入', '输出'],
    [
        ['1', '接收意图标签', '用户 query', '意图（HOW_TO/SEARCH/CHAT）'],
        ['2', '分配 Token 预算', '意图', '六层预算表'],
        ['3', '获取 L0-L2 固定层', 'System Prompt / 画像 / 工具', '固定层文本'],
        ['4', '获取 L3 RAG 结果', 'query + 预算', 'Top-K 文档'],
        ['5', '获取 L4 历史', '会话 ID + 预算', '历史消息列表'],
        ['6', '检查 L4 是否超预算', '历史 Token 数', '是否触发压缩'],
        ['7', '执行三级压缩', '历史 + 摘要服务', '压缩后历史'],
        ['8', '装配六层上下文', '所有层文本', 'messages 数组'],
        ['9', 'Token 计数校验', 'messages', '总 Token 数'],
        ['10', '超预算则淘汰 P3/P4', 'messages + 优先级', '裁剪后 messages'],
    ],
    col_widths=[1.5, 5, 4, 4])

doc.add_page_break()

# ==================== 四、核心代码 ====================
add_heading_styled(doc, '四、核心代码', level=1)

add_heading_styled(doc, '4.1 文件架构', level=2)

add_styled_table(doc,
    ['文件', '职责', '行数'],
    [
        ['ContextLayer.java', '六层枚举定义', '~30'],
        ['ContextAssembler.java', '上下文装配器（拼接六层）', '~120'],
        ['TokenBudgetAllocator.java', 'Token 预算分配（按意图）', '~80'],
        ['ContextCompressor.java', '三级压缩主控', '~150'],
        ['RollingSummaryService.java', '一级压缩：滚动摘要', '~100'],
        ['SlotFreezingService.java', '二级压缩：槽位冻结', '~90'],
        ['PinMessageService.java', '三级压缩：Pin Message', '~70'],
        ['TokenCounter.java', 'Token 计数（JTokkit）', '~50'],
        ['context_tables.sql', '上下文相关表结构', '~40'],
    ],
    col_widths=[6, 7, 2])

add_heading_styled(doc, '4.2 ContextLayer 枚举（六层定义）', level=2)

add_code_block(doc, '''package com.campushare.agent.context;

public enum ContextLayer {
    L0_SYSTEM("system", "System Prompt", 1500, true),
    L1_USER_PROFILE("user_profile", "用户画像", 200, true),
    L2_TOOLS("tools", "工具 Schema", 300, true),
    L3_RAG("context", "RAG 检索结果", 2000, false),
    L4_HISTORY("history", "对话历史", 3500, false),
    L5_CURRENT("current_input", "当前输入", 500, false);

    private final String tag;          // XML 标签名
    private final String description;   // 中文说明
    private final int defaultBudget;    //默认 Token 预算
    private final boolean cacheable;    // 是否可缓存

    ContextLayer(String tag, String description, int defaultBudget, boolean cacheable) {
        this.tag = tag;
        this.description = description;
        this.defaultBudget = defaultBudget;
        this.cacheable = cacheable;
    }

    public String wrap(String content) {
        return "<" + tag + ">\\n" + content + "\\n</" + tag + ">\\n";
    }

    // getter 省略
}''')

add_heading_styled(doc, '4.3 ContextAssembler（上下文装配器）', level=2)

add_code_block(doc, '''package com.campushare.agent.context;

import com.campushare.agent.prompt.PromptConstants;
import com.campushare.agent.dto.*;
import org.springframework.stereotype.Component;

@Component
public class ContextAssembler {

    private final TokenBudgetAllocator budgetAllocator;
    private final TokenCounter tokenCounter;

    public Context assemble(String intent, Long userId, Long sessionId,
                            String currentQuery, RagResult ragResult) {
        // 1. 分配 Token 预算
        BudgetTable budget = budgetAllocator.allocate(intent);

        // 2. 装配六层
        StringBuilder sb = new StringBuilder();

        // L0 System Prompt
        sb.append(ContextLayer.L0_SYSTEM.wrap(PromptConstants.PLATFORM_PROMPT));

        // L1 用户画像
        String profile = getUserProfile(userId);
        sb.append(ContextLayer.L1_USER_PROFILE.wrap(profile));

        // L2 工具 Schema（CHAT 无工具）
        if (!"CHAT".equals(intent)) {
            String tools = getToolSchemas(intent);
            sb.append(ContextLayer.L2_TOOLS.wrap(tools));
        }

        // L3 RAG 检索结果（CHAT 无 RAG）
        if (ragResult != null && ragResult.hasResults()) {
            String rag = ragResult.toPromptString(budget.getRagBudget());
            sb.append(ContextLayer.L3_RAG.wrap(rag));
        }

        // L4 对话历史（含压缩）
        String history = getHistory(sessionId, budget.getHistoryBudget());
        sb.append(ContextLayer.L4_HISTORY.wrap(history));

        // L5 当前输入
        sb.append(ContextLayer.L5_CURRENT.wrap(currentQuery));

        // 3. Token 计数校验
        int totalTokens = tokenCounter.count(sb.toString());
        if (totalTokens > budget.getMaxInput()) {
            // 超预算，触发淘汰策略
            sb = evict(sb, budget, totalTokens);
        }

        return Context.builder()
            .content(sb.toString())
            .totalTokens(tokenCounter.count(sb.toString()))
            .budget(budget)
            .build();
    }

    private String evict(StringBuilder sb, BudgetTable budget, int current) {
        // 按 P3 → P4 优先级淘汰
        // 1. 先减 RAG 数量（Top-5 → Top-3）
        // 2. 再删早期历史
        // 3. 最后省略工具 Schema（CHAT 已省略）
        // 实现略
        return sb;
    }
}''')

add_heading_styled(doc, '4.4 TokenBudgetAllocator（Token 预算分配）', level=2)

add_code_block(doc, '''package com.campushare.agent.context;

import org.springframework.stereotype.Component;

@Component
public class TokenBudgetAllocator {

    private static final int TOTAL_BUDGET = 8000;
    private static final int OUTPUT_RESERVE = 500;  // 预留输出

    public BudgetTable allocate(String intent) {
        return switch (intent) {
            case "HOW_TO" -> new BudgetTable(
                1500,  // L0 System
                200,   // L1 画像
                300,   // L2 工具
                1500,  // L3 RAG
                3500,  // L4 历史
                500,   // L5 当前
                OUTPUT_RESERVE  // 输出预留
            );
            case "SEARCH" -> new BudgetTable(
                1500, 200, 300, 3000, 2000, 500, OUTPUT_RESERVE
            );
            case "CHAT" -> new BudgetTable(
                1500, 200, 0, 0, 5000, 500, OUTPUT_RESERVE
            );
            default -> allocate("CHAT");
        };
    }

    public record BudgetTable(
        int system, int profile, int tools,
        int rag, int history, int current, int output
    ) {
        public int maxInput() {
            return TOTAL_BUDGET - output;
        }
        public int getRagBudget() { return rag; }
        public int getHistoryBudget() { return history; }
        public int getMaxInput() { return maxInput(); }
    }
}''')

add_heading_styled(doc, '4.5 ContextCompressor（三级压缩）', level=2)

add_code_block(doc, '''package com.campushare.agent.context;

import org.springframework.stereotype.Component;

@Component
public class ContextCompressor {

    private static final int L1_THRESHOLD = 2000;  // 一级触发
    private static final int L2_THRESHOLD = 3500;  // 二级触发
    private static final int L3_THRESHOLD = 5000;  // 三级触发

    private final RollingSummaryService summaryService;
    private final SlotFreezingService slotService;
    private final PinMessageService pinService;
    private final TokenCounter tokenCounter;

    public String compress(String history, Long sessionId) {
        int tokens = tokenCounter.count(history);

        if (tokens <= L1_THRESHOLD) {
            return history;  // 无需压缩
        }

        // 一级：Rolling Summary
        if (tokens > L1_THRESHOLD) {
            history = summaryService.compress(history, sessionId);
            tokens = tokenCounter.count(history);
        }

        // 二级：Slot Freezing
        if (tokens > L2_THRESHOLD) {
            history = slotService.freeze(history, sessionId);
            tokens = tokenCounter.count(history);
        }

        // 三级：Pin Message
        if (tokens > L3_THRESHOLD) {
            history = pinService.pinAndDrop(history, sessionId);
        }

        return history;
    }
}''')

add_heading_styled(doc, '4.6 RollingSummaryService（滚动摘要）', level=2)

add_code_block(doc, '''package com.campushare.agent.context;

import com.campushare.agent.client.DeepSeekClient;
import org.springframework.stereotype.Service;

@Service
public class RollingSummaryService {

    private final DeepSeekClient llmClient;
    private final TokenCounter tokenCounter;

    private static final String SUMMARY_PROMPT = """
            请将以下对话历史压缩成一段摘要，保留：
            1. 用户询问的核心问题
            2. Agent 给出的关键答案
            3. 提到的资源编号、学校、分类等关键信息
            4. 未解决的问题

            丢弃：
            1. 寒暄、客套话
            2. 重复信息
            3. 具体措辞细节

            输出格式：[摘要] 关键信息要点

            对话历史：
            %s
            """;

    public String compress(String history, Long sessionId) {
        // 找出最早 N 轮（保留最近 3 轮原文）
        String toCompress = extractOldTurns(history, keepRecent = 3);
        String recent = extractRecentTurns(history, keepRecent = 3);

        // 调用 LLM 生成摘要
        String prompt = String.format(SUMMARY_PROMPT, toCompress);
        String summary = llmClient.complete(prompt);

        // 保存摘要到数据库（供下次复用）
        saveSummary(sessionId, summary);

        return "[摘要] " + summary + "\\n\\n" + recent;
    }
}''')

add_heading_styled(doc, '4.7 数据库 Schema', level=2)

add_code_block(doc, '''-- context_summaries 表：一级压缩摘要
CREATE TABLE context_summaries (
    id BIGINT PRIMARY KEY AUTO_INCREMENT,
    session_id BIGINT NOT NULL COMMENT '会话 ID',
    summary_text TEXT NOT NULL COMMENT '摘要内容',
    covered_turn_ids VARCHAR(512) NOT NULL COMMENT '覆盖的轮次 ID 列表',
    token_count INT NOT NULL COMMENT '摘要 Token 数',
    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    INDEX idx_session (session_id),
    INDEX idx_created (created_at)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='上下文压缩摘要';

-- context_slots 表：二级压缩槽位
CREATE TABLE context_slots (
    id BIGINT PRIMARY KEY AUTO_INCREMENT,
    session_id BIGINT NOT NULL,
    slot_key VARCHAR(64) NOT NULL COMMENT '槽位名：school/subject/year等',
    slot_value VARCHAR(256) NOT NULL COMMENT '槽位值',
    frozen_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    UNIQUE KEY uk_session_slot (session_id, slot_key),
    INDEX idx_session (session_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='上下文槽位冻结';

-- pin_messages 表：三级压缩 Pin Message
CREATE TABLE pin_messages (
    id BIGINT PRIMARY KEY AUTO_INCREMENT,
    session_id BIGINT NOT NULL,
    turn_id BIGINT NOT NULL COMMENT '被 pin 的轮次 ID',
    pinned_by ENUM('USER', 'AGENT') NOT NULL,
    reason VARCHAR(256) COMMENT 'pin 原因',
    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    UNIQUE KEY uk_session_turn (session_id, turn_id),
    INDEX idx_session (session_id)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='Pin Message';''')

doc.add_page_break()

# ==================== 五、目标 ====================
add_heading_styled(doc, '五、目标：实现效果', level=1)

add_heading_styled(doc, '5.1 上下文装载目标', level=2)

add_styled_table(doc,
    ['指标', '定义', '目标值', '测量方式'],
    [
        ['六层装载完整率', '六层都正确装配的比例', '100%',
         '日志检查'],
        ['XML 标签正确率', '标签成对闭合的比例', '100%',
         'XML 解析器'],
        ['游离信息率', '未归入六层的信息占比', '0%',
         '人工抽检'],
    ],
    col_widths=[3, 5, 3, 4])

add_heading_styled(doc, '5.2 Token 预算目标', level=2)

add_styled_table(doc,
    ['指标', '定义', '目标值', '测量方式'],
    [
        ['总 Token 不超限', '总输入 ≤ 8K - 输出预留', '100%',
         'JTokkit 计数'],
        ['预算利用率', '实际 Token / 预算 Token', '70-90%',
         '统计'],
        ['超限触发淘汰率', '触发淘汰策略的比例', '≤ 5%',
         '日志统计'],
        ['预留输出充足率', 'LLM 输出未因空间不足被截断', '100%',
         '输出长度检查'],
    ],
    col_widths=[3, 5, 3, 4])

add_heading_styled(doc, '5.3 压缩质量目标', level=2)

add_styled_table(doc,
    ['指标', '定义', '目标值', '测量方式'],
    [
        ['摘要信息保留率', '压缩后保留的关键信息比例', '≥ 90%',
         'LLM-as-Judge'],
        ['摘要 Token 压缩比', '摘要 Token / 原文 Token', '≤ 20%',
         'Token 计数'],
        ['槽位准确率', '槽位提取正确的比例', '≥ 95%',
         '人工抽检'],
        ['Pin Message 命中率', '被 Pin 的消息在后续被引用的比例', '≥ 80%',
         '日志统计'],
        ['压缩后语义一致性', '压缩前后语义不变的程度', '≥ 4.5/5',
         'LLM-as-Judge'],
    ],
    col_widths=[3, 5, 3, 4])

add_heading_styled(doc, '5.4 性能与成本目标', level=2)

add_styled_table(doc,
    ['指标', '定义', '目标值', '测量方式'],
    [
        ['上下文装配延迟', '从意图到装配完成的时间', '≤ 200ms',
         '链路追踪'],
        ['压缩延迟', '三级压缩总耗时', '≤ 1.5s',
         '链路追踪'],
        ['Prefix Cache 命中率', 'L0-L2 固定层命中率', '≥ 95%',
         'API 返回'],
        ['平均输入 Token', '单次对话平均输入', '≤ 5000',
         '统计'],
        ['单次对话输入成本', '平均输入成本', '≤ 0.0005 元',
         'API 账单'],
    ],
    col_widths=[3, 5, 3, 4])

add_callout(doc,
    '成本对比：无上下文工程时，平均输入 ~15000 Token（全量历史），'
    '有上下文工程后 ~5000 Token（压缩 + 预算控制），成本降 67%。'
    '加上 Prefix Cache 命中 L0-L2，实际成本降 80%+。')

doc.add_page_break()

# ==================== 六、测试评估与验收 ====================
add_heading_styled(doc, '六、测试评估与验收', level=1)

add_paragraph_styled(doc,
    '上下文工程的测试评估与 System Prompt / RAG 不同——它不关注"角色对不对"或"检索准不准"，'
    '只关注"上下文组装得对不对、压缩丢没丢信息、Token 预算控没控住"。'
    '本章设计了一套上下文工程专属的评估体系。', bold=True)

add_heading_styled(doc, '6.1 评估指标体系', level=2)

add_styled_table(doc,
    ['类别', '指标', '公式', '目标值'],
    [
        ['装载正确性', '六层装载完整率', '六层完整装配数 / 总装配数', '100%'],
        ['装载正确性', 'XML 标签正确率', '标签成对闭合数 / 总标签数', '100%'],
        ['装载正确性', '游离信息率', '未归层信息数 / 总信息数', '0%'],
        ['Token 预算', '总 Token 超限率', '超限次数 / 总调用数', '0%'],
        ['Token 预算', '预算利用率', '实际 Token / 预算 Token', '70-90%'],
        ['Token 预算', '淘汰触发率', '触发淘汰次数 / 总调用数', '≤ 5%'],
        ['压缩质量', '摘要信息保留率', '保留关键信息数 / 原关键信息数', '≥ 90%'],
        ['压缩质量', '摘要压缩比', '摘要 Token / 原文 Token', '≤ 20%'],
        ['压缩质量', '槽位准确率', '正确槽位数 / 提取槽位数', '≥ 95%'],
        ['压缩质量', '压缩后语义一致性', 'LLM-as-Judge 评分', '≥ 4.5/5'],
        ['性能', '装配延迟', '装配耗时 P95', '≤ 200ms'],
        ['性能', '压缩延迟', '压缩耗时 P95', '≤ 1.5s'],
        ['成本', 'Prefix Cache 命中率', '缓存命中数 / 总调用数', '≥ 95%'],
        ['成本', '平均输入 Token', '总输入 Token / 调用数', '≤ 5000'],
    ],
    col_widths=[3, 4, 6, 2.5])

add_callout(doc,
    '与 System Prompt 文档的指标区别：System Prompt 关注"角色/格式/安全"；'
    '与 RAG 文档的区别：RAG 关注"召回率/MRR"；'
    '本文档只关注"装载/预算/压缩"，不涉及 Prompt 内容和检索算法。')

add_heading_styled(doc, '6.2 黄金测试集构建', level=2)

add_paragraph_styled(doc,
    '上下文工程的黄金测试集要覆盖"长对话触发压缩"的场景，共 150 条用例：')

add_styled_table(doc,
    ['类别', '用例数', '占比', '示例'],
    [
        ['短对话（≤5 轮）', '30', '20%',
         '不触发压缩，验证基本装载'],
        ['中对话（6-15 轮）', '40', '27%',
         '触发一级压缩，验证摘要质量'],
        ['长对话（16-30 轮）', '40', '27%',
         '触发二三级压缩，验证槽位+Pin'],
        ['超长对话（30+ 轮）', '20', '13%',
         '触发淘汰策略，验证优先级'],
        ['跨意图对话', '20', '13%',
         '中途切换意图，验证预算重分配'],
    ],
    col_widths=[3.5, 2.5, 2, 6.5])

add_heading_styled(doc, '6.3 评估流水线与 CI/CD', level=2)

add_code_block(doc, '''# .github/workflows/context-engineering-eval.yml
name: Context Engineering Evaluation

on:
  pull_request:
    paths:
      - 'agent-service/src/main/java/**/context/**'

jobs:
  evaluate:
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v3
      - name: Setup JDK 17
        uses: actions/setup-java@v3
        with:
          java-version: '17'
      
      - name: Run Assembly Test（六层装载）
        run: ./gradlew test --tests "*ContextAssemblyTest"
      
      - name: Run Budget Test（Token 预算）
        run: ./gradlew test --tests "*TokenBudgetTest"
      
      - name: Run Compression Test（三级压缩）
        run: ./gradlew test --tests "*CompressionTest"
      
      - name: Run Eviction Test（淘汰策略）
        run: ./gradlew test --tests "*EvictionTest"
      
      - name: LLM-as-Judge（压缩质量）
        run: python scripts/context_judge.py --threshold 4.5
      
      - name: Block on Regression
        run: python scripts/block_on_regression.py \\
          --max-token-overflow 0.0 \\
          --min-compression-retention 0.9''')

add_styled_table(doc,
    ['评估阶段', '用例数', '运行时长', '阻断阈值'],
    [
        ['六层装载测试', '30', '~5min', '完整率 < 100%'],
        ['Token 预算测试', '40', '~10min', '超限率 > 0%'],
        ['压缩质量测试', '50', '~20min', '保留率 < 90%'],
        ['淘汰策略测试', '20', '~5min', '淘汰错误率 > 5%'],
        ['LLM-as-Judge', '50', '~15min', '评分 < 4.5/5'],
    ],
    col_widths=[4, 2.5, 3, 5])

add_heading_styled(doc, '6.4 LLM-as-Judge 评估', level=2)

add_paragraph_styled(doc,
    '压缩质量的评估最难自动化——"摘要丢没丢信息"不能靠正则，只能靠 LLM 理解语义。'
    '用 LLM-as-Judge 对比"压缩前历史"和"压缩后历史"，打分信息保留度。')

add_styled_table(doc,
    ['评分维度', '说明', '分值'],
    [
        ['信息保留度', '关键信息（学校/资源/用户意图）是否保留', '0-5'],
        ['语义一致性', '压缩前后语义是否一致', '0-5'],
        ['简洁度', '摘要是否足够简洁', '0-5'],
        ['可读性', '摘要是否可读、不突兀', '0-5'],
    ],
    col_widths=[3, 8, 2.5])

add_code_block(doc, '''# context_judge.py - 压缩质量评估
JUDGE_PROMPT = """你是上下文压缩质量评估员。请对比压缩前后的对话历史，打分：

1. 信息保留度：关键信息（学校/科目/资源编号/用户意图）是否保留
2. 语义一致性：压缩前后语义是否一致
3. 简洁度：摘要是否足够简洁（Token 压缩比 ≤ 20%）
4. 可读性：摘要是否可读、不突兀

压缩前历史：
{original}

压缩后历史：
{compressed}

输出 JSON：
{{
  "info_retention": 0-5,
  "semantic_consistency": 0-5,
  "conciseness": 0-5,
  "readability": 0-5,
  "total": 0-20,
  "lost_info": ["丢失的关键信息1", "..."]
}}"""''')

add_callout(doc,
    '关键设计：评估脚本要输出 lost_info 字段——列出压缩丢失的关键信息。'
    '这比单纯打分更有用，可以指导摘要 Prompt 的迭代（"你丢失了学校信息"→'
    '在摘要 Prompt 里加"必须保留学校"）。')

add_heading_styled(doc, '6.5 测试用例设计', level=2)

add_styled_table(doc,
    ['测试类别', '目标', '用例数', '示例'],
    [
        ['装载测试', '六层正确装配', '30',
         '验证 XML 标签成对/六层都存在'],
        ['预算测试', 'Token 不超限', '40',
         '各意图预算分配正确/超限触发淘汰'],
        ['压缩测试', '压缩质量达标', '50',
         '一级摘要保留率/二级槽位准确/三级 Pin 命中'],
        ['淘汰测试', '优先级正确', '20',
         'P3 先淘汰/P0 永不淘汰'],
        ['跨意图测试', '中途切换意图', '10',
         'HOW_TO→SEARCH 预算重分配'],
    ],
    col_widths=[3, 3, 2, 7])

add_heading_styled(doc, '6.5.1 装载测试用例示例', level=3)

add_styled_table(doc,
    ['用例 ID', '场景', '期望行为', '验收点'],
    [
        ['A-001', 'HOW_TO 意图', '六层都装配，L2 含工具',
         'XML 标签 6 个'],
        ['A-002', 'CHAT 意图', 'L2/L3 省略，其余装配',
         'XML 标签 4 个'],
        ['A-003', '无 RAG 结果', 'L3 省略',
         '无 <context> 标签'],
        ['A-004', '新会话', 'L4 为空',
         '<history> 标签内空'],
    ],
    col_widths=[2, 4, 5, 4])

add_heading_styled(doc, '6.5.2 压缩测试用例示例', level=3)

add_styled_table(doc,
    ['用例 ID', '场景', '期望行为', '验收点'],
    [
        ['C-001', '历史 1500 Token', '不触发压缩',
         '原文保留'],
        ['C-002', '历史 2500 Token', '触发一级摘要',
         '含 [摘要] + 最近 3 轮'],
        ['C-003', '历史 4000 Token', '触发二级槽位',
         '含 [槽位] 字段'],
        ['C-004', '历史 6000 Token', '触发三级 Pin',
         '含 [Pin] 字段'],
        ['C-005', '摘要后仍超', '触发淘汰 P3',
         '早期历史被删'],
    ],
    col_widths=[2, 4, 5, 4])

doc.add_page_break()

# ==================== 6.6 压缩质量专项测试 ====================
add_heading_styled(doc, '6.6 压缩质量专项测试', level=2)

add_callout(doc,
    '压缩质量是上下文工程的核心难题——压缩太狠丢信息，压缩不够超预算。'
    '本节设计专项测试，验证三级压缩在"信息保留 vs Token 压缩"之间的平衡。',
    color='FFF3CD', border_color='FFC107')

add_heading_styled(doc, '6.6.1 一级压缩（Rolling Summary）测试', level=3)

add_styled_table(doc,
    ['测试维度', '测试方法', '通过标准'],
    [
        ['信息保留率', '压缩前后对比关键信息（学校/资源/意图）',
         '保留率 ≥ 90%'],
        ['压缩比', '摘要 Token / 原文 Token',
         '≤ 20%（即 5 倍压缩）'],
        ['摘要可读性', 'LLM-as-Judge 评分',
         '≥ 4.0/5'],
        ['摘要准确性', '摘要是否含幻觉（编造未提的信息）',
         '幻觉率 ≤ 2%'],
        ['最近 N 轮保留', '最近 3 轮是否原文保留',
         '100% 保留'],
    ],
    col_widths=[3, 6, 6])

add_heading_styled(doc, '6.6.2 二级压缩（Slot Freezing）测试', level=3)

add_styled_table(doc,
    ['测试维度', '测试方法', '通过标准'],
    [
        ['槽位提取准确率', '对比人工标注的槽位',
         '准确率 ≥ 95%'],
        ['槽位完整性', '关键槽位（学校/科目/年份/资源）是否齐全',
         '完整率 ≥ 90%'],
        ['槽位值正确性', '槽位值是否正确（如学校=清华）',
         '正确率 ≥ 95%'],
        ['槽位更新及时性', '用户修改条件后槽位是否更新',
         '更新率 100%'],
    ],
    col_widths=[3, 6, 6])

add_heading_styled(doc, '6.6.3 三级压缩（Pin Message）测试', level=3)

add_styled_table(doc,
    ['测试维度', '测试方法', '通过标准'],
    [
        ['Pin 命中率', '被 Pin 的消息在后续对话被引用的比例',
         '命中率 ≥ 80%'],
        ['Pin 误报率', '不该 Pin 的消息被 Pin 的比例',
         '误报率 ≤ 10%'],
        ['Pin 持久性', 'Pin 的消息在后续压缩中是否保留',
         '保留率 100%'],
        ['Agent 自动 Pin 准确率', 'Agent 自动 Pin 的消息是否合理',
         '准确率 ≥ 85%'],
    ],
    col_widths=[3, 6, 6])

add_heading_styled(doc, '6.6.4 压缩回归测试', level=3)

add_paragraph_styled(doc,
    '每次压缩算法变更，都要跑回归测试——确保新算法不退化。'
    '回归测试用 50 条固定长对话，对比新旧算法的压缩质量。')

add_styled_table(doc,
    ['回归维度', '基线', '新版要求', '阻断条件'],
    [
        ['信息保留率', '90%', '≥ 90%', '< 90%'],
        ['压缩比', '20%', '≤ 22%', '> 25%'],
        ['LLM-as-Judge 评分', '4.5/5', '≥ 4.5/5', '< 4.5'],
        ['幻觉率', '2%', '≤ 2%', '> 3%'],
    ],
    col_widths=[4, 3, 4, 4])

doc.add_page_break()

# ==================== 6.7 A/B 测试设计 ====================
add_heading_styled(doc, '6.7 A/B 测试设计', level=2)

add_paragraph_styled(doc,
    '上下文工程的 A/B 测试比"代码 A/B"复杂——LLM 输出有随机性，且压缩质量是"相对值"'
    '（压缩前后对比），需要双盲评估。', bold=True)

add_heading_styled(doc, '6.7.1 A/B 测试场景', level=3)

add_styled_table(doc,
    ['测试场景', '版本 A', '版本 B', '主要指标'],
    [
        ['预算分配策略', '固定预算', '按意图动态预算',
         'Token 利用率 / 准确率'],
        ['压缩触发阈值', '6K 触发', '7K 触发',
         '压缩频率 / 信息保留'],
        ['摘要 Prompt', '旧版摘要', '新版摘要（加槽位要求）',
         '信息保留率'],
        ['保留最近 N 轮', '保留 3 轮', '保留 5 轮',
         '准确率 / Token 成本'],
        ['Pin 策略', 'Agent 自动 Pin', '用户手动 Pin',
         'Pin 命中率'],
    ],
    col_widths=[3.5, 3.5, 3.5, 4])

add_heading_styled(doc, '6.7.2 样本量计算', level=3)

add_paragraph_styled(doc,
    '上下文工程的 A/B 测试样本量计算要考虑"长对话"——'
    '每个测试用例是"一段 20 轮对话"，不是"一条 query"。')

add_code_block(doc, '''# 样本量计算（连续型指标）
import math
from scipy import stats

def sample_size_continuous(sigma, mde, alpha=0.05, power=0.8):
    """连续型指标样本量（如信息保留率）"""
    z_alpha = stats.norm.ppf(1 - alpha / 2)
    z_beta = stats.norm.ppf(power)
    n = 2 * ((z_alpha + z_beta) * sigma / mde) ** 2
    return math.ceil(n)

# 示例：信息保留率标准差 0.05，希望检测 2pp 差异
n = sample_size_continuous(sigma=0.05, mde=0.02)
print(f"每组最小样本量: {n}")  # 输出: 每组 99

# 每个测试用例是 20 轮对话
print(f"需要的对话数: {n}")  # 输出: 99 段对话''')

add_styled_table(doc,
    ['指标', '标准差', 'MDE', '每组样本量'],
    [
        ['信息保留率', '0.05', '2pp', '99'],
        ['压缩比', '0.03', '3pp', '52'],
        ['Token 利用率', '0.08', '5pp', '201'],
        ['LLM-as-Judge 评分', '0.3', '0.2', '71'],
    ],
    col_widths=[4, 3, 3, 4])

add_heading_styled(doc, '6.7.3 灰度发布阶梯', level=3)

add_styled_table(doc,
    ['阶段', '流量', '持续时间', '准入指标', '回滚条件'],
    [
        ['Stage 1: 内测', '5%', '5 天',
         '压缩质量持平', '信息保留率下降 > 3pp'],
        ['Stage 2: 小流量', '10%', '7 天',
         'Token 成本持平或下降', '成本上升 > 10%'],
        ['Stage 3: 中流量', '30%', '7 天',
         '所有指标稳定', '任一指标退化'],
        ['Stage 4: 全量', '100%', '-',
         '所有指标稳定', '-'],
    ],
    col_widths=[3, 2.5, 2.5, 4, 4])

add_callout(doc,
    '上下文工程 A/B 的特殊性：要观察"长对话"——短对话看不出压缩差异。'
    '灰度至少 7 天，让用户有机会产生 20+ 轮的长对话，才能验证压缩质量。')

doc.add_page_break()

# ==================== 6.8 验收流程与准入准出 ====================
add_heading_styled(doc, '6.8 验收流程与准入准出', level=2)

add_heading_styled(doc, '6.8.1 四阶段验收流程', level=3)

add_styled_table(doc,
    ['阶段', '负责人', '检查项', '通过标准'],
    [
        ['1. 装载验收', '后端工程师',
         '六层装载 / XML 标签 / 游离信息',
         '完整率 100% / 标签 100% / 游离 0%'],
        ['2. 预算验收', '后端工程师',
         'Token 不超限 / 预算利用率',
         '超限率 0% / 利用率 70-90%'],
        ['3. 压缩验收', '算法工程师 + PM',
         '三级压缩质量 / LLM-as-Judge',
         '保留率 ≥ 90% / 评分 ≥ 4.5'],
        ['4. 准出观测', 'SRE',
         '全量后 7 天观测',
         '所有指标稳定 / 无成本飙升'],
    ],
    col_widths=[2.5, 3, 5, 5])

add_heading_styled(doc, '6.8.2 准入清单（PR 合并前）', level=3)

add_styled_table(doc,
    ['类别', '检查项', '阈值'],
    [
        ['装载', '六层装载完整率', '100%'],
        ['装载', 'XML 标签正确率', '100%'],
        ['装载', '游离信息率', '0%'],
        ['预算', '总 Token 超限率', '0%'],
        ['预算', '预算利用率', '70-90%'],
        ['预算', '淘汰触发率', '≤ 5%'],
        ['压缩', '摘要信息保留率', '≥ 90%'],
        ['压缩', '摘要压缩比', '≤ 20%'],
        ['压缩', '槽位准确率', '≥ 95%'],
        ['压缩', 'LLM-as-Judge 评分', '≥ 4.5/5'],
        ['性能', '装配延迟 P95', '≤ 200ms'],
        ['性能', '压缩延迟 P95', '≤ 1.5s'],
        ['成本', 'Prefix Cache 命中率', '≥ 95%'],
        ['成本', '平均输入 Token', '≤ 5000'],
        ['安全', '敏感词扫描', '0 命中'],
        ['文档', 'ADR 更新', '有变更必更新'],
    ],
    col_widths=[3, 7, 5])

add_heading_styled(doc, '6.8.3 准出清单（全量后 7 天观测）', level=3)

add_styled_table(doc,
    ['观测项', '观测窗口', '达标线', '不达标处理'],
    [
        ['Token 超限率', '7 天', '0%', '回滚'],
        ['压缩信息保留率', '7 天滑动', '≥ 90%', '优化摘要 Prompt'],
        ['Prefix Cache 命中率', '7 天均值', '≥ 95%', '检查 L0-L2 是否被改'],
        ['平均输入 Token', '7 天均值', '≤ 5000', '检查预算分配'],
        ['单次对话成本', '7 天均值', '≤ 基线', '回滚'],
        ['P0 Bug', '7 天', '0 起', '修复后重新灰度'],
    ],
    col_widths=[3.5, 3, 3, 5])

add_heading_styled(doc, '6.8.4 验收报告模板', level=3)

add_code_block(doc, '''# 《上下文工程 vX.X.X 验收报告》

## 1. 版本信息
- 版本号：v1.1.0
- 发布日期：2026-07-05
- 变更说明：增加二级槽位冻结压缩
- 负责人：@xxx

## 2. 评估结果
| 指标 | 基线 | 当前 | 变化 | 达标 |
|------|------|------|------|------|
| 六层装载完整率 | 100% | 100% | - | ✓ |
| Token 超限率 | 0% | 0% | - | ✓ |
| 摘要信息保留率 | 85% | 92% | +7pp | ✓ |
| 槽位准确率 | - | 96% | 新增 | ✓ |
| LLM-as-Judge 评分 | 4.3 | 4.6 | +0.3 | ✓ |
| 平均输入 Token | 6000 | 4800 | -1200 | ✓ |
| Prefix Cache 命中率 | 92% | 96% | +4pp | ✓ |

## 3. 压缩质量测试
- 一级摘要保留率：92%（基线 85%，+7pp）
- 二级槽位准确率：96%（新增）
- 三级 Pin 命中率：83%（基线 80%，+3pp）

## 4. 灰度结果
- Stage 1（5%）: 5 天，信息保留率 +5pp
- Stage 2（10%）: 7 天，信息保留率 +7pp
- Stage 3（30%）: 7 天，所有指标稳定

## 5. 结论
✅ 通过验收，建议全量发布。

## 6. 后续跟进
- 槽位提取的边界 case 纳入下版迭代
- 持续监控 7 天后关闭项目''')

doc.add_page_break()

# ==================== 6.9 持续监控与漂移检测 ====================
add_heading_styled(doc, '6.9 持续监控与漂移检测', level=2)

add_paragraph_styled(doc,
    '上下文工程上线后，监控的核心是"压缩质量漂移"和"Token 成本失控"——'
    'LLM 厂商静默更新模型可能导致摘要质量退化，用户行为变化可能导致长对话比例上升。', bold=True)

add_heading_styled(doc, '6.9.1 监控指标体系', level=3)

add_styled_table(doc,
    ['类别', '指标', '采集方式', '告警阈值'],
    [
        ['装载', '六层装载完整率', '日志采样', '< 100%'],
        ['装载', 'XML 标签正确率', 'XML 解析器', '< 100%'],
        ['预算', 'Token 超限率', '日志统计', '> 0%'],
        ['预算', '预算利用率', '统计', '< 60% 或 > 95%'],
        ['预算', '淘汰触发率', '日志统计', '> 10%'],
        ['压缩', '摘要信息保留率', 'LLM-as-Judge 抽样', '< 85%'],
        ['压缩', '槽位准确率', '人工抽检', '< 90%'],
        ['压缩', 'Pin 命中率', '日志统计', '< 75%'],
        ['性能', '装配延迟 P95', '链路追踪', '> 300ms'],
        ['性能', '压缩延迟 P95', '链路追踪', '> 2s'],
        ['成本', 'Prefix Cache 命中率', 'API 返回', '< 90%'],
        ['成本', '平均输入 Token', '统计', '> 6000'],
        ['成本', '单次对话成本', 'API 账单', '> 基线 × 1.2'],
    ],
    col_widths=[3, 4, 4, 4])

add_heading_styled(doc, '6.9.2 压缩质量漂移检测（影子评估）', level=3)

add_callout(doc,
    '⚠️ 压缩质量漂移是上下文工程最大的隐患。'
    '原因：(1) LLM 厂商静默更新模型，摘要能力可能退化；(2) 摘要 Prompt 微调；(3) 用户长对话比例变化。'
    '我们用"影子评估"持续检测——线上流量采样 1%，对比新旧压缩算法的质量差。',
    color='FDE7E9', border_color='E53935')

add_styled_table(doc,
    ['步骤', '动作', '频率', '目的'],
    [
        ['1', '线上长对话流量采样 1%（脱敏）', '实时', '获取真实长对话'],
        ['2', '并行跑当前压缩 + 上一稳定版', '实时', '对比压缩质量'],
        ['3', 'LLM-as-Judge 评估两个版本摘要', '实时', '计算信息保留率差'],
        ['4', '差值超阈值 → 触发漂移告警', '实时', '人工介入'],
        ['5', '归档到漂移日志，按周生成报告', '每周', '长期趋势'],
    ],
    col_widths=[1.5, 6, 2.5, 5])

add_heading_styled(doc, '6.9.3 Token 成本监控', level=3)

add_paragraph_styled(doc,
    'Token 成本是上下文工程最直接的"北极星指标"——成本上升说明压缩失效或预算失控。')

add_styled_table(doc,
    ['监控项', '期望值', '告警阈值', '排查方向'],
    [
        ['平均输入 Token', '≤ 5000', '> 6000',
         '压缩失效 / 长对话增加'],
        ['Prefix Cache 命中率', '≥ 95%', '< 90%',
         'L0-L2 被改 / 拼接顺序变化'],
        ['单次对话成本', '基线', '> 基线 × 1.2',
         'Token 超限 / 压缩失效'],
        ['压缩触发率', '10-20%', '> 30%',
         '长对话激增 / 阈值过低'],
    ],
    col_widths=[4, 3, 3, 5])

add_heading_styled(doc, '6.9.4 长对话比例监控', level=3)

add_paragraph_styled(doc,
    '长对话比例是上下文工程的"领先指标"——长对话增加意味着压缩压力增大，'
    '提前预警可以避免压缩质量退化。')

add_styled_table(doc,
    ['对话轮次分布', '期望占比', '告警阈值', '影响'],
    [
        ['1-5 轮', '60-70%', '< 50%', '短对话减少'],
        ['6-15 轮', '20-30%', '> 40%', '一级压缩压力'],
        ['16-30 轮', '5-10%', '> 20%', '二三级压缩压力'],
        ['30+ 轮', '< 5%', '> 10%', '淘汰策略压力'],
    ],
    col_widths=[3.5, 3, 3, 5.5])

add_callout(doc,
    '长对话比例突然上升时，要警惕——可能是产品功能变化（如新增"连续追问"功能）'
    '导致用户行为改变。这时要重新评估预算分配策略，而不是被动等压缩失效。')

doc.add_page_break()

# ==================== 七、总结与边界声明 ====================
add_heading_styled(doc, '七、总结与边界声明', level=1)

add_heading_styled(doc, '7.1 核心总结', level=2)

add_paragraph_styled(doc,
    '上下文工程是 Agent 的"工作记忆管理"——它决定每次调用 LLM 时喂什么、喂多少、太长了怎么压缩。'
    '本文档专注讨论这一个细小方向，核心要点：', bold=True)

add_styled_table(doc,
    ['维度', '核心决策', '关键 ADR'],
    [
        ['分层装载', 'L0-L5 六层（System/画像/工具/RAG/历史/当前）',
         'ADR-CTX-01'],
        ['Token 预算', '固定 8K，按意图动态分配',
         'ADR-CTX-02, 03'],
        ['压缩策略', '三级渐进（Summary + Slot + Pin）',
         'ADR-CTX-04'],
        ['分层标记', 'XML 标签（<context> <history>）',
         'ADR-CTX-05'],
        ['压缩触发', '6K 阈值（8K 的 75%）',
         'ADR-CTX-06'],
        ['质量评估', 'LLM-as-Judge（信息保留率）',
         'ADR-CTX-07'],
    ],
    col_widths=[3, 8, 4])

add_heading_styled(doc, '7.2 本文档与其他文档的关系', level=2)

add_styled_table(doc,
    ['顺序', '方向', '文档', '状态'],
    [
        ['1', 'System Prompt 工程',
         '《System Prompt 工程模块设计方案》', '✅ 已完成'],
        ['2', 'RAG 检索增强',
         '《RAG 检索增强生成模块设计方案》', '✅ 已完成'],
        ['3', '意图识别',
         '《意图识别模块设计方案》', '✅ 已完成'],
        ['4', '上下文工程（本文档）',
         '《上下文工程模块设计方案》', '✅ 本文档'],
        ['5', '对话编排',
         '《对话编排模块设计方案》', '⏳ 待规划'],
        ['6', '工具调用',
         '《工具调用模块设计方案》', '⏳ 待规划'],
        ['7', '长期记忆',
         '《长期记忆模块设计方案》', '⏳ 待规划'],
    ],
    col_widths=[1.5, 4, 6, 3.5])

add_callout(doc,
    '本文档与上下游的关系：'
    '上游——承接 System Prompt（L0 层）、RAG（L3 层）、意图识别（决定预算分配）；'
    '下游——为对话编排提供组装好的 messages 数组。'
    '本文档不展开上下游的实现，只讨论"如何组装"。')

add_heading_styled(doc, '7.3 演进路线', level=2)

add_styled_table(doc,
    ['阶段', '时间', '目标', '关键能力'],
    [
        ['Phase 1: MVP（当前）', '已完成',
         '能用', '滑动窗口截断'],
        ['Phase 2: 工程化', '2026 Q3',
         '可控预算', '六层 + 预算 + 三级压缩'],
        ['Phase 3: 智能化', '2026 Q4',
         '自适应压缩', '按对话内容动态调整压缩级别'],
        ['Phase 4: 个性化', '2027 Q1',
         '千人千面', '按用户画像调整预算分配'],
    ],
    col_widths=[3, 2, 3, 7])

add_heading_styled(doc, '7.4 结语', level=2)

add_paragraph_styled(doc,
    '上下文工程不是"塞满窗口"，而是"在有限预算内最大化信息密度"。'
    '它是 Agent 的"工作记忆管理"——压缩不是"删旧消息"，而是"保留信息密度"。'
    '把这一块做扎实，Agent 才能在长对话中保持连贯；做不扎实，第 10 轮就开始"失忆"。', bold=True)

add_callout(doc,
    '最后一句：你的 Agent 不是被 LLM 的窗口限制的，而是被你的上下文工程限制的。'
    'LLM 窗口是天花板，上下文工程是实际利用率——天花板再高，不会管理也是浪费。',
    color='E8F4FD', border_color='2196F3')

doc.add_page_break()

# ==================== 附录：ADR 摘要 ====================
add_heading_styled(doc, '附录：ADR 摘要', level=1)

add_paragraph_styled(doc,
    'ADR = Architecture Decision Record（架构决策记录），每条 ADR 包含'
    '"上下文 / 决策 / 后果"三段式。本附录列出本文档引用的所有 ADR 摘要。')

add_styled_table(doc,
    ['ADR 编号', '决策', '上下文', '后果'],
    [
        ['ADR-CTX-01',
         '采用 L0-L5 六层分层装载',
         '明确每层职责，避免信息混杂',
         '结构清晰，但需要严格归层，增加开发成本'],
        ['ADR-CTX-02',
         'Token 预算固定 8K',
         '平衡成本与质量；超 8K 后质量下降明显',
         '成本可控，但超长对话需要更激进压缩'],
        ['ADR-CTX-03',
         '按意图动态分配预算',
         'HOW_TO 多分历史，SEARCH 多分 RAG',
         '意图识别错误会导致预算分配错误'],
        ['ADR-CTX-04',
         '三级渐进压缩',
         '保留关键事实，渐进压缩而非一刀切',
         '工程量大，但信息保留率高'],
        ['ADR-CTX-05',
         '用 XML 标签分层',
         '让 LLM 明确每段性质，防隐式注入',
         '增加 ~100 Token 标签成本，但提升 LLM 理解'],
        ['ADR-CTX-06',
         '压缩触发阈值 6K',
         '提前压缩，避免到 8K 才被动截断',
         '压缩频率高，但避免被动截断的信息丢失'],
        ['ADR-CTX-07',
         '压缩质量用 LLM-as-Judge 评估',
         '摘要不能丢关键信息，必须可量化',
         '增加评估成本，但可发现压缩退化'],
    ],
    col_widths=[3, 4, 5, 5])

# ==================== 保存文档 ====================
doc.save(r'e:\workspace_work\CampusShare\docs\agent-design\上下文工程模块设计方案.docx')
print('文档已生成: 上下文工程模块设计方案.docx')
