# -*- coding: utf-8 -*-
"""
生成《CampusShare Agent 意图识别模块设计方案》Word 文档
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, color_hex):
    """设置单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_code_block(doc, code_text, language='java'):
    """添加代码块（灰色背景等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # 设置段落背景色
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    # 添加边框
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border)
    pPr.append(pBdr)

    lines = code_text.split('\n')
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if i < len(lines) - 1:
            run.add_break()
    return p


def add_callout(doc, text, color='FFF3CD', border_color='FFC107'):
    """添加提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), border_color)
        pBdr.append(border)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def style_table_header(row, bg_color='2563EB'):
    """设置表头样式"""
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_styled_table(doc, headers, rows, col_widths=None, header_bg='2563EB'):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    style_table_header(table.rows[0], header_bg)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
            # 隔行变色
            if r_idx % 2 == 1:
                set_cell_background(row_cells[c_idx], 'F8F9FA')

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            run.font.size = Pt(20)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            run.font.size = Pt(13)
    return h


def add_paragraph_styled(doc, text, bold=False, size=10.5, color=None, indent=None):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_bullet(doc, text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_numbered(doc, text):
    """添加编号列表"""
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


# ==================== 开始生成文档 ====================
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(10.5)

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ==================== 封面 ====================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CampusShare Agent\n意图识别模块设计方案')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('从平铺式 RAG 到意图驱动的智能路由')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

for _ in range(8):
    doc.add_paragraph()

# 文档信息表
info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('文档版本', 'v1.0'),
    ('文档日期', '2026-07-04'),
    ('文档状态', '设计中'),
    ('适用范围', 'campushare-agent 服务'),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    for p in info_table.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)
    for p in info_table.rows[i].cells[1].paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)
    set_cell_background(info_table.rows[i].cells[0], 'F0F4FF')

doc.add_page_break()

# ==================== 目录页 ====================
add_heading_styled(doc, '目录', level=1)

toc_items = [
    '一、场景：为什么需要意图识别',
    '    1.1 业务背景：当前 Agent 的工作方式',
    '    1.2 用户提问类型分析',
    '    1.3 平铺式 RAG 的三大痛点',
    '    1.4 意图识别的核心价值',
    '二、方案：业界做法与大厂思考',
    '    2.1 意图识别技术演进路线',
    '    2.2 六种技术方案对比',
    '    2.3 大厂实践案例分析',
    '    2.4 CampusShare 选型决策',
    '三、流程：如何搭建意图识别',
    '    3.1 前置条件',
    '    3.2 意图分类体系设计',
    '    3.3 整体架构与请求流转',
    '    3.4 详细实现步骤',
    '    3.5 关键设计决策（ADR）',
    '四、核心代码',
    '    4.1 文件架构',
    '    4.2 意图定义与 DTO',
    '    4.3 规则短路层',
    '    4.4 LLM 意图分类器',
    '    4.5 Prompt 模板',
    '    4.6 意图路由器',
    '    4.7 AgentChatService 改造',
    '    4.8 Redis 缓存层',
    '    4.9 监控指标',
    '五、目标：实现效果',
    '    5.1 功能目标',
    '    5.2 性能目标',
    '    5.3 体验目标',
    '    5.4 成本目标',
    '    5.5 评估指标与验收标准',
    '六、总结与展望',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    if not item.startswith('    '):
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ==================== 一、场景 ====================
add_heading_styled(doc, '一、场景：为什么需要意图识别', level=1)

add_heading_styled(doc, '1.1 业务背景：当前 Agent 的工作方式', level=2)

add_paragraph_styled(doc,
    'CampusShare AI 智能助手目前处于 MVP 阶段，已实现基于 DeepSeek-V3 的 SSE 流式对话、'
    'RAG 知识库检索（18 篇平台文档 + 帖子向量）、多轮会话管理等核心能力。')

add_paragraph_styled(doc,
    '但当前的工作方式是"平铺式 RAG"——无论用户问什么，都走同一条处理链路：', bold=True)

add_code_block(doc, """当前流程（无意图识别）：
用户消息
  → RetrievalService.retrieve()  // 三路并行检索（知识向量+关键词+帖子向量）
  → RRF 融合 Top-5
  → 注入系统提示词作为上下文
  → DeepSeekClient.chatCompletionStream()  // 流式生成
  → 持久化 agent_turns

问题：所有问题走同一条管线，"一刀切"处理""")

add_paragraph_styled(doc, '这种"一刀切"模式在 MVP 早期可以快速验证可用性，但随着用户量增长，问题逐渐暴露。')

add_heading_styled(doc, '1.2 用户提问类型分析', level=2)

add_paragraph_styled(doc, '通过对平台用户提问场景的调研，我们将用户问题归纳为 5 大类：')

add_styled_table(doc,
    ['意图', '占比预估', '典型示例', '理想处理方式'],
    [
        ['HOW_TO\n操作指引', '35%', '"怎么发帖？""如何修改密码？""单向消息是什么意思？"', '知识库检索 + 步骤生成'],
        ['SEARCH\n内容检索', '30%', '"求清华操作系统期末卷子""有没有考研面经"', '帖子向量检索 + 结构化过滤'],
        ['NAVIGATE\n页面导航', '15%', '"个人中心在哪？""我点赞的帖子怎么找"', '直接返回跳转卡片'],
        ['CLARIFY\n多轮澄清', '10%', '"那个有下载的""不对，我要讨论帖"', '上下文消解 + 筛选'],
        ['OUT_OF_SCOPE\n超范围', '10%', '"你好""今天天气怎么样""帮我发帖"', '模板回复 / 礼貌拒绝'],
    ],
    col_widths=[2.5, 1.8, 6, 4.5])

add_heading_styled(doc, '1.3 平铺式 RAG 的三大痛点', level=2)

add_heading_styled(doc, '痛点一：简单问题过度处理，响应慢、成本高', level=3)
add_paragraph_styled(doc,
    '用户问"你好"或"个人中心在哪"，本可以毫秒级模板回复，当前却要走完整的 RAG 检索 + LLM 流式生成。'
    '以"你好"为例：')
add_bullet(doc, '当前流程：Embedding 调用（~200ms）+ 三路 PG 检索（~50ms）+ DeepSeek 流式生成（~1.5s 首字）= 总耗时 ~1.8s')
add_bullet(doc, '理想流程：规则匹配"你好" → 直接返回"你好！我是 CampusShare AI 助手..." = 总耗时 <10ms')
add_bullet(doc, '浪费：1 次 Embedding API 调用 + 3 次 PG 查询 + 1 次 LLM 流式调用（约 200 token）')

add_heading_styled(doc, '痛点二：复杂问题检索不足，召回率低', level=3)
add_paragraph_styled(doc,
    '用户问"求清华操作系统期末卷子"，这是一个典型的 SEARCH/resource 意图，'
    '理想情况应该：')
add_bullet(doc, '抽取结构化槽位：school=清华, query=操作系统期末卷子, post_type=resource')
add_bullet(doc, '用槽位做结构化过滤（WHERE school_id=? AND post_type=?）+ 向量召回')
add_bullet(doc, '可选启用 HyDE（假设文档检索）提升短查询召回')
add_paragraph_styled(doc,
    '但当前流程把原始 query 直接向量化检索，没有结构化过滤，'
    '可能召回北大、复旦的操作系统资料，精准度不足。', bold=True)

add_heading_styled(doc, '痛点三：超范围问题浪费 LLM 成本', level=3)
add_paragraph_styled(doc,
    '约 10% 的提问是闲聊、开放域问题或写操作请求（"帮我发帖""帮我点赞"）。'
    '当前这些请求也走 LLM 生成，每次消耗约 200-500 token，按 DeepSeek 定价虽便宜但积少成多。'
    '更严重的是写操作请求——Agent 本应明确拒绝（平台规定不代替用户执行写操作），'
    '但 LLM 可能给出模糊回答，用户体验差。')

add_heading_styled(doc, '1.4 意图识别的核心价值', level=2)

add_paragraph_styled(doc, '意图识别是 Agent 从"能用"到"好用"的关键一跃，核心价值体现在三个维度：',
                     bold=True)

add_styled_table(doc,
    ['维度', '无意图识别（当前）', '有意图识别（目标）', '提升'],
    [
        ['响应延迟', '所有问题 ~1.8s 首字', '简单问题 <100ms，复杂问题 ~1.5s', '简单问题提速 18x'],
        ['LLM 成本', '每问必调 LLM', '约 25% 问题不调 LLM（NAVIGATE + OUT_OF_SCOPE）', '成本降低 ~30%'],
        ['回答精准度', '检索无差别召回', '按意图选择检索策略 + 结构化过滤', '召回精准度提升 ~40%'],
    ],
    col_widths=[2.5, 4, 4.5, 3.5])

add_callout(doc,
    '💡 核心洞察：意图识别的本质是"分类后路由"——先理解用户想干什么，'
    '再决定走哪条处理链路。这就像医院分诊台：发烧去内科，骨折去外科，'
    '而不是所有病人都去做全套检查。',
    color='E8F4FD', border_color='2196F3')

doc.add_page_break()

# ==================== 二、方案 ====================
add_heading_styled(doc, '二、方案：业界做法与大厂思考', level=1)

add_heading_styled(doc, '2.1 意图识别技术演进路线', level=2)

add_paragraph_styled(doc, '意图识别技术经历了三个阶段的演进：')

add_styled_table(doc,
    ['阶段', '时间', '代表技术', '典型精度', '维护成本'],
    [
        ['规则时代', '2015 前', '关键词匹配、正则、决策树', '60-75%', '高（规则爆炸）'],
        ['ML 时代', '2015-2022', 'SVM、BERT、TextCNN', '85-92%', '中（需标注数据）'],
        ['LLM 时代', '2023 至今', 'GPT/DeepSeek Prompt、Function Calling', '92-97%', '低（Prompt 即规则）'],
    ],
    col_widths=[2.2, 2, 4.5, 2, 2.5])

add_paragraph_styled(doc,
    'LLM 时代的核心变化：意图识别从"训练一个分类模型"变成"写好一个 Prompt"。'
    '这大幅降低了冷启动门槛——不再需要标注数千条训练数据，只需定义好意图体系和 Few-shot 示例即可。',
    bold=True)

add_heading_styled(doc, '2.2 六种技术方案对比', level=2)

add_paragraph_styled(doc, '我们调研了 6 种主流意图识别方案，从准确率、延迟、成本、维护性、冷启动 5 个维度对比：')

add_styled_table(doc,
    ['方案', '原理', '准确率', '延迟', '成本', '冷启动', '维护性'],
    [
        ['A. 规则匹配', '关键词 + 正则 + 优先级树', '70%', '<5ms', '极低', '即用', '差（规则爆炸）'],
        ['B. 传统 ML', 'BERT 微调分类头', '90%', '~50ms', '低（自部署）', '需 1k+ 标注', '中（需重训）'],
        ['C. Embedding 相似度', 'query 与意图描述向量余弦相似', '82%', '~80ms', '低', '即用', '好'],
        ['D. LLM Prompt 分类', 'LLM 输出 JSON {intent, confidence}', '94%', '~300ms', '中（API 调用）', '即用', '极好'],
        ['E. LLM Function Calling', 'LLM 隐式选择 tool 触发意图', '93%', '~400ms', '中', '即用', '极好'],
        ['F. 混合多层漏斗', '规则短路 + LLM 兜底 + Embedding 降级', '95%', '分层（5ms~300ms）', '中', '即用', '好'],
    ],
    col_widths=[2.8, 4, 1.5, 1.5, 1.5, 1.5, 1.5])

add_heading_styled(doc, '方案 A：基于规则的关键词匹配', level=3)
add_paragraph_styled(doc, '原理：维护关键词到意图的映射表，命中即分类。')
add_code_block(doc, """// 规则匹配示例
if (query.contains("怎么") || query.contains("如何")) return HOW_TO;
if (query.contains("求") || query.contains("找")) return SEARCH;
if (query.contains("你好") || query.contains("谢谢")) return OUT_OF_SCOPE;""")
add_paragraph_styled(doc, '优点：极快（<5ms）、零成本、可解释。')
add_paragraph_styled(doc, '缺点：用户表达多样（"咋/咋整/咋关"），规则爆炸，维护噩梦，准确率仅 70%。')

add_heading_styled(doc, '方案 B：基于传统 ML 模型', level=3)
add_paragraph_styled(doc, '原理：用 BERT 等预训练模型，在标注数据上微调一个分类头。')
add_paragraph_styled(doc, '优点：准确率高（90%+）、延迟低（~50ms 自部署）。')
add_paragraph_styled(doc,
    '缺点：需要标注 1000+ 条训练数据（冷启动难）；模型迭代需重新训练；'
    '部署需 GPU 或量化；多意图体系调整需重训。')

add_heading_styled(doc, '方案 C：基于 Embedding 语义相似度', level=3)
add_paragraph_styled(doc,
    '原理：为每个意图写一段描述（如"用户想找资料、求资源、找帖子"），'
    '将用户 query 和意图描述都向量化，取余弦相似度最高的意图。')
add_code_block(doc, """// Embedding 相似度示例
intent_descriptions = {
    HOW_TO: "用户询问怎么使用功能、如何操作、在哪设置",
    SEARCH: "用户想找资料、求资源、找帖子、有没有",
    NAVIGATE: "用户想知道某个功能在哪、入口在哪、怎么找自己的内容",
    ...
}
query_vec = embeddingClient.embed(query)
scores = {intent: cosine(query_vec, embed(desc)) for intent, desc in intent_descriptions}
best_intent = max(scores, key=scores.get)""")
add_paragraph_styled(doc, '优点：无需训练、即用、可动态调整意图描述。')
add_paragraph_styled(doc, '缺点：准确率偏低（~82%），意图间语义重叠时易混淆（HOW_TO vs NAVIGATE）。')

add_heading_styled(doc, '方案 D：基于 LLM Prompt 的意图分类', level=3)
add_paragraph_styled(doc, '原理：用 LLM 理解用户意图，输出结构化 JSON。')
add_code_block(doc, """// LLM Prompt 分类示例
System: 你是意图分类器。判断用户问题的意图，输出 JSON：
{"intent": "HOW_TO|SEARCH|NAVIGATE|CLARIFY|OUT_OF_SCOPE",
 "sub_intent": "...",
 "confidence": 0.0-1.0,
 "rewritten_query": "改写后的查询"}

User: 求清华操作系统期末卷子
Assistant: {"intent":"SEARCH","sub_intent":"resource","confidence":0.95,"rewritten_query":"操作系统 期末 卷子","slots":{"school":"清华","post_type":"resource"}}""")
add_paragraph_styled(doc, '优点：准确率高（94%+）、零冷启动、Prompt 即规则易迭代、可同时做查询改写。')
add_paragraph_styled(doc, '缺点：延迟较高（~300ms）、有 API 成本、依赖 LLM 可用性。')

add_heading_styled(doc, '方案 E：基于 LLM Function Calling', level=3)
add_paragraph_styled(doc,
    '原理：把每个意图定义为 LLM 的 tool/function，让 LLM 隐式选择 tool 触发对应意图。'
    '这是 OpenAI 推荐的 Agentic 模式。')
add_paragraph_styled(doc,
    '优点：与 ReAct 循环天然契合、可执行实际动作。'
    '缺点：当前 DeepSeek Function Calling 稳定性待验证；'
    'MVP 阶段无工具执行框架；延迟更高（~400ms+）。')

add_heading_styled(doc, '方案 F：混合多层漏斗（推荐）', level=3)
add_paragraph_styled(doc,
    '原理：多层漏斗，逐层过滤。规则层处理高确定性请求（短路省成本），'
    'LLM 层处理复杂分类，Embedding 层作为 LLM 降级方案。',
    bold=True)

add_code_block(doc, """混合漏斗架构：
请求 → [Layer 1: 规则短路]  → 命中? → 直接返回意图（<5ms）
         ↓ 未命中
       [Layer 2: LLM 分类]  → confidence ≥ 0.6? → 返回意图（~300ms）
         ↓ confidence < 0.6 或 LLM 不可用
       [Layer 3: Embedding 兜底] → 返回最相似意图（~80ms）
         ↓ 全部失败
       [Default: SEARCH]  → 最通用管线兜底""")

add_heading_styled(doc, '2.3 大厂实践案例分析', level=2)

add_heading_styled(doc, '案例一：阿里小蜜——3 层漏斗架构', level=3)
add_paragraph_styled(doc,
    '阿里小蜜日均承接千万级对话，采用经典的 3 层漏斗：')
add_bullet(doc, 'Layer 1 规则层（30% 流量）：高频确定性意图（催发货、查物流）直接命中，<10ms')
add_bullet(doc, 'Layer 2 浅层模型（40% 流量）：FastText + CNN 做粗分类，~30ms')
add_bullet(doc, 'Layer 3 深度模型（30% 流量）：BERT + 任务型对话管理，~80ms')
add_paragraph_styled(doc,
    '启示：高流量场景必须分层，规则层能省 30% 的模型调用量。'
    'CampusShare 借鉴此思路，用规则短路处理"你好/帮我发帖/我点赞的"等高频确定性请求。',
    bold=True)

add_heading_styled(doc, '案例二：腾讯客服——BERT 分类 + 槽位抽取', level=3)
add_paragraph_styled(doc,
    '腾讯智能客服采用 BERT 做意图分类，同时用 BERT 做槽位抽取（NER），'
    '两者共享 BERT 编码器。准确率达 92%+，但需要标注 5 万+ 条数据。')
add_paragraph_styled(doc,
    '启示：槽位抽取与意图分类可合并（CampusShare 在 LLM 方案中也合并，但用 Prompt 而非双 BERT）。')

add_heading_styled(doc, '案例三：字节豆包——LLM 原生意图识别', level=3)
add_paragraph_styled(doc,
    '豆包采用 LLM 原生方案：不做显式意图分类，直接用 LLM 的 Function Calling 决定走哪个工具。'
    '优势是灵活（无需预定义意图体系），劣势是延迟高、成本高、不可控。')
add_paragraph_styled(doc,
    '启示：纯 LLM 方案适合通用助手，但垂直领域（如 CampusShare）'
    '显式意图分类更可控、更省成本。', bold=True)

add_heading_styled(doc, '案例四：OpenAI ChatGPT——隐式意图 + 路由', level=3)
add_paragraph_styled(doc,
    'GPT-4 内部完成意图理解，通过 Function Calling 选择工具（浏览、代码执行、DALL-E）。'
    '用户感知不到"意图分类"这一步，但底层就是意图路由。')
add_paragraph_styled(doc,
    '启示：对终端用户而言，意图识别应当是"隐形"的——用户只感受到"回答更快更准"，'
    '不需要知道背后做了分类。CampusShare 的设计也遵循此原则。')

add_heading_styled(doc, '2.4 CampusShare 选型决策', level=2)

add_heading_styled(doc, '业务约束分析', level=3)
add_styled_table(doc,
    ['约束维度', 'CampusShare 现状', '影响'],
    [
        ['数据量', '冷启动，无标注数据', '排除方案 B（需 1k+ 标注）'],
        ['性能要求', '首字 P95 ≤ 1.5s', 'LLM 分类需 <400ms，规则层需 <10ms'],
        ['成本敏感', '日均 1000 问，月成本 ≤ ¥120', '25% 流量不调 LLM'],
        ['LLM 可用', '已接入 DeepSeek-V3', '复用现有 WebClient + 熔断器'],
        ['维护能力', '1-2 人小团队', '避免重训模型，Prompt 即规则'],
        ['意图体系', '5 大意图 + 14 子意图', 'LLM 可处理，规则难覆盖全'],
    ],
    col_widths=[2.5, 5, 5.5])

add_heading_styled(doc, '最终选型：方案 F（混合多层漏斗）', level=3)

add_paragraph_styled(doc,
    '综合大厂经验与业务约束，CampusShare 选择方案 F——混合多层漏斗架构：',
    bold=True)

add_styled_table(doc,
    ['层级', '技术', '处理流量', '延迟', '职责'],
    [
        ['Layer 1\n规则短路', '关键词 + 正则', '~25%', '<5ms', '处理高确定性请求（你好/帮我发帖/我点赞的）'],
        ['Layer 2\nLLM 分类', 'DeepSeek Prompt', '~70%', '~300ms', '主分类器，输出 JSON（意图+子意图+置信度+改写query+槽位）'],
        ['Layer 3\nEmbedding 兜底', 'BGE-M3 相似度', '~5%', '~80ms', 'LLM 降级时兜底（熔断开路或超时）'],
        ['Default\nSEARCH 兜底', '直接路由', '<1%', '0ms', '全部失败时走最通用管线'],
    ],
    col_widths=[2.5, 3, 1.8, 1.5, 5])

add_callout(doc,
    '🎯 选型核心逻辑：用规则层过滤 25% 的"不需要思考"的请求（省成本），'
    '用 LLM 层处理 70% 的"需要理解"的请求（保精度），'
    '用 Embedding 层兜底 5% 的"LLM 不可用"场景（保可用）。',
    color='E8F5E9', border_color='4CAF50')

doc.add_page_break()

# ==================== 三、流程 ====================
add_heading_styled(doc, '三、流程：如何搭建意图识别', level=1)

add_heading_styled(doc, '3.1 前置条件', level=2)

add_paragraph_styled(doc, '意图识别模块不是从零开始，而是嵌入现有 AgentChatService 编排链路。前置条件：')

add_styled_table(doc,
    ['前置组件', '当前状态', '用途', '文件位置'],
    [
        ['AgentChatService', '✅ 已实现', '编排核心，意图识别将插入此处', 'service/AgentChatService.java'],
        ['RetrievalService', '✅ 已实现', 'RAG 三路检索，按意图选择性调用', 'service/RetrievalService.java'],
        ['DeepSeekClient', '✅ 已实现', 'LLM 调用，复用于意图分类', 'llm/DeepSeekClient.java'],
        ['EmbeddingClient', '✅ 已实现', '向量化，用于 Embedding 兜底层', 'llm/EmbeddingClient.java'],
        ['AgentRateLimiter', '✅ 已实现', 'Redis 限流，复用', 'service/AgentRateLimiter.java'],
        ['ResilienceConfig', '✅ 已实现', '熔断器，为意图分类新增 1 个', 'config/ResilienceConfig.java'],
        ['PromptConstants', '✅ 基础版', '系统提示词，将扩展意图分类 Prompt', 'llm/PromptConstants.java'],
        ['Redis 连接', '✅ 已配置', '缓存意图分类结果', 'application.yml'],
    ],
    col_widths=[3.5, 2, 5, 4])

add_heading_styled(doc, '3.2 意图分类体系设计', level=2)

add_paragraph_styled(doc, '基于用户场景调研，设计 5 大意图 + 14 子意图的两级分类体系：')

add_styled_table(doc,
    ['L1 意图', 'L2 子意图', '触发特征', '处理管线', 'LLM 调用数'],
    [
        ['HOW_TO\n操作指引', 'feature_help', '"怎么/如何/在哪设置" + 功能词', '知识库检索 → 生成步骤', '1 次'],
        ['HOW_TO', 'rule_explain', '"为什么/什么意思" + 规则词', '知识库检索 → 解释', '1 次'],
        ['SEARCH\n内容检索', 'resource', '"求/找/有没有" + 资源词 + 学校', '帖子检索（资源类）+ 槽位过滤', '2-4 次（ReAct）'],
        ['SEARCH', 'discussion', '"讨论/聊聊/怎么看" + 话题', '帖子检索（讨论类）', '2-4 次'],
        ['SEARCH', 'content_qa', '"那个帖子说了什么" + 指代', '定位帖子 → 内容提取', '2-4 次'],
        ['NAVIGATE\n页面导航', 'feature_loc', '"在哪/入口" + 功能名', '返回跳转卡片', '0-1 次'],
        ['NAVIGATE', 'section_loc', '"板块/分类在哪" + 分类名', '返回分类卡片', '0-1 次'],
        ['NAVIGATE', 'my_list', '"我xxx的帖子" + 个人列表词', '返回 /profile/:type 跳转', '0 次'],
        ['CLARIFY\n多轮澄清', 'coreference', '"那个/它/上面那个" + 指代', '上下文消解 → 直接答/筛选', '0-1 次'],
        ['CLARIFY', 'refine', '"我说的是xxx/不对" + 修正', '修正槽位 → 重走原管线', '同原管线'],
        ['CLARIFY', 'followup', '"那xxx呢/接着问" + 追问', '上下文 + 新检索', '同原管线'],
        ['OUT_OF_SCOPE\n超范围', 'chitchat', '闲聊问候', '礼貌引导模板', '0 次'],
        ['OUT_OF_SCOPE', 'open_domain', '开放域知识', '拒绝 + 引导', '0 次'],
        ['OUT_OF_SCOPE', 'write_action', '"帮我发/帮我点赞/帮我改"', '拒绝 + 操作指引', '0 次'],
        ['OUT_OF_SCOPE', 'sensitive', '政治/医疗/法律', '拒绝模板', '0 次'],
    ],
    col_widths=[2, 2, 3.5, 3.5, 1.8])

add_callout(doc,
    '⚠️ 关键规则：当 query 包含指代词（那个/它/上面那个）时，强制走 CLARIFY/coreference，'
    '无论置信度多高。这是为了避免指代词被误判为 SEARCH 后重新检索，丢失上一轮上下文。',
    color='FFF3CD', border_color='FFC107')

add_heading_styled(doc, '3.3 整体架构与请求流转', level=2)

add_paragraph_styled(doc, '意图识别插入后的完整请求流转（红色为新增部分）：')

add_code_block(doc, """请求流转（含意图识别）：

用户消息
  ↓
[1] 规则短路层 RuleShortCircuitFilter          ← 新增
    ├─ 命中"你好/谢谢"         → OUT_OF_SCOPE/chitchat    → 模板回复（0 LLM）
    ├─ 命中"帮我发/帮我点赞"   → OUT_OF_SCOPE/write_action → 拒绝模板（0 LLM）
    ├─ 命中"我点赞的/我收藏的" → NAVIGATE/my_list          → 跳转卡片（0 LLM）
    └─ 未命中                   → 进入 Layer 2
  ↓
[2] LLM 意图分类 IntentClassifier              ← 新增
    ├─ 调用 DeepSeek（temperature=0, max_tokens=200）
    ├─ 输出 JSON: {intent, sub_intent, confidence, rewritten_query, slots}
    ├─ 熔断器保护（新增 intentClassifierCircuitBreaker）
    ├─ Redis 缓存（key=query hash, TTL=1h）     ← 新增
    └─ LLM 降级 → 进入 Layer 3
  ↓
[3] Embedding 兜底 EmbeddingIntentFallback     ← 新增
    ├─ query 向量化 → 与意图描述余弦相似度
    └─ 返回最相似意图
  ↓
[4] 意图路由 IntentRouter                       ← 新增
    ├─ OUT_OF_SCOPE  → 模板回复（0 LLM）
    ├─ NAVIGATE      → 跳转卡片（0-1 LLM）
    ├─ HOW_TO        → 知识库 RAG（1 LLM）
    ├─ CLARIFY       → 上下文消解（0-1 LLM）
    └─ SEARCH        → ReAct 主循环（2-4 LLM）   ← Advanced 阶段
  ↓
[5] AgentChatService 编排（已有，按路由结果执行）
  ↓
流式响应 → 持久化""")

add_heading_styled(doc, '3.4 详细实现步骤', level=2)

add_heading_styled(doc, 'Step 1：定义意图枚举与 DTO', level=3)
add_paragraph_styled(doc, '创建意图枚举类和分类结果 DTO，作为整个模块的数据基础。')
add_bullet(doc, '文件：dto/IntentResult.java（新增）')
add_bullet(doc, '文件：dto/SlotResult.java（新增，槽位抽取结果）')

add_heading_styled(doc, 'Step 2：实现规则短路层', level=3)
add_paragraph_styled(doc, '用关键词 + 正则处理高确定性请求，避免 LLM 调用。')
add_bullet(doc, '文件：service/RuleShortCircuitFilter.java（新增）')
add_bullet(doc, '规则配置：维护在 application.yml，支持热更新')

add_heading_styled(doc, 'Step 3：实现 LLM 意图分类器', level=3)
add_paragraph_styled(doc, '核心分类器，复用 DeepSeekClient，新增分类专用 Prompt。')
add_bullet(doc, '文件：service/IntentClassifier.java（新增）')
add_bullet(doc, '文件：llm/PromptConstants.java（扩展，新增意图分类 Prompt）')
add_bullet(doc, '新增熔断器：intentClassifierCircuitBreaker')

add_heading_styled(doc, 'Step 4：实现 Embedding 兜底', level=3)
add_paragraph_styled(doc, 'LLM 降级时用 Embedding 相似度兜底，保证可用性。')
add_bullet(doc, '文件：service/EmbeddingIntentFallback.java（新增）')
add_bullet(doc, '意图描述维护：IntentDescription枚举')

add_heading_styled(doc, 'Step 5：实现意图路由器', level=3)
add_paragraph_styled(doc, '根据意图分发到不同处理链路。')
add_bullet(doc, '文件：service/IntentRouter.java（新增）')
add_bullet(doc, '路由策略表：见 3.2 节')

add_heading_styled(doc, 'Step 6：改造 AgentChatService', level=3)
add_paragraph_styled(doc, '在现有编排链路中插入意图识别环节。')
add_bullet(doc, '文件：service/AgentChatService.java（修改）')
add_bullet(doc, '改造点：prepareContext() 前插入意图识别，按路由结果选择处理链路')

add_heading_styled(doc, 'Step 7：添加 Redis 缓存', level=3)
add_paragraph_styled(doc, '缓存意图分类结果，相同 query 不重复调 LLM。')
add_bullet(doc, '缓存 key：agent:intent:{md5(query)}')
add_bullet(doc, 'TTL：1 小时')
add_bullet(doc, '缓存命中：直接返回 IntentResult，跳过 LLM 调用')

add_heading_styled(doc, 'Step 8：添加监控指标', level=3)
add_paragraph_styled(doc, '用 Micrometer 暴露意图分类相关指标。')
add_bullet(doc, 'agent.intent.classification.total（标签：intent, layer, result）')
add_bullet(doc, 'agent.intent.classification.duration（标签：layer）')
add_bullet(doc, 'agent.intent.cache.hit.rate')

add_heading_styled(doc, '3.5 关键设计决策（ADR）', level=2)

add_paragraph_styled(doc, '本设计涉及 7 个关键 ADR（架构决策记录）：')

add_styled_table(doc,
    ['ADR', '决策', '理由'],
    [
        ['ADR-009', '意图分类用 LLM 而非纯规则', '用户表达多样（咋/咋整/咋关），规则爆炸；LLM 泛化好。高频确定性模式用规则短路省成本'],
        ['ADR-010', 'confidence < 0.6 兜底为 SEARCH', 'SEARCH 管线最通用（检索+生成），分错也能给有用结果；NAVIGATE/HOW_TO 分错会完全跑偏'],
        ['ADR-011', '意图分类与查询改写合并为一次 LLM 调用', '省 ~500ms 延迟和 1 次 API 成本；两者输入相同，输出可结构化合并。风险：多任务 Prompt 质量略降，用 JSON Schema + Few-shot 缓解'],
        ['ADR-012', 'HyDE 条件触发', 'HyDE 对短模糊查询有益，对长清晰查询浪费。触发条件：SEARCH 类 + query < 15 字 + 无明确实体'],
        ['ADR-013', '简单意图走快路径', 'OUT_OF_SCOPE/NAVIGATE 占 25%+ 流量，快路径省 60%+ 成本和延迟。风险：误路由用置信度阈值 + SEARCH 兜底缓解'],
        ['ADR-014', '短期记忆用 Redis', '高频读写、TTL 自动过期、低延迟。重要会话异步持久化到 MySQL'],
        ['ADR-015', '指代词强制 CLARIFY', '指代词误判为 SEARCH 会丢失上轮上下文。宁可过度 CLARIFY（LLM 判断真假指代），不可漏判'],
    ],
    col_widths=[1.5, 4.5, 8.5])

doc.add_page_break()

# ==================== 四、核心代码 ====================
add_heading_styled(doc, '四、核心代码', level=1)

add_heading_styled(doc, '4.1 文件架构', level=2)

add_paragraph_styled(doc, '意图识别模块新增/修改的文件清单：')

add_code_block(doc, """backend/campushare-agent/src/main/java/com/campushare/agent/
├── controller/
│   └── AgentController.java                    # [修改] 无需改动，流式端点不变
├── service/
│   ├── AgentChatService.java                   # [修改] 插入意图识别环节
│   ├── IntentClassifier.java                   # [新增] ⭐ LLM 意图分类器（核心）
│   ├── RuleShortCircuitFilter.java             # [新增] ⭐ 规则短路层
│   ├── EmbeddingIntentFallback.java            # [新增] Embedding 兜底层
│   ├── IntentRouter.java                       # [新增] ⭐ 意图路由器
│   └── IntentCacheService.java                 # [新增] Redis 缓存
├── llm/
│   ├── PromptConstants.java                    # [修改] 新增意图分类 Prompt
│   └── DeepSeekClient.java                     # [修改] 新增分类专用方法
├── config/
│   └── ResilienceConfig.java                   # [修改] 新增 intentClassifierCircuitBreaker
├── dto/
│   ├── IntentResult.java                       # [新增] ⭐ 意图分类结果 DTO
│   ├── SlotResult.java                         # [新增] 槽位抽取结果
│   └── RouteDecision.java                      # [新增] 路由决策
└── enums/
    └── Intent.java                             # [新增] ⭐ 意图枚举（5 大 + 14 子）""")

add_heading_styled(doc, '4.2 意图定义与 DTO', level=2)

add_paragraph_styled(doc, '文件：enums/Intent.java（新增）', bold=True)
add_code_block(doc, """package com.campushare.agent.enums;

import lombok.Getter;

@Getter
public enum Intent {
    // L1 意图
    HOW_TO("操作指引", "平台功能操作指引", 1),
    SEARCH("内容检索", "资源/讨论帖检索", 2),
    NAVIGATE("页面导航", "功能/板块定位", 3),
    CLARIFY("多轮澄清", "多轮澄清/追问/指代", 4),
    OUT_OF_SCOPE("超范围", "闲聊/开放域/写操作/敏感", 5);

    private final String label;
    private final String description;
    private final int code;

    Intent(String label, String description, int code) {
        this.label = label;
        this.description = description;
        this.code = code;
    }

    // 子意图（14 个，用 String 而非枚举，便于扩展）
    public static final class SubIntent {
        public static final String FEATURE_HELP = "feature_help";
        public static final String RULE_EXPLAIN = "rule_explain";
        public static final String RESOURCE = "resource";
        public static final String DISCUSSION = "discussion";
        public static final String CONTENT_QA = "content_qa";
        public static final String FEATURE_LOC = "feature_loc";
        public static final String SECTION_LOC = "section_loc";
        public static final String MY_LIST = "my_list";
        public static final String COREFERENCE = "coreference";
        public static final String REFINE = "refine";
        public static final String FOLLOWUP = "followup";
        public static final String CHITCHAT = "chitchat";
        public static final String OPEN_DOMAIN = "open_domain";
        public static final String WRITE_ACTION = "write_action";
        public static final String SENSITIVE = "sensitive";
    }
}""")

add_paragraph_styled(doc, '文件：dto/IntentResult.java（新增）', bold=True)
add_code_block(doc, """package com.campushare.agent.dto;

import com.campushare.agent.enums.Intent;
import lombok.Builder;
import lombok.Data;

@Data
@Builder
public class IntentResult {
    private Intent intent;              // L1 意图
    private String subIntent;           // L2 子意图
    private double confidence;          // 置信度 0.0-1.0
    private String rewrittenQuery;      // 改写后的查询
    private SlotResult slots;           // 抽取的槽位
    private String hydeDoc;             // HyDE 假设文档（可选）
    private String classifyLayer;       // 分类层级：RULE / LLM / EMBEDDING / DEFAULT

    // 槽位
    @Data
    @Builder
    public static class SlotResult {
        private String school;          // 清华/北大/...
        private String category;        // 音乐/游戏/面经/...
        private String postType;        // resource/discussion
        private String sort;            // 最新/最热
    }

    public boolean isHighConfidence() {
        return confidence >= 0.6;
    }

    public boolean isLowConfidence() {
        return confidence < 0.6;
    }
}""")

add_heading_styled(doc, '4.3 规则短路层', level=2)

add_paragraph_styled(doc, '文件：service/RuleShortCircuitFilter.java（新增）', bold=True)
add_paragraph_styled(doc, '处理高确定性请求，避免 LLM 调用。三层规则：写操作、闲聊、个人列表。')

add_code_block(doc, """package com.campushare.agent.service;

import com.campushare.agent.dto.IntentResult;
import com.campushare.agent.enums.Intent;
import lombok.extern.slf4j.Slf4j;
import org.springframework.stereotype.Service;

import java.util.List;
import java.util.Optional;
import java.util.regex.Pattern;

@Service
@Slf4j
public class RuleShortCircuitFilter {

    // 规则 1：写操作请求（帮我发帖/帮我点赞/帮我关注/帮我改密码）
    private static final List<Pattern> WRITE_ACTION_PATTERNS = List.of(
        Pattern.compile("帮我(发|发布|写).*帖"),
        Pattern.compile("帮我(点|赞|收藏|关注|取消)"),
        Pattern.compile("帮我(改|修改|删除|编辑)"),
        Pattern.compile("代替我(发|点|改|删)")
    );

    // 规则 2：闲聊问候
    private static final List<Pattern> CHITCHAT_PATTERNS = List.of(
        Pattern.compile("^(你好|您好|hi|hello|嗨|哈喽|hey).*", Pattern.CASE_INSENSITIVE),
        Pattern.compile("^(谢谢|感谢|thx|thanks).*", Pattern.CASE_INSENSITIVE),
        Pattern.compile("^(你是谁|你叫什么|你能做什么).*"),
        Pattern.compile("^(早上好|晚上好|中午好|晚安).*")
    );

    // 规则 3：个人列表（我点赞的/我收藏的/我回复的/我的浏览历史）
    private static final List<Pattern> MY_LIST_PATTERNS = List.of(
        Pattern.compile("我(点赞|赞过|喜欢).*帖"),
        Pattern.compile("我(收藏|存|关注).*帖"),
        Pattern.compile("我(回复|评论)过.*"),
        Pattern.compile("我(浏览|看过)历史"),
        Pattern.compile("我(关注|粉丝|互关)列表")
    );

    // 规则 4：指代词强制 CLARIFY
    private static final Pattern COREFERENCE_PATTERN = Pattern.compile(
        ".*(那个|它|上面那个|刚才那个|第几个|有下载的|带图的|最热的).*"
    );

    /**
     * 规则短路：命中返回 IntentResult，未命中返回 empty
     */
    public Optional<IntentResult> filter(String query) {
        if (query == null || query.isBlank()) {
            return Optional.empty();
        }
        String trimmed = query.trim();

        // 优先级 1：指代词 → 强制 CLARIFY（ADR-015）
        if (COREFERENCE_PATTERN.matcher(trimmed).matches()) {
            log.debug("Rule matched: COREFERENCE for query='{}'", trimmed);
            return Optional.of(IntentResult.builder()
                    .intent(Intent.CLARIFY)
                    .subIntent(Intent.SubIntent.COREFERENCE)
                    .confidence(0.95)  // 规则命中视为高置信
                    .rewrittenQuery(trimmed)
                    .classifyLayer("RULE")
                    .build());
        }

        // 优先级 2：写操作 → OUT_OF_SCOPE/write_action
        for (Pattern p : WRITE_ACTION_PATTERNS) {
            if (p.matcher(trimmed).find()) {
                log.debug("Rule matched: WRITE_ACTION for query='{}'", trimmed);
                return Optional.of(IntentResult.builder()
                        .intent(Intent.OUT_OF_SCOPE)
                        .subIntent(Intent.SubIntent.WRITE_ACTION)
                        .confidence(0.99)
                        .rewrittenQuery(trimmed)
                        .classifyLayer("RULE")
                        .build());
            }
        }

        // 优先级 3：闲聊 → OUT_OF_SCOPE/chitchat
        for (Pattern p : CHITCHAT_PATTERNS) {
            if (p.matcher(trimmed).matches()) {
                log.debug("Rule matched: CHITCHAT for query='{}'", trimmed);
                return Optional.of(IntentResult.builder()
                        .intent(Intent.OUT_OF_SCOPE)
                        .subIntent(Intent.SubIntent.CHITCHAT)
                        .confidence(0.99)
                        .rewrittenQuery(trimmed)
                        .classifyLayer("RULE")
                        .build());
            }
        }

        // 优先级 4：个人列表 → NAVIGATE/my_list
        for (Pattern p : MY_LIST_PATTERNS) {
            if (p.matcher(trimmed).find()) {
                log.debug("Rule matched: MY_LIST for query='{}'", trimmed);
                return Optional.of(IntentResult.builder()
                        .intent(Intent.NAVIGATE)
                        .subIntent(Intent.SubIntent.MY_LIST)
                        .confidence(0.95)
                        .rewrittenQuery(trimmed)
                        .classifyLayer("RULE")
                        .build());
            }
        }

        return Optional.empty();
    }
}""")

add_heading_styled(doc, '4.4 LLM 意图分类器', level=2)

add_paragraph_styled(doc, '文件：service/IntentClassifier.java（新增，核心组件）', bold=True)
add_paragraph_styled(doc,
    '核心分类器，调用 DeepSeek 输出结构化 JSON。'
    '同时完成意图分类 + 查询改写 + 槽位抽取（ADR-011 合并调用）。')

add_code_block(doc, """package com.campushare.agent.service;

import com.campushare.agent.dto.IntentResult;
import com.campushare.agent.enums.Intent;
import com.campushare.agent.llm.DeepSeekClient;
import com.campushare.agent.llm.PromptConstants;
import com.fasterxml.jackson.databind.JsonNode;
import com.fasterxml.jackson.databind.ObjectMapper;
import io.github.resilience4j.circuitbreaker.CircuitBreaker;
import io.github.resilience4j.reactor.circuitbreaker.operator.CircuitBreakerOperator;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.beans.factory.annotation.Qualifier;
import org.springframework.stereotype.Service;
import reactor.core.publisher.Mono;

import java.time.Duration;

@Service
@RequiredArgsConstructor
@Slf4j
public class IntentClassifier {

    private final DeepSeekClient deepSeekClient;
    private final IntentCacheService cacheService;
    private final EmbeddingIntentFallback embeddingFallback;
    private final ObjectMapper objectMapper;

    @Qualifier("intentClassifierCircuitBreaker")
    private final CircuitBreaker circuitBreaker;

    private static final Duration CLASSIFY_TIMEOUT = Duration.ofSeconds(3);
    private static final double CONFIDENCE_THRESHOLD = 0.6;

    /**
     * 意图分类主入口
     * @param query 用户原始查询
     * @param sessionId 会话 ID（用于多轮上下文）
     * @return IntentResult 分类结果
     */
    public Mono<IntentResult> classify(String query, String sessionId) {
        // 1. 查 Redis 缓存
        return cacheService.get(query)
                .map(cached -> {
                    log.debug("Intent cache hit for query='{}'", query);
                    return cached;
                })
                .switchIfEmpty(Mono.defer(() -> 
                    // 2. 缓存未命中，调用 LLM 分类
                    classifyByLLM(query)
                        .timeout(CLASSIFY_TIMEOUT)
                        .transform(CircuitBreakerOperator.of(circuitBreaker))
                        .onErrorResume(e -> {
                            // 3. LLM 降级 → Embedding 兜底
                            log.warn("LLM classify failed, fallback to embedding: {}", e.getMessage());
                            return embeddingFallback.classify(query);
                        })
                        .doOnNext(result -> {
                            // 4. 写入缓存
                            cacheService.put(query, result).subscribe();
                        })
                ));
    }

    /**
     * LLM 分类：调用 DeepSeek 输出 JSON
     */
    private Mono<IntentResult> classifyByLLM(String query) {
        String prompt = PromptConstants.buildIntentClassificationPrompt(query);

        return deepSeekClient.chatCompletionJson(prompt)  // 非流式 JSON 调用
                .map(response -> parseIntentResult(query, response))
                .doOnNext(result -> {
                    // ADR-010：低置信度兜底为 SEARCH
                    if (result.isLowConfidence()) {
                        log.info("Low confidence ({}), fallback to SEARCH for query='{}'",
                                result.getConfidence(), query);
                        result.setIntent(Intent.SEARCH);
                        result.setSubIntent(Intent.SubIntent.RESOURCE);
                    }
                    log.info("Intent classified: query='{}' → {} / {} (conf={}, layer=LLM)",
                            query, result.getIntent(), result.getSubIntent(),
                            result.getConfidence());
                });
    }

    /**
     * 解析 LLM 返回的 JSON
     */
    private IntentResult parseIntentResult(String query, String llmResponse) {
        try {
            JsonNode root = objectMapper.readTree(llmResponse);
            Intent intent = Intent.valueOf(root.path("intent").asText("SEARCH"));
            String subIntent = root.path("sub_intent").asText(Intent.SubIntent.RESOURCE);
            double confidence = root.path("confidence").asDouble(0.5);
            String rewrittenQuery = root.path("rewritten_query").asText(query);

            // 解析槽位
            IntentResult.SlotResult slots = null;
            JsonNode slotsNode = root.path("slots");
            if (!slotsNode.isMissingNode()) {
                slots = IntentResult.SlotResult.builder()
                        .school(slotsNode.path("school").asText(null))
                        .category(slotsNode.path("category").asText(null))
                        .postType(slotsNode.path("post_type").asText(null))
                        .sort(slotsNode.path("sort").asText(null))
                        .build();
            }

            return IntentResult.builder()
                    .intent(intent)
                    .subIntent(subIntent)
                    .confidence(confidence)
                    .rewrittenQuery(rewrittenQuery)
                    .slots(slots)
                    .hydeDoc(root.path("hyde_doc").asText(null))
                    .classifyLayer("LLM")
                    .build();
        } catch (Exception e) {
            log.error("Failed to parse intent JSON: {}", llmResponse, e);
            // 解析失败兜底
            return IntentResult.builder()
                    .intent(Intent.SEARCH)
                    .subIntent(Intent.SubIntent.RESOURCE)
                    .confidence(0.5)
                    .rewrittenQuery(query)
                    .classifyLayer("DEFAULT")
                    .build();
        }
    }
}""")

add_heading_styled(doc, '4.5 Prompt 模板', level=2)

add_paragraph_styled(doc, '文件：llm/PromptConstants.java（扩展）', bold=True)
add_paragraph_styled(doc,
    '意图分类专用 Prompt，包含意图定义、Few-shot 示例、JSON 输出格式约束。'
    '温度设为 0 保证稳定性。')

add_code_block(doc, '''package com.campushare.agent.llm;

public class PromptConstants {

    /**
     * 意图分类 Prompt（ADR-011：分类+改写+槽位合并）
     */
    public static final String INTENT_CLASSIFICATION_SYSTEM = """
            你是 CampusShare 校园资源共享平台的意图分类器。
            你的任务是判断用户问题的意图，并输出结构化 JSON。

            ## 意图体系（5 大 + 14 子）

            ### HOW_TO - 操作指引
            - feature_help: "怎么/如何/在哪设置" + 平台功能词
            - rule_explain: "为什么/什么意思" + 平台规则词

            ### SEARCH - 内容检索
            - resource: "求/找/有没有" + 资源词 + 学校/科目
            - discussion: "讨论/聊聊/怎么看" + 话题
            - content_qa: "那个帖子说了什么" + 指代

            ### NAVIGATE - 页面导航
            - feature_loc: "在哪/入口" + 功能名
            - section_loc: "板块/分类在哪" + 分类名
            - my_list: "我xxx的帖子" + 个人列表词

            ### CLARIFY - 多轮澄清
            - coreference: "那个/它/上面那个" + 指代
            - refine: "我说的是xxx/不对" + 修正
            - followup: "那xxx呢/接着问" + 追问

            ### OUT_OF_SCOPE - 超范围
            - chitchat: 闲聊问候
            - open_domain: 开放域知识（天气/新闻/百科）
            - write_action: "帮我发/帮我点赞/帮我改"
            - sensitive: 政治/医疗/法律

            ## 输出格式（严格 JSON，不要 Markdown 代码块）

            {
              "intent": "HOW_TO|SEARCH|NAVIGATE|CLARIFY|OUT_OF_SCOPE",
              "sub_intent": "feature_help|resource|...",
              "confidence": 0.0-1.0,
              "rewritten_query": "改写后的查询（规范化+同义词扩展）",
              "slots": {
                "school": "清华|北大|复旦|...|null",
                "category": "音乐|游戏|面经|...|null",
                "post_type": "resource|discussion|null",
                "sort": "最新|最热|null"
              },
              "hyde_doc": "假设文档（仅 SEARCH+短query 时生成）|null"
            }

            ## 判定原则
            1. confidence < 0.6 时，intent 设为 SEARCH（最通用兜底）
            2. 含指代词（那个/它/上面那个）时，强制 CLARIFY/coreference
            3. rewritten_query 需做：全角转半角、繁转简、同义词扩展
            4. hyde_doc 仅在 SEARCH 类 + query < 15 字 + 无明确实体时生成

            ## Few-shot 示例

            用户：怎么发帖？
            输出：{"intent":"HOW_TO","sub_intent":"feature_help","confidence":0.95,"rewritten_query":"如何发布帖子","slots":null,"hyde_doc":null}

            用户：求清华操作系统期末卷子
            输出：{"intent":"SEARCH","sub_intent":"resource","confidence":0.92,"rewritten_query":"操作系统 期末 卷子","slots":{"school":"清华","category":null,"post_type":"resource","sort":null},"hyde_doc":"这是一份清华大学操作系统期末复习卷子，包含进程管理、内存管理、文件系统..."}

            用户：个人中心在哪
            输出：{"intent":"NAVIGATE","sub_intent":"feature_loc","confidence":0.94,"rewritten_query":"个人中心 入口","slots":null,"hyde_doc":null}

            用户：那个有下载的
            输出：{"intent":"CLARIFY","sub_intent":"coreference","confidence":0.90,"rewritten_query":"在上一轮结果中筛选有文件下载的","slots":null,"hyde_doc":null}

            用户：帮我发帖
            输出：{"intent":"OUT_OF_SCOPE","sub_intent":"write_action","confidence":0.99,"rewritten_query":"帮我发帖","slots":null,"hyde_doc":null}

            用户：今天天气怎么样
            输出：{"intent":"OUT_OF_SCOPE","sub_intent":"open_domain","confidence":0.97,"rewritten_query":"今天天气","slots":null,"hyde_doc":null}
            """;

    public static String buildIntentClassificationPrompt(String userQuery) {
        return INTENT_CLASSIFICATION_SYSTEM + "\\n\\n用户：" + userQuery + "\\n输出：";
    }
}''')

add_heading_styled(doc, '4.6 意图路由器', level=2)

add_paragraph_styled(doc, '文件：service/IntentRouter.java（新增）', bold=True)
add_paragraph_styled(doc,
    '根据意图分发到不同处理链路。每个意图对应一个处理方法，返回 Flux<ChatEvent> 统一流式协议。')

add_code_block(doc, """package com.campushare.agent.service;

import com.campushare.agent.dto.IntentResult;
import com.campushare.agent.enums.Intent;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.stereotype.Service;
import reactor.core.publisher.Flux;

@Service
@RequiredArgsConstructor
@Slf4j
public class IntentRouter {

    private final RetrievalService retrievalService;
    private final DeepSeekClient deepSeekClient;
    private final SessionContextService sessionContextService;  // 多轮上下文

    /**
     * 按意图路由到对应处理链路
     */
    public Flux<ChatEvent> route(IntentResult intent, String sessionId, String userId) {
        return switch (intent.getIntent()) {
            case OUT_OF_SCOPE -> handleOutOfScope(intent);
            case NAVIGATE     -> handleNavigate(intent, sessionId);
            case HOW_TO       -> handleHowTo(intent, sessionId);
            case CLARIFY      -> handleClarify(intent, sessionId);
            case SEARCH       -> handleSearch(intent, sessionId, userId);
        };
    }

    // ========== OUT_OF_SCOPE：模板回复，0 次 LLM ==========
    private Flux<ChatEvent> handleOutOfScope(IntentResult intent) {
        String reply = switch (intent.getSubIntent()) {
            case Intent.SubIntent.CHITCHAT ->
                "你好！我是 CampusShare AI 助手，可以帮你找资料、教你怎么用功能～有什么可以帮你的？";
            case Intent.SubIntent.WRITE_ACTION ->
                "抱歉，我无法代替你执行操作（发帖/点赞/收藏等）。你可以参考以下指引自行操作：\\n"
                + "1. 发帖：进入对应板块 → 点击右下角 + 按钮\\n"
                + "2. 点赞/收藏：在帖子详情页点击爱心或星标";
            case Intent.SubIntent.OPEN_DOMAIN ->
                "抱歉，我是 CampusShare 平台助手，只能回答平台相关问题（找资料、用功能、看帖子）。"
                + "关于其他问题，建议你咨询相关专业渠道哦～";
            case Intent.SubIntent.SENSITIVE ->
                "抱歉，这个问题我无法回答。如果你有 CampusShare 平台相关问题，欢迎继续提问。";
            default -> "抱歉，我无法理解你的问题。你可以试试问我：怎么发帖、找清华考研资料、个人中心在哪。";
        };

        return Flux.just(new ChatEvent("delta", reply));
    }

    // ========== NAVIGATE：跳转卡片，0-1 次 LLM ==========
    private Flux<ChatEvent> handleNavigate(IntentResult intent, String sessionId) {
        if (Intent.SubIntent.MY_LIST.equals(intent.getSubIntent())) {
            // 个人列表：直接返回跳转卡片
            String route = mapMyListRoute(intent.getRewrittenQuery());
            String card = String.format(
                "你可以在这里查看：\\n[点击跳转 →](%s)", route);
            return Flux.just(new ChatEvent("delta", card));
        }
        // feature_loc / section_loc：可选 1 次 LLM 生成引导语
        return Flux.just(new ChatEvent("delta", "正在为你定位入口..."));
        // 实际实现：查知识库/分类表 → 返回跳转卡片
    }

    // ========== HOW_TO：知识库 RAG，1 次 LLM ==========
    private Flux<ChatEvent> handleHowTo(IntentResult intent, String sessionId) {
        return retrievalService.retrieveKnowledge(intent.getRewrittenQuery())
                .flatMapMany(context -> {
                    String prompt = PromptConstants.buildHowToPrompt(
                            intent.getRewrittenQuery(), context);
                    return deepSeekClient.chatCompletionStream(prompt)
                            .map(chunk -> new ChatEvent("delta", chunk.content()));
                });
    }

    // ========== CLARIFY：上下文消解，0-1 次 LLM ==========
    private Flux<ChatEvent> handleClarify(IntentResult intent, String sessionId) {
        return sessionContextService.resolveCoreference(sessionId, intent)
                .flatMapMany(resolved -> {
                    if (resolved.isDirectAnswer()) {
                        // 直接答（如"那个有下载的" → 筛选上轮结果）
                        return Flux.just(new ChatEvent("delta", resolved.getAnswer()));
                    }
                    // 重走原管线
                    return route(resolved.getNextIntent(), sessionId, null);
                });
    }

    // ========== SEARCH：ReAct 主循环，2-4 次 LLM（Advanced 阶段） ==========
    private Flux<ChatEvent> handleSearch(IntentResult intent, String sessionId, String userId) {
        // MVP 阶段：仍走 RAG 检索 + 生成（与当前实现一致）
        // Advanced 阶段：接入 ReAct 循环 + 工具调用
        return retrievalService.retrieve(intent.getRewrittenQuery())
                .flatMapMany(context -> {
                    String prompt = PromptConstants.buildSearchPrompt(
                            intent.getRewrittenQuery(), context, intent.getSlots());
                    return deepSeekClient.chatCompletionStream(prompt)
                            .map(chunk -> new ChatEvent("delta", chunk.content()));
                });
    }

    private String mapMyListRoute(String query) {
        if (query.contains("点赞") || query.contains("赞过")) return "/profile/liked";
        if (query.contains("收藏")) return "/profile/starred";
        if (query.contains("回复") || query.contains("评论")) return "/profile/comments";
        if (query.contains("浏览") || query.contains("历史")) return "/profile/history";
        if (query.contains("关注")) return "/profile/following";
        if (query.contains("粉丝")) return "/profile/followers";
        return "/profile";
    }
}""")

add_heading_styled(doc, '4.7 AgentChatService 改造', level=2)

add_paragraph_styled(doc, '文件：service/AgentChatService.java（修改）', bold=True)
add_paragraph_styled(doc,
    '在现有 chat() 方法中插入意图识别环节，按路由结果分发。'
    '改造点：prepareContext() 前插入意图分类，用 IntentRouter 替代直接 RAG。')

add_code_block(doc, """// AgentChatService.chat() 改造后（核心片段）

public Flux<ChatEvent> chat(String userId, ChatRequest request) {
    return Mono.fromCallable(() -> prepareSession(userId, request))
            .subscribeOn(Schedulers.boundedElastic())
            .flatMap(session -> {
                String query = request.getMessage();

                // ===== 新增：意图识别 =====
                return ruleShortCircuitFilter.filter(query)          // Layer 1: 规则短路
                    .map(Mono::just)
                    .orElseGet(() -> intentClassifier.classify(query, session.getId())  // Layer 2/3: LLM + Embedding
                            .map(Optional::of)
                            .defaultIfEmpty(Optional.empty()))
                    .flatMap(intentOpt -> {
                        IntentResult intent = intentOpt.orElseGet(() ->
                                IntentResult.builder()
                                        .intent(Intent.SEARCH)        // Default: SEARCH 兜底
                                        .subIntent(Intent.SubIntent.RESOURCE)
                                        .confidence(0.0)
                                        .rewrittenQuery(query)
                                        .classifyLayer("DEFAULT")
                                        .build());

                        // 记录意图分类指标
                        metricsService.recordIntentClassification(intent);

                        // ===== 新增：按意图路由 =====
                        return intentRouter.route(intent, session.getId(), userId)
                                .doOnNext(event -> accumulateAssistantContent(event, session))
                                .doFinally(signal -> persistTurn(session, intent, signal));
                    });
            });
}""")

add_heading_styled(doc, '4.8 Redis 缓存层', level=2)

add_paragraph_styled(doc, '文件：service/IntentCacheService.java（新增）', bold=True)
add_code_block(doc, """package com.campushare.agent.service;

import com.campushare.agent.dto.IntentResult;
import com.fasterxml.jackson.databind.ObjectMapper;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.data.redis.core.StringRedisTemplate;
import org.springframework.stereotype.Service;
import reactor.core.publisher.Mono;

import java.time.Duration;
import java.util.DigestUtils;

@Service
@RequiredArgsConstructor
@Slf4j
public class IntentCacheService {

    private final StringRedisTemplate redis;
    private final ObjectMapper objectMapper;

    private static final String KEY_PREFIX = "agent:intent:";
    private static final Duration TTL = Duration.ofHours(1);

    public Mono<IntentResult> get(String query) {
        String key = buildKey(query);
        return Mono.fromCallable(() -> {
            String json = redis.opsForValue().get(key);
            if (json == null) return null;
            return objectMapper.readValue(json, IntentResult.class);
        })
        .onErrorResume(e -> {
            log.warn("Intent cache read failed: {}", e.getMessage());
            return Mono.empty();
        })
        .map(Mono::just)
        .flatMap(m -> m != null ? m : Mono.empty());
    }

    public Mono<Void> put(String query, IntentResult result) {
        String key = buildKey(query);
        return Mono.fromRunnable(() -> {
            try {
                String json = objectMapper.writeValueAsString(result);
                redis.opsForValue().set(key, json, TTL);
            } catch (Exception e) {
                log.warn("Intent cache write failed: {}", e.getMessage());
            }
        }).then();
    }

    private String buildKey(String query) {
        return KEY_PREFIX + DigestUtils.md5Hex(query.trim().toLowerCase());
    }
}""")

add_heading_styled(doc, '4.9 监控指标', level=2)

add_paragraph_styled(doc, '文件：config/IntentMetricsConfig.java（新增）', bold=True)
add_code_block(doc, """package com.campushare.agent.config;

import io.micrometer.core.instrument.Counter;
import io.micrometer.core.instrument.MeterRegistry;
import io.micrometer.core.instrument.Timer;
import lombok.RequiredArgsConstructor;
import org.springframework.stereotype.Component;

import java.time.Duration;

@Component
@RequiredArgsConstructor
public class IntentMetricsConfig {

    private final MeterRegistry registry;

    // 意图分类总数（标签：intent, sub_intent, layer, result）
    public void recordClassification(String intent, String subIntent,
                                      String layer, String result) {
        Counter.builder("agent.intent.classification.total")
                .tag("intent", intent)
                .tag("sub_intent", subIntent)
                .tag("layer", layer)        // RULE / LLM / EMBEDDING / DEFAULT
                .tag("result", result)      // SUCCESS / FALLBACK / ERROR
                .register(registry)
                .increment();
    }

    // 分类耗时（标签：layer）
    public void recordDuration(String layer, Duration duration) {
        Timer.builder("agent.intent.classification.duration")
                .tag("layer", layer)
                .register(registry)
                .record(duration);
    }

    // 缓存命中率
    public void recordCacheHit(boolean hit) {
        Counter.builder("agent.intent.cache.total")
                .tag("result", hit ? "HIT" : "MISS")
                .register(registry)
                .increment();
    }
}""")

add_paragraph_styled(doc, 'Grafana 仪表盘建议监控指标：')
add_styled_table(doc,
    ['指标', '类型', '说明', '告警阈值'],
    [
        ['agent.intent.classification.total', 'Counter', '意图分类总数（按意图/层级/结果）', '—'],
        ['agent.intent.classification.duration', 'Timer', '分类耗时 P95', 'P95 > 500ms'],
        ['agent.intent.cache.total{result=HIT}', 'Counter', '缓存命中数', '命中率 < 30%'],
        ['LLM 降级率', 'Computed', 'EMBEDDING+DEFAULT / total', '> 10%'],
        ['SEARCH 兜底率', 'Computed', 'confidence<0.6 / total', '> 20%'],
    ],
    col_widths=[5, 2, 5, 3])

doc.add_page_break()

# ==================== 五、目标 ====================
add_heading_styled(doc, '五、目标：实现效果', level=1)

add_heading_styled(doc, '5.1 功能目标', level=2)

add_styled_table(doc,
    ['功能点', '当前（无意图识别）', '目标（有意图识别）', '验收方式'],
    [
        ['意图分类', '❌ 无', '✅ 5 大意图 + 14 子意图', 'Golden Set 准确率 ≥ 92%'],
        ['规则短路', '❌ 无', '✅ 写操作/闲聊/个人列表/指代词', '命中 25% 流量'],
        ['查询改写', '❌ 无', '✅ 规范化+同义词+槽位抽取', '改写 Recall@20 提升 ≥ 15%'],
        ['HyDE 检索', '❌ 无', '✅ SEARCH+短query 条件触发', '短 query Top-5 命中率 ≥ 80%'],
        ['意图路由', '❌ 一条管线', '✅ 5 条差异化管线', '按 ADR-013 路由表'],
        ['多轮指代', '❌ 无', '✅ CLARIFY/coreference', "指代词 100% 走 CLARIFY"],
        ['缓存优化', '❌ 无', '✅ Redis 缓存 1h', '相同 query 0 LLM 调用'],
    ],
    col_widths=[2.5, 3, 4, 4])

add_heading_styled(doc, '5.2 性能目标', level=2)

add_styled_table(doc,
    ['指标', '当前（无意图识别）', '目标（有意图识别）', '提升幅度'],
    [
        ['首字延迟 P50', '~1.5s', '简单问题 <100ms / 复杂问题 ~1.2s', '简单问题提速 15x'],
        ['首字延迟 P95', '~2.5s', '简单问题 <200ms / 复杂问题 ~1.8s', 'P95 降低 ~30%'],
        ['规则短路延迟', '—', '<5ms', '—'],
        ['LLM 分类延迟', '—', 'P95 < 400ms', '—'],
        ['缓存命中延迟', '—', '<10ms', '—'],
        ['Embedding 兜底延迟', '—', '<100ms', '—'],
    ],
    col_widths=[3.5, 3, 5, 2.5])

add_heading_styled(doc, '5.3 体验目标', level=2)

add_paragraph_styled(doc, '意图识别对用户体验的改善：')

add_bullet(doc, '简单问题秒回：用户问"你好"，从等 1.5s 变成毫秒级回复，体感"快了"')
add_bullet(doc, '检索更精准：用户问"求清华操作系统资料"，不再混入北大、复旦的资料')
add_bullet(doc, '导航直达：用户问"个人中心在哪"，直接给跳转链接，不用等 LLM 生成')
add_bullet(doc, '写操作明确拒绝：用户问"帮我发帖"，明确告知不能代操作 + 给操作指引，不再含糊')
add_bullet(doc, '多轮连贯：用户问"那个有下载的"，能理解指代上一轮结果，不再重新检索')

add_heading_styled(doc, '5.4 成本目标', level=2)

add_styled_table(doc,
    ['成本项', '当前（无意图识别）', '目标（有意图识别）', '节省'],
    [
        ['LLM 调用次数/日', '~1000 次（每问必调）', '~750 次（25% 不调 LLM）', '25%'],
        ['LLM 成本/日', '~¥4', '~¥3', '25%'],
        ['Embedding 调用/日', '~1000 次', '~800 次（25% 规则短路省）', '20%'],
        ['月度成本', '~¥120', '~¥90', '25%'],
        ['缓存节省', '0', '约 15% 重复 query 命中缓存', '额外 15%'],
    ],
    col_widths=[3, 3.5, 4, 2])

add_heading_styled(doc, '5.5 评估指标与验收标准', level=2)

add_heading_styled(doc, 'Golden Set 评估', level=3)
add_paragraph_styled(doc,
    '构建 200 条标注数据（每意图 40 条），评估分类准确率：')

add_styled_table(doc,
    ['意图', '样本数', '准确率目标', '召回率目标', 'F1 目标'],
    [
        ['HOW_TO', '40', '≥ 90%', '≥ 92%', '≥ 91%'],
        ['SEARCH', '40', '≥ 92%', '≥ 90%', '≥ 91%'],
        ['NAVIGATE', '40', '≥ 90%', '≥ 88%', '≥ 89%'],
        ['CLARIFY', '40', '≥ 85%', '≥ 90%', '≥ 87%'],
        ['OUT_OF_SCOPE', '40', '≥ 95%', '≥ 95%', '≥ 95%'],
        ['总体', '200', '≥ 92%', '≥ 92%', '≥ 92%'],
    ],
    col_widths=[3, 2, 2.5, 2.5, 2.5])

add_callout(doc,
    '⚠️ 特殊要求：OUT_OF_SCOPE 召回率 ≥ 95%（宁可拒勿错答）。'
    '把 OUT_OF_SCOPE 误判为其他意图，会导致 LLM 胡说八道；'
    '把其他意图误判为 OUT_OF_SCOPE，只是少回答一次，风险低。',
    color='FFEBEE', border_color='F44336')

add_heading_styled(doc, '线上 A/B 测试', level=3)
add_styled_table(doc,
    ['指标', '对照组（无意图识别）', '实验组（有意图识别）', '显著性'],
    [
        ['首字延迟 P95', '~2.5s', '~1.8s', 'p < 0.05'],
        ['用户 👍 率', '—', '提升 ≥ 10%', 'p < 0.05'],
        ['会话轮次', '—', '减少 ≥ 15%', 'p < 0.05'],
        ['LLM 成本/日', '~¥4', '~¥3', '—'],
    ],
    col_widths=[3, 3.5, 4, 2])

add_heading_styled(doc, '上线验收清单', level=3)
add_bullet(doc, '✅ Golden Set 总体准确率 ≥ 92%')
add_bullet(doc, '✅ OUT_OF_SCOPE 召回率 ≥ 95%')
add_bullet(doc, '✅ 规则短路命中率 ≥ 20%')
add_bullet(doc, '✅ LLM 分类 P95 < 400ms')
add_bullet(doc, '✅ 熔断器 + Embedding 兜底验证通过')
add_bullet(doc, '✅ Redis 缓存命中率 ≥ 15%')
add_bullet(doc, '✅ Grafana 仪表盘指标正常')
add_bullet(doc, '✅ A/B 测试首字延迟 P95 降低 ≥ 25%')

doc.add_page_break()

# ==================== 六、总结 ====================
add_heading_styled(doc, '六、总结与展望', level=1)

add_heading_styled(doc, '6.1 设计总结', level=2)

add_paragraph_styled(doc,
    '本方案围绕"意图识别"这一核心方向，从场景、方案、流程、代码、目标五个维度'
    '完整阐述了 CampusShare Agent 意图识别模块的设计：')

add_numbered(doc, '场景：当前平铺式 RAG 存在"简单问题过度处理、复杂问题检索不足、超范围浪费成本"三大痛点')
add_numbered(doc, '方案：对比 6 种技术方案 + 4 家大厂实践，选定"规则短路 + LLM 分类 + Embedding 兜底"混合漏斗架构')
add_numbered(doc, '流程：5 大意图 + 14 子意图体系，三层漏斗逐层过滤，7 个 ADR 记录关键决策')
add_numbered(doc, '代码：新增 7 个类（IntentClassifier/RuleShortCircuitFilter/IntentRouter 等），改造 AgentChatService')
add_numbered(doc, '目标：准确率 ≥ 92%，简单问题提速 15x，LLM 成本降低 25%')

add_heading_styled(doc, '6.2 与现有架构的关系', level=2)

add_paragraph_styled(doc,
    '意图识别不是推翻现有 RAG 实现，而是在其上层增加"智能路由层"。'
    '现有 RetrievalService、DeepSeekClient、AgentChatService 全部复用，'
    '只是 AgentChatService 的编排逻辑从"一刀切 RAG"变为"按意图差异化处理"。')

add_heading_styled(doc, '6.3 后续演进方向', level=2)

add_styled_table(doc,
    ['阶段', '能力', '说明'],
    [
        ['MVP（本方案）', '意图分类 + 路由', '5 意图 + 规则短路 + LLM 分类 + Embedding 兜底'],
        ['Advanced', 'ReAct + 工具调用', 'SEARCH 意图接入 ReAct 循环，支持多轮工具调用'],
        ['Advanced', 'Reranker', 'bge-reranker-v2-m3 重排序，提升检索精度'],
        ['Advanced', '长期记忆', 'user_memory 表，记住用户偏好（学校/分类）'],
        ['Future', '多 Agent 编排', 'Planner + Specialist + Critic 多 Agent 协作'],
        ['Future', 'Function Calling', 'DeepSeek Function Calling 替代显式意图分类'],
    ],
    col_widths=[3, 4, 7])

add_heading_styled(doc, '6.4 风险与应对', level=2)

add_styled_table(doc,
    ['风险', '概率', '影响', '应对'],
    [
        ['LLM 分类延迟超 500ms', '中', '首字延迟劣化', '熔断器 + Embedding 兜底 + Redis 缓存'],
        ['规则短路误判', '低', '用户感知"答非所问"', '仅对高确定性模式短路，置信度 0.95+'],
        ['JSON 解析失败', '中', '降级为 SEARCH', 'try-catch 兜底 + Few-shot 示例约束'],
        ['DeepSeek 服务不可用', '低', 'LLM 分类全失效', 'Embedding 兜底 + SEARCH 默认管线'],
        ['意图体系需扩展', '中', '代码改动', '子意图用 String 而非枚举，便于扩展'],
    ],
    col_widths=[4, 1.5, 3.5, 5])

# ==================== 附录 ====================
doc.add_page_break()
add_heading_styled(doc, '附录：参考文档', level=1)

add_paragraph_styled(doc, '本方案参考了以下项目内文档与外部资料：')

add_heading_styled(doc, '项目内文档', level=2)
add_bullet(doc, 'docs/agent-assistant/04-intent-understanding/intent-taxonomy.md（意图分类体系设计）')
add_bullet(doc, 'docs/agent-assistant/04-intent-understanding/routing-strategy.md（意图路由策略）')
add_bullet(doc, 'docs/agent-assistant/04-intent-understanding/query-rewriting.md（查询改写设计）')
add_bullet(doc, 'docs/agent-assistant/04-intent-understanding/multi-turn-context.md（多轮上下文）')
add_bullet(doc, 'docs/agent-assistant/architecture-overview.md（架构总览，含实现状态）')
add_bullet(doc, 'docs/agent-architecture.md（Agent 模块架构文档）')

add_heading_styled(doc, '技术栈学习文档', level=2)
add_bullet(doc, 'docs/agent技术栈/03-agent-invocation/（Agent 调用与 Function Calling）')
add_bullet(doc, 'docs/agent技术栈/04-RAG检索增强生成/（RAG 与查询改写）')
add_bullet(doc, 'docs/agent技术栈/09-LLM基础原理与推理范式/（LLM 推理范式）')
add_bullet(doc, 'docs/agent技术栈/11-Prompt Engineering/（Prompt 工程最佳实践）')

add_heading_styled(doc, '外部参考', level=2)
add_bullet(doc, 'DeepSeek API 文档：https://platform.deepseek.com/api-docs')
add_bullet(doc, 'OpenAI Function Calling 指南')
add_bullet(doc, '阿里小蜜技术架构公开分享')
add_bullet(doc, 'RAG 基础架构：Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks')

# 保存
output_path = r'e:\workspace_work\CampusShare\docs\agent-design\意图识别模块设计方案.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
