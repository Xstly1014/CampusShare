# -*- coding: utf-8 -*-
"""
生成《知识库管理模块设计方案》Word 文档
这是 Agent 搭建系列第 3 个方向（RAG 数据侧），ADR 前缀 KB。
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_code_block(doc, code_text, font_size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border)
    pPr.append(pBdr)
    lines = code_text.split('\n')
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if i < len(lines) - 1:
            run.add_break()
    return p


def add_callout(doc, text, color='FFF3CD', border_color='FFC107'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), border_color)
        pBdr.append(border)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def style_table_header(row, bg_color='2563EB'):
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_styled_table(doc, headers, rows, col_widths=None, header_bg='2563EB'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    style_table_header(table.rows[0], header_bg)
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_background(row_cells[c_idx], 'F8F9FA')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            run.font.size = Pt(20)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            run.font.size = Pt(13)
    return h


def add_paragraph_styled(doc, text, bold=False, size=10.5, color=None, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


# ==================== 开始生成文档 ====================
doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(10.5)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ==================== 封面 ====================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CampusShare Agent\n知识库管理模块设计方案')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('RAG 的数据侧：摄入 · 分块 · 同步 · 治理')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

for _ in range(8):
    doc.add_paragraph()

info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('文档版本', 'v1.0'),
    ('文档日期', '2026-07-05'),
    ('文档状态', '设计中'),
    ('适用范围', 'campushare-agent 服务 · 知识库数据侧'),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    for p in info_table.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)
    for p in info_table.rows[i].cells[1].paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)
    set_cell_background(info_table.rows[i].cells[0], 'F0F4FF')

doc.add_page_break()

# ==================== 本文档范围声明 ====================
add_heading_styled(doc, '本文档范围声明', level=1)

add_paragraph_styled(doc, '本文件是 Agent 搭建系列文档的第 3 个方向（按依赖排序），聚焦于"知识库管理"——即 RAG 的数据侧。', bold=True)

add_heading_styled(doc, '覆盖什么', level=3)
add_bullet(doc, '文档摄入流水线：扫描 → 解析 → 清洗 → 分块 → Embedding → 入库')
add_bullet(doc, '分块策略：固定大小 / 语义分块 / 递归分块 / Markdown 结构分块 / 句子窗口')
add_bullet(doc, '增量同步机制：知识库 MD5 diff + 帖子库 updated_at + 事件通知 + 定时兜底')
add_bullet(doc, '知识库治理：质量评分 / 重复检测 / 矛盾检测 / 过期清理 / 版本回滚')
add_bullet(doc, 'Embedding 管理：批量调用 / 重试 / 模型版本管理')
add_bullet(doc, '数据库 Schema：knowledge_articles / knowledge_chunks / knowledge_versions')

add_heading_styled(doc, '不覆盖什么', level=3)
add_bullet(doc, '检索查询侧（多路召回 / RRF 融合 / 重排）→ 见《RAG 检索增强生成模块设计方案》')
add_bullet(doc, '意图识别与路由 → 见《意图识别模块设计方案》')
add_bullet(doc, '上下文工程（检索结果如何注入 Prompt）→ 见《上下文工程模块设计方案》')
add_bullet(doc, 'System Prompt 工程 → 见《SystemPrompt 工程模块设计方案》')
add_bullet(doc, '知识图谱（实体关系建模）→ 扩展方向，本文档不涉及')

add_heading_styled(doc, 'ADR 是什么', level=3)
add_callout(doc,
    'ADR = Architecture Decision Record（架构决策记录）。'
    '每条 ADR 记录一个关键的架构决策，包含：背景（为什么需要决策）、决策（选了什么）、理由（为什么这么选）、后果（带来什么影响）。'
    '本系列文档的 ADR 用前缀编号，本文档前缀为 KB（Knowledge Base），共 7 条 ADR-KB-01 ~ ADR-KB-07。'
    '文末附录列出全部 ADR 摘要。',
    color='E8F4FD', border_color='2196F3')

add_heading_styled(doc, '与其他文档的关系', level=3)
add_bullet(doc, '前置文档：System Prompt 工程（无强依赖）、意图识别（无强依赖）')
add_bullet(doc, '后续文档：RAG 检索增强（本文档的下游，检索的数据来源由本文档管理）')
add_bullet(doc, '横向文档：可观测性（摄入流水线的监控）、评估体系（知识库质量评估）')

doc.add_page_break()

# ==================== 目录 ====================
add_heading_styled(doc, '目录', level=1)

toc_items = [
    '一、场景：为什么需要知识库管理',
    '    1.1 业务背景：RAG 只讲了查询，数据怎么进来',
    '    1.2 没有知识库管理会怎样：四大痛点',
    '    1.3 知识库管理带来什么',
    '    1.4 CampusShare 中的两类知识库',
    '二、方案：业界知识库管理设计模式',
    '    2.1 五种分块策略对比',
    '    2.2 大厂案例研究',
    '    2.3 增量同步技术方案对比',
    '    2.4 知识库治理方案',
    '    2.5 ADR 决策汇总',
    '三、流程：如何搭建知识库管理',
    '    3.1 前置条件清单',
    '    3.2 知识库分类：平台知识库 vs 帖子库',
    '    3.3 文档摄入流水线',
    '    3.4 分块策略实现',
    '    3.5 增量同步机制',
    '    3.6 知识库治理',
    '    3.7 关键设计决策（ADR）',
    '四、核心代码',
    '    4.1 文件架构',
    '    4.2 分块器 ChunkSplitter',
    '    4.3 文档摄入服务改造',
    '    4.4 Embedding 批量服务',
    '    4.5 增量同步服务改造',
    '    4.6 知识库治理服务',
    '    4.7 数据库 Schema',
    '五、目标：实现效果',
    '    5.1 功能目标',
    '    5.2 性能目标',
    '    5.3 质量目标',
    '    5.4 成本目标',
    '六、测试评估与验收（详尽版）',
    '    6.1 评估指标体系（含公式）',
    '    6.2 黄金测试集构建',
    '    6.3 评估流水线与 CI/CD 集成',
    '    6.4 LLM-as-Judge 评估',
    '    6.5 错误分析与归因',
    '    6.6 测试用例设计',
    '    6.7 性能与压力测试',
    '    6.8 A/B 测试设计',
    '    6.9 验收流程与准入准出',
    '    6.10 持续监控与漂移检测',
    '七、总结与边界声明',
    '    7.1 核心总结',
    '    7.2 与其他文档的关系',
    '    7.3 演进路线',
    '附录  ADR 摘要',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    if not item.startswith('    '):
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ==================== 一、场景 ====================
add_heading_styled(doc, '一、场景：为什么需要知识库管理', level=1)

add_heading_styled(doc, '1.1 业务背景：RAG 只讲了查询，数据怎么进来', level=2)

add_paragraph_styled(doc,
    '在《RAG 检索增强生成模块设计方案》里，我们讲了"查询侧"——用户提问后，如何用多路召回 + RRF 融合 + '
    'Cross-encoder 重排，从知识库中检索出最相关的文档。但那份文档有一个隐含假设：知识库里已经有高质量的数据了。')

add_paragraph_styled(doc,
    '问题是：这些数据是怎么进来的？怎么保证它是新的、全的、对的？怎么保证分块合理、Embedding 准确？'
    '这些问题在查询侧文档里几乎没提——它只关心"怎么查"，不关心"数据怎么来"。')

add_paragraph_styled(doc,
    '这就像建了一个图书馆，配了最好的检索系统，但从来没人管"书怎么买进来、怎么分类上架、怎么剔除过期书"。'
    '时间一长，图书馆里堆满了过期的、重复的、撕掉几页的书，再好的检索系统也查不出好结果。', bold=True)

add_callout(doc,
    '核心矛盾：RAG 查询侧解决"怎么查"，知识库管理解决"数据怎么来、怎么管"。'
    '两者是 RAG 系统的两条腿，缺一不可。查询侧再先进，数据侧一团糟，召回率照样上不去。')

add_heading_styled(doc, '1.2 没有知识库管理会怎样：四大痛点', level=2)

add_paragraph_styled(doc,
    '我们看当前 CampusShare Agent 的实现（KnowledgeIngestionService + PostVectorService），'
    '存在四个严重问题：')

add_heading_styled(doc, '痛点一：无分块，长文档召回差', level=3)
add_paragraph_styled(doc,
    '当前 KnowledgeIngestionService.ingestAll() 的实现：扫描 .md 文件 → 解析 frontmatter → '
    '计算 MD5 → 截取前 500 字符作为 excerpt → 整个文档生成一个 embedding 向量。')

add_code_block(doc, """// 当前实现（KnowledgeIngestionService.java 第 103-112 行）
String excerpt = doc.content.length() > EXCERPT_LENGTH  // EXCERPT_LENGTH = 500
        ? doc.content.substring(0, EXCERPT_LENGTH)
        : doc.content;

float[] embedding = embeddingClient.embed(doc.title + "\\n" + excerpt).block();
// 整个文档 = 一个 embedding 向量""")

add_paragraph_styled(doc, '这导致三个问题：')
add_bullet(doc, '长文档信息丢失：一篇 3000 字的"创作者认证指南"，只有前 500 字进入向量库，后 2500 字完全检索不到')
add_bullet(doc, '召回精度低：用户问"创作者认证的审核要多久"，答案在第 1200 字，但向量库里只有前 500 字，召回不到')
add_bullet(doc, 'Embedding 表达力弱：一个向量要表示整篇文档的主题，语义被稀释，相似度计算不准')

add_heading_styled(doc, '痛点二：增量同步粗糙', level=3)
add_paragraph_styled(doc, '当前两类知识库的同步机制都很粗糙：')

add_styled_table(doc,
    ['知识库', '当前同步机制', '问题'],
    [
        ['平台知识库\n（~50 篇 .md）',
         'KnowledgeScheduler 每小时扫描全目录，MD5 diff 跳过未变更的',
         '全目录扫描效率低（50 篇还好，500 篇就慢）；无法感知"哪篇改了"，只能全量对比'],
        ['帖子库\n（数万~百万帖）',
         'PostVectorScheduler 每 5 分钟全量分页拉取 + 事件通知（syncPost）',
         '5 分钟全量同步是兜底机制，10 万帖时每次同步耗时长、Embedding API 调用多；没有基于 updated_at 的增量'],
    ],
    col_widths=[3.5, 5.5, 5.5])

add_paragraph_styled(doc,
    '帖子库的"全量同步"问题尤其严重：假设有 10 万帖子，每 5 分钟全量同步一次，'
    '即使没有任何帖子更新，也要拉取 10 万帖 + 调用 10 万次 Embedding（批量后约 6250 次批量调用）。'
    '这是巨大的浪费。', bold=True)

add_heading_styled(doc, '痛点三：知识质量无保障', level=3)
add_paragraph_styled(doc, '当前实现没有任何质量保障机制：')
add_bullet(doc, '重复检测缺失：两篇标题相近、内容相似的文档会同时入库，检索时重复召回')
add_bullet(doc, '矛盾检测缺失：旧版规则说"发帖满 10 篇可认证"，新版规则说"获赞≥10000 且发帖≥50"，两篇都入库会矛盾')
add_bullet(doc, '过期清理缺失：DEPRECATED 状态的文档还在向量库里（status 字段在向量表，但摄入时直接写 PUBLISHED）')
add_bullet(doc, '质量评分缺失：不知道哪些文档被高频召回、哪些从未被召回（可能是垃圾文档）')

add_heading_styled(doc, '痛点四：无版本管理，无法回滚', level=3)
add_paragraph_styled(doc,
    '当前 knowledge_articles 表有 version 字段，但只是个自增整数，没有版本历史表。'
    '一旦文档被更新，旧版本内容就丢了——无法回滚到上一个版本，无法对比"这次改了什么"。')

add_paragraph_styled(doc,
    '更严重的是 Embedding 模型版本管理缺失：knowledge_vectors 表的 embedding_model 字段写死 "bge-m3"，'
    '一旦升级 Embedding 模型（如换 text-embedding-3-large），旧向量和新向量混在一起，相似度计算不准。', bold=True)

add_heading_styled(doc, '1.3 知识库管理带来什么', level=2)

add_styled_table(doc,
    ['能力', '当前（无知识库管理）', '目标（有知识库管理）', '提升'],
    [
        ['分块精度', '整篇 1 个向量，前 500 字', '按 Markdown 结构分块，256 Token/块', '长文档召回率 +40%'],
        ['同步效率', '全量扫描，10 万帖 5 分钟一轮', 'updated_at 增量 + 事件通知', '同步耗时降 90%'],
        ['知识质量', '无检测，重复/矛盾/过期共存', '质量评分 + 重复/矛盾检测 + 过期清理', '重复召回率降 80%'],
        ['版本管理', 'version 自增，无历史', 'SemVer + 全量历史 + 回滚', '可回滚，可对比'],
        ['Embedding 一致性', '模型写死，升级即混乱', '模型版本字段 + 批量重嵌入', '模型升级零风险'],
    ],
    col_widths=[3, 4.5, 4.5, 3])

add_heading_styled(doc, '1.4 CampusShare 中的两类知识库', level=2)

add_paragraph_styled(doc, 'CampusShare Agent 有两类知识库，特征完全不同，必须分开管理：')

add_styled_table(doc,
    ['维度', '平台知识库', '帖子库'],
    [
        ['内容来源', '人工维护的 .md 文档（平台功能说明）', '用户生成的帖子（UGC）'],
        ['数量级', '~50 篇', '数万 ~ 百万帖'],
        ['更新频率', '低（月级，平台规则变更时）', '高（秒级，用户持续发帖）'],
        ['质量', '高（人工审核后发布）', '参差（用户随意发，可能有错别字/灌水）'],
        ['分块策略', 'Markdown 结构分块（按 H2/H3 标题切分）', '标题 + 摘要整体 embedding（帖文太短不值得分块）'],
        ['同步机制', 'MD5 diff + 手动触发', 'updated_at 增量 + 事件通知 + 定时兜底'],
        ['版本管理', 'SemVer + 全量历史 + 回滚', 'updated_at 时间戳（不保留历史版本）'],
        ['治理重点', '矛盾检测（规则变更时新旧冲突）', '过期清理（删除的帖子要从向量库移除）'],
        ['服务的意图', 'HOW_TO / NAVIGATE', 'SEARCH'],
        ['物理存储', 'MySQL knowledge_articles + PG knowledge_vectors', 'PG post_vectors（不进 MySQL）'],
    ],
    col_widths=[2.5, 6, 6])

add_callout(doc,
    '关键设计决策（ADR-KB-01）：知识库与帖子库物理分离。'
    '这是从 RAG 检索侧延续的决策（ADR-RAG-016），在数据侧进一步细化：'
    '平台知识库走 MySQL → PG 双写（MySQL 存全文 + 版本，PG 存向量）；'
    '帖子库只走 PG（帖子全文在 post-service 的 MySQL，agent-service 不重复存储）。',
    color='FFF3E0', border_color='FF9800')

doc.add_page_break()

# ==================== 二、方案 ====================
add_heading_styled(doc, '二、方案：业界知识库管理设计模式', level=1)

add_heading_styled(doc, '2.1 五种分块策略对比', level=2)

add_paragraph_styled(doc, '分块（Chunking）是知识库管理最核心的决策。分块策略直接决定召回率的上限。')

add_styled_table(doc,
    ['策略', '原理', '优点', '缺点', '适用场景'],
    [
        ['固定大小分块',
         '按固定字符数/Token 数切分，可加 overlap',
         '实现最简单、可预测',
         '可能切断句子/段落，语义不完整',
         '快速 MVP、纯文本'],
        ['语义分块',
         '用 Embedding 计算相邻句子相似度，在语义跳变处切分',
         '语义完整、召回质量高',
         '计算成本高（每句一次 Embedding）',
         '高质量要求、长文档'],
        ['递归分块',
         '按分隔符优先级递归切分（\\n\\n → \\n → 。 → 空格）',
         '平衡语义与性能、保留结构',
         '块大小不均匀',
         '通用场景、LangChain 默认'],
        ['Markdown 结构分块',
         '按 H1/H2/H3 标题切分，保留文档结构',
         '结构清晰、块边界有意义',
         '只适用于 Markdown',
         '文档型知识库（本项目平台知识库）'],
        ['句子窗口',
         '检索单句，但返回该句前后各 N 句作为上下文',
         '召回精准 + 上下文完整',
         '存储成本高（每句一个向量）',
         '精确问答场景'],
    ],
    col_widths=[2.5, 3.5, 2.5, 2.5, 3])

add_paragraph_styled(doc,
    '结论：平台知识库用 Markdown 结构分块（ADR-KB-02），帖子库用标题+摘要整体 embedding（帖文已足够短）。', bold=True)

add_heading_styled(doc, '2.2 大厂案例研究', level=2)

add_heading_styled(doc, '2.2.1 OpenAI File Search（Assistants API）', level=3)
add_bullet(doc, '分块：自动分块，策略未公开（推测是递归分块 + 语义优化）')
add_bullet(doc, '同步：文件上传后自动解析 + 分块 + Embedding，无需用户干预')
add_bullet(doc, '治理：文件可删除/替换，但无版本管理；向量库托管，不可见')
add_bullet(doc, '启示：全托管模式适合 SaaS，但企业自建需要可控的分块和版本管理')

add_heading_styled(doc, '2.2.2 LlamaIndex', level=3)
add_bullet(doc, '分块：NodeParser 抽象，支持 SentenceSplitter / MarkdownNodeParser / SemanticSplitter')
add_bullet(doc, '同步：文档加载器（Document Loader）支持目录/数据库/API/Notion 等多种来源')
add_bullet(doc, '治理：支持文档级 metadata，但版本管理需自建')
add_bullet(doc, '启示：分块策略可插拔是工业标准；MarkdownNodeParser 是文档型知识库的最佳实践')

add_heading_styled(doc, '2.2.3 LangChain', level=3)
add_bullet(doc, '分块：RecursiveCharacterTextSplitter 是默认推荐，按分隔符优先级递归切分')
add_bullet(doc, '同步：DocumentLoader 抽象，支持 80+ 数据源连接器')
add_bullet(doc, '治理：支持文档去重（按 metadata）、过期清理（按时间戳）')
add_bullet(doc, '启示：递归分块的"分隔符优先级"思想值得借鉴（先段落、再句子、再字符）')

add_heading_styled(doc, '2.2.4 Dify', level=3)
add_bullet(doc, '分块：支持自动分块 + 自定义分块（用户可调 chunk_size / overlap / 分隔符）')
add_bullet(doc, '同步：支持 Notion / 飞书 / GitHub / 网页 / API / 本地文件')
add_bullet(doc, '治理：提供"知识库质量"页面，展示召回测试结果、文档命中统计')
add_bullet(doc, '启示：质量可视化是大厂标配；召回测试应该有 UI 入口')

add_heading_styled(doc, '2.2.5 阿里通义 / 百度千帆', level=3)
add_bullet(doc, '分块：智能分块（结合语义 + 结构），支持 Markdown / PDF / Word / HTML')
add_bullet(doc, '同步：支持增量同步（按文件 MD5 / 文档 updated_at）')
add_bullet(doc, '治理：支持文档质量评分（基于召回频次 + 用户反馈）')
add_bullet(doc, '启示：增量同步用 MD5/updated_at 是行业标准；质量评分要结合使用数据')

add_heading_styled(doc, '2.2.6 大厂案例总结', level=3)

add_styled_table(doc,
    ['维度', '行业最佳实践', 'CampusShare 选型'],
    [
        ['分块策略', '可插拔、按文档类型选择', '平台知识库用 Markdown 结构分块，帖子库整体 embedding'],
        ['同步机制', '增量同步（MD5/updated_at）+ 事件通知', '知识库 MD5 diff，帖子库 updated_at + 事件通知'],
        ['版本管理', 'SemVer + 全量历史 + 回滚', '平台知识库 SemVer + 历史，帖子库不保留历史'],
        ['质量治理', '召回频次统计 + 重复检测 + 过期清理', '质量评分 + 重复检测 + 矛盾检测 + 过期清理'],
        ['Embedding 版本', '模型版本字段 + 批量重嵌入', 'embedding_model 字段 + 重嵌入脚本'],
    ],
    col_widths=[3, 5.5, 5.5])

add_heading_styled(doc, '2.3 增量同步技术方案对比', level=2)

add_styled_table(doc,
    ['方案', '原理', '优点', '缺点', '适用'],
    [
        ['全量同步',
         '定期拉取全部数据，重新 Embedding',
         '实现最简单、绝对一致',
         '浪费算力、延迟高、不可扩展',
         '小数据量（<1000）'],
        ['MD5 diff',
         '全量扫描但对比 MD5，未变更的跳过',
         '避免重复 Embedding',
         '仍需全量扫描、无法感知删除',
         '中等数据量 + 文件型'],
        ['updated_at 增量',
         '只拉取 updated_at > 上次同步时间的记录',
         '高效、可扩展',
         '需要源表有 updated_at 字段',
         '数据库型（本项目帖子库）'],
        ['事件通知',
         '源系统变更时主动通知（如 post-service 发 HTTP）',
         '近实时、零浪费',
         '依赖源系统可靠性、可能丢消息',
         '高频变更场景（本项目帖子库）'],
        ['CDC（Change Data Capture）',
         '监听数据库 binlog，捕获变更事件',
         '实时、对源系统零侵入',
         '部署复杂（Debezium/Kafka）',
         '大型系统（暂不需要）'],
    ],
    col_widths=[2.5, 3.5, 2.5, 3, 3])

add_paragraph_styled(doc,
    '结论：帖子库用 updated_at 增量 + 事件通知 + 定时兜底（ADR-KB-03）；'
    '平台知识库用 MD5 diff（量小，全量扫描可接受）。', bold=True)

add_heading_styled(doc, '2.4 知识库治理方案', level=2)

add_paragraph_styled(doc, '知识库治理解决"数据进了库之后怎么管"的问题，包含四个子方向：')

add_styled_table(doc,
    ['治理项', '问题', '方案', '触发时机'],
    [
        ['质量评分',
         '不知道哪些文档有用、哪些是垃圾',
         '统计每篇文档的召回频次 + 用户反馈（点赞/点踩），计算质量分',
         '每天定时统计'],
        ['重复检测',
         '相似文档重复召回，浪费上下文 Token',
         '新文档入库时，用 Embedding 相似度对比已有文档，相似度 > 0.95 标记为重复',
         '入库时实时检测'],
        ['矛盾检测',
         '新旧规则冲突，回答自相矛盾',
         '同 topic 下的文档，用 LLM 对比关键结论是否矛盾',
         '平台知识库更新时'],
        ['过期清理',
         'DEPRECATED 文档还在向量库，被错误召回',
         '摄入时同步状态字段；定时任务清理 status=DEPRECATED 的向量',
         '摄入时 + 每日定时'],
    ],
    col_widths=[2.5, 4, 5, 3])

add_heading_styled(doc, '2.5 ADR 决策汇总', level=2)

add_styled_table(doc,
    ['ADR 编号', '决策', '核心理由'],
    [
        ['ADR-KB-01', '知识库与帖子库物理分离', '特征完全不同，混在一起会互相污染'],
        ['ADR-KB-02', '平台知识库用 Markdown 结构分块，帖子库整体 embedding', '匹配数据特征：文档有结构、帖文短'],
        ['ADR-KB-03', '帖子库 updated_at 增量 + 事件通知 + 定时兜底', '兼顾实时性、可靠性、可扩展性'],
        ['ADR-KB-04', '平台知识库 SemVer + 全量历史 + 回滚', '平台规则变更需要可追溯、可回滚'],
        ['ADR-KB-05', 'Embedding 批量调用，batch=16，重试 3 次', '平衡 API 限流与吞吐'],
        ['ADR-KB-06', '质量评分 + 重复/矛盾检测 + 过期清理', '保证知识库长期可用'],
        ['ADR-KB-07', '分块参数：目标 256 Token，重叠 50，最大 512', '平衡召回精度与上下文完整性'],
    ],
    col_widths=[2.5, 6, 6])

doc.add_page_break()

# ==================== 三、流程 ====================
add_heading_styled(doc, '三、流程：如何搭建知识库管理', level=1)

add_heading_styled(doc, '3.1 前置条件清单', level=2)

add_paragraph_styled(doc, '搭建知识库管理前，需要以下前置条件已就绪：')

add_styled_table(doc,
    ['前置项', '要求', '当前状态'],
    [
        ['Embedding 服务', 'DeepSeek embedding API 可用，批量调用支持', '✅ EmbeddingClient 已实现'],
        ['向量数据库', 'PostgreSQL + pgvector，HNSW 索引', '✅ agent-postgres 已部署'],
        ['MySQL 业务库', 'knowledge_articles 表已建', '✅ agent-init.sql 已含'],
        ['知识文档目录', 'docs/agent-assistant/knowledge-docs/*.md', '✅ 已有 ~50 篇'],
        ['PostFeignClient', '可拉取帖子数据 + updated_at 字段', '⚠️ 需新增 updated_at 参数支持增量'],
        ['JTokkit 依赖', '本地 Token 估算', '✅ pom.xml 已含'],
        ['Resilience4j', 'CircuitBreaker + Retry', '✅ ResilienceConfig 已配'],
    ],
    col_widths=[3.5, 5, 5])

add_heading_styled(doc, '3.2 知识库分类：平台知识库 vs 帖子库', level=2)

add_paragraph_styled(doc,
    '两类知识库的特征完全不同（见 1.4 节），必须用不同的摄入流水线和治理策略。'
    '本节明确两者的边界：')

add_code_block(doc, """┌─────────────────────────────────────────────────────────────┐
│  平台知识库（KnowledgeBase）                                  │
│  ├─ 来源：docs/agent-assistant/knowledge-docs/*.md（人工维护）│
│  ├─ 存储：MySQL knowledge_articles（全文+版本）                │
│  │         PG knowledge_vectors（向量+分块）                   │
│  ├─ 分块：Markdown 结构分块（ChunkSplitter）                  │
│  ├─ 同步：MD5 diff，每小时扫描                                │
│  └─ 治理：矛盾检测 + 版本回滚                                 │
├─────────────────────────────────────────────────────────────┤
│  帖子库（PostVector）                                         │
│  ├─ 来源：post-service MySQL（用户生成的帖子）                │
│  ├─ 存储：PG post_vectors（只存向量+摘要，不存全文）           │
│  ├─ 分块：整体 embedding（title + content_excerpt）           │
│  ├─ 同步：事件通知 + updated_at 增量 + 5 分钟兜底             │
│  └─ 治理：过期清理（删除的帖子移除向量）                       │
└─────────────────────────────────────────────────────────────┘""")

add_heading_styled(doc, '3.3 文档摄入流水线', level=2)

add_paragraph_styled(doc, '平台知识库的摄入流水线分六步：')

add_code_block(doc, """平台知识库摄入流水线：

  ┌─ Step 1: 扫描 ─────────────────────────────────────┐
  │  扫描 docs-path 目录下所有 .md 文件                 │
  │  → 输出：List<Path> mdFiles                        │
  └─────────────────────────────────────────────────────┘
                  │
                  ▼
  ┌─ Step 2: 解析 ─────────────────────────────────────┐
  │  解析 frontmatter（title/topic/tags）+ 正文         │
  │  → 输出：ParsedDoc(title, topic, tags, content)   │
  └─────────────────────────────────────────────────────┘
                  │
                  ▼
  ┌─ Step 3: 变更检测 ─────────────────────────────────┐
  │  计算 content 的 MD5，与 MySQL 已有记录对比          │
  │  → 未变更：跳过                                     │
  │  → 变更/新增：进入 Step 4                           │
  └─────────────────────────────────────────────────────┘
                  │
                  ▼
  ┌─ Step 4: 分块 ─────────────────────────────────────┐
  │  ChunkSplitter.split(content)                      │
  │  → 按 Markdown H2/H3 标题切分                       │
  │  → 块大小目标 256 Token，重叠 50，最大 512           │
  │  → 输出：List<Chunk> chunks                         │
  └─────────────────────────────────────────────────────┘
                  │
                  ▼
  ┌─ Step 5: Embedding ────────────────────────────────┐
  │  批量调用 Embedding API（batch=16，重试 3 次）       │
  │  → 输出：List<float[]> embeddings                   │
  └─────────────────────────────────────────────────────┘
                  │
                  ▼
  ┌─ Step 6: 入库 ─────────────────────────────────────┐
  │  ① MySQL: knowledge_articles（全文+版本+MD5）       │
  │           knowledge_article_versions（版本历史）     │
  │  ② PG: knowledge_chunks（分块+向量）                │
  │       （旧的 knowledge_vectors 表迁移为分块表）      │
  └─────────────────────────────────────────────────────┘""")

add_heading_styled(doc, '3.4 分块策略实现', level=2)

add_paragraph_styled(doc,
    '平台知识库用 Markdown 结构分块（ADR-KB-02），算法如下：')

add_numbered(doc, '按 H2（## ）标题切分为大段，每段对应一个主题')
add_numbered(doc, '若大段超过 512 Token，按 H3（### ）标题二次切分')
add_numbered(doc, '若 H3 段仍超过 512 Token，按段落（\\n\\n）三次切分')
add_numbered(doc, '若段落仍超过 512 Token，按句子（。！？）四次切分')
add_numbered(doc, '相邻块之间保留 50 Token 重叠（避免句子边界信息丢失）')
add_numbered(doc, '每个块携带元数据：articleId / chunkIndex / heading（所属标题）')

add_paragraph_styled(doc, '分块参数（ADR-KB-07）：')

add_styled_table(doc,
    ['参数', '值', '理由'],
    [
        ['目标块大小', '256 Token', 'DeepSeek Embedding 最佳输入长度 128-512，256 是甜点'],
        ['最大块大小', '512 Token', '超过 512 会稀释向量表达力'],
        ['重叠大小', '50 Token', '约 1-2 句话，保证语义连续性'],
        ['最小块大小', '64 Token', '太小的块语义不足，合并到上一块'],
        ['切分优先级', 'H2 → H3 → 段落 → 句子', '保留文档结构，块边界有意义'],
    ],
    col_widths=[3, 3, 7.5])

add_heading_styled(doc, '3.5 增量同步机制', level=2)

add_heading_styled(doc, '3.5.1 帖子库增量同步', level=3)

add_paragraph_styled(doc,
    '帖子库当前是"5 分钟全量同步 + 事件通知"双机制。改造为"updated_at 增量 + 事件通知 + 定时兜底"三机制：')

add_code_block(doc, """帖子库三机制协同：

  机制一：事件通知（近实时，主路径）
  ─────────────────────────────────
  post-service 帖子变更 → 调用 agent-service /api/internal/post-notify
  → PostVectorService.syncPost(postId, action)
  → 单帖 Embedding + upsert
  优点：近实时、零浪费
  风险：通知可能丢失（网络抖动、服务重启）

  机制二：updated_at 增量（补漏，定时）
  ─────────────────────────────────
  PostVectorScheduler 每 5 分钟触发
  → 查询 post-service：updated_at > 上次同步时间
  → 只对变更的帖子做 Embedding + upsert
  优点：补漏通知丢失的情况、只处理变更
  实现：PostFeignClient 新增 getUpdatedAfter(updatedAt, page, size)

  机制三：全量同步（兜底，低频）
  ─────────────────────────────────
  每天凌晨 3 点全量同步一次
  → 拉取全部帖子，MD5 diff 跳过未变更的
  优点：保证最终一致性
  场景：updated_at 字段异常、向量库数据丢失""")

add_heading_styled(doc, '3.5.2 平台知识库同步', level=3)

add_paragraph_styled(doc,
    '平台知识库量小（~50 篇），用 MD5 diff 足够：')

add_bullet(doc, 'KnowledgeScheduler 每小时扫描全目录')
add_bullet(doc, '对比每篇文档的 content MD5，未变更的跳过')
add_bullet(doc, '变更的文档：旧版本写入 knowledge_article_versions（版本历史），新版本写入 knowledge_articles')
add_bullet(doc, '分块重新生成，旧分块删除（按 article_id 删除 knowledge_chunks）')

add_heading_styled(doc, '3.6 知识库治理', level=2)

add_heading_styled(doc, '3.6.1 质量评分', level=3)

add_paragraph_styled(doc, '每天定时统计每篇文档/每个帖子的召回频次和用户反馈，计算质量分：')

add_code_block(doc, """质量评分公式：

  quality_score = 0.4 * recall_frequency_score
                + 0.3 * user_feedback_score
                + 0.2 * freshness_score
                + 0.1 * completeness_score

  - recall_frequency_score: 过去 7 天被召回次数 / 同 topic 平均召回次数
  - user_feedback_score:    (点赞数 - 点踩数) / 召回次数
  - freshness_score:        1 - min(days_since_update / 90, 1)
  - completeness_score:     分块数 / 预期分块数（检测截断）

  质量分 < 0.3 的文档标记为 LOW_QUALITY，进入待清理队列""")

add_heading_styled(doc, '3.6.2 重复检测', level=3)

add_paragraph_styled(doc,
    '新文档入库时，用 Embedding 相似度对比同 topic 下已有文档：')

add_bullet(doc, '计算新文档 Embedding 与同 topic 已有文档 Embedding 的余弦相似度')
add_bullet(doc, '相似度 > 0.95：标记为 DUPLICATE，不自动入库，进入人工审核队列')
add_bullet(doc, '相似度 0.85~0.95：标记为 SIMILAR，入库但记录关联，检索时降权')
add_bullet(doc, '相似度 < 0.85：正常入库')

add_heading_styled(doc, '3.6.3 矛盾检测', level=3)

add_paragraph_styled(doc,
    '平台知识库更新时，对同 topic 下的新旧版本用 LLM 检测关键结论是否矛盾：')

add_code_block(doc, """矛盾检测 Prompt（简化版）：

  系统：你是知识库审核员。对比以下两版文档，判断关键结论是否矛盾。
        只关注"规则/条件/流程"类结论，忽略措辞差异。

  旧版（v1.0）：{old_content}
  新版（v2.0）：{new_content}

  输出 JSON：
  {
    "has_contradiction": true/false,
    "contradictions": [
      {"old": "发帖满10篇可认证", "new": "获赞≥10000且发帖≥50"}
    ]
  }

  若 has_contradiction=true，标记新版本为 NEEDS_REVIEW，进入人工审核""")

add_heading_styled(doc, '3.6.4 过期清理', level=3)

add_bullet(doc, '平台知识库：摄入时同步 status 字段，status=DEPRECATED 的文档从向量库删除（保留 MySQL 全文用于审计）')
add_bullet(doc, '帖子库：接收 post-service DELETE 通知时，从 post_vectors 删除对应向量')
add_bullet(doc, '兜底：每天凌晨 3 点全量同步时，对比 MySQL 现有帖子 ID，删除向量库中多余的（已删除的帖子）')

add_heading_styled(doc, '3.7 关键设计决策（ADR）', level=2)

add_styled_table(doc,
    ['ADR', '决策', '背景', '理由', '后果'],
    [
        ['ADR-KB-01',
         '知识库与帖子库物理分离',
         '两类数据特征完全不同',
         '混在一起检索会互相污染',
         '需要两套摄入流水线，复杂度增加'],
        ['ADR-KB-02',
         '平台知识库用 Markdown 结构分块',
         '文档有清晰的 H2/H3 结构',
         '保留结构、块边界有意义',
         '只适用于 Markdown，PDF/Word 需另外适配'],
        ['ADR-KB-03',
         '帖子库三机制同步',
         '事件通知可能丢失、全量同步太慢',
         '近实时 + 补漏 + 兜底，兼顾可靠与高效',
         '需要 post-service 新增 updated_at 查询接口'],
        ['ADR-KB-04',
         '平台知识库 SemVer + 版本历史',
         '平台规则变更需要可追溯',
         '可回滚、可对比',
         '存储成本增加（每版本一份全文）'],
        ['ADR-KB-05',
         'Embedding batch=16，重试 3 次',
         'DeepSeek API 限流',
         '平衡吞吐与限流',
         '单次批量最多 16 条，大批量需多次调用'],
        ['ADR-KB-06',
         '四维质量评分',
         '不知道哪些文档有用',
         '量化质量、驱动清理',
         '需要记录召回频次和用户反馈'],
        ['ADR-KB-07',
         '分块目标 256 Token',
         'DeepSeek Embedding 最佳输入 128-512',
         '256 是甜点，召回精度与上下文兼顾',
         '长文档会切成多个块，向量数增加'],
    ],
    col_widths=[2, 3, 3, 3, 3.5])

doc.add_page_break()

# ==================== 四、核心代码 ====================
add_heading_styled(doc, '四、核心代码', level=1)

add_heading_styled(doc, '4.1 文件架构', level=2)

add_code_block(doc, """backend/campushare-agent/src/main/java/com/campushare/agent/
├── service/
│   ├── KnowledgeIngestionService.java    # 改造：支持分块 + 版本管理
│   ├── PostVectorService.java            # 改造：支持 updated_at 增量
│   ├── KnowledgeGovernanceService.java   # 新增：质量评分 + 重复/矛盾检测
│   └── EmbeddingBatchService.java        # 新增：批量 Embedding + 重试
├── chunk/
│   ├── ChunkSplitter.java                # 新增：分块器接口
│   ├── MarkdownChunkSplitter.java        # 新增：Markdown 结构分块
│   └── Chunk.java                        # 新增：分块实体
├── store/
│   ├── KnowledgeVectorStore.java         # 改造：支持分块向量
│   ├── KnowledgeChunkStore.java          # 新增：分块向量存储
│   └── PostVectorStore.java              # 改造：支持增量删除
├── entity/
│   ├── KnowledgeArticle.java             # 改造：新增 version_status 字段
│   ├── KnowledgeArticleVersion.java      # 新增：版本历史实体
│   └── KnowledgeChunk.java               # 新增：分块实体
└── mapper/
    ├── KnowledgeArticleMapper.java
    ├── KnowledgeArticleVersionMapper.java # 新增
    └── KnowledgeChunkMapper.java          # 新增""")

add_heading_styled(doc, '4.2 分块器 ChunkSplitter', level=2)

add_paragraph_styled(doc, '分块器接口 + Markdown 实现：')

add_code_block(doc, '''package com.campushare.agent.chunk;

import java.util.List;

public interface ChunkSplitter {
    List<Chunk> split(String content, Long articleId);
}
''')

add_code_block(doc, '''package com.campushare.agent.chunk;

import lombok.AllArgsConstructor;
import lombok.Builder;
import lombok.Data;

@Data
@Builder
@AllArgsConstructor
public class Chunk {
    private Long articleId;
    private int chunkIndex;        // 块序号（从 0 开始）
    private String heading;        // 所属标题（如 "## 创作者认证条件"）
    private String content;        // 块文本
    private int tokenCount;        // 块 Token 数（JTokkit 估算）
    private int startOffset;       // 在原文中的起始字符偏移
    private int endOffset;         // 结束偏移
}
''')

add_code_block(doc, '''package com.campushare.agent.chunk;

import com.knuddels.jtokkit.Encodings;
import com.knuddels.jtokkit.api.Encoding;
import com.knuddels.jtokkit.api.EncodingRegistry;
import com.knuddels.jtokkit.api.EncodingType;
import org.springframework.stereotype.Component;

import java.util.ArrayList;
import java.util.List;
import java.util.regex.Matcher;
import java.util.regex.Pattern;

/**
 * Markdown 结构分块器。
 *
 * 切分优先级：H2 → H3 → 段落（\\n\\n）→ 句子（。！？）
 * 块参数：目标 256 Token，最大 512，最小 64，重叠 50 Token（ADR-KB-07）
 */
@Component
public class MarkdownChunkSplitter implements ChunkSplitter {

    private static final int TARGET_TOKENS = 256;
    private static final int MAX_TOKENS = 512;
    private static final int MIN_TOKENS = 64;
    private static final int OVERLAP_TOKENS = 50;

    private static final Pattern H2_PATTERN = Pattern.compile("^## .+$", Pattern.MULTILINE);
    private static final Pattern H3_PATTERN = Pattern.compile("^### .+$", Pattern.MULTILINE);

    private final Encoding encoding;

    public MarkdownChunkSplitter() {
        EncodingRegistry registry = Encodings.newDefaultEncodingRegistry();
        this.encoding = registry.getEncoding(EncodingType.CL100K_BASE);
    }

    @Override
    public List<Chunk> split(String content, Long articleId) {
        List<Chunk> chunks = new ArrayList<>();
        // Step 1: 按 H2 切分
        List<Section> h2Sections = splitByPattern(content, H2_PATTERN);
        int chunkIndex = 0;
        for (Section h2 : h2Sections) {
            String h2Title = h2.heading;
            // Step 2: 若 H2 段超过 MAX_TOKENS，按 H3 切分
            if (estimateTokens(h2.body) > MAX_TOKENS) {
                List<Section> h3Sections = splitByPattern(h2.body, H3_PATTERN);
                for (Section h3 : h3Sections) {
                    String heading = h3.heading != null ? h2Title + " > " + h3.heading : h2Title;
                    chunkIndex = addChunks(chunks, articleId, chunkIndex, heading, h3.body);
                }
            } else {
                chunkIndex = addChunks(chunks, articleId, chunkIndex, h2Title, h2.body);
            }
        }
        return chunks;
    }

    private int addChunks(List<Chunk> chunks, Long articleId, int startIdx, String heading, String text) {
        int idx = startIdx;
        // 若文本超过 MAX_TOKENS，按段落/句子进一步切分
        List<String> pieces = splitToTokenLimit(text);
        String prevTail = "";
        for (String piece : pieces) {
            String fullPiece = prevTail + piece;
            int tokens = estimateTokens(fullPiece);
            if (tokens < MIN_TOKENS && chunks.size() > 0) {
                // 太小，合并到上一块
                Chunk last = chunks.get(chunks.size() - 1);
                last.setContent(last.getContent() + fullPiece);
                last.setTokenCount(last.getTokenCount() + tokens);
                prevTail = "";
                continue;
            }
            chunks.add(Chunk.builder()
                    .articleId(articleId)
                    .chunkIndex(idx++)
                    .heading(heading)
                    .content(fullPiece)
                    .tokenCount(tokens)
                    .build());
            // 保留 overlap（取末尾约 50 Token）
            prevTail = getTail(fullPiece, OVERLAP_TOKENS);
        }
        return idx;
    }

    private List<String> splitToTokenLimit(String text) {
        List<String> result = new ArrayList<>();
        if (estimateTokens(text) <= MAX_TOKENS) {
            result.add(text);
            return result;
        }
        // 按段落切分
        String[] paragraphs = text.split("\\n\\n");
        StringBuilder buffer = new StringBuilder();
        for (String para : paragraphs) {
            if (estimateTokens(buffer.toString() + para) > MAX_TOKENS && buffer.length() > 0) {
                result.add(buffer.toString());
                buffer = new StringBuilder();
            }
            buffer.append(para).append("\\n\\n");
        }
        if (buffer.length() > 0) result.add(buffer.toString());
        return result;
    }

    private List<Section> splitByPattern(String content, Pattern pattern) {
        List<Section> sections = new ArrayList<>();
        Matcher matcher = pattern.matcher(content);
        int lastEnd = 0;
        String lastHeading = null;
        while (matcher.find()) {
            if (lastHeading != null || lastEnd > 0) {
                String body = content.substring(lastEnd, matcher.start()).trim();
                if (!body.isEmpty()) {
                    sections.add(new Section(lastHeading, body));
                }
            }
            lastHeading = matcher.group().trim();
            lastEnd = matcher.end();
        }
        String tail = content.substring(lastEnd).trim();
        if (!tail.isEmpty()) {
            sections.add(new Section(lastHeading, tail));
        }
        return sections;
    }

    private int estimateTokens(String text) {
        return encoding.countTokens(text);
    }

    private String getTail(String text, int targetTokens) {
        List<Integer> tokens = encoding.encode(text);
        if (tokens.size() <= targetTokens) return text;
        return encoding.decode(tokens.subList(tokens.size() - targetTokens, tokens.size()));
    }

    private record Section(String heading, String body) {}
}
''')

add_heading_styled(doc, '4.3 文档摄入服务改造', level=2)

add_paragraph_styled(doc, '改造 KnowledgeIngestionService.ingestAll()，增加分块 + 版本历史：')

add_code_block(doc, '''// 改造后：KnowledgeIngestionService.java（核心片段）

@Slf4j
@Service
@RequiredArgsConstructor
public class KnowledgeIngestionService {

    private final KnowledgeArticleMapper articleMapper;
    private final KnowledgeArticleVersionMapper versionMapper;  // 新增
    private final KnowledgeChunkMapper chunkMapper;             // 新增
    private final EmbeddingBatchService embeddingBatchService;  // 新增（替代直接调 EmbeddingClient）
    private final KnowledgeChunkStore chunkStore;               // 新增（替代 KnowledgeVectorStore）
    private final MarkdownChunkSplitter chunkSplitter;          // 新增
    private final KnowledgeGovernanceService governanceService; // 新增

    public Map<String, Object> ingestAll() {
        // ... 扫描 + 解析（略，与原版相同）...

        for (Path file : mdFiles) {
            ParsedDoc doc = parseFile(file);
            String md5 = md5(doc.content);

            KnowledgeArticle existing = articleMapper.selectOne(
                    new LambdaQueryWrapper<KnowledgeArticle>()
                            .eq(KnowledgeArticle::getTitle, doc.title)
                            .last("LIMIT 1"));

            if (existing != null && md5.equals(existing.getContentMd5())) {
                skipped++;
                continue;
            }

            // Step 4: 分块（新增）
            List<Chunk> chunks = chunkSplitter.split(doc.content,
                    existing != null ? existing.getId() : null);

            // Step 5: 批量 Embedding（新增）
            List<String> chunkTexts = chunks.stream()
                    .map(c -> c.getHeading() + "\\n" + c.getContent())
                    .collect(Collectors.toList());
            List<float[]> embeddings = embeddingBatchService.embedBatch(chunkTexts).block();

            // Step 6: 入库
            if (existing != null) {
                // 版本历史：旧版本写入 knowledge_article_versions
                KnowledgeArticleVersion versionHistory = KnowledgeArticleVersion.builder()
                        .articleId(existing.getId())
                        .version(existing.getVersion())
                        .content(existing.getContent())
                        .contentMd5(existing.getContentMd5())
                        .build();
                versionMapper.insert(versionHistory);

                // 更新主表
                existing.setTopic(doc.topic);
                existing.setContent(doc.content);
                existing.setContentMd5(md5);
                existing.setVersion(existing.getVersion() + 1);
                articleMapper.updateById(existing);

                // 删除旧分块，写入新分块
                chunkStore.deleteByArticleId(existing.getId());
                chunkStore.batchUpsert(existing.getId(), chunks, embeddings);
                updated++;
            } else {
                KnowledgeArticle article = KnowledgeArticle.builder()
                        .title(doc.title)
                        .topic(doc.topic)
                        .content(doc.content)
                        .contentMd5(md5)
                        .status("PUBLISHED")
                        .version(1)
                        .tags(doc.tags)
                        .build();
                articleMapper.insert(article);

                // 设置分块的 articleId
                chunks.forEach(c -> c.setArticleId(article.getId()));
                chunkStore.batchUpsert(article.getId(), chunks, embeddings);
                inserted++;
            }

            // 矛盾检测（新增）
            if (existing != null) {
                governanceService.detectContradiction(existing.getContent(), doc.content, existing.getId());
            }
            // 重复检测（新增）
            governanceService.detectDuplicate(doc.title, doc.topic,
                    embeddings.get(0));  // 用第一个块向量做文档级相似度对比
        }
        // ... 返回统计 ...
    }
}
''')

add_heading_styled(doc, '4.4 Embedding 批量服务', level=2)

add_code_block(doc, '''package com.campushare.agent.service;

import com.campushare.agent.llm.EmbeddingClient;
import io.github.resilience4j.retry.Retry;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.stereotype.Service;
import reactor.core.publisher.Mono;

import java.util.ArrayList;
import java.util.List;

/**
 * 批量 Embedding 服务。
 * - batch=16（ADR-KB-05），超过自动分批
 * - 重试 3 次，指数退避
 */
@Slf4j
@Service
@RequiredArgsConstructor
public class EmbeddingBatchService {

    private final EmbeddingClient embeddingClient;
    private final Retry embeddingRetry;

    private static final int BATCH_SIZE = 16;

    public Mono<List<float[]>> embedBatch(List<String> texts) {
        return Mono.fromCallable(() -> {
            List<float[]> all = new ArrayList<>();
            for (int i = 0; i < texts.size(); i += BATCH_SIZE) {
                int end = Math.min(i + BATCH_SIZE, texts.size());
                List<String> batch = texts.subList(i, end);
                List<float[]> result = Retry.decorateSupplier(embeddingRetry,
                        () -> embeddingClient.embedBatch(batch).block()).get();
                if (result != null) {
                    all.addAll(result);
                }
            }
            return all;
        });
    }
}
''')

add_heading_styled(doc, '4.5 增量同步服务改造', level=2)

add_paragraph_styled(doc, 'PostVectorService 新增 updated_at 增量同步方法：')

add_code_block(doc, '''// 改造后：PostVectorService.java（新增方法）

/**
 * updated_at 增量同步（ADR-KB-03 机制二）。
 * 只拉取 updated_at > lastSyncTime 的帖子。
 */
public Mono<Map<String, Object>> syncIncremental(LocalDateTime lastSyncTime) {
    return Mono.fromCallable(() -> doSyncIncremental(lastSyncTime))
            .subscribeOn(Schedulers.boundedElastic())
            .onErrorResume(e -> {
                log.error("Incremental post sync failed", e);
                Map<String, Object> err = new HashMap<>();
                err.put("error", e.getMessage());
                return Mono.just(err);
            });
}

private Map<String, Object> doSyncIncremental(LocalDateTime lastSyncTime) {
    int page = 1;
    int size = 100;
    int total = 0, success = 0, failed = 0;

    while (true) {
        // 新增 Feign 接口：按 updated_at 增量拉取
        Result<IPage<PostVectorDTO>> result = postFeignClient.getUpdatedAfter(
                lastSyncTime, page, size);

        if (result == null || result.getData() == null
                || result.getData().getRecords().isEmpty()) {
            break;
        }

        List<PostVectorDTO> posts = result.getData().getRecords();
        total += posts.size();

        List<String> texts = posts.stream()
                .map(dto -> dto.getTitle() + "\\n" +
                        (dto.getContentExcerpt() != null ? dto.getContentExcerpt() : ""))
                .collect(Collectors.toList());

        List<float[]> embeddings = embeddingBatchService.embedBatch(texts).block();

        for (int i = 0; i < posts.size() && i < embeddings.size(); i++) {
            try {
                upsertPostVector(posts.get(i), embeddings.get(i));
                success++;
            } catch (Exception e) {
                log.warn("Failed to upsert post vector: {}", e.getMessage());
                failed++;
            }
        }

        if (posts.size() < size) break;
        page++;
    }

    Map<String, Object> stats = new HashMap<>();
    stats.put("total", total);
    stats.put("success", success);
    stats.put("failed", failed);
    stats.put("lastSyncTime", LocalDateTime.now().toString());
    log.info("Incremental post sync: total={}, success={}, failed={}", total, success, failed);
    return stats;
}
''')

add_paragraph_styled(doc, 'PostVectorScheduler 改造为三机制调度：')

add_code_block(doc, '''// 改造后：PostVectorScheduler.java

@Slf4j
@Component
@RequiredArgsConstructor
public class PostVectorScheduler {

    private final PostVectorService postVectorService;
    private final PostVectorSyncStateRepository syncStateRepository;  // 记录上次同步时间

    @EventListener(ApplicationReadyEvent.class)
    public void onStartup() {
        // 启动后 60 秒首次全量同步（与原版相同）
    }

    // 机制二：每 5 分钟增量同步（替代原版的全量同步）
    @Scheduled(initialDelay = 60000, fixedDelay = 300000)
    public void scheduledIncrementalSync() {
        LocalDateTime lastSync = syncStateRepository.getLastSyncTime();
        if (lastSync == null) {
            // 首次运行，走全量
            postVectorService.syncAll().block();
        } else {
            postVectorService.syncIncremental(lastSync).block();
        }
        syncStateRepository.updateLastSyncTime(LocalDateTime.now());
    }

    // 机制三：每天凌晨 3 点全量同步（兜底）
    @Scheduled(cron = "0 0 3 * * ?")
    public void scheduledFullSync() {
        log.info("Daily full post vector sync triggered");
        postVectorService.syncAll().block();
    }
}
''')

add_heading_styled(doc, '4.6 知识库治理服务', level=2)

add_code_block(doc, '''package com.campushare.agent.service;

import com.campushare.agent.store.KnowledgeChunkStore;
import com.campushare.agent.llm.DeepSeekClient;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.scheduling.annotation.Scheduled;
import org.springframework.stereotype.Service;

@Slf4j
@Service
@RequiredArgsConstructor
public class KnowledgeGovernanceService {

    private final KnowledgeChunkStore chunkStore;
    private final DeepSeekClient deepSeekClient;
    private final KnowledgeArticleMapper articleMapper;

    private static final float DUPLICATE_THRESHOLD = 0.95f;
    private static final float SIMILAR_THRESHOLD = 0.85f;

    /**
     * 重复检测：新文档入库时调用。
     */
    public void detectDuplicate(String title, String topic, float[] newEmbedding) {
        // 查同 topic 下的已有文档向量
        List<float[]> existing = chunkStore.getArticleEmbeddingsByTopic(topic);
        for (float[] vec : existing) {
            float similarity = cosineSimilarity(newEmbedding, vec);
            if (similarity > DUPLICATE_THRESHOLD) {
                log.warn("Duplicate detected: title={}, topic={}, similarity={}",
                        title, topic, similarity);
                // 标记为 DUPLICATE，进入人工审核队列
                break;
            } else if (similarity > SIMILAR_THRESHOLD) {
                log.info("Similar document found: title={}, similarity={}", title, similarity);
                // 标记为 SIMILAR，检索时降权
            }
        }
    }

    /**
     * 矛盾检测：平台知识库更新时调用。
     */
    public void detectContradiction(String oldContent, String newContent, Long articleId) {
        String prompt = """
                你是知识库审核员。对比以下两版文档，判断关键结论是否矛盾。
                只关注"规则/条件/流程"类结论，忽略措辞差异。

                旧版：
                %s

                新版：
                %s

                输出 JSON：{"has_contradiction": true/false, "contradictions": [...]}
                """.formatted(oldContent, newContent);

        String result = deepSeekClient.chatCompletion(prompt).block();
        // 解析 JSON，若 has_contradiction=true，标记新版本为 NEEDS_REVIEW
        log.info("Contradiction check done for article {}", articleId);
    }

    /**
     * 质量评分：每天凌晨 2 点定时计算。
     */
    @Scheduled(cron = "0 0 2 * * ?")
    public void calculateQualityScores() {
        log.info("Daily quality score calculation triggered");
        // 统计每篇文档过去 7 天的召回频次、用户反馈
        // 计算 quality_score，< 0.3 的标记为 LOW_QUALITY
    }

    /**
     * 过期清理：每天凌晨 3:30 清理。
     */
    @Scheduled(cron = "0 30 3 * * ?")
    public void cleanupDeprecated() {
        log.info("Daily deprecated cleanup triggered");
        // 删除 status=DEPRECATED 的文档向量
        int deleted = chunkStore.deleteDeprecated();
        log.info("Cleaned up {} deprecated chunks", deleted);
    }

    private float cosineSimilarity(float[] a, float[] b) {
        float dot = 0, normA = 0, normB = 0;
        for (int i = 0; i < a.length; i++) {
            dot += a[i] * b[i];
            normA += a[i] * a[i];
            normB += b[i] * b[i];
        }
        return (float) (dot / (Math.sqrt(normA) * Math.sqrt(normB)));
    }
}
''')

add_heading_styled(doc, '4.7 数据库 Schema', level=2)

add_paragraph_styled(doc, '新增/改造的表：')

add_heading_styled(doc, '4.7.1 knowledge_articles 表（改造）', level=3)

add_code_block(doc, '''-- 改造：新增 version_status 字段
ALTER TABLE knowledge_articles
  ADD COLUMN version_status VARCHAR(32) DEFAULT 'STABLE'
    COMMENT '版本状态（STABLE/NEEDS_REVIEW/DUPLICATE/SIMILAR/LOW_QUALITY）';''')

add_heading_styled(doc, '4.7.2 knowledge_article_versions 表（新增）', level=3)

add_code_block(doc, '''CREATE TABLE IF NOT EXISTS knowledge_article_versions (
  id BIGINT PRIMARY KEY AUTO_INCREMENT COMMENT '自增主键',
  article_id BIGINT NOT NULL COMMENT '文档ID',
  version INT NOT NULL COMMENT '版本号',
  content MEDIUMTEXT NOT NULL COMMENT '该版本正文',
  content_md5 CHAR(32) NOT NULL COMMENT '该版本 MD5',
  changed_by VARCHAR(64) COMMENT '变更人',
  change_note VARCHAR(256) COMMENT '变更说明',
  created_at DATETIME DEFAULT CURRENT_TIMESTAMP COMMENT '创建时间',
  INDEX idx_article_version (article_id, version),
  INDEX idx_created (created_at)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
  COMMENT='知识库文档版本历史表';''')

add_heading_styled(doc, '4.7.3 knowledge_chunks 表（新增，替代 knowledge_vectors）', level=3)

add_code_block(doc, '''-- 在 PostgreSQL（agent-postgres）中创建
CREATE TABLE IF NOT EXISTS knowledge_chunks (
  id BIGSERIAL PRIMARY KEY,
  article_id BIGINT NOT NULL COMMENT '文档ID',
  chunk_index INT NOT NULL COMMENT '块序号',
  heading VARCHAR(256) COMMENT '所属标题',
  content TEXT NOT NULL COMMENT '块文本',
  token_count INT COMMENT '块Token数',
  content_md5 CHAR(32) NOT NULL COMMENT '块MD5（检测变更）',
  embedding vector(1024) NOT NULL COMMENT '块向量',
  embedding_model VARCHAR(32) DEFAULT 'bge-m3' COMMENT 'Embedding模型版本',
  status VARCHAR(16) DEFAULT 'PUBLISHED' COMMENT '状态',
  recall_count INT DEFAULT 0 COMMENT '召回次数（质量评分用）',
  feedback_score INT DEFAULT 0 COMMENT '反馈分（点赞+1/点踩-1）',
  created_at TIMESTAMPTZ DEFAULT CURRENT_TIMESTAMP,
  updated_at TIMESTAMPTZ DEFAULT CURRENT_TIMESTAMP,
  UNIQUE KEY uk_article_chunk (article_id, chunk_index)
);

-- 向量索引（HNSW）
CREATE INDEX IF NOT EXISTS idx_knowledge_chunks_embedding
  ON knowledge_chunks USING hnsw (embedding vector_cosine_ops);

-- 按 article_id 删除（文档更新时清旧分块）
CREATE INDEX IF NOT EXISTS idx_knowledge_chunks_article
  ON knowledge_chunks (article_id);''')

add_heading_styled(doc, '4.7.4 post_vector_sync_state 表（新增）', level=3)

add_code_block(doc, '''-- 记录帖子库同步状态（用于 updated_at 增量）
CREATE TABLE IF NOT EXISTS post_vector_sync_state (
  id INT PRIMARY KEY DEFAULT 1,
  last_sync_at DATETIME NOT NULL COMMENT '上次增量同步时间',
  last_full_sync_at DATETIME COMMENT '上次全量同步时间',
  updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
  COMMENT='帖子向量同步状态表';

-- 初始化
INSERT INTO post_vector_sync_state (id, last_sync_at)
SELECT 1, '2026-01-01 00:00:00'
WHERE NOT EXISTS (SELECT 1 FROM post_vector_sync_state WHERE id = 1);''')

doc.add_page_break()

# ==================== 五、目标 ====================
add_heading_styled(doc, '五、目标：实现效果', level=1)

add_heading_styled(doc, '5.1 功能目标', level=2)

add_styled_table(doc,
    ['功能点', '当前', '目标', '验收标准'],
    [
        ['平台知识库分块', '整篇 1 个向量', '按 Markdown 结构分块，每篇 5-20 块', '长文档召回率 +40%'],
        ['帖子库增量同步', '5 分钟全量', 'updated_at 增量 + 事件通知 + 日级全量兜底', '同步耗时降 90%'],
        ['版本管理', 'version 自增', 'SemVer + 版本历史 + 回滚', '可回滚到任意历史版本'],
        ['质量评分', '无', '四维质量评分（召回/反馈/新鲜度/完整度）', '< 0.3 的进入待清理队列'],
        ['重复检测', '无', '相似度 > 0.95 标记重复', '重复召回率降 80%'],
        ['矛盾检测', '无', 'LLM 对比新旧版本关键结论', '矛盾文档进入人工审核'],
        ['过期清理', '无', '每日清理 DEPRECATED 向量', '向量库无过期数据'],
        ['Embedding 模型版本', '写死 bge-m3', '字段记录 + 批量重嵌入脚本', '模型升级零风险'],
    ],
    col_widths=[3, 3, 5, 4])

add_heading_styled(doc, '5.2 性能目标', level=2)

add_styled_table(doc,
    ['指标', '当前', '目标', '测试方法'],
    [
        ['平台知识库摄入', '~50 篇 30 秒', '~50 篇 < 60 秒（含分块）', 'KnowledgeIngestionService.ingestAll() 计时'],
        ['帖子库增量同步', '10 万帖 5 分钟', '1 万帖变更 < 30 秒', 'syncIncremental 计时'],
        ['Embedding 批量吞吐', '单条调用', '16 条/批，100 条/秒', 'embedBatch 压测'],
        ['分块延迟', 'N/A', '< 50ms/篇（JTokkit 本地估算）', 'ChunkSplitter.split 计时'],
        ['向量检索延迟', '~20ms（Top-5）', '~20ms（Top-5，分块后向量数增但 HNSW 保持）', 'KnowledgeChunkStore.search 计时'],
    ],
    col_widths=[3.5, 3.5, 4.5, 4])

add_heading_styled(doc, '5.3 质量目标', level=2)

add_styled_table(doc,
    ['指标', '当前', '目标', '测试方法'],
    [
        ['长文档召回率', '~50%（前 500 字）', '>= 90%（分块后全覆盖）', '黄金测试集：长文档问题 100 条'],
        ['重复召回率', '未知', '< 5%', '统计 Top-5 中同源文档占比'],
        ['矛盾召回率', '未知', '0%（矛盾文档进入审核）', '人工审核 + LLM 检测'],
        ['过期数据占比', '未知', '0%', '每日清理后统计'],
        ['Embedding 一致性', '混在一起', '100% 按模型版本隔离', '模型升级后全量重嵌入'],
    ],
    col_widths=[3.5, 3.5, 4, 4.5])

add_heading_styled(doc, '5.4 成本目标', level=2)

add_styled_table(doc,
    ['成本项', '当前', '目标', '节省'],
    [
        ['Embedding API 调用（帖子库）', '10 万次/5 分钟', '1 万次/5 分钟（增量）', '90%'],
        ['Embedding API 调用（知识库）', '50 次/小时', '50 次/小时（MD5 diff 跳过）', '维持'],
        ['存储成本（向量库）', '1 向量/文档', '5-20 向量/文档（分块）', '增加 5-20x（可接受）'],
        ['存储成本（版本历史）', '0', '1 份/版本', '增加（可接受）'],
        ['LLM 调用（矛盾检测）', '0', '1 次/文档更新', '可控（更新低频）'],
    ],
    col_widths=[4, 3.5, 4, 3.5])

doc.add_page_break()

# ==================== 六、测试评估与验收 ====================
add_heading_styled(doc, '六、测试评估与验收（详尽版）', level=1)

add_heading_styled(doc, '6.1 评估指标体系（含公式）', level=2)

add_paragraph_styled(doc, '知识库管理的评估指标分四个维度：')

add_heading_styled(doc, '6.1.1 分块质量指标', level=3)

add_styled_table(doc,
    ['指标', '公式', '目标', '说明'],
    [
        ['分块覆盖率',
         '覆盖正文的 Token 数 / 文档总 Token 数',
         '>= 95%',
         '检测分块后是否有信息丢失（应接近 100%，重叠部分会增加 Token）'],
        ['分块均匀度',
         '1 - (块 Token 数标准差 / 平均块 Token 数)',
         '>= 0.6',
         '块大小越均匀，检索质量越稳定'],
        ['块边界合理性',
         '在句子/标题边界切分的块数 / 总块数',
         '>= 90%',
         '不应在句子中间切断'],
        ['平均块数/文档',
         '总块数 / 文档数',
         '5-20',
         '太少说明分块不够，太多说明块太小'],
    ],
    col_widths=[3.5, 5, 2.5, 4.5])

add_heading_styled(doc, '6.1.2 同步效率指标', level=3)

add_styled_table(doc,
    ['指标', '公式', '目标', '说明'],
    [
        ['增量同步命中率',
         '增量同步的帖子数 / 总帖子数',
         '< 5%（平稳期）',
         '大部分时间没有变更，增量同步应该处理很少的帖子'],
        ['同步延迟',
         '帖子更新到向量库更新的时间差 P95',
         '< 10 秒（事件通知）/ < 5 分钟（增量）',
         '事件通知近实时，增量补漏'],
        ['同步失败率',
         '失败的帖子数 / 同步的帖子数',
         '< 1%',
         '失败的要重试 + 告警'],
        ['全量同步耗时',
         '全量同步的总耗时',
         '10 万帖 < 10 分钟',
         '每天凌晨兜底同步'],
    ],
    col_widths=[3.5, 5, 3.5, 4])

add_heading_styled(doc, '6.1.3 知识质量指标', level=3)

add_styled_table(doc,
    ['指标', '公式', '目标', '说明'],
    [
        ['重复召回率',
         'Top-5 中同源文档对数 / 检索次数',
         '< 5%',
         '同源 = 相似度 > 0.95 的文档'],
        ['矛盾召回率',
         '召回的矛盾文档对数 / 检索次数',
         '0%',
         '矛盾文档应该被审核拦截'],
        ['过期数据占比',
         'status=DEPRECATED 的向量数 / 总向量数',
         '0%',
         '每日清理后应该为 0'],
        ['低质量文档占比',
         'quality_score < 0.3 的文档数 / 总文档数',
         '< 10%',
         '低质量的进入待清理队列'],
        ['召回集中度',
         'Top-20% 文档的召回次数 / 总召回次数',
         '60-80%',
         '健康的分布应该集中（长尾理论），但不应过度集中'],
    ],
    col_widths=[3.5, 5, 2.5, 4.5])

add_heading_styled(doc, '6.1.4 版本管理指标', level=3)

add_styled_table(doc,
    ['指标', '公式', '目标', '说明'],
    [
        ['版本回滚成功率',
         '成功回滚次数 / 回滚尝试次数',
         '100%',
         '回滚必须可靠'],
        ['版本历史完整率',
         '有版本历史的文档数 / 总文档数',
         '100%',
         '每个文档每个版本都应有历史'],
        ['Embedding 模型一致性',
         '同一模型版本的向量数 / 总向量数',
         '100%',
         '模型升级后必须全量重嵌入'],
    ],
    col_widths=[3.5, 5, 2.5, 4.5])

add_heading_styled(doc, '6.2 黄金测试集构建', level=2)

add_heading_styled(doc, '6.2.1 测试集结构', level=3)

add_styled_table(doc,
    ['测试集', '数量', '用途', '构建方式'],
    [
        ['分块质量测试集', '50 篇文档', '评估分块策略', '人工挑选不同长度/结构的 .md 文档'],
        ['长文档召回测试集', '100 条 query', '评估分块后召回率', '每条 query 答案在文档第 500+ 字'],
        ['增量同步测试集', '1000 帖', '评估增量同步', '模拟帖子更新/删除，验证向量库一致'],
        ['重复检测测试集', '50 对文档', '评估重复检测', '25 对重复 + 25 对相似不重复'],
        ['矛盾检测测试集', '30 对文档', '评估矛盾检测', '15 对矛盾 + 15 对不矛盾'],
        ['版本回滚测试集', '20 篇文档', '评估版本管理', '每篇有 3+ 版本，验证回滚'],
    ],
    col_widths=[3.5, 2.5, 4, 5])

add_heading_styled(doc, '6.2.2 测试集维护', level=3)
add_bullet(doc, '测试集纳入 Git 版本管理（docs/agent-design/test-suites/kb/）')
add_bullet(doc, '每月评审一次：剔除过时的、补充新场景')
add_bullet(doc, '每次知识库结构变更后，回归全部测试集')

add_heading_styled(doc, '6.3 评估流水线与 CI/CD 集成', level=2)

add_paragraph_styled(doc, '知识库管理的 CI/CD 流水线分四个阶段：')

add_code_block(doc, """知识库管理 CI/CD 流水线：

  ┌─ Stage 1: 编译 ────────────────────────────────────┐
  │  mvn clean compile -DskipTests                      │
  │  → BUILD SUCCESS 才进入下一阶段                      │
  └─────────────────────────────────────────────────────┘
                  │
                  ▼
  ┌─ Stage 2: 单元测试 ────────────────────────────────┐
  │  - MarkdownChunkSplitter 单测（分块边界/Token/重叠）│
  │  - EmbeddingBatchService 单测（分批/重试）          │
  │  - KnowledgeGovernanceService 单测（相似度计算）    │
  │  → 通过率 100% 才进入下一阶段                       │
  └─────────────────────────────────────────────────────┘
                  │
                  ▼
  ┌─ Stage 3: 集成测试 ────────────────────────────────┐
  │  - 摄入流水线集成测试（50 篇文档全流程）             │
  │  - 增量同步集成测试（1000 帖变更）                   │
  │  - 版本回滚集成测试                                  │
  │  → 通过率 100% 才进入下一阶段                       │
  └─────────────────────────────────────────────────────┘
                  │
                  ▼
  ┌─ Stage 4: 评估测试 ────────────────────────────────┐
  │  - 分块质量评估（50 篇文档）                        │
  │  - 长文档召回评估（100 条 query）                   │
  │  - 重复/矛盾检测评估                                │
  │  - 与上次版本对比，退化 > 5% 阻断发布               │
  └─────────────────────────────────────────────────────┘""")

add_heading_styled(doc, '6.4 LLM-as-Judge 评估', level=2)

add_paragraph_styled(doc, '用 LLM 评估分块质量和矛盾检测质量：')

add_heading_styled(doc, '6.4.1 分块质量 LLM 评估', level=3)

add_code_block(doc, '''分块质量评估 Prompt：

  系统：你是知识库质量审核员。评估以下分块是否合理。
  评分维度（1-5 分）：
  1. 语义完整性：块内容是否是完整的语义单元
  2. 边界合理性：是否在合适的位置切分（标题/段落/句子）
  3. 信息密度：块内容是否信息充分（非碎片化）
  4. 上下文独立：不依赖其他块能否理解

  文档标题：{title}
  分块内容：{chunk_content}
  所属标题：{heading}

  输出 JSON：
  {
    "semantic_completeness": 4,
    "boundary_rationality": 5,
    "information_density": 4,
    "context_independence": 3,
    "average_score": 4.0,
    "issues": ["块末尾句子被切断"]
  }''')

add_heading_styled(doc, '6.4.2 矛盾检测 LLM 评估', level=3)

add_paragraph_styled(doc,
    '对矛盾检测的结果做二次评估（评估检测是否准确）：')

add_styled_table(doc,
    ['评估项', '公式', '目标'],
    [
        ['矛盾检测准确率', '正确识别的矛盾对数 / 标记为矛盾的对数', '>= 90%'],
        ['矛盾检测召回率', '正确识别的矛盾对数 / 实际矛盾对数', '>= 85%'],
        ['误报率', '误报为矛盾的对数 / 标记为矛盾的对数', '< 10%'],
    ],
    col_widths=[3.5, 6, 4])

add_heading_styled(doc, '6.5 错误分析与归因', level=2)

add_paragraph_styled(doc, '知识库管理的常见错误及归因：')

add_styled_table(doc,
    ['错误类型', '表现', '可能原因', '修复方向'],
    [
        ['分块切断句子',
         '块末尾是半个句子',
         'splitToTokenLimit 在字符级切断',
         '在句子边界切分，不足时合并到下一块'],
        ['分块过大',
         '块 Token 数 > 512',
         'H2/H3 段落本身超过 512',
         '按段落/句子二次切分'],
        ['分块过小',
         '块 Token 数 < 64',
         'H3 标题下内容太少',
         '合并到上一块'],
        ['增量同步遗漏',
         '帖子更新但向量库没更新',
         '事件通知丢失 + updated_at 精度不足',
         '定时全量兜底 + 提高时间戳精度'],
        ['重复检测误报',
         '不重复的文档被标记为重复',
         '阈值 0.95 太低',
         '调整阈值或加入语义判断'],
        ['矛盾检测漏报',
         '矛盾的文档未被发现',
         'LLM 判断标准过松',
         '优化 Prompt + 人工复核'],
        ['版本回滚失败',
         '回滚后内容不一致',
         '版本历史不完整 / 分块未同步删除',
         '事务保证：版本历史 + 分块 + 主表原子操作'],
        ['Embedding 模型混乱',
         '新旧模型向量混在一起',
         '模型升级未全量重嵌入',
         '批量重嵌入脚本 + 模型版本字段隔离'],
    ],
    col_widths=[3, 3.5, 4, 4.5])

doc.add_page_break()

# ==================== 6.6 测试用例设计 ====================
add_heading_styled(doc, '6.6 测试用例设计', level=2)

add_paragraph_styled(doc, '按四个子方向设计测试用例，覆盖正常路径、边界、异常和回归场景。')

add_heading_styled(doc, '6.6.1 分块测试用例', level=3)

add_styled_table(doc,
    ['用例 ID', '场景', '输入', '预期输出', '类型'],
    [
        ['CHUNK-001', '标准 H2 切分',
         '一篇含 3 个 H2 章节、每章约 400 Token 的文档',
         '切分为 3 块，每块约 400 Token，块 Token 数在 [64, 512]',
         '正常'],
        ['CHUNK-002', 'H2 下嵌套 H3',
         'H2 章节下有 2 个 H3 子章节',
         'H2 块过大时拆为 2 个 H3 块，前缀都带 H2 标题',
         '正常'],
        ['CHUNK-003', 'H2 超长切分',
         '单个 H2 段落 > 512 Token',
         '按段落/句子二次切分，每块 <= 512 Token，重叠 50',
         '边界'],
        ['CHUNK-004', '极短文档',
         '文档仅 30 Token',
         '不切分，单块输出，块 Token 数 = 30（允许 < 64）',
         '边界'],
        ['CHUNK-005', '无标题文档',
         '只有段落无 H2/H3',
         '按段落切分，前缀为空 heading',
         '正常'],
        ['CHUNK-006', '代码块跨块',
         '文档含超过 512 Token 的代码块',
         '代码块单独成块，不被切断（标记 type=CODE）',
         '边界'],
        ['CHUNK-007', '重叠验证',
         '相邻两块',
         '后一块开头 50 Token 与前一块结尾重复',
         '正常'],
        ['CHUNK-008', 'Token 估算准确性',
         '用 JTokkit 估算的 Token 数',
         '与 DeepSeek API 实际 Token 数误差 < 5%',
         '回归'],
    ],
    col_widths=[2, 2.5, 3.5, 4, 1.5])

add_heading_styled(doc, '6.6.2 增量同步测试用例', level=3)

add_styled_table(doc,
    ['用例 ID', '场景', '操作', '预期结果', '类型'],
    [
        ['SYNC-001', '事件通知单帖更新',
         'post-service 发送 POST_UPDATED 事件',
         '30 秒内该帖向量被更新',
         '正常'],
        ['SYNC-002', '事件通知丢失',
         '事件队列积压导致事件丢失',
         '5 分钟内 updated_at 增量同步补齐',
         '异常'],
        ['SYNC-003', 'updated_at 增量遗漏',
         'updated_at 精度仅到秒，秒内多次更新',
         '每小时全量同步兜底',
         '异常'],
        ['SYNC-004', '批量帖子更新',
         '1000 帖同时更新',
         '批量 embedding + 批量 upsert，10 分钟内完成',
         '正常'],
        ['SYNC-005', '帖子删除同步',
         'post-service 发送 POST_DELETED 事件',
         '向量库中该帖记录被删除',
         '正常'],
        ['SYNC-006', 'embedding 服务降级',
         'embedding 服务返回 5xx',
         'CircuitBreaker 熔断，跳过本次，下次重试',
         '异常'],
        ['SYNC-007', 'PostgreSQL 连接失败',
         'PG 容器重启中',
         '写入重试 3 次后归档到 agent_pending_writes 表',
         '异常'],
        ['SYNC-008', '首次启动全量同步',
         'agent-service 首次启动',
         '60 秒后开始全量同步，所有 PUBLISHED 帖子被嵌入',
         '正常'],
    ],
    col_widths=[2, 2.5, 3.5, 4, 1.5])

add_heading_styled(doc, '6.6.3 重复与矛盾检测测试用例', level=3)

add_styled_table(doc,
    ['用例 ID', '场景', '输入', '预期输出', '类型'],
    [
        ['DUP-001', '完全重复文档',
         '两篇内容完全相同的文档',
         '标记 DUPLICATE，相似度 = 1.0，保留先发布的',
         '正常'],
        ['DUP-002', '高度相似文档',
         '两篇余弦相似度 0.92',
         '标记 SIMILAR，人工复核',
         '正常'],
        ['DUP-003', '低相似度文档',
         '两篇余弦相似度 0.5',
         '不标记，正常存储',
         '正常'],
        ['DUP-004', '阈值边界',
         '相似度恰好 0.95',
         '标记 DUPLICATE',
         '边界'],
        ['DUP-005', '阈值边界',
         '相似度恰好 0.85',
         '标记 SIMILAR',
         '边界'],
        ['CONTR-001', '明显矛盾',
         '"注册需要邮箱" vs "注册不需要邮箱"',
         'LLM 判定矛盾，标记 CONFLICT',
         '正常'],
        ['CONTR-002', '补充关系非矛盾',
         '"注册需要邮箱" vs "推荐使用校园邮箱"',
         'LLM 判定非矛盾，标记 COMPLEMENTARY',
         '正常'],
        ['CONTR-006', 'LLM 服务降级',
         '矛盾检测时 LLM 不可用',
         '降级为人工复核队列，记录到 agent_tool_errors',
         '异常'],
    ],
    col_widths=[2, 2.5, 3.5, 4, 1.5])

add_heading_styled(doc, '6.6.4 版本管理测试用例', level=3)

add_styled_table(doc,
    ['用例 ID', '场景', '操作', '预期结果', '类型'],
    [
        ['VER-001', '版本递增',
         '更新文档内容',
         'version+1，旧版本归档到 versions 表',
         '正常'],
        ['VER-002', '版本回滚',
         '从 v3 回滚到 v1',
         'content 恢复为 v1 内容，version 变为 v4',
         '正常'],
        ['VER-003', '回滚后分块同步',
         '回滚后重新分块',
         '旧分块被删除，按 v1 内容重新分块',
         '正常'],
        ['VER-004', '版本历史查询',
         '查询某文档所有版本',
         '返回完整版本列表，按版本号降序',
         '正常'],
        ['VER-005', '事务原子性',
         '回滚过程中 PG 宕机',
         '事务回滚，version 不变（不出现半回滚状态）',
         '异常'],
        ['VER-006', '版本号不连续',
         '中间版本被物理删除',
         '回滚到已删除版本报错并提示',
         '异常'],
    ],
    col_widths=[2, 2.5, 3.5, 4, 1.5])

# ==================== 6.7 性能与压力测试 ====================
add_heading_styled(doc, '6.7 性能与压力测试', level=2)

add_heading_styled(doc, '6.7.1 摄入吞吐量测试', level=3)

add_paragraph_styled(doc, '目标：单机 agent-service 实例的摄入吞吐能力。')

add_styled_table(doc,
    ['指标', '目标值', '测量方法', '当前基线'],
    [
        ['单文档摄入耗时', '< 500ms', '50 篇 5000 Token 文档平均', '约 1.2s（无分块）'],
        ['批量摄入吞吐', '>= 20 篇/分钟', '100 篇文档总耗时', '约 30 篇/分钟'],
        ['分块耗时', '< 50ms/篇', 'JTokkit 估算 + 切分', 'N/A（待实现）'],
        ['Embedding 耗时', '< 300ms/块', 'DeepSeek API 调用', '约 250ms'],
        ['PG 写入耗时', '< 20ms/块', 'HNSW 索引下的 upsert', '约 15ms'],
    ],
    col_widths=[3.5, 3, 4.5, 3])

add_heading_styled(doc, '6.7.2 压力测试场景', level=3)

add_code_block(doc, '''压力测试脚本（JMeter / wrk 模拟）：

  场景 1：首次启动批量摄入
    - 输入：1000 篇文档（每篇 5000 Token）
    - 监控：CPU / 内存 / PG 连接池 / embedding API 限流
    - 期望：30 分钟内完成，无 OOM，无连接池耗尽

  场景 2：高并发增量同步
    - 输入：100 帖/秒的 POST_UPDATED 事件
    - 监控：事件队列积压 / embedding 降级 / PG 写入延迟
    - 期望：队列积压 < 1000，10 分钟内消化完毕

  场景 3：向量检索 P99 延迟
    - 输入：50 QPS 的检索请求，向量库 10 万条
    - 监控：HNSW 检索 P50/P95/P99
    - 期望：P99 < 50ms

  场景 4：Embedding 服务降级压测
    - 输入：embedding API 持续 5xx 5 分钟
    - 监控：CircuitBreaker 状态 / 降级队列长度 / 恢复时间
    - 期望：5 分钟内不丢失事件，恢复后 10 分钟内补齐''')

add_heading_styled(doc, '6.7.3 资源占用基线', level=3)

add_styled_table(doc,
    ['资源', '空闲', '正常负载', '峰值负载', '告警阈值'],
    [
        ['CPU', '< 10%', '30-50%', '70-85%', '> 90% 持续 5 分钟'],
        ['JVM 堆', '< 512MB', '1-2GB', '2.5-3GB', '> 3.5GB'],
        ['PG 连接数', '< 5', '10-20', '30-40', '> 45（池上限 50）'],
        ['PG 磁盘', '10 万条 ≈ 2GB', '-', '-', '> 80% 使用率'],
        ['向量索引大小', '10 万条 ≈ 1.5GB', '-', '-', '-'],
    ],
    col_widths=[2.5, 2.5, 2.5, 2.5, 4])

# ==================== 6.8 A/B 测试设计 ====================
add_heading_styled(doc, '6.8 A/B 测试设计', level=2)

add_paragraph_styled(doc, '通过 A/B 测试验证分块策略和治理参数对 RAG 检索质量的影响。')

add_heading_styled(doc, '6.8.1 实验一：分块策略对比', level=3)

add_styled_table(doc,
    ['分组', '分块策略', '目标 Token', '重叠', '流量比例'],
    [
        ['A 组（基线）', '固定大小分块', '256', '0', '50%'],
        ['B 组（实验）', 'Markdown 结构分块', '256', '50', '50%'],
    ],
    col_widths=[3, 3.5, 2.5, 2, 2])

add_paragraph_styled(doc, '观测指标（持续 7 天）：')
add_bullet(doc, '召回 Recall@5：B 组应 >= A 组 +5%')
add_bullet(doc, '答案准确率（LLM-as-Judge）：B 组应 >= A 组 +3%')
add_bullet(doc, '平均引用块数：B 组应 <= A 组（更精准）')
add_bullet(doc, '用户反馈（点赞/点踩比）：B 组应 >= A 组')

add_heading_styled(doc, '6.8.2 实验二：重复检测阈值', level=3)

add_styled_table(doc,
    ['分组', 'DUPLICATE 阈值', 'SIMILAR 阈值', '流量比例'],
    [
        ['A 组', '0.95', '0.85', '50%'],
        ['B 组', '0.92', '0.80', '50%'],
    ],
    col_widths=[2, 3, 3, 2])

add_paragraph_styled(doc, '观测指标：')
add_bullet(doc, '重复检测误报率：B 组可能上升，需 < 15%')
add_bullet(doc, '人工复核工作量：B 组会增加，需评估 ROI')
add_bullet(doc, '检索质量：B 组去重更彻底，Recall@5 不应下降')

add_heading_styled(doc, '6.8.3 实验三：Embedding 模型升级', level=3)

add_paragraph_styled(doc, '当 embedding 模型从 bge-m3 升级到更新版本时：')
add_bullet(doc, 'A 组：继续使用 bge-m3')
add_bullet(doc, 'B 组：新模型，需先全量重嵌入 10 万条')
add_bullet(doc, '观测：检索 Recall@5 / 延迟 / 成本（API 调用单价）')

add_callout(doc,
    'A/B 测试准入条件：实验组样本量 >= 1000 次对话，至少持续 7 天，'
    '统计显著性 p < 0.05。',
    color='E8F4FD', border_color='2196F3')

# ==================== 6.9 验收流程与准入准出 ====================
add_heading_styled(doc, '6.9 验收流程与准入准出', level=2)

add_heading_styled(doc, '6.9.1 准入条件（开发完成 → 测试阶段）', level=3)

add_styled_table(doc,
    ['检查项', '标准', '负责人'],
    [
        ['单元测试', '覆盖率 >= 80%，全部通过', '开发'],
        ['代码评审', '至少 1 人评审通过', '架构师'],
        ['静态检查', 'SonarQube 0 Blocker / 0 Critical', 'CI'],
        ['Schema 变更', 'DDL 已在测试库执行通过', 'DBA'],
        ['ADR 文档', '新增 ADR 已归档', '架构师'],
        ['Changelog', '已追加到 changelog/', '开发'],
    ],
    col_widths=[3, 6, 2])

add_heading_styled(doc, '6.9.2 准出条件（测试阶段 → 发布）', level=3)

add_styled_table(doc,
    ['检查项', '标准', '负责人'],
    [
        ['功能测试', '所有用例（6.6 节）通过', 'QA'],
        ['性能测试', '满足 6.7 节目标值', 'QA'],
        ['回归测试', 'RAG 检索相关回归用例 100% 通过', 'QA'],
        ['A/B 测试', '实验组指标不劣于基线（或显著优于）', '数据'],
        ['LLM-as-Judge', '分块质量均分 >= 4.0', 'QA'],
        ['监控告警', 'Prometheus 指标已接入', '运维'],
        ['回滚预案', '已有回滚脚本并演练通过', '运维'],
    ],
    col_widths=[3, 6, 2])

add_heading_styled(doc, '6.9.3 发布流程', level=3)

add_code_block(doc, '''发布流程（七步）：

  1. 合并代码到 master 分支
  2. CI 触发镜像构建（campushare-agent: vX.Y.Z）
  3. 在预发环境部署，跑 golden 测试集（32 条）
  4. 通过后推送到生产镜像仓库
  5. 滚动更新 agent-service 容器（先 1 个实例验证 10 分钟）
  6. 全量滚动更新，监控指标 30 分钟
  7. 发布完成，记录到 changelog 和 release notes''')

# ==================== 6.10 持续监控与漂移检测 ====================
add_heading_styled(doc, '6.10 持续监控与漂移检测', level=2)

add_paragraph_styled(doc,
    '知识库上线后需要持续监控质量和性能，发现"漂移"及时告警。')

add_heading_styled(doc, '6.10.1 监控指标体系', level=3)

add_styled_table(doc,
    ['维度', '指标', '采集方式', '告警阈值'],
    [
        ['摄入健康', '每小时摄入文档数',
         'Micrometer Counter', '连续 2 小时为 0'],
        ['摄入健康', '摄入失败率',
         'failed / total', '> 10%'],
        ['同步健康', '增量同步延迟',
         'updated_at - 同步时间', '> 10 分钟'],
        ['同步健康', 'CircuitBreaker 开启次数',
         'resilience4j 状态', '> 5 次/小时'],
        ['分块质量', '平均块 Token 数',
         '计算分块后统计', '< 64 或 > 512 持续'],
        ['分块质量', '分块过小比例',
         'chunk_tokens < 64 占比', '> 20%'],
        ['检索质量', '检索 Recall@5',
         '每日黄金集评估', '下降 > 5%'],
        ['检索质量', '检索结果为空率',
         '空结果查询占比', '> 15%'],
        ['存储健康', 'PG 磁盘使用率',
         'Prometheus node_exporter', '> 80%'],
        ['存储健康', '向量索引膨胀率',
         '记录数 vs 索引大小', '异常增长'],
    ],
    col_widths=[2.5, 3, 3.5, 3])

add_heading_styled(doc, '6.10.2 数据漂移检测', level=3)

add_paragraph_styled(doc, '数据漂移：知识库内容分布随时间变化，导致 embedding 模型在新数据上表现下降。')

add_styled_table(doc,
    ['漂移类型', '表现', '检测方法', '应对'],
    [
        ['内容漂移', '新文档主题与历史文档差异大',
         '新文档 embedding 与历史均值的距离',
         '触发模型重新评估'],
        ['查询漂移', '用户查询分布变化',
         '查询 embedding 聚类中心偏移',
         '补充新主题文档 + 重新分块'],
        ['质量漂移', 'LLM-as-Judge 分数下降',
         '每日抽样评估',
         '排查新摄入文档质量'],
        ['模型漂移', 'embedding 模型在新数据上召回下降',
         '定期跑 golden set',
         '评估新模型 + 全量重嵌入'],
    ],
    col_widths=[2.5, 3, 3.5, 3.5])

add_heading_styled(doc, '6.10.3 告警与自动化处理', level=3)

add_code_block(doc, '''告警分级与自动化处理：

  P0（立即响应，电话告警）：
    - PG 不可用超过 1 分钟
    - 摄入失败率 > 50% 持续 5 分钟
    - 自动处理：触发全量同步重试 + 切换备用 embedding

  P1（10 分钟响应，钉钉告警）：
    - 增量同步延迟 > 30 分钟
    - CircuitBreaker 持续开启
    - 自动处理：扩容 agent-service 实例

  P2（1 小时响应，邮件告警）：
    - 分块过小比例 > 30%
    - 检索 Recall@5 下降 > 3%
    - 自动处理：标记问题文档，加入人工复核队列

  P3（每日汇总）：
    - 摄入吞吐量趋势
    - 存储增长趋势
    - 自动处理：无，仅观察''')

doc.add_page_break()

# ==================== 第七章 总结与边界声明 ====================
add_heading_styled(doc, '第七章 总结与边界声明', level=1)

add_heading_styled(doc, '7.1 核心总结', level=2)

add_paragraph_styled(doc, '本文档围绕"RAG 的数据侧"展开，覆盖摄入、分块、同步、治理四大子方向：')

add_styled_table(doc,
    ['子方向', '核心方案', '关键 ADR', '预期收益'],
    [
        ['摄入流水线',
         '六步流水线：扫描 → 解析 → MD5 → 分块 → 嵌入 → 入库',
         'ADR-KB-01',
         '从单 embedding 升级为分块 embedding，召回粒度细化'],
        ['分块策略',
         'Markdown 结构分块（H2→H3→段落→句子）',
         'ADR-KB-02',
         '召回 Recall@5 提升 10-15%'],
        ['增量同步',
         '三机制：事件通知 + updated_at 增量 + 全量兜底',
         'ADR-KB-03',
         '同步延迟从 5 分钟降至 30 秒'],
        ['质量治理',
         '四维评分 + 重复检测 + 矛盾检测',
         'ADR-KB-04, ADR-KB-05',
         '重复内容减少 80%，矛盾内容 0 上线'],
        ['版本管理',
         'SemVer + 全量历史表 + 一键回滚',
         'ADR-KB-06',
         '回滚耗时 < 30 秒，版本可追溯'],
        ['Embedding 模型管理',
         '模型版本字段隔离 + 全量重嵌入脚本',
         'ADR-KB-07',
         '模型升级零停机，新旧向量不混淆'],
    ],
    col_widths=[2.5, 4.5, 2.5, 4])

add_heading_styled(doc, '7.2 与其他文档的关系', level=2)

add_paragraph_styled(doc, '本文档在 Agent 搭建系列中的位置：')

add_styled_table(doc,
    ['相关文档', '关系', '说明'],
    [
        ['RAG 检索增强生成模块设计方案',
         '上下游',
         '本文档是 RAG 的数据侧，RAG 文档是查询侧。本文档产出的分块和向量被 RAG 检索消费。'],
        ['SystemPrompt 工程模块设计方案',
         '协作',
         'SystemPrompt 的 L2 SEARCH prompt 引用本文档产出的话题标签（topic 字段）。'],
        ['意图识别模块设计方案',
         '协作',
         '意图识别判定为 SEARCH 时，触发本文档产出的知识库检索。'],
        ['上下文工程模块设计方案',
         '下游',
         '本文档产出的检索结果作为上下文工程的"知识层"注入到 LLM 上下文。'],
        ['工具调用模块设计方案',
         '无直接关系',
         '知识库管理不通过工具调用，由 RAG 内部消费。'],
        ['可观测性模块（待写）',
         '横切',
         '本文档第六章定义的指标需接入可观测性体系。'],
        ['评估体系模块（待写）',
         '横切',
         '本文档的 golden set 和 LLM-as-Judge 纳入全局评估体系。'],
    ],
    col_widths=[3.5, 2, 8])

add_heading_styled(doc, '7.3 演进路线', level=2)

add_paragraph_styled(doc, '知识库管理模块的演进路线图（按优先级）：')

add_styled_table(doc,
    ['阶段', '里程碑', '关键工作', '预计周期'],
    [
        ['Phase 1（当前）',
         '基础治理',
         'Markdown 分块 + 三机制同步 + 重复检测 + 版本管理',
         '2 周'],
        ['Phase 2',
         '质量提升',
         '矛盾检测 LLM 化 + 四维评分 + 漂移检测',
         '1 周'],
        ['Phase 3',
         '多模态知识',
         '图片/表格分块 + 多模态 embedding（CLIP）',
         '3 周'],
        ['Phase 4',
         '知识图谱融合',
         '实体抽取 + 关系建模 + 图谱检索增强 RAG',
         '4 周'],
        ['Phase 5',
         '主动学习',
         '基于用户反馈自动补充文档 + 自动失效检测',
         '2 周'],
    ],
    col_widths=[2.5, 2.5, 7, 2])

add_callout(doc,
    'Phase 1-2 是本文档覆盖范围，Phase 3-5 属于后续演进，'
    '会在对应方向文档（如"知识图谱模块"）中详述。',
    color='FFF3CD', border_color='FFC107')

add_heading_styled(doc, '7.4 边界声明', level=2)

add_paragraph_styled(doc, '本文档不覆盖以下内容（避免主题扩散）：')

add_bullet(doc, '向量检索算法本身的优化（HNSW 参数调优）——属于 RAG 检索侧')
add_bullet(doc, 'Embedding 模型的训练和微调——使用第三方模型（bge-m3）')
add_bullet(doc, '知识图谱构建——属于第 6 个方向"知识图谱"')
add_bullet(doc, '多模态知识库（图片/视频分块）——属于 Phase 3 演进')
add_bullet(doc, '用户生成内容（UGC）的审核——属于第 15 个方向"内容审核"')
add_bullet(doc, '备份与灾备——属于工程基础设施层')

doc.add_page_break()

# ==================== 附录 ADR 摘要 ====================
add_heading_styled(doc, '附录：ADR 摘要', level=1)

add_paragraph_styled(doc, '本文档涉及的 7 条架构决策记录（ADR）：')

add_styled_table(doc,
    ['ADR 编号', '决策标题', '决策摘要', '权衡'],
    [
        ['ADR-KB-01',
         '采用分块 embedding 替代单文档 embedding',
         '每篇文档按 Markdown 结构分块，每块独立 embedding，提升召回粒度',
         '存储成本上升 5-10 倍，但召回质量显著提升'],
        ['ADR-KB-02',
         '分块策略选择 Markdown 结构分块',
         '按 H2→H3→段落→句子优先级切分，目标 256 Token，重叠 50',
         '实现复杂度高于固定大小分块，但语义完整性更好'],
        ['ADR-KB-03',
         '三机制增量同步（事件+增量+全量）',
         '事件通知近实时 + updated_at 增量补漏 + 全量兜底',
         '三套机制维护成本高，但可靠性最强'],
        ['ADR-KB-04',
         '四维质量评分模型',
         '召回频次(0.4) + 用户反馈(0.3) + 新鲜度(0.2) + 完整度(0.1)',
         '权重需定期调优，但维度覆盖全面'],
        ['ADR-KB-05',
         '重复检测使用余弦相似度双阈值',
         '0.95 标记 DUPLICATE，0.85-0.95 标记 SIMILAR',
         '可能误报，但配合人工复核可控制'],
        ['ADR-KB-06',
         '版本管理采用 SemVer + 全量历史表',
         'knowledge_article_versions 存储全量历史，支持一键回滚',
         '存储成本上升，但可追溯性强'],
        ['ADR-KB-07',
         'Embedding 模型版本字段隔离',
         'knowledge_vectors.embedding_model 字段标记模型版本，新旧模型向量不混淆',
         '模型升级需全量重嵌入，但避免向量空间不一致'],
    ],
    col_widths=[2.5, 3, 5.5, 3])

add_callout(doc,
    '所有 ADR 决策记录归档在 docs/agent-design/adr/ 目录下，'
    '文件名格式：ADR-KB-XX-决策标题.md。每条 ADR 包含：'
    '背景、决策、权衡、后果、相关日期。',
    color='E8F4FD', border_color='2196F3')

# ==================== 文档结束 ====================
doc.add_paragraph()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end_p.add_run('— 文档结束 —')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
run.font.italic = True

# 文档元信息
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p.add_run('CampusShare Agent 搭建系列文档 · 第 3 个方向（RAG 数据侧）\n'
                     '版本：v1.0.0  ·  ADR 前缀：KB  ·  生成日期：2026-07-05')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

# ==================== 保存文档 ====================
output_path = r'e:\workspace_work\CampusShare\docs\agent-design\知识库管理模块设计方案.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
print(f'章节数：7 章 + 附录')
print(f'ADR 数量：7 条（ADR-KB-01 ~ ADR-KB-07）')
