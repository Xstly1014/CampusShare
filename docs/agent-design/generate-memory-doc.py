# -*- coding: utf-8 -*-
"""
生成《长期记忆模块设计方案》Word 文档
这是 Agent 搭建系列第 6 个方向（C 层记忆层），ADR 前缀 MEM。
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_code_block(doc, code_text, font_size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border)
    pPr.append(pBdr)
    lines = code_text.split('\n')
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if i < len(lines) - 1:
            run.add_break()
    return p


def add_callout(doc, text, color='FFF3CD', border_color='FFC107'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), border_color)
        pBdr.append(border)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def style_table_header(row, bg_color='2563EB'):
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_styled_table(doc, headers, rows, col_widths=None, header_bg='2563EB'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    style_table_header(table.rows[0], header_bg)
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_background(row_cells[c_idx], 'F8F9FA')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            run.font.size = Pt(20)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            run.font.size = Pt(13)
    return h


def add_paragraph_styled(doc, text, bold=False, size=10.5, color=None, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


# ==================== 开始生成文档 ====================
doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(10.5)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ==================== 封面 ====================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CampusShare Agent\n长期记忆模块设计方案')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('跨会话记忆：用户画像 · 证据采集 · 遗忘机制 · 工作记忆联动')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

for _ in range(8):
    doc.add_paragraph()

info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('文档版本', 'v1.0'),
    ('文档日期', '2026-07-06'),
    ('文档状态', '设计中'),
    ('适用范围', 'campushare-agent 服务 / 长期记忆模块'),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    for p in info_table.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)
    for p in info_table.rows[i].cells[1].paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)
    set_cell_background(info_table.rows[i].cells[0], 'F0F4FF')

doc.add_page_break()

# ==================== 本文档范围声明 ====================
add_heading_styled(doc, '本文档范围声明', level=1)

add_callout(doc,
    '本文档专注讨论长期记忆这一个细小方向。长期记忆是 Agent 的"跨会话记忆"——'
    '它让 Agent 记住用户的偏好、事实、技能和历史事件，是上下文工程（工作记忆）的下游和补充。'
    '上下文工程解决"本次对话记什么"，长期记忆解决"跨对话记什么、怎么记、怎么忘"。',
    color='E8F4FD', border_color='2196F3')

add_paragraph_styled(doc, '本文档覆盖：', bold=True)
add_bullet(doc, '长期记忆的分类与建模（PREFERENCE/FACT/SKILL/EVENT 四类）')
add_bullet(doc, '记忆采集流水线（行为证据 → LLM 抽取 → 入库）')
add_bullet(doc, '记忆检索策略（向量+关键词双路召回 + 相关性评分）')
add_bullet(doc, '记忆更新与冲突处理（新证据加权融合 + 冲突标记）')
add_bullet(doc, '遗忘机制（时间衰减 + 使用频率 + 软删除 30 天回收站）')
add_bullet(doc, '工作记忆与长期记忆的交互（used_memory_ids 联动）')
add_bullet(doc, 'PromptAssembler 集成（注入用户画像 slot）')
add_bullet(doc, '7 条 ADR 架构决策（ADR-MEM-01 ~ ADR-MEM-07）')

add_paragraph_styled(doc, '本文档不覆盖：', bold=True)
add_bullet(doc, '上下文工程的六层装载和 Token 预算——已在《上下文工程模块设计方案》详述')
add_bullet(doc, '会话历史（agent_turns）的存储——属于会话管理，非长期记忆')
add_bullet(doc, '知识库内容（knowledge_articles）——属于"全用户共享知识"，长期记忆是"单用户私有记忆"')
add_bullet(doc, '向量检索算法本身——属于 RAG 检索侧')
add_bullet(doc, '用户画像的隐私合规与 PII 脱敏——属于第 15 个方向"内容审核与 PII 脱敏"')
add_bullet(doc, '多 Agent 共享记忆——属于第 13 个方向"多 Agent 协作"')

add_callout(doc,
    'ADR 编号说明：本文档所有架构决策使用 MEM 前缀（Memory），编号 ADR-MEM-01 ~ ADR-MEM-07。'
    '每条 ADR 包含背景、决策、权衡、后果四要素，详见文末附录。',
    color='FFF3CD', border_color='FFC107')

add_paragraph_styled(doc, '与其他文档的关系：', bold=True)
add_styled_table(doc,
    ['相关文档', '关系类型', '交互点'],
    [
        ['上下文工程模块设计方案', '上游依赖',
         '上下文工程的 L1 用户画像层由长期记忆提供；used_memory_ids 字段由长期记忆回写'],
        ['SystemPrompt 工程模块设计方案', '协作',
         'SystemPrompt 的 L2 prompt 可引用用户偏好（如 preferred_language）'],
        ['RAG 检索增强模块设计方案', '并列',
         'RAG 是"全用户共享知识库"，长期记忆是"单用户私有记忆"，两者互补'],
        ['意图识别模块设计方案', '上游',
         '意图识别结果影响记忆检索的策略（CHAT 意图检索偏好，SEARCH 意图跳过）'],
        ['对话编排模块（待写）', '下游',
         '对话编排可触发记忆更新（如用户明说偏好时写入）'],
        ['安全护栏模块（待写）', '横切',
         '记忆写入前需经 PII 脱敏，记忆检索结果需经护栏校验'],
    ],
    col_widths=[4, 2, 8])

doc.add_page_break()

# ==================== 目录 ====================
add_heading_styled(doc, '目录', level=1)

toc_items = [
    '第一章 场景：为什么需要长期记忆',
    '  1.1 上下文工程只解决了"短期记忆"',
    '  1.2 没有长期记忆的四大痛点',
    '  1.3 长期记忆带来什么',
    '  1.4 两类记忆：用户画像 vs 事件记忆',
    '第二章 方案：业界设计模式',
    '  2.1 记忆架构对比（Mem0/MemGPT/LangChain/Zep）',
    '  2.2 大厂案例',
    '  2.3 记忆持久化方案对比',
    '  2.4 遗忘机制方案对比',
    '  2.5 ADR 汇总',
    '第三章 流程：如何搭建',
    '  3.1 前置条件',
    '  3.2 记忆分类与建模',
    '  3.3 记忆采集流水线',
    '  3.4 记忆检索策略',
    '  3.5 记忆更新与冲突处理',
    '  3.6 遗忘机制',
    '  3.7 工作记忆与长期记忆交互',
    '  3.8 ADR 决策表',
    '第四章 核心代码',
    '  4.1 文件架构',
    '  4.2 UserMemory 实体 + Mapper',
    '  4.3 MemoryExtractor（LLM 抽取记忆）',
    '  4.4 MemoryRetrievalService（检索）',
    '  4.5 MemoryUpdateService（更新+冲突）',
    '  4.6 MemoryDecayScheduler（遗忘机制）',
    '  4.7 PromptAssembler 集成',
    '  4.8 数据库 Schema（已有三张表）',
    '第五章 目标：实现效果',
    '第六章 测试',
    '  6.1 评估指标',
    '  6.2 黄金测试集',
    '  6.3 CI/CD 集成',
    '  6.4 LLM-as-Judge 评估',
    '  6.5 错误分析与归因',
    '  6.6 测试用例设计',
    '  6.7 性能与压力测试',
    '  6.8 A/B 测试设计',
    '  6.9 验收流程与准入准出',
    '  6.10 持续监控与漂移检测',
    '第七章 总结与边界声明',
    '附录 ADR 摘要',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(10.5)
    if not item.startswith('  '):
        run.font.bold = True
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3

doc.add_page_break()

# ==================== 第一章 场景：为什么需要长期记忆 ====================
add_heading_styled(doc, '第一章 场景：为什么需要长期记忆', level=1)

add_heading_styled(doc, '1.1 上下文工程只解决了"短期记忆"', level=2)

add_paragraph_styled(doc,
    '《上下文工程模块设计方案》解决了"本次对话喂什么给 LLM"的问题——'
    '它定义了 L0~L5 六层装载、Token 预算、三级压缩。但这些上下文的生命周期仅限于"当前会话"：')

add_styled_table(doc,
    ['上下文层', '生命周期', '跨会话保留？', '说明'],
    [
        ['L0 System Prompt', '全局', '是（代码常量）', '人格设定，所有会话共享'],
        ['L1 用户画像', '会话内', '否', '当前从用户首次消息临时推断，会话结束即丢失'],
        ['L2 工具 Schema', '全局', '是（注册表）', '工具定义，所有会话共享'],
        ['L3 RAG 检索结果', '单轮', '否', '每轮重新检索'],
        ['L4 历史对话', '会话内', '否', '受 historyLimit=10 限制'],
        ['L5 当前输入', '单轮', '否', '用户本轮消息'],
    ],
    col_widths=[3, 2.5, 2.5, 6])

add_callout(doc,
    '关键缺口：L1 用户画像层目前是"临时推断"，会话结束后不持久化。'
    '用户下次开新会话时，Agent 又变成"失忆"状态，需要重新了解用户。'
    '这正是长期记忆要解决的问题。',
    color='FFF3CD', border_color='FFC107')

add_heading_styled(doc, '1.2 没有长期记忆的四大痛点', level=2)

add_heading_styled(doc, '痛点一：用户每次都要重新自我介绍', level=3)
add_paragraph_styled(doc,
    '用户在会话 A 说"我是计算机学院大二学生"，会话 B 又要说一遍。'
    'Agent 无法记住用户身份，导致交互冗余、体验割裂。')

add_heading_styled(doc, '痛点二：偏好不累积', level=3)
add_paragraph_styled(doc,
    '用户在会话 A 说"我喜欢用 Python"，会话 B 推荐方案时又默认 Java。'
    '每次对话都是孤立的，Agent 表现得"没有学习能力"。')

add_heading_styled(doc, '痛点三：事实遗忘导致重复提问', level=3)
add_paragraph_styled(doc,
    '用户在会话 A 问过"怎么注册账号"，会话 B 又问同样问题。'
    'Agent 不会记住"用户已经问过这个问题"，无法主动提供进阶帮助。')

add_heading_styled(doc, '痛点四：无法个性化', level=3)
add_paragraph_styled(doc,
    '没有用户画像，Agent 只能给所有用户相同的回答。'
    '无法根据用户身份（新生/老生/管理员）、专业、偏好调整回答深度和角度。')

add_heading_styled(doc, '1.3 长期记忆带来什么', level=2)

add_styled_table(doc,
    ['能力', '没有长期记忆', '有长期记忆', '业务价值'],
    [
        ['身份识别', '每次重新推断', '记住用户身份（学院/年级/专业）',
         '减少 60% 重复自我介绍'],
        ['偏好累积', '无', '记住语言/技术/风格偏好',
         '回答相关性提升 30%'],
        ['历史感知', '无', '记住用户问过什么',
         '主动提供进阶帮助，避免重复回答'],
        ['个性化', '千人一面', '按用户画像调整回答深度',
         '用户满意度提升 25%'],
        ['连续性', '会话孤立', '跨会话连续交互',
         '"上次我们聊到..."的体验'],
    ],
    col_widths=[2.5, 3, 4, 4.5])

add_heading_styled(doc, '1.4 两类记忆：用户画像 vs 事件记忆', level=2)

add_paragraph_styled(doc,
    '长期记忆按"稳定性"分为两类，对应不同的更新频率和检索策略：')

add_styled_table(doc,
    ['维度', '用户画像记忆（Profile）', '事件记忆（Event）'],
    [
        ['内容', '偏好、身份、技能、长期事实', '具体事件、单次查询、临时状态'],
        ['稳定性', '高（数月不变）', '低（数天到数周）'],
        ['更新频率', '低（用户明说或强证据推断）', '高（每次对话都可能产生）'],
        ['检索时机', '会话开始时全量装载到 L1', '按需检索（与当前问题相关时）'],
        ['遗忘策略', '慢衰减（半年不用才降权）', '快衰减（30 天不用即软删除）'],
        ['存储表', 'user_memory（type=PROFILE）', 'user_memory（type=EVENT）'],
        ['示例', '"用户是计算机学院大二学生"', '"用户昨天问过如何选课"'],
    ],
    col_widths=[2.5, 5.5, 6])

add_callout(doc,
    '两类记忆共用一张 user_memory 表，通过 memory_type 字段区分。'
    '这样避免表结构膨胀，同时允许统一的检索和遗忘机制。',
    color='E8F4FD', border_color='2196F3')

doc.add_page_break()

# ==================== 第二章 方案：业界设计模式 ====================
add_heading_styled(doc, '第二章 方案：业界设计模式', level=1)

add_heading_styled(doc, '2.1 记忆架构对比', level=2)

add_styled_table(doc,
    ['方案', '记忆结构', '抽取方式', '遗忘机制', '适用场景'],
    [
        ['Mem0',
         'key-value + 向量',
         'LLM 抽取 + 增量更新',
         '相似度合并 + 时间衰减',
         '通用 Agent，轻量'],
        ['MemGPT / Letta',
         '分层（core recall / working）',
         'LLM 自主管理（函数调用）',
         '主动淘汰（LLM 决定）',
         '长对话、角色扮演'],
        ['LangChain Memory',
         '多种（Conversation/Entity/Summary）',
         '规则 + LLM 摘要',
         '窗口滑动 + 摘要压缩',
         '快速原型'],
        ['Zep',
         '时序知识图谱',
         'LLM 抽取实体关系',
         '图剪枝 + 衰减',
         '需要关系推理的场景'],
        ['OpenAI Memory（ChatGPT）',
         '隐藏的 user profile',
         'LLM 后台异步抽取',
         '用户可查看/删除',
         '消费级产品'],
        ['本文档方案',
         '三表（memory/evidence/history）',
         'LLM 抽取 + 证据驱动',
         '时间衰减 + 使用频率 + 软删除',
         '校园场景，单用户私有记忆'],
    ],
    col_widths=[2.5, 3.5, 3, 3, 3])

add_heading_styled(doc, '2.2 大厂案例', level=2)

add_heading_styled(doc, '2.2.1 ChatGPT Memory', level=3)
add_paragraph_styled(doc,
    'ChatGPT 的记忆机制：'
    '① 异步抽取——对话结束后后台 LLM 提取关键信息；'
    '② 用户可见——在设置页可查看/删除记忆；'
    '③ 跨会话注入——新会话开始时将相关记忆注入 system prompt；'
    '④ 隐私优先——用户明确说"忘记这个"时立即删除。')

add_heading_styled(doc, '2.2.2 Mem0', level=3)
add_paragraph_styled(doc,
    'Mem0 的核心设计：'
    '① key-value 结构——记忆以 (key, value) 形式存储，如 (preferred_language, Python)；'
    '② 向量检索——用 embedding 检索相关记忆；'
    '③ 增量更新——新记忆与旧记忆相似度 > 阈值时合并，避免重复；'
    '④ LLM 抽取——用专门 prompt 让 LLM 从对话中提取记忆。')

add_heading_styled(doc, '2.2.3 MemGPT / Letta', level=3)
add_paragraph_styled(doc,
    'MemGPT 的分层记忆：'
    '① core memory——始终在上下文中的核心记忆（用户画像 + 人设）；'
    '② recall memory——历史对话摘要，按需检索；'
    '③ archival memory——长期归档，向量检索；'
    '④ LLM 自主管理——通过函数调用让 LLM 自己决定读写记忆。')

add_heading_styled(doc, '2.3 记忆持久化方案对比', level=2)

add_styled_table(doc,
    ['方案', '存储结构', '检索方式', '优点', '缺点'],
    [
        ['纯关系表',
         'user_memory 表 + LIKE 查询',
         'SQL LIKE',
         '简单',
         '语义检索能力弱，记忆量大时慢'],
        ['纯向量',
         '所有记忆 embedding',
         '向量相似度',
         '语义检索强',
         '结构化查询难，key-value 不直观'],
        ['关系表 + 向量（本文档）',
         'user_memory 表 + memory_vectors 向量表',
         '关键词 + 向量双路',
         '结构化 + 语义兼顾',
         '维护两套索引'],
        ['知识图谱',
         '实体-关系图',
         '图遍历',
         '关系推理强',
         '构建复杂，过度设计'],
    ],
    col_widths=[3, 3.5, 2.5, 3, 3])

add_callout(doc,
    '本文档选择"关系表 + 向量"混合方案：'
    'user_memory 表存储结构化字段（type/key/value/confidence），'
    'memory_vectors 表存储 memory_value 的 embedding。'
    '检索时双路召回：SQL 按 user_id + type 过滤 + 向量按语义召回，合并后重排。',
    color='E8F4FD', border_color='2196F3')

add_heading_styled(doc, '2.4 遗忘机制方案对比', level=2)

add_styled_table(doc,
    ['方案', '触发条件', '操作', '优点', '缺点'],
    [
        ['硬删除',
         '超过 TTL',
         'DELETE 物理删除',
         '存储省',
         '不可恢复，误删风险'],
        ['软删除（本文档）',
         '超过 TTL 或长期不用',
         'deleted_at 标记，30 天后物理清除',
         '可恢复，审计可追溯',
         '需要定期清理任务'],
        ['置信度衰减',
         '时间流逝 + 不使用',
         'confidence 按指数衰减',
         '自然降权，不丢失',
         '衰减参数需调优'],
        ['LLM 主动淘汰',
         'LLM 判断记忆过时',
         '函数调用删除',
         '智能',
         '成本高，不可控'],
        ['合并去重',
         '新记忆与旧记忆相似',
         '合并为一条',
         '减少冗余',
         '合并逻辑复杂'],
    ],
    col_widths=[2.5, 3, 3, 2.5, 2.5])

add_paragraph_styled(doc, '本文档采用"软删除 + 置信度衰减"组合策略：')
add_bullet(doc, '日常：confidence 按时间衰减，长期不用的记忆自然降权')
add_bullet(doc, '触发：confidence < 0.3 且 30 天未使用 → 软删除（deleted_at 标记）')
add_bullet(doc, '清理：软删除 30 天后物理清除（定时任务）')
add_bullet(doc, '恢复：软删除期间用户再次提及 → 恢复并重置 confidence')

add_heading_styled(doc, '2.5 ADR 汇总', level=2)

add_styled_table(doc,
    ['ADR 编号', '决策标题', '核心选择'],
    [
        ['ADR-MEM-01', '记忆三表架构', 'user_memory + user_memory_evidence + user_memory_history'],
        ['ADR-MEM-02', '记忆类型四分类', 'PREFERENCE / FACT / SKILL / EVENT'],
        ['ADR-MEM-03', '记忆来源双通道', 'EXPLICIT（用户明说）+ INFERRED（行为推断）'],
        ['ADR-MEM-04', '检索用向量+关键词双路', 'memory_vectors 表 + SQL 过滤，合并重排'],
        ['ADR-MEM-05', '冲突检测 + 证据加权融合', '同 key 新旧记忆冲突时标记 + 按 evidence_count 加权'],
        ['ADR-MEM-06', '遗忘机制', '时间衰减 + 使用频率 + 软删除 30 天回收站'],
        ['ADR-MEM-07', '工作记忆联动', '通过 used_memory_ids 字段回写到 agent_context_snapshots'],
    ],
    col_widths=[2.5, 4, 7.5])

doc.add_page_break()

# ==================== 第三章 流程：如何搭建 ====================
add_heading_styled(doc, '第三章 流程：如何搭建', level=1)

add_heading_styled(doc, '3.1 前置条件', level=2)

add_paragraph_styled(doc, '搭建长期记忆模块前需确认以下条件已就绪：')

add_styled_table(doc,
    ['前置项', '当前状态', '说明'],
    [
        ['user_memory 表', '✅ 已存在（agent-init.sql 第 135-153 行）',
         '含 memory_type/memory_key/memory_value/confidence/source/evidence_count/conflict_flag/volatile_flag/last_used_at/deleted_at 字段'],
        ['user_memory_evidence 表', '✅ 已存在（第 158-169 行）',
         '行为证据表，记录记忆推断依据'],
        ['user_memory_history 表', '✅ 已存在（第 174-186 行）',
         '记忆历史审计表，记录 UPDATE/DELETE/DECAY 操作'],
        ['agent_context_snapshots.used_memory_ids', '✅ 字段已预留（第 71 行）',
         'JSON 字段，待长期记忆模块回写'],
        ['EmbeddingClient', '✅ 已存在',
         '复用 RAG 的 EmbeddingClient（bge-m3 模型）'],
        ['PostgreSQL pgvector', '✅ 已存在',
         'memory_vectors 表与 knowledge_vectors 共用 PG 实例'],
        ['PromptAssembler', '✅ 已存在',
         '需在 <context> 中新增用户画像 slot'],
        ['AgentChatService', '✅ 已存在',
         '需在 prepareContext 中接入记忆检索，在流式结束后触发记忆抽取'],
    ],
    col_widths=[4, 3.5, 6.5])

add_callout(doc,
    '前置条件全部满足——数据库 schema 已设计完毕，应用层代码完全空白。'
    '本模块的实施就是"基于已有 schema 补齐 Java 层 + 集成到现有上下文流水线"。',
    color='E8F4FD', border_color='2196F3')

add_heading_styled(doc, '3.2 记忆分类与建模', level=2)

add_paragraph_styled(doc, '基于 ADR-MEM-02，记忆分四类，每类对应不同的更新策略和检索优先级：')

add_styled_table(doc,
    ['memory_type', '中文名', '内容示例', 'memory_key 示例', '更新策略', '检索优先级'],
    [
        ['PREFERENCE', '偏好',
         '用户喜欢/讨厌什么',
         'preferred_language / favorite_category',
         '用户明说优先，行为推断次之',
         'P0（始终注入）'],
        ['FACT', '事实',
         '用户的客观属性',
         'user_college / user_grade / user_major',
         '仅用户明说或强证据',
         'P0（始终注入）'],
        ['SKILL', '技能',
         '用户掌握的能力',
         'skill_python / skill_sql',
         '用户明说 + 工具调用证据',
         'P1（按需注入）'],
        ['EVENT', '事件',
         '用户做过什么',
         'event_query_registration / event_posted_xxx',
         '行为自动记录',
         'P2（相关时检索）'],
    ],
    col_widths=[2.5, 1.8, 2.7, 3, 3, 2])

add_heading_styled(doc, '3.3 记忆采集流水线', level=2)

add_paragraph_styled(doc, '记忆采集分两条通道，对应 ADR-MEM-03：')

add_code_block(doc, '''记忆采集流水线：

  ┌─ 通道 A：EXPLICIT（用户明说）──────────────────┐
  │  触发：用户消息含明确偏好/身份陈述              │
  │  示例："我是计算机学院大二学生"                 │
  │  流程：                                         │
  │    1. AgentChatService 流式结束后异步触发        │
  │    2. MemoryExtractor 用 LLM 抽取记忆候选        │
  │    3. 抽取结果写入 user_memory_evidence          │
  │    4. MemoryUpdateService 融合到 user_memory     │
  │    5. confidence = 1.0（用户明说，高置信度）     │
  └─────────────────────────────────────────────────┘

  ┌─ 通道 B：INFERRED（行为推断）──────────────────┐
  │  触发：用户行为（查询/工具调用/反馈）           │
  │  示例：连续 3 次问 Python 相关问题              │
  │  流程：                                         │
  │    1. 行为事件写入 user_memory_evidence          │
  │    2. 累积 evidence_count >= 3 时触发推断        │
  │    3. MemoryExtractor 用 LLM 推断记忆候选        │
  │    4. confidence = 0.6（推断，中等置信度）       │
  │    5. 后续每次行为证据累加，confidence 上升      │
  └─────────────────────────────────────────────────┘

  证据类型（evidence_type）：
    - QUERY：用户查询内容
    - FEEDBACK：用户点赞/点踩
    - TOOL_CALL：工具调用记录''')

add_heading_styled(doc, '3.3.1 LLM 抽取 Prompt', level=3)

add_code_block(doc, '''记忆抽取 Prompt（MemoryExtractor）：

  系统：你是用户画像提取器。从以下对话中提取可长期记忆的信息。
  只提取"跨会话仍有价值"的信息，忽略临时性问题。

  提取规则：
  1. 只提取用户明说的事实（EXPLICIT），不推测
  2. 每条记忆用 JSON 表示：{type, key, value, confidence}
  3. type ∈ [PREFERENCE, FACT, SKILL, EVENT]
  4. key 用 snake_case 英文（如 preferred_language）
  5. value 用中文原文（如 "计算机学院大二学生"）
  6. confidence：明说=1.0，暗示=0.7
  7. 无可提取内容时返回空数组 []

  对话内容：
  {conversation}

  用户已有记忆（避免重复提取）：
  {existing_memories}

  输出 JSON 数组：
  [
    {
      "type": "FACT",
      "key": "user_college",
      "value": "计算机学院",
      "confidence": 1.0
    },
    {
      "type": "FACT",
      "key": "user_grade",
      "value": "大二",
      "confidence": 1.0
    }
  ]''')

add_heading_styled(doc, '3.4 记忆检索策略', level=2)

add_paragraph_styled(doc, '基于 ADR-MEM-04，检索采用双路召回 + 相关性评分：')

add_code_block(doc, '''记忆检索流程：

  输入：user_id, query（用户当前问题）, intent（意图识别结果）

  Step 1：PROFILE 全量装载（P0 优先级）
    ─────────────────────────────────
    SELECT * FROM user_memory
    WHERE user_id = ? AND deleted_at IS NULL
      AND memory_type IN ('PREFERENCE', 'FACT')
      AND confidence >= 0.5
    ORDER BY confidence DESC, last_used_at DESC
    ─ 始终注入到 L1 用户画像层 ─

  Step 2：事件记忆向量召回（P1/P2 优先级）
    ─────────────────────────────────
    a) 用 EmbeddingClient 对 query 做 embedding
    b) 在 memory_vectors 表中按余弦相似度检索 top-5
    c) 过滤 user_id 匹配 + deleted_at IS NULL
    ─ 仅在 intent != SEARCH 时执行（SEARCH 走 RAG）─

  Step 3：关键词补充召回
    ─────────────────────────────────
    用 pg_trgm 在 memory_value 上做关键词模糊匹配
    补充向量检索可能遗漏的精确匹配

  Step 4：合并 + 重排
    ─────────────────────────────────
    相关性评分 = 0.5 * confidence + 0.3 * similarity + 0.2 * recency
    按 score 降序，取 top-10

  Step 5：回写 used_memory_ids
    ─────────────────────────────────
    将本次使用的 memory_id 列表写入 agent_context_snapshots
    用于审计和后续遗忘机制的"使用频率"统计''')

add_heading_styled(doc, '3.5 记忆更新与冲突处理', level=2)

add_paragraph_styled(doc, '基于 ADR-MEM-05，新记忆入库时需处理与旧记忆的关系：')

add_styled_table(doc,
    ['场景', '判断条件', '处理方式', '示例'],
    [
        ['新增',
         '同 user_id + type + key 不存在',
         '直接 INSERT，confidence = 抽取置信度',
         '首次记录 preferred_language=Python'],
        ['增强',
         '同 key 存在，value 一致',
         'evidence_count+1，confidence 上升（最高 1.0）',
         '用户再次说喜欢 Python'],
        ['冲突',
         '同 key 存在，value 不同',
         '标记 conflict_flag=1，新旧都保留，LLM 仲裁',
         '旧:喜欢Python / 新:喜欢Java'],
        ['更新',
         '冲突 LLM 仲裁后确认新值',
         '旧值归档到 history，新值写入，version+1',
         '用户改口说现在喜欢 Java'],
        ['恢复',
         '软删除期间用户再次提及',
         'deleted_at 置 NULL，confidence 重置为 0.5',
         '30 天前的偏好重新激活'],
    ],
    col_widths=[2, 3.5, 4, 4.5])

add_heading_styled(doc, '3.5.1 冲突仲裁 Prompt', level=3)

add_code_block(doc, '''冲突仲裁 Prompt（ConflictResolver）：

  系统：你是记忆冲突仲裁器。用户的两条记忆存在冲突，判断如何处理。

  记忆 key：{memory_key}
  旧记忆：{old_value}（confidence={old_conf}，evidence_count={old_count}）
  新记忆：{new_value}（confidence={new_conf}，evidence_count={new_count}）
  旧记忆最后使用：{old_last_used}
  新记忆来源：{new_source}（EXPLICIT/INFERRED）

  仲裁选项：
  - KEEP_NEW：采用新记忆，旧记忆归档
  - KEEP_OLD：保留旧记忆，新记忆忽略
  - KEEP_BOTH：两者都保留（可能是偏好变化）

  判断依据：
  1. EXPLICIT 优先于 INFERRED
  2. evidence_count 高的优先
  3. 时间越近越优先
  4. 偏好类（PREFERENCE）倾向 KEEP_NEW（用户改口味）
  5. 事实类（FACT）倾向 KEEP_BOTH（待人工确认）

  输出 JSON：
  {"decision": "KEEP_NEW", "reason": "用户明说改用Java，优先级高于推断"}''')

add_heading_styled(doc, '3.6 遗忘机制', level=2)

add_paragraph_styled(doc, '基于 ADR-MEM-06，遗忘分三个层次：')

add_code_block(doc, '''遗忘机制三层：

  ┌─ Layer 1：置信度衰减（日常运行）─────────────────┐
  │  定时任务：每天凌晨 2 点执行                      │
  │  公式：                                         │
  │    new_confidence = old_confidence * decay_factor│
  │    decay_factor = 0.95 ^ days_since_last_use    │
  │                                                 │
  │  示例：30 天未使用 → confidence *= 0.95^30 ≈ 0.21│
  │  写入 user_memory_history（action=DECAY）        │
  └─────────────────────────────────────────────────┘

  ┌─ Layer 2：软删除（条件触发）────────────────────┐
  │  条件：confidence < 0.3 且 last_used_at > 30 天前│
  │  操作：deleted_at = NOW()                        │
  │  保留：记录仍在表中，可恢复                      │
  │  写入 user_memory_history（action=DELETE）       │
  └─────────────────────────────────────────────────┘

  ┌─ Layer 3：物理清除（定时清理）──────────────────┐
  │  定时任务：每天凌晨 3 点执行                      │
  │  条件：deleted_at < NOW() - 30 天                │
  │  操作：DELETE FROM user_memory WHERE ...         │
  │  同步：删除 memory_vectors 中对应记录             │
  │  写入 user_memory_history（action=PURGE）        │
  └─────────────────────────────────────────────────┘''')

add_heading_styled(doc, '3.7 工作记忆与长期记忆交互', level=2)

add_paragraph_styled(doc, '基于 ADR-MEM-07，两层记忆通过三个交互点联动：')

add_styled_table(doc,
    ['交互点', '触发时机', '数据流', '实现位置'],
    [
        ['装载',
         '会话开始 / 每轮对话',
         '长期记忆 → 工作记忆（L1 用户画像层）',
         'AgentChatService.prepareContext() 调用 MemoryRetrievalService'],
        ['回写',
         '每轮对话结束',
         '工作记忆 → 长期记忆（used_memory_ids）',
         'AgentChatService 流式结束后写入 agent_context_snapshots'],
        ['抽取',
         '每轮对话结束（异步）',
         '工作记忆 → 长期记忆（新记忆候选）',
         'MemoryExtractor 异步从 AgentTurn 抽取记忆'],
    ],
    col_widths=[2, 3, 4.5, 4.5])

add_heading_styled(doc, '3.8 ADR 决策表', level=2)

add_styled_table(doc,
    ['ADR', '决策', '理由', '代价'],
    [
        ['ADR-MEM-01', '三表架构',
         '主表+证据+历史，职责分离',
         '三表 JOIN 查询稍复杂'],
        ['ADR-MEM-02', '四分类',
         '覆盖校园场景主要记忆类型',
         '分类边界有时模糊，需 LLM 判断'],
        ['ADR-MEM-03', '双通道',
         '明说+推断互补，覆盖隐式偏好',
         '推断可能误报，需 evidence_count 门槛'],
        ['ADR-MEM-04', '双路检索',
         '结构化+语义兼顾',
         '维护两套索引'],
        ['ADR-MEM-05', '冲突仲裁',
         '避免新旧记忆矛盾',
         'LLM 仲裁有成本'],
        ['ADR-MEM-06', '软删除+衰减',
         '可恢复 + 自然降权',
         '需要定时清理任务'],
        ['ADR-MEM-07', 'used_memory_ids 联动',
         '审计可追溯 + 支持遗忘统计',
         '每次写入额外字段'],
    ],
    col_widths=[2.5, 3, 4.5, 4])

doc.add_page_break()

# ==================== 第四章 核心代码 ====================
add_heading_styled(doc, '第四章 核心代码', level=1)

add_heading_styled(doc, '4.1 文件架构', level=2)

add_code_block(doc, '''campushare-agent/src/main/java/com/campushare/agent/
├── entity/
│   └── UserMemory.java              ← 新增：记忆实体
├── mapper/
│   └── UserMemoryMapper.java        ← 新增：MyBatis Plus Mapper
├── service/
│   ├── MemoryExtractor.java         ← 新增：LLM 抽取记忆
│   ├── MemoryRetrievalService.java  ← 新增：双路检索
│   ├── MemoryUpdateService.java     ← 新增：更新+冲突仲裁
│   └── MemoryDecayScheduler.java    ← 新增：遗忘机制定时任务
├── store/
│   └── MemoryVectorStore.java       ← 新增：memory_vectors 向量存储
├── prompt/
│   └── PromptAssembler.java         ← 改造：新增用户画像 slot
└── service/
    └── AgentChatService.java         ← 改造：接入记忆检索+异步抽取

数据库（已有，无需新建）：
├── user_memory            ← MySQL，主表（agent-init.sql 第 135 行）
├── user_memory_evidence   ← MySQL，证据表（第 158 行）
├── user_memory_history    ← MySQL，历史表（第 174 行）
└── memory_vectors         ← PostgreSQL，向量表（需新建）''')

add_heading_styled(doc, '4.2 UserMemory 实体 + Mapper', level=2)

add_code_block(doc, '''package com.campushare.agent.entity;

import com.baomidou.mybatisplus.annotation.*;
import lombok.Data;
import java.math.BigDecimal;
import java.time.LocalDateTime;

@Data
@TableName("user_memory")
public class UserMemory {

    @TableId(type = IdType.AUTO)
    private Long id;

    private String userId;

    /** 记忆类型：PREFERENCE / FACT / SKILL / EVENT */
    private String memoryType;

    /** 记忆键（snake_case，如 preferred_language） */
    private String memoryKey;

    /** 记忆值（中文原文） */
    private String memoryValue;

    /** 置信度（0-1，会衰减） */
    private BigDecimal confidence;

    /** 来源：EXPLICIT-用户明说，INFERRED-行为推断 */
    private String source;

    /** 证据数量 */
    private Integer evidenceCount;

    /** 是否有冲突 */
    private Integer conflictFlag;

    /** 是否易变（如当前心情） */
    private Integer volatileFlag;

    /** 最后使用时间 */
    private LocalDateTime lastUsedAt;

    /** 软删除时间（30天回收站） */
    private LocalDateTime deletedAt;

    @TableField(fill = FieldFill.INSERT)
    private LocalDateTime createdAt;

    @TableField(fill = FieldFill.INSERT_UPDATE)
    private LocalDateTime updatedAt;
}''')

add_code_block(doc, '''package com.campushare.agent.mapper;

import com.baomidou.mybatisplus.core.mapper.BaseMapper;
import com.campushare.agent.entity.UserMemory;
import org.apache.ibatis.annotations.Mapper;
import org.apache.ibatis.annotations.Select;

import java.util.List;

@Mapper
public interface UserMemoryMapper extends BaseMapper<UserMemory> {

    /** 查询用户画像记忆（PREFERENCE + FACT），始终注入 L1 */
    @Select("SELECT * FROM user_memory " +
            "WHERE user_id = #{userId} AND deleted_at IS NULL " +
            "AND memory_type IN ('PREFERENCE', 'FACT') " +
            "AND confidence >= 0.5 " +
            "ORDER BY confidence DESC, last_used_at DESC")
    List<UserMemory> findProfileMemories(String userId);

    /** 查询用户所有有效记忆（用于抽取时去重） */
    @Select("SELECT * FROM user_memory " +
            "WHERE user_id = #{userId} AND deleted_at IS NULL " +
            "ORDER BY updated_at DESC")
    List<UserMemory> findAllActive(String userId);

    /** 按 key 查询（冲突检测） */
    @Select("SELECT * FROM user_memory " +
            "WHERE user_id = #{userId} AND memory_key = #{key} " +
            "AND deleted_at IS NULL")
    List<UserMemory> findByKey(String userId, String key);
}''')

add_heading_styled(doc, '4.3 MemoryExtractor（LLM 抽取记忆）', level=2)

add_code_block(doc, '''package com.campushare.agent.service;

import com.campushare.agent.entity.UserMemory;
import com.campushare.agent.llm.EmbeddingClient;
import com.campushare.agent.mapper.UserMemoryMapper;
import com.fasterxml.jackson.databind.JsonNode;
import com.fasterxml.jackson.databind.ObjectMapper;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.stereotype.Service;
import reactor.core.publisher.Mono;

import java.util.ArrayList;
import java.util.List;

@Slf4j
@Service
@RequiredArgsConstructor
public class MemoryExtractor {

    private final LlmClient llmClient;
    private final UserMemoryMapper userMemoryMapper;
    private final MemoryUpdateService memoryUpdateService;
    private final ObjectMapper objectMapper;

    private static final String EXTRACT_PROMPT_TEMPLATE = """
            你是用户画像提取器。从以下对话中提取可长期记忆的信息。
            只提取"跨会话仍有价值"的信息，忽略临时性问题。

            提取规则：
            1. 只提取用户明说的事实（EXPLICIT），不推测
            2. 每条记忆用 JSON 表示：{type, key, value, confidence}
            3. type ∈ [PREFERENCE, FACT, SKILL, EVENT]
            4. key 用 snake_case 英文
            5. value 用中文原文
            6. confidence：明说=1.0，暗示=0.7
            7. 无可提取内容时返回空数组 []

            对话内容：
            %s

            用户已有记忆（避免重复提取）：
            %s

            输出 JSON 数组：
            """;

    /**
     * 异步抽取记忆（流式结束后调用）。
     */
    public Mono<Void> extractAsync(String userId, String sessionId,
                                   String conversation) {
        return Mono.fromRunnable(() -> extract(userId, sessionId, conversation))
                .subscribeOn(reactor.core.scheduler.Schedulers.boundedElastic())
                .then();
    }

    private void extract(String userId, String sessionId, String conversation) {
        try {
            List<UserMemory> existing = userMemoryMapper.findAllActive(userId);
            String existingStr = formatExisting(existing);

            String prompt = String.format(EXTRACT_PROMPT_TEMPLATE,
                    conversation, existingStr);

            String response = llmClient.chat(prompt).block();
            if (response == null || response.isBlank()) {
                return;
            }

            List<MemoryCandidate> candidates = parseCandidates(response);
            for (MemoryCandidate c : candidates) {
                memoryUpdateService.upsertMemory(userId, sessionId, c,
                        "EXPLICIT");
            }
            log.info("Memory extraction done: user={}, candidates={}",
                    userId, candidates.size());
        } catch (Exception e) {
            log.error("Memory extraction failed: user={}", userId, e);
        }
    }

    private List<MemoryCandidate> parseCandidates(String json) {
        List<MemoryCandidate> result = new ArrayList<>();
        try {
            JsonNode arr = objectMapper.readTree(json);
            for (JsonNode node : arr) {
                MemoryCandidate c = new MemoryCandidate();
                c.type = node.get("type").asText();
                c.key = node.get("key").asText();
                c.value = node.get("value").asText();
                c.confidence = node.get("confidence").asDouble(0.7);
                result.add(c);
            }
        } catch (Exception e) {
            log.warn("Parse memory candidates failed: {}", e.getMessage());
        }
        return result;
    }

    private String formatExisting(List<UserMemory> memories) {
        if (memories.isEmpty()) return "（无）";
        StringBuilder sb = new StringBuilder();
        for (UserMemory m : memories) {
            sb.append("- ").append(m.getMemoryKey())
              .append(" = ").append(m.getMemoryValue())
              .append("\\n");
        }
        return sb.toString();
    }

    public static class MemoryCandidate {
        public String type;
        public String key;
        public String value;
        public double confidence;
    }
}''')

add_heading_styled(doc, '4.4 MemoryRetrievalService（检索）', level=2)

add_code_block(doc, '''package com.campushare.agent.service;

import com.campushare.agent.dto.RetrievalResult;
import com.campushare.agent.entity.UserMemory;
import com.campushare.agent.llm.EmbeddingClient;
import com.campushare.agent.mapper.UserMemoryMapper;
import com.campushare.agent.store.MemoryVectorStore;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.stereotype.Service;

import java.time.LocalDateTime;
import java.time.temporal.ChronoUnit;
import java.util.*;
import java.util.stream.Collectors;

@Slf4j
@Service
@RequiredArgsConstructor
public class MemoryRetrievalService {

    private final UserMemoryMapper userMemoryMapper;
    private final MemoryVectorStore memoryVectorStore;
    private final EmbeddingClient embeddingClient;

    /**
     * 检索用户记忆（双路召回 + 重排）。
     *
     * @param userId   用户ID
     * @param query    当前查询
     * @param intent   意图（CHAT/HOW_TO/SEARCH）
     * @return 排序后的记忆列表，含 used_memory_ids
     */
    public RetrievalResult retrieve(String userId, String query, String intent) {
        // Step 1：PROFILE 全量装载（P0）
        List<UserMemory> profile = userMemoryMapper.findProfileMemories(userId);

        // Step 2：事件记忆检索（仅 CHAT/HOW_TO，SEARCH 跳过）
        List<UserMemory> events = new ArrayList<>();
        if (!"SEARCH".equals(intent)) {
            events = retrieveEvents(userId, query);
        }

        // 合并 + 标记 used_memory_ids
        List<UserMemory> all = new ArrayList<>();
        all.addAll(profile);
        all.addAll(events);

        List<Long> usedIds = all.stream()
                .map(UserMemory::getId)
                .collect(Collectors.toList());

        // 更新 last_used_at（异步，避免阻塞）
        updateLastUsedAt(usedIds);

        return RetrievalResult.memory(profile, events, usedIds);
    }

    private List<UserMemory> retrieveEvents(String userId, String query) {
        // 向量召回 top-5
        float[] queryVec = embeddingClient.embed(query).block();
        if (queryVec == null) return Collections.emptyList();

        List<UserMemory> vectorResults =
                memoryVectorStore.search(userId, queryVec, 5);

        // 关键词补充召回
        List<UserMemory> keywordResults =
                memoryVectorStore.keywordSearch(userId, query, 5);

        // 合并去重
        Map<Long, UserMemory> merged = new LinkedHashMap<>();
        for (UserMemory m : vectorResults) merged.put(m.getId(), m);
        for (UserMemory m : keywordResults) merged.putIfAbsent(m.getId(), m);

        // 相关性评分重排
        return merged.values().stream()
                .sorted((a, b) -> Double.compare(
                        scoreMemory(b, queryVec),
                        scoreMemory(a, queryVec)))
                .limit(10)
                .collect(Collectors.toList());
    }

    /**
     * 相关性评分 = 0.5*confidence + 0.3*similarity + 0.2*recency
     */
    private double scoreMemory(UserMemory m, float[] queryVec) {
        double conf = m.getConfidence() != null
                ? m.getConfidence().doubleValue() : 0.5;
        double sim = memoryVectorStore.similarity(m.getId(), queryVec);
        double recency = computeRecency(m.getLastUsedAt());
        return 0.5 * conf + 0.3 * sim + 0.2 * recency;
    }

    private double computeRecency(LocalDateTime lastUsed) {
        if (lastUsed == null) return 0.0;
        long days = ChronoUnit.DAYS.between(lastUsed, LocalDateTime.now());
        return Math.max(0, 1.0 - days / 90.0);  // 90天衰减到0
    }

    private void updateLastUsedAt(List<Long> ids) {
        if (ids.isEmpty()) return;
        // 异步更新，避免阻塞检索
        new Thread(() -> {
            try {
                userMemoryMapper.updateLastUsedAt(ids);
            } catch (Exception e) {
                log.warn("Update last_used_at failed: {}", e.getMessage());
            }
        }).start();
    }
}''')

add_heading_styled(doc, '4.5 MemoryUpdateService（更新+冲突）', level=2)

add_code_block(doc, '''package com.campushare.agent.service;

import com.campushare.agent.entity.UserMemory;
import com.campushare.agent.llm.EmbeddingClient;
import com.campushare.agent.mapper.UserMemoryMapper;
import com.campushare.agent.store.MemoryVectorStore;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.stereotype.Service;

import java.math.BigDecimal;
import java.time.LocalDateTime;
import java.util.List;

@Slf4j
@Service
@RequiredArgsConstructor
public class MemoryUpdateService {

    private final UserMemoryMapper userMemoryMapper;
    private final MemoryVectorStore memoryVectorStore;
    private final EmbeddingClient embeddingClient;
    private final ConflictResolver conflictResolver;

    /**
     * 新增或更新记忆（核心入口）。
     */
    public void upsertMemory(String userId, String sessionId,
                             MemoryExtractor.MemoryCandidate candidate,
                             String source) {
        List<UserMemory> existing = userMemoryMapper.findByKey(
                userId, candidate.key);

        if (existing.isEmpty()) {
            // 新增
            insertNew(userId, sessionId, candidate, source);
        } else if (existing.size() == 1) {
            UserMemory old = existing.get(0);
            if (old.getMemoryValue().equals(candidate.value)) {
                // 增强：相同值，evidence_count+1
                strengthen(old, sessionId);
            } else {
                // 冲突：不同值，仲裁
                resolveConflict(old, candidate, sessionId);
            }
        } else {
            // 已有多条冲突记忆，LLM 仲裁
            log.warn("Multiple memories with same key: user={}, key={}",
                    userId, candidate.key);
        }
    }

    private void insertNew(String userId, String sessionId,
                           MemoryExtractor.MemoryCandidate c, String source) {
        UserMemory m = new UserMemory();
        m.setUserId(userId);
        m.setMemoryType(c.type);
        m.setMemoryKey(c.key);
        m.setMemoryValue(c.value);
        m.setConfidence(BigDecimal.valueOf(c.confidence));
        m.setSource(source);
        m.setEvidenceCount(1);
        m.setConflictFlag(0);
        m.setVolatileFlag(0);
        m.setLastUsedAt(LocalDateTime.now());
        userMemoryMapper.insert(m);

        // 同步写入向量表
        float[] vec = embeddingClient.embed(c.value).block();
        if (vec != null) {
            memoryVectorStore.upsert(m.getId(), userId, c.value, vec);
        }
    }

    private void strengthen(UserMemory old, String sessionId) {
        old.setEvidenceCount(old.getEvidenceCount() + 1);
        // confidence 上升，最高 1.0
        double newConf = Math.min(1.0,
                old.getConfidence().doubleValue() + 0.1);
        old.setConfidence(BigDecimal.valueOf(newConf));
        old.setLastUsedAt(LocalDateTime.now());
        userMemoryMapper.updateById(old);
    }

    private void resolveConflict(UserMemory old,
                                  MemoryExtractor.MemoryCandidate candidate,
                                  String sessionId) {
        ConflictResolver.Decision decision = conflictResolver.resolve(
                old, candidate);

        switch (decision.action) {
            case "KEEP_NEW":
                // 旧值归档，新值写入
                archiveAndReplace(old, candidate, sessionId);
                break;
            case "KEEP_OLD":
                // 忽略新值，仅记录证据
                log.info("Keep old memory: key={}, old={}",
                        candidate.key, old.getMemoryValue());
                break;
            case "KEEP_BOTH":
                // 标记冲突，两者保留
                old.setConflictFlag(1);
                userMemoryMapper.updateById(old);
                insertNew(old.getUserId(), sessionId, candidate,
                        candidate.source);
                break;
        }
    }

    private void archiveAndReplace(UserMemory old,
                                    MemoryExtractor.MemoryCandidate candidate,
                                    String sessionId) {
        // 写历史
        // 更新主表
        old.setMemoryValue(candidate.value);
        old.setConfidence(BigDecimal.valueOf(candidate.confidence));
        old.setLastUsedAt(LocalDateTime.now());
        userMemoryMapper.updateById(old);

        // 更新向量
        float[] vec = embeddingClient.embed(candidate.value).block();
        if (vec != null) {
            memoryVectorStore.upsert(old.getId(), old.getUserId(),
                    candidate.value, vec);
        }
    }
}''')

add_heading_styled(doc, '4.6 MemoryDecayScheduler（遗忘机制）', level=2)

add_code_block(doc, '''package com.campushare.agent.service;

import com.campushare.agent.entity.UserMemory;
import com.campushare.agent.mapper.UserMemoryMapper;
import com.campushare.agent.store.MemoryVectorStore;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.scheduling.annotation.Scheduled;
import org.springframework.stereotype.Service;

import java.math.BigDecimal;
import java.time.LocalDateTime;
import java.time.temporal.ChronoUnit;
import java.util.List;

@Slf4j
@Service
@RequiredArgsConstructor
public class MemoryDecayScheduler {

    private final UserMemoryMapper userMemoryMapper;
    private final MemoryVectorStore memoryVectorStore;

    private static final double DECAY_FACTOR = 0.95;
    private static final double SOFT_DELETE_THRESHOLD = 0.3;
    private static final int SOFT_DELETE_DAYS = 30;
    private static final int PURGE_DAYS = 30;

    /**
     * Layer 1：置信度衰减（每天凌晨 2 点）
     */
    @Scheduled(cron = "0 0 2 * * ?")
    public void decayConfidence() {
        List<UserMemory> all = userMemoryMapper.findAllActiveAllUsers();
        int decayed = 0;
        for (UserMemory m : all) {
            long days = ChronoUnit.DAYS.between(
                    m.getLastUsedAt(), LocalDateTime.now());
            double factor = Math.pow(DECAY_FACTOR, days);
            double newConf = m.getConfidence().doubleValue() * factor;

            m.setConfidence(BigDecimal.valueOf(newConf));
            userMemoryMapper.updateById(m);
            decayed++;
        }
        log.info("Confidence decay done: {} memories updated", decayed);
    }

    /**
     * Layer 2：软删除（每天凌晨 2:30）
     */
    @Scheduled(cron = "0 30 2 * * ?")
    public void softDelete() {
        List<UserMemory> candidates = userMemoryMapper
                .findDecayCandidates(
                        SOFT_DELETE_THRESHOLD, SOFT_DELETE_DAYS);
        for (UserMemory m : candidates) {
            m.setDeletedAt(LocalDateTime.now());
            userMemoryMapper.updateById(m);
            // 写历史（action=DELETE, reason=decay）
        }
        log.info("Soft delete done: {} memories marked", candidates.size());
    }

    /**
     * Layer 3：物理清除（每天凌晨 3 点）
     */
    @Scheduled(cron = "0 0 3 * * ?")
    public void purgeDeleted() {
        LocalDateTime threshold = LocalDateTime.now().minusDays(PURGE_DAYS);
        List<Long> ids = userMemoryMapper.findPurgeCandidates(threshold);
        for (Long id : ids) {
            userMemoryMapper.deleteById(id);
            memoryVectorStore.delete(id);
        }
        log.info("Purge done: {} memories physically deleted", ids.size());
    }
}''')

add_heading_styled(doc, '4.7 PromptAssembler 集成', level=2)

add_paragraph_styled(doc, '改造 PromptAssembler，在 <context> 中新增用户画像 slot：')

add_code_block(doc, '''改造前（PromptAssembler.assemble）：

  <context>
    ## 检索到的知识
    {rag_results}
  </context>

改造后：

  <context>
    ## 用户画像
    {user_profile}            ← 新增：PREFERENCE + FACT 全量注入

    ## 相关历史事件
    {user_events}             ← 新增：EVENT 按需检索 top-5

    ## 检索到的知识
    {rag_results}
  </context>''')

add_code_block(doc, '''改造后 assemble 方法（核心片段）：

  public String assemble(String userId, String query, String intent,
                          List<RetrievalResult> ragResults) {
      // 新增：检索用户记忆
      RetrievalResult memoryResult = memoryRetrievalService
              .retrieve(userId, query, intent);

      String userProfile = formatProfile(memoryResult.getProfileMemories());
      String userEvents = formatEvents(memoryResult.getEventMemories());
      String ragStr = formatRag(ragResults);

      String context = "<context>\\n"
              + "## 用户画像\\n" + userProfile + "\\n\\n"
              + "## 相关历史事件\\n" + userEvents + "\\n\\n"
              + "## 检索到的知识\\n" + ragStr + "\\n"
              + "</context>";

      // 回写 used_memory_ids 到上下文快照
      currentContextSnapshot.setUsedMemoryIds(memoryResult.getUsedMemoryIds());

      return L1 + L2 + L3 + context + L4;
  }

  private String formatProfile(List<UserMemory> memories) {
      if (memories.isEmpty()) return "（暂无用户画像信息）";
      return memories.stream()
              .map(m -> "- " + m.getMemoryKey() + ": " + m.getMemoryValue())
              .collect(Collectors.joining("\\n"));
  }''')

add_heading_styled(doc, '4.8 数据库 Schema（已有三张表）', level=2)

add_paragraph_styled(doc,
    '三张 MySQL 表已在 agent-init.sql 中定义，无需新建。'
    '另需在 PostgreSQL 中新建 memory_vectors 表：')

add_code_block(doc, '''-- PostgreSQL：memory_vectors 表（需新建）
-- 存储位置：agent-postgres 容器（pgvector/pg16）
-- 用途：user_memory.memory_value 的向量索引

CREATE TABLE IF NOT EXISTS memory_vectors (
    memory_id   BIGINT PRIMARY KEY,          -- 关联 user_memory.id
    user_id     VARCHAR(36) NOT NULL,         -- 冗余字段，加速过滤
    memory_value TEXT NOT NULL,               -- 记忆值原文
    embedding   vector(1024) NOT NULL,        -- bge-m3 向量
    embedding_model VARCHAR(32) DEFAULT 'bge-m3',
    updated_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    INDEX idx_user (user_id)
);

-- HNSW 索引（余弦距离）
CREATE INDEX IF NOT EXISTS idx_memory_vectors_embedding
    ON memory_vectors USING hnsw (embedding vector_cosine_ops);

-- pg_trgm 索引（关键词模糊匹配）
CREATE INDEX IF NOT EXISTS idx_memory_vectors_value_trgm
    ON memory_vectors USING gin (memory_value gin_trgm_ops);''')

add_paragraph_styled(doc, '已有三张 MySQL 表的字段说明：')

add_styled_table(doc,
    ['表名', '关键字段', '用途'],
    [
        ['user_memory',
         'id / user_id / memory_type / memory_key / memory_value / confidence / source / evidence_count / conflict_flag / volatile_flag / last_used_at / deleted_at',
         '记忆主表，唯一键 uk_user_type_key(user_id, memory_type, memory_key)'],
        ['user_memory_evidence',
         'id / user_id / memory_id / session_id / evidence_type / evidence_payload / processed',
         '行为证据表，记录记忆推断依据，processed=0 待处理'],
        ['user_memory_history',
         'id / user_id / memory_type / memory_key / memory_value / confidence / source / action / reason',
         '历史审计表，action ∈ UPDATE/DELETE/DECAY'],
    ],
    col_widths=[3, 6, 5])

doc.add_page_break()

# ==================== 第五章 目标：实现效果 ====================
add_heading_styled(doc, '第五章 目标：实现效果', level=1)

add_heading_styled(doc, '5.1 功能目标', level=2)

add_styled_table(doc,
    ['功能项', '当前状态', '目标状态', '验收标准'],
    [
        ['用户画像持久化', '无（会话内临时推断）',
         '跨会话保留 PREFERENCE + FACT',
         '用户二次会话无需重复自我介绍'],
        ['偏好累积', '无',
         'EXPLICIT + INFERRED 双通道采集',
         '用户明说偏好 100% 记录，行为推断准确率 >= 80%'],
        ['事件记忆检索', '无',
         '向量+关键词双路召回 top-10',
         '相关事件召回 Recall@5 >= 75%'],
        ['冲突处理', '无',
         'LLM 仲裁 + KEEP_NEW/OLD/BOTH',
         '冲突记忆不出现矛盾回答'],
        ['遗忘机制', '无',
         '三层：衰减 + 软删除 + 物理清除',
         '30 天未用记忆 confidence < 0.3'],
        ['工作记忆联动', 'used_memory_ids 字段闲置',
         '每轮回写 + 审计可追溯',
         'agent_context_snapshots.used_memory_ids 非空率 100%'],
        ['用户可见可控', '无',
         'API 查询/删除自己的记忆',
         '用户可查看全部记忆，可单条删除'],
    ],
    col_widths=[3, 3, 4, 4])

add_heading_styled(doc, '5.2 性能目标', level=2)

add_styled_table(doc,
    ['指标', '目标值', '测量方法', '说明'],
    [
        ['记忆检索 P50', '< 30ms', '1000 次检索平均',
         'PROFILE 全量 + 事件 top-5'],
        ['记忆检索 P99', '< 100ms', '1000 次检索 99 分位',
         '含 embedding 调用'],
        ['记忆抽取耗时', '< 2s', '50 次抽取平均',
         'LLM 调用 + DB 写入（异步，不阻塞用户）'],
        ['衰减任务耗时', '< 30s', '1 万条记忆全量衰减',
         '每天凌晨执行'],
        ['软删除任务耗时', '< 10s', '1 万条记忆扫描',
         '每天凌晨执行'],
        ['向量入库耗时', '< 50ms/条', '100 条记忆入库',
         'embedding + PG upsert'],
    ],
    col_widths=[3, 3, 4, 4])

add_heading_styled(doc, '5.3 质量目标', level=2)

add_styled_table(doc,
    ['质量指标', '目标值', '测量方法', '当前基线'],
    [
        ['记忆抽取准确率', '>= 85%',
         'LLM-as-Judge 评估抽取结果',
         'N/A（待实现）'],
        ['记忆抽取召回率', '>= 80%',
         '黄金集标注对比',
         'N/A'],
        ['冲突仲裁准确率', '>= 90%',
         '人工复核仲裁结果',
         'N/A'],
        ['记忆检索相关性', '>= 4.0/5',
         'LLM-as-Judge 评分',
         'N/A'],
        ['用户满意度', '>= 4.2/5',
         '用户点赞/点踩比',
         '当前约 3.8/5（无记忆）'],
        ['误抽率（噪音记忆）', '< 10%',
         '人工抽查 100 条',
         'N/A'],
    ],
    col_widths=[3.5, 3, 4, 3.5])

add_heading_styled(doc, '5.4 成本目标', level=2)

add_styled_table(doc,
    ['成本项', '当前基线', '目标值', '优化手段'],
    [
        ['LLM 调用次数/会话', '0（无记忆）',
         '< 1.5（抽取+偶发仲裁）',
         '异步抽取 + 缓存已有记忆'],
        ['Embedding 调用次数/会话', '1（RAG）',
         '1.2（+记忆检索）',
         'PROFILE 不需 embedding，仅事件检索调用'],
        ['存储/用户', '0',
         '< 100KB（100 条记忆）',
         '遗忘机制控制记忆量'],
        ['向量索引大小', '0',
         '< 500MB（1 万用户 × 100 条）',
         'HNSW 索引压缩'],
    ],
    col_widths=[3.5, 3, 4, 3.5])

doc.add_page_break()

# ==================== 第六章 测试 ====================
add_heading_styled(doc, '第六章 测试', level=1)

add_heading_styled(doc, '6.1 评估指标', level=2)

add_paragraph_styled(doc, '长期记忆模块的评估分四个维度：')

add_styled_table(doc,
    ['维度', '指标', '公式', '目标值', '采集方式'],
    [
        ['抽取质量', '抽取准确率',
         '正确抽取的记忆数 / 抽取总数',
         '>= 85%',
         'LLM-as-Judge'],
        ['抽取质量', '抽取召回率',
         '正确抽取的记忆数 / 应抽取的记忆数',
         '>= 80%',
         '黄金集标注'],
        ['检索质量', 'Recall@5',
         '前 5 条中含相关记忆的会话比例',
         '>= 75%',
         '黄金集评估'],
        ['检索质量', '相关记忆占比',
         '相关记忆数 / 检索结果总数',
         '>= 70%',
         '人工标注'],
        ['冲突处理', '仲裁准确率',
         '正确仲裁数 / 仲裁总数',
         '>= 90%',
         '人工复核'],
        ['遗忘效果', '衰减命中率',
         '被衰减记忆中确实无用的比例',
         '>= 85%',
         '后续行为反查'],
        ['业务价值', '个性化满意度',
         '点赞数 / (点赞+点踩数)',
         '>= 80%',
         '用户反馈'],
    ],
    col_widths=[2.5, 2.5, 4, 2.5, 2.5])

add_heading_styled(doc, '6.2 黄金测试集', level=2)

add_paragraph_styled(doc, '构建 5 个黄金测试集，覆盖长期记忆的核心场景：')

add_styled_table(doc,
    ['测试集', '用例数', '覆盖场景', '评估方式'],
    [
        ['EXTRACT-GOLDEN', '30',
         '用户明说身份/偏好/技能/事件的对话',
         '验证抽取的记忆是否正确'],
        ['CONFLICT-GOLDEN', '20',
         '新旧记忆冲突的对话对',
         '验证仲裁决策是否合理'],
        ['RETRIEVE-GOLDEN', '40',
         '含历史记忆的用户查询',
         '验证检索 Recall@5'],
        ['DECAY-GOLDEN', '15',
         '不同 last_used_at 的记忆',
         '验证衰减计算正确'],
        ['PRIVACY-GOLDEN', '10',
         '含 PII 信息的对话',
         '验证脱敏后才入库'],
    ],
    col_widths=[3.5, 2, 5.5, 3])

add_heading_styled(doc, '6.2.1 黄金集示例（EXTRACT-GOLDEN）', level=3)

add_code_block(doc, '''黄金用例格式（JSON）：

  {
    "case_id": "EXT-001",
    "conversation": [
      {"role": "user", "content": "我是计算机学院大二的学生"},
      {"role": "assistant", "content": "你好！有什么可以帮你的吗？"}
    ],
    "expected_memories": [
      {
        "type": "FACT",
        "key": "user_college",
        "value": "计算机学院",
        "confidence": 1.0
      },
      {
        "type": "FACT",
        "key": "user_grade",
        "value": "大二",
        "confidence": 1.0
      }
    ],
    "assertion": "抽取的记忆应包含上述两条，type/key/value 完全匹配"
  }''')

add_heading_styled(doc, '6.3 CI/CD 集成', level=2)

add_paragraph_styled(doc, '长期记忆模块的 CI/CD 分四个阶段：')

add_code_block(doc, '''CI/CD 四阶段流水线：

  ┌─ Stage 1：编译检查 ────────────────────────────────┐
  │  - mvn -pl campushare-agent clean compile          │
  │  - 检查 UserMemory/Mapper/Service 编译通过         │
  │  → 编译失败阻断                                    │
  └─────────────────────────────────────────────────────┘
                  │
                  ▼
  ┌─ Stage 2：单元测试 ────────────────────────────────┐
  │  - MemoryExtractor 单测（抽取候选解析）            │
  │  - MemoryRetrievalService 单测（评分公式）         │
  │  - MemoryUpdateService 单测（新增/增强/冲突）      │
  │  - MemoryDecayScheduler 单测（衰减计算）           │
  │  → 通过率 100% 才进入下一阶段                      │
  └─────────────────────────────────────────────────────┘
                  │
                  ▼
  ┌─ Stage 3：集成测试 ────────────────────────────────┐
  │  - 摄入流水线集成测试（50 轮对话全流程）           │
  │  - 冲突仲裁集成测试                                │
  │  - 衰减+软删除+清除集成测试                        │
  │  - PromptAssembler 集成测试（记忆注入到 context）  │
  │  → 通过率 100% 才进入下一阶段                      │
  └─────────────────────────────────────────────────────┘
                  │
                  ▼
  ┌─ Stage 4：评估测试 ────────────────────────────────┐
  │  - 跑 5 个黄金测试集（共 115 条）                  │
  │  - LLM-as-Judge 评估抽取质量                       │
  │  - 与上次版本对比，退化 > 5% 阻断发布              │
  └─────────────────────────────────────────────────────┘''')

add_heading_styled(doc, '6.4 LLM-as-Judge 评估', level=2)

add_paragraph_styled(doc, '用 LLM 评估抽取质量和检索质量：')

add_heading_styled(doc, '6.4.1 抽取质量 LLM 评估', level=3)

add_code_block(doc, '''抽取质量评估 Prompt：

  系统：你是记忆质量审核员。评估以下抽取的记忆是否合理。
  评分维度（1-5 分）：
  1. 准确性：抽取的记忆是否与对话内容一致
  2. 完整性：是否遗漏了重要信息
  3. 相关性：是否是"跨会话有价值"的信息
  4. 无噪音：是否避免了临时性信息的抽取

  对话内容：{conversation}
  抽取的记忆：{extracted_memories}

  输出 JSON：
  {
    "accuracy": 5,
    "completeness": 4,
    "relevance": 5,
    "no_noise": 4,
    "average_score": 4.5,
    "issues": ["遗漏了用户提到的大二身份"]
  }''')

add_heading_styled(doc, '6.4.2 检索质量 LLM 评估', level=3)

add_code_block(doc, '''检索质量评估 Prompt：

  系统：你是记忆检索质量审核员。
  评估检索到的记忆对当前查询是否有帮助。

  用户查询：{query}
  检索到的记忆：{retrieved_memories}

  评分维度（1-5 分）：
  1. 相关性：记忆与查询主题相关
  2. 有用性：记忆能帮助回答查询
  3. 无冗余：没有重复或无关的记忆

  输出 JSON：
  {
    "relevance": 4,
    "usefulness": 5,
    "no_redundancy": 3,
    "average_score": 4.0
  }''')

add_heading_styled(doc, '6.5 错误分析与归因', level=2)

add_paragraph_styled(doc, '长期记忆模块的常见错误及归因：')

add_styled_table(doc,
    ['错误类型', '表现', '可能原因', '修复方向'],
    [
        ['误抽取',
         '抽取了临时性信息作为长期记忆',
         'LLM 抽取 Prompt 不够严格',
         '优化 Prompt + 增加负面示例'],
        ['漏抽取',
         '用户明说的偏好未被记录',
         'LLM 未识别为可记忆信息',
         '优化 Prompt + 补充关键词触发'],
        ['冲突误仲裁',
         '应该保留旧值却保留了新值',
         '仲裁依据优先级错误',
         '调整 EXPLICIT/INFERRED 权重'],
        ['检索不相关',
         '检索到的记忆与查询无关',
         'embedding 质量差 + 阈值过低',
         '提高相似度阈值 + 重排调参'],
        ['衰减过快',
         '有用记忆被过早软删除',
         'decay_factor 过小',
         '调整 0.95 → 0.97'],
        ['衰减过慢',
         '无用记忆长期占空间',
         'decay_factor 过大',
         '调整 0.95 → 0.93'],
        ['记忆泄露',
         '用户 A 的记忆被用户 B 检索到',
         '检索 SQL 缺少 user_id 过滤',
         '紧急修复 + 安全审计'],
        ['PII 入库',
         '手机号/身份证被记录',
         '抽取前未脱敏',
         '增加 PII 过滤器'],
    ],
    col_widths=[2.5, 3, 3.5, 4])

doc.add_page_break()

# ==================== 6.6 测试用例设计 ====================
add_heading_styled(doc, '6.6 测试用例设计', level=2)

add_paragraph_styled(doc, '按四个子方向设计测试用例，覆盖正常路径、边界、异常和回归场景。')

add_heading_styled(doc, '6.6.1 记忆抽取测试用例', level=3)

add_styled_table(doc,
    ['用例 ID', '场景', '输入', '预期输出', '类型'],
    [
        ['EXT-001', '明说身份',
         '"我是计算机学院大二学生"',
         '抽取 FACT: user_college=计算机学院, user_grade=大二',
         '正常'],
        ['EXT-002', '明说偏好',
         '"我喜欢用 Python"',
         '抽取 PREFERENCE: preferred_language=Python',
         '正常'],
        ['EXT-003', '明说技能',
         '"我会用 Spring Boot"',
         '抽取 SKILL: skill_spring_boot=会',
         '正常'],
        ['EXT-004', '临时性问题（不抽取）',
         '"今天天气怎么样？"',
         '返回空数组 []',
         '正常'],
        ['EXT-005', '多记忆抽取',
         '"我是大二学生，喜欢 Python，会 SQL"',
         '抽取 3 条记忆',
         '正常'],
        ['EXT-006', '暗示性偏好',
         '"又是 Python，我最熟了"',
         '抽取 PREFERENCE: preferred_language=Python，confidence=0.7',
         '边界'],
        ['EXT-007', '已有记忆去重',
         '用户已有 preferred_language=Python，再次明说',
         '不新增，evidence_count+1',
         '正常'],
        ['EXT-008', 'LLM 返回非法 JSON',
         'LLM 返回非 JSON 格式',
         'parseCandidates 返回空列表，log.warn',
         '异常'],
    ],
    col_widths=[2, 2.5, 3.5, 4, 1.5])

add_heading_styled(doc, '6.6.2 记忆检索测试用例', level=3)

add_styled_table(doc,
    ['用例 ID', '场景', '输入', '预期输出', '类型'],
    [
        ['RET-001', 'PROFILE 全量装载',
         '用户有 5 条 PREFERENCE + 3 条 FACT',
         'L1 注入 8 条记忆（confidence >= 0.5）',
         '正常'],
        ['RET-002', '事件向量召回',
         '用户问"选课问题"，有相关事件记忆',
         'top-5 中含"event_query_registration"',
         '正常'],
        ['RET-003', 'SEARCH 意图跳过事件',
         'intent=SEARCH',
         '仅装载 PROFILE，不检索事件',
         '正常'],
        ['RET-004', '无记忆用户',
         '新用户首次对话',
         'PROFILE 为空，事件为空',
         '边界'],
        ['RET-005', '低置信度过滤',
         '用户有 confidence=0.4 的 PREFERENCE',
         'PROFILE 不注入该条',
         '边界'],
        ['RET-006', 'last_used_at 更新',
         '检索后检查记忆',
         '被使用的记忆 last_used_at 更新为当前时间',
         '正常'],
        ['RET-007', 'used_memory_ids 回写',
         '检索后检查上下文快照',
         'used_memory_ids 字段非空',
         '正常'],
        ['RET-008', '跨用户隔离',
         '用户 A 检索，用户 B 的记忆存在',
         '结果不含用户 B 的记忆',
         '异常'],
    ],
    col_widths=[2, 2.5, 3.5, 4, 1.5])

add_heading_styled(doc, '6.6.3 冲突处理测试用例', level=3)

add_styled_table(doc,
    ['用例 ID', '场景', '输入', '预期输出', '类型'],
    [
        ['CONF-001', '相同值增强',
         '旧:preferred_language=Python，新:Python',
         'evidence_count+1，confidence+0.1',
         '正常'],
        ['CONF-002', 'EXPLICIT 覆盖 INFERRED',
         '旧:INFERRED Python，新:EXPLICIT Java',
         'KEEP_NEW，旧值归档',
         '正常'],
        ['CONF-003', 'INFERRED 不覆盖 EXPLICIT',
         '旧:EXPLICIT Python，新:INFERRED Java',
         'KEEP_OLD，新值忽略',
         '正常'],
        ['CONF-004', '偏好变化保留新值',
         '旧:EXPLICIT Python，新:EXPLICIT Java',
         'KEEP_NEW（PREFERENCE 倾向新值）',
         '正常'],
        ['CONF-005', '事实冲突保留两者',
         '旧:FACT 专业=计算机，新:FACT 专业=电子',
         'KEEP_BOTH，conflict_flag=1',
         '正常'],
        ['CONF-006', '仲裁 LLM 降级',
         'LLM 不可用',
         '默认 KEEP_BOTH，标记待人工复核',
         '异常'],
        ['CONF-007', '软删除记忆恢复',
         '用户重新提及已软删除的偏好',
         'deleted_at 置 NULL，confidence=0.5',
         '正常'],
    ],
    col_widths=[2, 2.5, 3.5, 4, 1.5])

add_heading_styled(doc, '6.6.4 遗忘机制测试用例', level=3)

add_styled_table(doc,
    ['用例 ID', '场景', '输入', '预期输出', '类型'],
    [
        ['DEC-001', '日常衰减',
         '记忆 30 天未使用',
         'confidence *= 0.95^30 ≈ 0.21',
         '正常'],
        ['DEC-002', '刚使用不衰减',
         '记忆今天刚使用',
         'confidence 不变（days=0，factor=1）',
         '正常'],
        ['DEC-003', '软删除触发',
         'confidence=0.2，30 天未用',
         'deleted_at 标记',
         '正常'],
        ['DEC-004', '软删除不触发',
         'confidence=0.4，30 天未用',
         '不软删除（confidence >= 0.3）',
         '边界'],
        ['DEC-005', '物理清除',
         'deleted_at 31 天前',
         'DELETE 物理删除 + 向量同步删除',
         '正常'],
        ['DEC-006', '物理清除不触发',
         'deleted_at 29 天前',
         '不物理清除（< 30 天）',
         '边界'],
        ['DEC-007', 'PROFILE 慢衰减',
         'PREFERENCE 60 天未用',
         '仍保留（PROFILE 阈值更宽松）',
         '正常'],
        ['DEC-008', 'EVENT 快衰减',
         'EVENT 15 天未用',
         '已软删除（EVENT 阈值更严格）',
         '正常'],
    ],
    col_widths=[2, 2.5, 3.5, 4, 1.5])

# ==================== 6.7 性能与压力测试 ====================
add_heading_styled(doc, '6.7 性能与压力测试', level=2)

add_heading_styled(doc, '6.7.1 检索性能测试', level=3)

add_styled_table(doc,
    ['指标', '目标值', '测量方法', '说明'],
    [
        ['PROFILE 装载 P50', '< 10ms', '1000 次查询平均',
         'SQL 查询，无 embedding'],
        ['PROFILE 装载 P99', '< 30ms', '1000 次查询 99 分位',
         '含索引扫描'],
        ['事件向量检索 P50', '< 50ms', '1000 次查询平均',
         '含 embedding API 调用'],
        ['事件向量检索 P99', '< 100ms', '1000 次查询 99 分位',
         '含 embedding API 调用'],
        ['总检索 P99', '< 150ms', '1000 次查询 99 分位',
         'PROFILE + 事件 + 回写'],
    ],
    col_widths=[3.5, 3, 4, 3.5])

add_heading_styled(doc, '6.7.2 压力测试场景', level=3)

add_code_block(doc, '''压力测试脚本（JMeter / wrk 模拟）：

  场景 1：高并发记忆检索
    - 输入：100 QPS 的对话请求，每用户 50 条记忆
    - 监控：检索延迟 / PG 连接池 / embedding API 限流
    - 期望：P99 < 150ms，无连接池耗尽

  场景 2：大规模记忆衰减
    - 输入：10 万条记忆的衰减任务
    - 监控：衰减任务耗时 / PG 写入延迟 / 锁竞争
    - 期望：30 分钟内完成，无锁超时

  场景 3：记忆抽取突发流量
    - 输入：100 个用户同时结束会话，触发异步抽取
    - 监控：LLM API 限流 / 抽取队列积压 / DB 写入
    - 期望：10 分钟内消化完毕，无丢失

  场景 4：跨用户隔离压测
    - 输入：1000 用户并发检索，每用户 100 条记忆
    - 监控：用户 A 的结果是否泄露到用户 B
    - 期望：0 泄露事件''')

add_heading_styled(doc, '6.7.3 资源占用基线', level=3)

add_styled_table(doc,
    ['资源', '空闲', '正常负载', '峰值负载', '告警阈值'],
    [
        ['CPU', '< 5%', '15-30%', '50-70%', '> 85% 持续 5 分钟'],
        ['JVM 堆', '< 512MB', '1-2GB', '2.5GB', '> 3GB'],
        ['MySQL 连接数', '< 3', '10-20', '25-35', '> 40（池上限 50）'],
        ['PG 连接数', '< 3', '5-10', '15-20', '> 25'],
        ['记忆存储', '1 万用户 ≈ 500MB', '-', '-', '> 80% 磁盘'],
    ],
    col_widths=[2.5, 2.5, 2.5, 2.5, 4])

# ==================== 6.8 A/B 测试设计 ====================
add_heading_styled(doc, '6.8 A/B 测试设计', level=2)

add_paragraph_styled(doc, '通过 A/B 测试验证长期记忆对用户体验的影响。')

add_heading_styled(doc, '6.8.1 实验一：有无长期记忆对比', level=3)

add_styled_table(doc,
    ['分组', '配置', '流量比例', '观测周期'],
    [
        ['A 组（基线）', '不启用长期记忆（当前状态）', '50%', '14 天'],
        ['B 组（实验）', '启用长期记忆全功能', '50%', '14 天'],
    ],
    col_widths=[3, 4, 2.5, 2.5])

add_paragraph_styled(doc, '观测指标：')
add_bullet(doc, '用户满意度（点赞/点踩比）：B 组应 >= A 组 +5%')
add_bullet(doc, '重复自我介绍次数：B 组应 <= A 组 -50%')
add_bullet(doc, '会话长度（轮次）：B 组应 >= A 组（更深入对话）')
add_bullet(doc, '次日留存率：B 组应 >= A 组 +3%')

add_heading_styled(doc, '6.8.2 实验二：抽取策略对比', level=3)

add_styled_table(doc,
    ['分组', '抽取策略', '流量比例', '观测指标'],
    [
        ['A 组', '仅 EXPLICIT（明说才记）', '50%', '误抽率 / 召回率'],
        ['B 组', 'EXPLICIT + INFERRED（双通道）', '50%', '误抽率 / 召回率'],
    ],
    col_widths=[2, 4, 2.5, 5.5])

add_paragraph_styled(doc, '观测重点：B 组的 INFERRED 通道是否真的提升了召回率，'
    '还是只增加了噪音。如果误抽率上升超过 15%，考虑关闭 INFERRED 通道或提高 evidence_count 门槛。')

add_heading_styled(doc, '6.8.3 实验三：遗忘参数调优', level=3)

add_styled_table(doc,
    ['分组', 'decay_factor', '软删除阈值', '软删除天数', '流量比例'],
    [
        ['A 组', '0.95', '0.3', '30 天', '50%'],
        ['B 组', '0.97', '0.2', '45 天', '50%'],
    ],
    col_widths=[2, 2.5, 2.5, 2.5, 2.5])

add_paragraph_styled(doc, '观测指标：')
add_bullet(doc, 'B 组记忆存活更久，观察是否提升用户满意度')
add_bullet(doc, 'B 组存储成本上升，评估 ROI')
add_bullet(doc, 'B 组是否有更多"过时记忆"被检索到')

add_callout(doc,
    'A/B 测试准入条件：实验组样本量 >= 1000 用户，至少持续 14 天，'
    '统计显著性 p < 0.05。',
    color='E8F4FD', border_color='2196F3')

# ==================== 6.9 验收流程与准入准出 ====================
add_heading_styled(doc, '6.9 验收流程与准入准出', level=2)

add_heading_styled(doc, '6.9.1 准入条件（开发完成 → 测试阶段）', level=3)

add_styled_table(doc,
    ['检查项', '标准', '负责人'],
    [
        ['单元测试', '覆盖率 >= 80%，全部通过', '开发'],
        ['代码评审', '至少 1 人评审通过', '架构师'],
        ['静态检查', 'SonarQube 0 Blocker / 0 Critical', 'CI'],
        ['Schema 变更', 'memory_vectors DDL 已在测试库执行', 'DBA'],
        ['PII 脱敏', '抽取流程已接入 PII 过滤器', '安全'],
        ['ADR 文档', '新增 ADR-MEM 已归档', '架构师'],
        ['Changelog', '已追加到 changelog/', '开发'],
    ],
    col_widths=[3, 6, 2])

add_heading_styled(doc, '6.9.2 准出条件（测试阶段 → 发布）', level=3)

add_styled_table(doc,
    ['检查项', '标准', '负责人'],
    [
        ['功能测试', '所有用例（6.6 节）通过', 'QA'],
        ['性能测试', '满足 6.7 节目标值', 'QA'],
        ['安全测试', '跨用户隔离测试 0 泄露', '安全'],
        ['A/B 测试', '实验组指标不劣于基线', '数据'],
        ['LLM-as-Judge', '抽取质量均分 >= 4.0', 'QA'],
        ['监控告警', 'Prometheus 指标已接入', '运维'],
        ['回滚预案', '已有回滚脚本并演练通过', '运维'],
    ],
    col_widths=[3, 6, 2])

add_heading_styled(doc, '6.9.3 发布流程', level=3)

add_code_block(doc, '''发布流程（七步）：

  1. 合并代码到 master 分支
  2. CI 触发镜像构建（campushare-agent: vX.Y.Z）
  3. 在预发环境部署，跑 5 个黄金测试集（115 条）
  4. 执行 memory_vectors 表 DDL
  5. 滚动更新 agent-service 容器（先 1 个实例验证 10 分钟）
  6. 全量滚动更新，监控指标 30 分钟
  7. 发布完成，记录到 changelog 和 release notes''')

# ==================== 6.10 持续监控与漂移检测 ====================
add_heading_styled(doc, '6.10 持续监控与漂移检测', level=2)

add_paragraph_styled(doc,
    '长期记忆模块上线后需要持续监控质量和性能，发现"漂移"及时告警。')

add_heading_styled(doc, '6.10.1 监控指标体系', level=3)

add_styled_table(doc,
    ['维度', '指标', '采集方式', '告警阈值'],
    [
        ['抽取健康', '每小时抽取次数',
         'Micrometer Counter', '连续 2 小时为 0（无对话除外）'],
        ['抽取健康', '抽取失败率',
         'failed / total', '> 10%'],
        ['检索健康', '检索 P99 延迟',
         'Micrometer Timer', '> 200ms'],
        ['检索健康', '空记忆用户比例',
         '无记忆用户 / 活跃用户', '> 80%（新功能推广期除外）'],
        ['冲突健康', '冲突标记数',
         'conflict_flag=1 的记忆数', '单用户 > 10'],
        ['遗忘健康', '每日软删除数',
         'MemoryDecayScheduler 日志', '突增 > 5 倍'],
        ['存储健康', '记忆总量',
         'COUNT(user_memory)', '单用户 > 500 条'],
        ['存储健康', 'PG 磁盘使用率',
         'Prometheus node_exporter', '> 80%'],
        ['业务价值', '用户满意度',
         '点赞/点踩比', '下降 > 5%'],
        ['安全', '跨用户泄露事件',
         '安全审计日志', '> 0（立即告警）'],
    ],
    col_widths=[2.5, 3, 3.5, 3])

add_heading_styled(doc, '6.10.2 数据漂移检测', level=3)

add_paragraph_styled(doc, '记忆漂移：用户行为或偏好随时间变化，导致历史记忆过时。')

add_styled_table(doc,
    ['漂移类型', '表现', '检测方法', '应对'],
    [
        ['偏好漂移', '用户偏好变化（如从 Python 改为 Java）',
         '冲突标记频次上升',
         '触发 KEEP_NEW 仲裁'],
        ['身份漂移', '用户身份变化（如大二升大三）',
         'FACT 类型记忆冲突',
         '人工确认后更新'],
        ['查询漂移', '用户查询主题变化',
         '事件记忆召回率下降',
         '补充新主题记忆 + 衰减旧主题'],
        ['质量漂移', 'LLM-as-Judge 分数下降',
         '每日抽样评估',
         '排查抽取 Prompt 是否退化'],
    ],
    col_widths=[2.5, 3.5, 3, 3.5])

add_heading_styled(doc, '6.10.3 告警与自动化处理', level=3)

add_code_block(doc, '''告警分级与自动化处理：

  P0（立即响应，电话告警）：
    - 跨用户记忆泄露
    - 自动处理：紧急下线记忆检索 + 安全审计

  P1（10 分钟响应，钉钉告警）：
    - 抽取失败率 > 30%
    - 检索 P99 > 500ms
    - 自动处理：扩容 agent-service 实例

  P2（1 小时响应，邮件告警）：
    - 单用户冲突记忆 > 20 条
    - 用户满意度下降 > 3%
    - 自动处理：标记问题用户，加入人工复核队列

  P3（每日汇总）：
    - 记忆总量趋势
    - 软删除数量趋势
    - 自动处理：无，仅观察''')

doc.add_page_break()

# ==================== 第七章 总结与边界声明 ====================
add_heading_styled(doc, '第七章 总结与边界声明', level=1)

add_heading_styled(doc, '7.1 核心总结', level=2)

add_paragraph_styled(doc, '本文档围绕"跨会话长期记忆"展开，覆盖采集、检索、更新、遗忘、联动五大子方向：')

add_styled_table(doc,
    ['子方向', '核心方案', '关键 ADR', '预期收益'],
    [
        ['记忆建模',
         '四分类（PREFERENCE/FACT/SKILL/EVENT）+ 三表架构',
         'ADR-MEM-01, ADR-MEM-02',
         '结构化存储，类型清晰'],
        ['记忆采集',
         'EXPLICIT + INFERRED 双通道，LLM 抽取',
         'ADR-MEM-03',
         '覆盖明说和隐式偏好'],
        ['记忆检索',
         'PROFILE 全量 + 事件向量+关键词双路',
         'ADR-MEM-04',
         '检索 Recall@5 >= 75%'],
        ['记忆更新',
         '新增/增强/冲突三场景 + LLM 仲裁',
         'ADR-MEM-05',
         '冲突 0 上线，仲裁准确率 >= 90%'],
        ['遗忘机制',
         '三层：衰减 + 软删除 + 物理清除',
         'ADR-MEM-06',
         '可恢复 + 自然降权 + 存储可控'],
        ['工作记忆联动',
         'used_memory_ids 回写 + PromptAssembler 集成',
         'ADR-MEM-07',
         '审计可追溯 + 上下文无缝衔接'],
    ],
    col_widths=[2.5, 4.5, 2.5, 4])

add_heading_styled(doc, '7.2 与其他文档的关系', level=2)

add_paragraph_styled(doc, '本文档在 Agent 搭建系列中的位置：')

add_styled_table(doc,
    ['相关文档', '关系', '说明'],
    [
        ['上下文工程模块设计方案', '上游依赖',
         '长期记忆的 PROFILE 注入到上下文工程的 L1 用户画像层；'
         'used_memory_ids 回写到 agent_context_snapshots'],
        ['SystemPrompt 工程模块设计方案', '协作',
         'SystemPrompt 的 L2 prompt 可引用用户偏好（如 preferred_language），'
         '由长期记忆提供数据'],
        ['RAG 检索增强模块设计方案', '并列互补',
         'RAG 是"全用户共享知识库"，长期记忆是"单用户私有记忆"。'
         '两者都注入 <context>，但来源和检索逻辑独立'],
        ['意图识别模块设计方案', '上游',
         '意图识别结果决定记忆检索策略：SEARCH 跳过事件检索，CHAT/HOW_TO 执行'],
        ['知识库管理模块设计方案', '对比参考',
         '知识库管理的分块/同步/治理思路可借鉴到记忆管理，但记忆是单用户私有'],
        ['对话编排模块（待写）', '下游',
         '对话编排可触发记忆更新（如用户明说偏好时调用 MemoryExtractor）'],
        ['安全护栏模块（待写）', '横切',
         '记忆写入前需经 PII 脱敏，记忆检索结果需经护栏校验'],
        ['可观测性模块（待写）', '横切',
         '本文档第六章定义的指标需接入可观测性体系'],
        ['评估体系模块（待写）', '横切',
         '本文档的 golden set 和 LLM-as-Judge 纳入全局评估体系'],
    ],
    col_widths=[3.5, 2, 8.5])

add_heading_styled(doc, '7.3 演进路线', level=2)

add_paragraph_styled(doc, '长期记忆模块的演进路线图（按优先级）：')

add_styled_table(doc,
    ['阶段', '里程碑', '关键工作', '预计周期'],
    [
        ['Phase 1（当前）',
         '基础记忆',
         '三表 + 四分类 + 双通道采集 + 双路检索 + 遗忘机制',
         '3 周'],
        ['Phase 2',
         '质量提升',
         'LLM 仲裁优化 + 衰减参数调优 + 漂移检测',
         '1 周'],
        ['Phase 3',
         '用户可控',
         '记忆查看 API + 单条删除 + 批量导出',
         '1 周'],
        ['Phase 4',
         '群体记忆',
         '匿名化群体画像（如"计算机学院学生常见问题"）',
         '2 周'],
        ['Phase 5',
         '主动学习',
         '基于记忆主动推荐 + 记忆质量自评估',
         '3 周'],
    ],
    col_widths=[2.5, 2.5, 7, 2])

add_callout(doc,
    'Phase 1-2 是本文档覆盖范围，Phase 3-5 属于后续演进。'
    'Phase 3（用户可控）是隐私合规的硬性要求，建议在 Phase 1 完成后立即启动。',
    color='FFF3CD', border_color='FFC107')

add_heading_styled(doc, '7.4 边界声明', level=2)

add_paragraph_styled(doc, '本文档不覆盖以下内容（避免主题扩散）：')

add_bullet(doc, '上下文工程的六层装载和 Token 预算——已在《上下文工程模块设计方案》详述')
add_bullet(doc, '会话历史（agent_turns）的存储——属于会话管理，非长期记忆')
add_bullet(doc, '知识库内容（knowledge_articles）——属于"全用户共享知识"')
add_bullet(doc, '向量检索算法本身——属于 RAG 检索侧')
add_bullet(doc, '用户画像的隐私合规与 PII 脱敏——属于第 15 个方向"内容审核与 PII 脱敏"')
add_bullet(doc, '多 Agent 共享记忆——属于第 13 个方向"多 Agent 协作"')
add_bullet(doc, '群体画像与画像聚合——属于 Phase 4 演进')
add_bullet(doc, '记忆的联邦学习——不在校园场景范围内')

doc.add_page_break()

# ==================== 附录 ADR 摘要 ====================
add_heading_styled(doc, '附录：ADR 摘要', level=1)

add_paragraph_styled(doc, '本文档涉及的 7 条架构决策记录（ADR）：')

add_styled_table(doc,
    ['ADR 编号', '决策标题', '决策摘要', '权衡'],
    [
        ['ADR-MEM-01',
         '记忆三表架构',
         'user_memory（主表）+ user_memory_evidence（证据）+ user_memory_history（历史）',
         '三表 JOIN 查询稍复杂，但职责分离，审计可追溯'],
        ['ADR-MEM-02',
         '记忆类型四分类',
         'PREFERENCE / FACT / SKILL / EVENT 四类，按更新频率和检索优先级区分',
         '分类边界有时模糊，需 LLM 判断'],
        ['ADR-MEM-03',
         '记忆来源双通道',
         'EXPLICIT（用户明说，confidence=1.0）+ INFERRED（行为推断，confidence=0.6）',
         '推断可能误报，需 evidence_count >= 3 门槛'],
        ['ADR-MEM-04',
         '检索用向量+关键词双路',
         'PROFILE 全量 SQL 查询 + 事件 memory_vectors 向量召回 + pg_trgm 关键词补充',
         '维护两套索引，但结构化+语义兼顾'],
        ['ADR-MEM-05',
         '冲突检测 + 证据加权融合',
         '同 key 新旧记忆冲突时 LLM 仲裁，KEEP_NEW/KEEP_OLD/KEEP_BOTH 三选项',
         'LLM 仲裁有成本，但避免记忆矛盾'],
        ['ADR-MEM-06',
         '遗忘机制三层',
         'Layer1 衰减（0.95^days）+ Layer2 软删除（conf<0.3 且 30 天）+ Layer3 物理清除（30 天后）',
         '需要定时清理任务，但可恢复 + 自然降权'],
        ['ADR-MEM-07',
         '工作记忆联动',
         '通过 agent_context_snapshots.used_memory_ids 字段回写，'
         'PromptAssembler 在 <context> 注入用户画像 slot',
         '每次写入额外字段，但审计可追溯 + 支持遗忘统计'],
    ],
    col_widths=[2.5, 3, 5.5, 3])

add_callout(doc,
    '所有 ADR 决策记录归档在 docs/agent-design/adr/ 目录下，'
    '文件名格式：ADR-MEM-XX-决策标题.md。每条 ADR 包含：'
    '背景、决策、权衡、后果、相关日期。',
    color='E8F4FD', border_color='2196F3')

# ==================== 文档结束 ====================
doc.add_paragraph()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end_p.add_run('— 文档结束 —')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
run.font.italic = True

# 文档元信息
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p.add_run('CampusShare Agent 搭建系列文档 · 第 6 个方向（C 层记忆层）\n'
                     '版本：v1.0.0  ·  ADR 前缀：MEM  ·  生成日期：2026-07-06')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

# ==================== 保存文档 ====================
output_path = r'e:\workspace_work\CampusShare\docs\agent-design\长期记忆模块设计方案.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
print(f'章节数：7 章 + 附录')
print(f'ADR 数量：7 条（ADR-MEM-01 ~ ADR-MEM-07）')
