# -*- coding: utf-8 -*-
"""
可观测性模块设计方案文档生成脚本。
ADR 前缀：OBS
输出：可观测性模块设计方案.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ==================== 样式工具函数 ====================

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell, color_hex='BFBFBF', size='4'):
    """设置单元格边框"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color_hex)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_heading_styled(doc, text, level=1, color_hex='1F4E79'):
    """添加带颜色的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
        if level == 0:
            run.font.size = Pt(26)
        elif level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(12)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading


def add_body_paragraph(doc, text, bold=False, italic=False, color=None, size=11, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """添加项目符号段落"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.name = '微软雅黑'
        run_b._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def add_code_block(doc, code_text, language='java'):
    """添加代码块（灰色背景等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.15
    # 设置段落背景色
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    # 添加边框
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'D0D0D0')
        border.set(qn('w:space'), '4')
        pBdr.append(border)
    pPr.append(pBdr)

    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.color.rgb = RGBColor.from_string('333333')
    return p


def add_table_styled(doc, headers, rows, col_widths=None, header_color='1F4E79'):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string('FFFFFF')
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, header_color)
        set_cell_border(cell)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            set_cell_border(cell)
            # 隔行变色
            if row_idx % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')

    return table


def add_callout_box(doc, title, content, color_hex='FFF4CE', border_color='FFC000'):
    """添加提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)

    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:color'), border_color)
        border.set(qn('w:space'), '6')
        pBdr.append(border)
    pPr.append(pBdr)

    run_t = p.add_run(title)
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.name = '微软雅黑'
    run_t._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run_t.font.color.rgb = RGBColor.from_string('7F6000')

    run_c = p.add_run(content)
    run_c.font.size = Pt(10)
    run_c.font.name = '微软雅黑'
    run_c._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def add_adr_box(doc, adr_id, title, context, decision, rationale, consequences):
    """添加 ADR 决策框"""
    # ADR 标题
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{adr_id}：{title}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string('1F4E79')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    rows = [
        ['背景', context],
        ['决策', decision],
        ['理由', rationale],
        ['后果', consequences],
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for cell in table.columns[0].cells:
        cell.width = Cm(2.0)
    for cell in table.columns[1].cells:
        cell.width = Cm(15.0)
    for i, (label, text) in enumerate(rows):
        row_cells = table.rows[i].cells
        # 标签列
        cell_label = row_cells[0]
        cell_label.text = ''
        p_label = cell_label.paragraphs[0]
        run_label = p_label.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(9.5)
        run_label.font.name = '微软雅黑'
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run_label.font.color.rgb = RGBColor.from_string('FFFFFF')
        set_cell_shading(cell_label, '4472C4')
        set_cell_border(cell_label)
        # 内容列
        cell_text = row_cells[1]
        cell_text.text = ''
        p_text = cell_text.paragraphs[0]
        run_text = p_text.add_run(text)
        run_text.font.size = Pt(9.5)
        run_text.font.name = '微软雅黑'
        run_text._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_border(cell_text)
        set_cell_shading(cell_text, 'F7F9FC')
    doc.add_paragraph()


def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()


# ==================== 创建文档 ====================

doc = Document()

# 设置默认样式
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


# ==================== 封面 ====================

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('可观测性模块设计方案')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor.from_string('1F4E79')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Observability Design — 全链路追踪 / Metrics 指标 / 结构化日志 / 告警规则 / Grafana 仪表盘')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string('666666')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CampusShare Agent 搭建系列文档 · 第 13 篇')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string('999999')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ADR 前缀：OBS')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string('999999')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('版本 v1.0.0    日期：2026-07-11')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor.from_string('AAAAAA')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

add_page_break(doc)


# ==================== 本文档范围声明 ====================

add_heading_styled(doc, '本文档范围声明', level=1, color_hex='C00000')

add_body_paragraph(doc, '本文档是 CampusShare Agent 搭建系列文档的第 13 篇，聚焦"可观测性（Observability）"方向。'
    '可观测性是 Agent 系统的"眼睛"——没有可观测性，Agent 在线上就是黑盒：你不知道哪个请求慢、不知道成本花在哪、不知道错误率是多少、不知道用户是否满意。'
    '本方案在 CampusShare agent-service 现有基础上，设计一套覆盖 Logs / Metrics / Trace 三支柱的可观测性体系。', indent=True)

add_heading_styled(doc, '覆盖内容', level=3, color_hex='2E75B6')
add_bullet(doc, '网关生成 TraceId → HTTP Header 传递 → agent-service MDC → 日志/DB/响应全链路贯通', bold_prefix='全链路 TraceId 传递：')
add_bullet(doc, 'agent_trace_spans 表记录意图识别 / RAG 检索 / LLM 调用 / 工具调用各阶段耗时，支持单请求多阶段下钻', bold_prefix='Span 级链路追踪：')
add_bullet(doc, '6 大维度 Metrics（延迟 / Token / 成本 / 错误率 / 工具调用率 / 缓存命中率），Micrometer + Prometheus 采集', bold_prefix='统一 Metrics 体系：')
add_bullet(doc, 'logback-spring.xml + LogstashEncoder 输出 JSON 结构化日志，便于 ELK / Loki 采集和检索', bold_prefix='结构化 JSON 日志：')
add_bullet(doc, '5 类告警（延迟 P99 / 错误率 / 成本 / 可用率 / 熔断），Prometheus AlertRule + Alertmanager', bold_prefix='告警规则体系：')
add_bullet(doc, '4 大面板（概览 / 链路 / 成本 / 错误），Grafana Dashboard JSON 即开即用', bold_prefix='Grafana 仪表盘：')
add_bullet(doc, '从 agent_turns 表 status=ERROR 记录自动采集 BadCase → 分诊 → 供评估体系使用', bold_prefix='BadCase 自动采集：')

add_heading_styled(doc, '不覆盖内容', level=3, color_hex='2E75B6')
add_bullet(doc, '不覆盖 SLO 目标定义与延迟预算分配——属于"性能 SLO 工程"方向（ADR 前缀 SLO）')
add_bullet(doc, '不覆盖评估指标体系与黄金测试集——属于"评估体系"方向（ADR 前缀 EVAL）')
add_bullet(doc, '不覆盖分布式追踪后端部署（Jaeger/Zipkin/Tempo）——CampusShare V1.0 用 DB 表 + Grafana 即可，不引入额外组件')
add_bullet(doc, '不覆盖日志平台部署（ELK/Loki）——V1.0 用文件日志 + Grafana Loki（可选），不强制全量 ELK')
add_bullet(doc, '不覆盖 APM 商业方案（SkyWalking/Pinpoint）——CampusShare 是开源项目，用开源方案')

add_heading_styled(doc, 'ADR 解释', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'ADR（Architecture Decision Record，架构决策记录）是一种记录架构决策的文档格式。'
    '每条 ADR 包含四个部分：背景（为什么要做这个决策）、决策（做了什么决策）、理由（为什么这样决策）、后果（决策带来的影响）。'
    '本文档所有 ADR 使用前缀 OBS 编号（ADR-OBS-01 ~ ADR-OBS-07），完整摘要见文末附录。', indent=True)

add_heading_styled(doc, '与其他文档的关系', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '可观测性是横切关注点，贯穿 Agent 搭建的所有方向：', indent=True)
add_bullet(doc, '前置文档：安全护栏（SEC）、LLM 网关（GW）——SEC 的安全审计日志、GW 的成本追踪都需要可观测性基础设施')
add_bullet(doc, '后续文档：性能 SLO 工程（SLO）依赖本文档的 Metrics 和告警；评估体系（EVAL）依赖本文档的 BadCase 采集')
add_bullet(doc, '并行文档：分层部署（DEPLOY）——在线/异步/离线分层后，各层的可观测性策略不同')

add_callout_box(doc, '核心认知：',
    '可观测性不是"加几个日志"，而是"回答任何问题的能力"。'
    '一个好的可观测性系统应该能在 30 秒内回答：哪个请求慢？慢在哪一步？为什么慢？影响多少用户？花了多少钱？',
    color_hex='E8F0FE', border_color='4285F4')

add_page_break(doc)


# ==================== 目录 ====================

add_heading_styled(doc, '目录', level=1, color_hex='1F4E79')

toc_items = [
    ('一、场景：为什么需要可观测性', ''),
    ('  1.1 没有可观测性会怎样', ''),
    ('  1.2 带来什么', ''),
    ('  1.3 CampusShare 场景', ''),
    ('  1.4 现状评估', ''),
    ('二、方案：业界可观测性设计模式', ''),
    ('  2.1 三支柱模型（Logs / Metrics / Trace）', ''),
    ('  2.2 大厂案例', ''),
    ('  2.3 选型决策', ''),
    ('  2.4 ADR 汇总', ''),
    ('三、流程：如何搭建可观测性体系', ''),
    ('  3.1 现状评估', ''),
    ('  3.2 全链路 TraceId 传递', ''),
    ('  3.3 Span 级链路追踪表', ''),
    ('  3.4 统一 Metrics 体系（6 大维度）', ''),
    ('  3.5 结构化 JSON 日志', ''),
    ('  3.6 告警规则体系（5 类告警）', ''),
    ('  3.7 Grafana 仪表盘（4 大面板）', ''),
    ('  3.8 BadCase 自动采集', ''),
    ('  3.9 ADR 决策表', ''),
    ('四、核心代码', ''),
    ('  4.1 文件架构', ''),
    ('  4.2 TraceIdFilter — 网关生成 + 传递', ''),
    ('  4.3 TraceContext — agent-service 接收', ''),
    ('  4.4 AgentTraceSpan 实体 + Mapper', ''),
    ('  4.5 ObservabilityMetricsConfig — 6 大维度指标', ''),
    ('  4.6 logback-spring.xml — 结构化 JSON 日志', ''),
    ('  4.7 告警规则 + Grafana 仪表盘 JSON', ''),
    ('  4.8 AgentChatService 集成改造', ''),
    ('  4.9 application.yml 配置', ''),
    ('五、目标：实现效果', ''),
    ('六、测试评估与验收', ''),
    ('七、总结与边界声明', ''),
    ('附录：ADR 摘要', ''),
]
for item, _ in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if not item.startswith('  '):
        run.bold = True
        run.font.color.rgb = RGBColor.from_string('1F4E79')

add_page_break(doc)


# ==================== 第一章 场景 ====================

add_heading_styled(doc, '一、场景：为什么需要可观测性', level=1)

add_heading_styled(doc, '1.1 没有可观测性会怎样', level=2)

add_body_paragraph(doc, 'Agent 系统比传统 Web 服务更迫切地需要可观测性，因为 Agent 的行为不确定性远高于传统服务：', indent=True)

add_heading_styled(doc, '问题一：无法定位慢请求', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '用户反馈"Agent 回复很慢"，但你看不到以下信息：', indent=True)
add_bullet(doc, '慢在哪一步？意图识别 300ms？RAG 检索 800ms？LLM 推理 2s？工具调用 500ms？')
add_bullet(doc, '是偶发还是必现？是所有用户都慢还是个别用户？')
add_bullet(doc, '慢请求的 traceId 是什么？日志里能否串起来？')
add_body_paragraph(doc, 'CampusShare 现状：AgentChatService.prepareContext 中生成了 MDC traceId，'
    '但只在 agent-service 单次请求内有效，无法跨服务追踪。'
    '更关键的是——没有 Span 级记录，你只能看到"整个请求耗时 3s"，看不到"3s 里 RAG 占 1.5s"。', indent=True)

add_heading_styled(doc, '问题二：无法归因成本', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'DeepSeek API 按 Token 计费，但你看不到以下信息：', indent=True)
add_bullet(doc, '哪个用户消耗 Token 最多？哪个会话？哪个意图？')
add_bullet(doc, 'RAG 检索结果是否有效？如果检索了一堆无关内容，Token 白花了')
add_bullet(doc, '快路径（模板回复）占比多少？如果 80% 请求都走快路径，成本应该很低')
add_body_paragraph(doc, 'CampusShare 现状：AgentTurn 表记录了 inputTokens/outputTokens/totalTokens，'
    '但没有按用户/意图/模型维度汇总，也没有成本单价换算，无法回答"今天花了多少钱"。', indent=True)

add_heading_styled(doc, '问题三：无法发现异常', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'Agent 上线后可能出现的异常：', indent=True)
add_bullet(doc, '错误率突增（LLM API 限流 / 超时 / 熔断打开）')
add_bullet(doc, '延迟突增（P99 从 2s 飙到 8s）')
add_bullet(doc, '成本突增（有人刷接口 / Prompt 注入导致大量 Token 消耗）')
add_bullet(doc, '工具调用失败率突增（下游服务挂了）')
add_body_paragraph(doc, '没有告警规则，这些问题只能靠用户投诉发现——等你收到投诉，已经影响了大量用户。', indent=True)

add_heading_styled(doc, '问题四：无法度量质量', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'Agent 回复质量怎么衡量？', indent=True)
add_bullet(doc, '用户点赞率（LIKE）/ 踩率（DISLIKE）是多少？')
add_bullet(doc, '哪些回复被用户点了踩？BadCase 长什么样？')
add_bullet(doc, '意图识别准确率是多少？路由决策正确率？')
add_body_paragraph(doc, 'CampusShare 现状：AgentTurn 表有 feedback 字段（LIKE/DISLIKE/NONE），'
    '但没有自动采集 BadCase 的机制——ERROR 状态的记录散落在表中，无人分析和利用。', indent=True)

add_callout_box(doc, '一句话总结：',
    '没有可观测性的 Agent 系统 = 黑盒运行。你不知道它好不好、快不快、贵不贵、稳不稳。'
    '可观测性的目标是让 Agent 系统从"黑盒"变成"玻璃盒"。',
    color_hex='FFF4CE', border_color='FFC000')

add_heading_styled(doc, '1.2 带来什么', level=2)

add_body_paragraph(doc, '建设可观测性体系后，Agent 系统获得四项核心能力：', indent=True)

add_table_styled(doc,
    headers=['能力', '描述', '回答的问题', '对应组件'],
    rows=[
        ['全链路追踪', '一个请求从网关到 LLM 的每一步都可见', '哪个请求慢？慢在哪一步？', 'TraceId + Span 表'],
        ['实时指标监控', '延迟/Token/成本/错误率实时可查', '今天花了多少钱？错误率多少？', 'Micrometer + Prometheus'],
        ['异常自动告警', '异常发生时主动通知，不等用户投诉', '什么时候出问题？影响多大？', 'AlertRule + Alertmanager'],
        ['BadCase 采集', '线上失败案例自动归集，供评估改进', '哪些回复不好？为什么不好？', 'BadCase 采集器'],
    ],
    col_widths=[2.5, 4.5, 5.0, 4.0])

add_heading_styled(doc, '1.3 CampusShare 场景', level=2)

add_body_paragraph(doc, 'CampusShare 是校园资源共享平台，Agent 作为智能助手回答用户关于校园生活的问题。'
    '具体可观测性场景包括：', indent=True)

add_bullet(doc, '学生反馈"问 Agent 问题半天才回复"——需要看到 RAG 检索耗时、LLM 推理耗时，定位瓶颈',
    bold_prefix='场景一：慢请求定位  ')
add_bullet(doc, '运营想知道"Agent 每天消耗多少 Token，月成本多少"——需要按天/用户/意图汇总成本',
    bold_prefix='场景二：成本归因  ')
add_bullet(doc, 'DeepSeek API 突然限流（429），错误率飙升——需要告警 + 自动降级到备用模型',
    bold_prefix='场景三：异常告警  ')
add_bullet(doc, '用户给 Agent 回复点了踩——需要自动采集 DISLIKE 的 turn，供评估体系分析改进',
    bold_prefix='场景四：BadCase 采集  ')
add_bullet(doc, '意图识别准确率下降（新学期开学后问题类型变了）——需要监控意图分布变化',
    bold_prefix='场景五：质量监控  ')

add_heading_styled(doc, '1.4 现状评估', level=2)

add_body_paragraph(doc, '调研 agent-service 现有可观测性代码，评估结果如下：', indent=True)

add_table_styled(doc,
    headers=['维度', '现有实现', '缺失项', '建设优先级'],
    rows=[
        ['日志',
         'application.yml 配置 logging pattern，含 MDC traceId 占位符 [%X{traceId:-}]；AgentChatService.prepareContext 中 MDC.put("traceId", UUID 前 8 位)',
         '无结构化 JSON 日志；traceId 仅在 agent-service 内部，不跨服务；无全局 TraceIdFilter',
         '高'],
        ['Metrics',
         'KnowledgeMetricsConfig（6 项知识库指标）；IntentMetricsConfig（4 项意图指标）；AgentChatService 中 2 项 Prompt 指标（violation/injection）',
         '无 RAG 检索耗时指标；无 LLM 调用耗时/Token 指标；无工具调用指标；无成本指标；无缓存命中率指标',
         '高'],
        ['Trace',
         'AgentTurn 表记录每轮对话（responseTimeMs/intent/inputTokens/outputTokens/status）',
         '无 Span 级追踪表（看不到 RAG/LLM/工具各阶段耗时）；无跨服务 traceId 传递；无分布式追踪',
         '高'],
        ['告警',
         '无',
         '无任何告警规则',
         '中'],
        ['仪表盘',
         '无',
         '无 Grafana 仪表盘',
         '中'],
        ['BadCase',
         'AgentTurn 表 status=ERROR 和 feedback=DISLIKE 字段存在',
         '无自动采集机制；无分诊标注流程',
         '中'],
    ],
    col_widths=[2.0, 5.5, 6.0, 2.5])

add_callout_box(doc, '现状结论：',
    'CampusShare 已有可观测性"零件"（MDC traceId、Micrometer Metrics、AgentTurn 表），'
    '但缺少"组装"——traceId 不跨服务、Metrics 不覆盖全链路、没有 Span 表、没有告警和仪表盘。'
    '本方案的任务是把这些零件组装成完整的可观测性体系。',
    color_hex='E8F0FE', border_color='4285F4')

add_page_break(doc)


# ==================== 第二章 方案 ====================

add_heading_styled(doc, '二、方案：业界可观测性设计模式', level=1)

add_heading_styled(doc, '2.1 三支柱模型（Logs / Metrics / Trace）', level=2)

add_body_paragraph(doc, '业界公认的可观测性三支柱模型（Three Pillars of Observability），'
    '由 CNCF（Cloud Native Computing Foundation）推广：', indent=True)

add_table_styled(doc,
    headers=['支柱', '定义', '特点', '适用场景', 'CampusShare 方案'],
    rows=[
        ['Logs（日志）',
         '离散事件记录，包含时间戳 + 上下文',
         '高基数（每条日志不同）；可搜索；信息量大',
         '排查具体问题、审计追踪',
         'logback JSON 结构化日志 + MDC traceId'],
        ['Metrics（指标）',
         '聚合数值，按时间序列存储',
         '低基数（标签有限）；可聚合；可告警',
         '监控趋势、告警触发、容量规划',
         'Micrometer + Prometheus 6 大维度'],
        ['Trace（追踪）',
         '一个请求经过的所有服务/组件的因果链',
         '单请求粒度；Span 树结构；可下钻',
         '定位慢请求、跨服务排查',
         'agent_trace_spans 表 + traceId 传递'],
    ],
    col_widths=[2.0, 3.0, 3.5, 3.0, 4.5])

add_body_paragraph(doc, '三支柱的关系：', indent=True)
add_bullet(doc, 'Metrics 告诉你"出问题了"（错误率飙升）→ 低延迟，适合告警')
add_bullet(doc, 'Trace 告诉你"问题在哪"（RAG 慢了）→ 单请求下钻，适合定位')
add_bullet(doc, 'Logs 告诉你"为什么出问题"（RAG 检索 SQL 超时）→ 最细粒度，适合根因分析')

add_callout_box(doc, '关键认知：',
    '三支柱不是三个独立系统，而是一个整体——traceId 贯穿三支柱，'
    '让你从 Metrics（告警）→ Trace（下钻）→ Logs（根因）形成完整排查链路。',
    color_hex='E8F0FE', border_color='4285F4')

add_heading_styled(doc, '2.2 大厂案例', level=2)

add_heading_styled(doc, '案例一：OpenAI — LLM 可观测性标杆', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'OpenAI 内部 observability 团队为 ChatGPT 构建了专门的 LLM 可观测性体系：', indent=True)
add_bullet(doc, '每条请求记录 model / prompt_tokens / completion_tokens / latency / ttft（Time To First Token）')
add_bullet(doc, '按用户/模型/region 三维度聚合成本，实时仪表盘展示日消耗')
add_bullet(doc, 'TTFT P99 作为核心 SLO，超阈值自动告警')
add_bullet(doc, 'BadCase 自动采集：用户点踩 → 自动归集 → 人工标注 → 微调数据集')

add_heading_styled(doc, '案例二：字节跳动 — Coze 平台 Agent 可观测性', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '字节跳动 Coze（扣子）平台的 Agent 可观测性方案：', indent=True)
add_bullet(doc, '全链路 Span：Bot 编排 → 意图识别 → 知识库检索 → LLM 调用 → 工具调用，每一步记录耗时和状态')
add_bullet(doc, '成本看板：按 Bot / 用户 / 模型维度展示 Token 消耗和成本')
add_bullet(doc, '质量监控：对话满意度、任务完成率、转人工率')
add_bullet(doc, '告警体系：LLM 调用失败率 / 平均延迟 / Token 消耗突增告警')

add_heading_styled(doc, '案例三：阿里云 — 通义千问平台', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '阿里云通义千问平台可观测性特点：', indent=True)
add_bullet(doc, 'ARMS（应用实时监控服务）集成 LLM 调用链路追踪')
add_bullet(doc, '按 prompt 模板维度聚合指标，监控不同 Prompt 的效果差异')
add_bullet(doc, 'SLS（日志服务）结构化日志 + SQL 查询分析')

add_heading_styled(doc, '2.3 选型决策', level=2)

add_body_paragraph(doc, '基于大厂案例和 CampusShare 实际情况，选型决策如下：', indent=True)

add_table_styled(doc,
    headers=['组件', '选项', '选型', '理由'],
    rows=[
        ['Metrics 采集',
         'Micrometer / OpenTelemetry SDK / 自研',
         'Micrometer',
         'Spring Boot 原生集成，现有代码已用，学习成本低'],
        ['Metrics 存储',
         'Prometheus / InfluxDB / VictoriaMetrics',
         'Prometheus',
         'CNCF 标准，Grafana 原生支持，社区成熟'],
        ['Trace 后端',
         'Jaeger / Zipkin / Tempo / DB 表',
         'DB 表（agent_trace_spans）',
         'V1.0 不引入额外组件，DB 表 + Grafana 查询即可；后续可升级 Tempo'],
        ['日志格式',
         '纯文本 / JSON 结构化',
         'JSON 结构化',
         '便于 ELK/Loki 采集和检索，字段可查询'],
        ['日志采集',
         'Filebeat / Fluentd / Promtail',
         'Promtail（可选）',
         '与 Grafana Loki 配套，轻量级；V1.0 可先用文件日志'],
        ['告警',
         'Prometheus AlertManager / Grafana Alerts',
         'Prometheus AlertManager',
         '与 Prometheus 配套，支持多窗口多燃烧率'],
        ['仪表盘',
         'Grafana / Kibana / 自研',
         'Grafana',
         '开源标准，支持 Prometheus + MySQL 双数据源'],
        ['TraceId 传递',
         'W3C Trace Context / B3 / 自定义 Header',
         '自定义 X-Trace-Id Header',
         '简单直接，不依赖追踪框架；后续可升级 W3C'],
    ],
    col_widths=[2.5, 3.5, 3.0, 7.0])

add_heading_styled(doc, '2.4 ADR 汇总', level=2)

add_body_paragraph(doc, '本方案共 7 条 ADR，涵盖 TraceId 传递、Span 追踪、Metrics 体系、结构化日志、告警、仪表盘、BadCase 采集：', indent=True)

add_table_styled(doc,
    headers=['ADR 编号', '决策标题', '核心选择'],
    rows=[
        ['ADR-OBS-01', '全链路 TraceId 传递', '网关生成 → X-Trace-Id Header → agent-service MDC → 日志/DB/响应'],
        ['ADR-OBS-02', 'Span 级链路追踪表', 'agent_trace_spans 表记录各阶段耗时，不引入 Jaeger/Zipkin'],
        ['ADR-OBS-03', '统一 Metrics 体系', '6 大维度（延迟/Token/成本/错误率/工具调用/缓存），Micrometer + Prometheus'],
        ['ADR-OBS-04', '结构化 JSON 日志', 'logback LogstashEncoder 输出 JSON，traceId 贯穿'],
        ['ADR-OBS-05', '告警规则体系', '5 类告警（延迟P99/错误率/成本/可用率/熔断），Prometheus AlertRule'],
        ['ADR-OBS-06', 'Grafana 仪表盘', '4 大面板（概览/链路/成本/错误），JSON 即开即用'],
        ['ADR-OBS-07', 'BadCase 自动采集', 'ERROR + DISLIKE 记录自动归集 → 分诊 → 供 EVAL 使用'],
    ],
    col_widths=[2.5, 4.5, 9.0])

add_page_break(doc)


# ==================== 第三章 流程 ====================

add_heading_styled(doc, '三、流程：如何搭建可观测性体系', level=1)

add_body_paragraph(doc, '本章按"先追踪、再指标、后告警"的顺序，逐步搭建可观测性体系。'
    '每一步包含：设计目标 + 核心思路 + 关键决策（ADR）+ 数据流。', indent=True)

# ---- 3.1 现状评估 ----
add_heading_styled(doc, '3.1 现状评估', level=2)

add_body_paragraph(doc, '在搭建前，先梳理 agent-service 现有可观测性代码，明确改造起点：', indent=True)

add_heading_styled(doc, '3.1.1 日志现状', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'application.yml 中的日志配置：', indent=True)
add_code_block(doc, '''logging:
  level:
    root: INFO
    '[com.campushare]': DEBUG
  pattern:
    console: '%d{yyyy-MM-dd HH:mm:ss} [%thread] [%X{traceId:-}] %-5level %logger{36} - %msg%n'

management:
  endpoints:
    web:
      exposure:
        include: health,info,prometheus,metrics,loggers
  prometheus:
    metrics:
      export:
        enabled: true
  metrics:
    distribution:
      percentiles-histogram:
        '[http.server.requests]': true
      percentiles:
        '[http.server.requests]':
          - 0.5
          - 0.75
          - 0.9
          - 0.95
          - 0.99''')

add_body_paragraph(doc, 'AgentChatService 中的 MDC traceId（prepareContext 方法）：', indent=True)
add_code_block(doc, '''private ChatContext prepareContext(String userId, ChatRequest request) {
    // MDC 链路追踪：traceId 贯穿意图识别->检索->prompt 装配日志（ADR-030）
    MDC.put("traceId", UUID.randomUUID().toString().substring(0, 8));
    try {
        AgentSession session = getOrCreateSession(userId, request);
        // ... 意图识别、检索、Prompt 装配 ...
        return new ChatContext(...);
    } finally {
        MDC.remove("traceId");
    }
}''')

add_callout_box(doc, '现状问题：',
    '1. traceId 仅在 agent-service 内部生成，网关没有生成 traceId，无法跨服务追踪。\n'
    '2. 日志是纯文本格式，无法被 ELK/Loki 结构化检索。\n'
    '3. MDC traceId 在 Reactor 异步线程间不自动传递（WebFlux 场景下可能丢失）。',
    color_hex='FCE4E4', border_color='E06666')

add_heading_styled(doc, '3.1.2 Metrics 现状', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '现有 3 个 Metrics 配置类，共 12 项指标：', indent=True)

add_table_styled(doc,
    headers=['配置类', '指标数', '指标列表', '覆盖范围'],
    rows=[
        ['KnowledgeMetricsConfig', '6',
         'agent.knowledge.ingest.total / ingest.duration / chunks.perDoc / embedding.batchSize / retrieval.recallCount / duplicate.detected',
         '知识库摄入'],
        ['IntentMetricsConfig', '4',
         'agent.intent.classification.total / classification.duration / cache.total / route.total',
         '意图识别'],
        ['AgentChatService (initCounters)', '2',
         'agent.prompt.violation / agent.prompt.injection.detected',
         'Prompt 安全'],
    ],
    col_widths=[4.5, 1.5, 7.0, 3.0])

add_body_paragraph(doc, '缺失的指标：', indent=True)
add_bullet(doc, 'LLM 调用耗时 / TTFT（Time To First Token）—— 看不到 LLM 响应速度')
add_bullet(doc, 'LLM Token 消耗（input/output/total）—— 看不到实时 Token 用量')
add_bullet(doc, '成本指标（按用户/意图/模型维度）—— 看不到花了多少钱')
add_bullet(doc, 'RAG 检索耗时 / 召回数 —— 看不到检索性能')
add_bullet(doc, '工具调用次数 / 成功率 / 耗时 —— 看不到工具调用情况')
add_bullet(doc, '缓存命中率（意图缓存 / 检索缓存）—— 看不到缓存效果')
add_bullet(doc, '整体错误率 / 可用率 —— 看不到系统健康度')

add_heading_styled(doc, '3.1.3 Trace 现状', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'AgentTurn 表记录每轮对话的汇总信息，但缺少阶段级 Span：', indent=True)
add_code_block(doc, '''// AgentTurn 实体（agent_turns 表）
public class AgentTurn {
    private String id;
    private String sessionId;
    private Integer turnNumber;
    private String userMessage;
    private String assistantMessage;
    private Integer tokensUsed;        // 总 Token
    private String modelName;
    private String retrievalContext;   // RAG 结果 JSON
    private String toolsUsed;          // 工具使用 JSON
    private Integer responseTimeMs;    // 总耗时
    private String status;             // COMPLETED / ERROR
    private String errorMessage;
    private String intent;             // L1 意图
    private BigDecimal intentConfidence;
    private Integer inputTokens;
    private Integer outputTokens;
    private String feedback;           // LIKE / DISLIKE / NONE
    private Long contextSnapshotId;
    private Integer interrupted;
    private LocalDateTime createdAt;
}''')

add_callout_box(doc, '现状问题：',
    'AgentTurn 记录的是"整轮汇总"——你能看到总耗时 3s，但看不到 3s 里意图识别用了多少、RAG 用了多少、LLM 用了多少。'
    '缺少 Span 级追踪表，无法下钻到单请求的多阶段耗时。',
    color_hex='FCE4E4', border_color='E06666')

add_page_break(doc)

# ---- 3.2 全链路 TraceId 传递 ----
add_heading_styled(doc, '3.2 全链路 TraceId 传递', level=2)

add_body_paragraph(doc, '设计目标：一个请求从网关到 agent-service，traceId 贯穿全程，日志/DB/响应中都能看到同一个 traceId。', indent=True)

add_heading_styled(doc, '3.2.1 数据流', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'TraceId 传递链路：', indent=True)

add_table_styled(doc,
    headers=['步骤', '位置', '动作', '产出'],
    rows=[
        ['1', 'gateway-service', 'TraceIdFilter 拦截请求，检查 X-Trace-Id Header',
         '无则生成 32 位 UUID，放入 MDC + Response Header'],
        ['2', 'gateway-service → agent-service', 'WebClient 转发请求',
         '自动携带 X-Trace-Id Header'],
        ['3', 'agent-service', 'TraceIdFilter 拦截请求，读取 X-Trace-Id Header',
         '放入 MDC，无则生成新 traceId'],
        ['4', 'agent-service', 'AgentChatService 使用 MDC 中的 traceId',
         '所有日志自动输出 traceId'],
        ['5', 'agent-service', 'AgentTraceSpan 记录 traceId',
         'DB 表中每条 Span 都关联 traceId'],
        ['6', 'agent-service → 前端', '响应 Header 返回 X-Trace-Id',
         '前端可用于问题反馈关联'],
    ],
    col_widths=[1.2, 3.5, 5.0, 6.3])

add_heading_styled(doc, '3.2.2 关键设计', level=3, color_hex='2E75B6')
add_bullet(doc, '网关生成优先：gateway-service 作为入口，负责生成 traceId（如果客户端没传），确保所有下游服务共享同一 traceId',
    bold_prefix='设计一：')
add_bullet(doc, '下游兜底：agent-service 如果没收到 X-Trace-Id Header，自己生成一个，保证 MDC 不为空',
    bold_prefix='设计二：')
add_bullet(doc, 'WebFlux 兼容：使用 Reactor Context 传递 traceId，而非 ThreadLocal（MDC 在异步线程间不自动传递）',
    bold_prefix='设计三：')
add_bullet(doc, '响应回传：响应 Header 返回 X-Trace-Id，前端可在问题反馈时附上 traceId，便于排查',
    bold_prefix='设计四：')

add_adr_box(doc,
    'ADR-OBS-01',
    '全链路 TraceId 传递',
    'Agent 系统涉及网关 → agent-service → LLM API 多跳调用，当前 traceId 仅在 agent-service 内部生成，无法跨服务追踪。'
    '用户反馈"慢"时，无法在日志中用同一 traceId 串联网关和 agent-service 的日志。',
    'gateway-service 的 TraceIdFilter 生成 32 位 UUID traceId，放入 MDC + X-Trace-Id 响应 Header；'
    '转发到 agent-service 时携带 X-Trace-Id 请求 Header；agent-service 的 TraceIdFilter 读取该 Header 放入 MDC。'
    'WebFlux 场景使用 Reactor Context 传递，避免 ThreadLocal 丢失。',
    '1. 自定义 Header 比引入 W3C Trace Context 更简单，不依赖追踪框架；'
    '2. 网关生成保证唯一入口，下游兜底保证不丢；'
    '3. Reactor Context 是 WebFlux 异步传递的标准方案。',
    '正面：traceId 贯穿全链路，日志可串联。'
    '负面：自定义 Header 不兼容 OpenTelemetry 生态，后续升级到 OTel 时需要适配。')

add_page_break(doc)

# ---- 3.3 Span 级链路追踪表 ----
add_heading_styled(doc, '3.3 Span 级链路追踪表', level=2)

add_body_paragraph(doc, '设计目标：一个请求拆成多个 Span（阶段），每个 Span 记录耗时和状态，支持单请求多阶段下钻。', indent=True)

add_heading_styled(doc, '3.3.1 Span 定义', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'Agent 请求的典型阶段（Span）：', indent=True)

add_table_styled(doc,
    headers=['Span 名称', '阶段类型', '记录内容', '典型耗时'],
    rows=[
        ['INTENT_CLASSIFICATION', '意图识别', 'L1 意图 / 置信度 / 分类层级（RULE/LLM/EMBEDDING）', '5-300ms'],
        ['RETRIEVAL', 'RAG 检索', '召回数 / 三路耗时 / 重排耗时', '100-800ms'],
        ['CONTEXT_ASSEMBLY', '上下文装配', 'Token 数 / 装配耗时', '5-50ms'],
        ['LLM_CALL', 'LLM 调用', '模型名 / TTFT / 总耗时 / Token 用量', '500-3000ms'],
        ['TOOL_CALL', '工具调用', '工具名 / 参数 / 结果 / 耗时', '100-2000ms'],
        ['OUTPUT_VALIDATION', '输出验证', 'Constitutional AI 检查结果', '1-10ms'],
        ['MEMORY_WRITE', '记忆写入', 'Redis 写入耗时', '1-20ms'],
    ],
    col_widths=[3.8, 2.5, 6.5, 3.2])

add_heading_styled(doc, '3.3.2 agent_trace_spans 表设计', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '每条 Span 记录一个阶段的执行信息：', indent=True)

add_table_styled(doc,
    headers=['字段', '类型', '说明'],
    rows=[
        ['id', 'BIGINT AUTO_INCREMENT', '主键'],
        ['trace_id', 'VARCHAR(32)', '链路 ID，同一请求的所有 Span 共享'],
        ['session_id', 'VARCHAR(32)', '会话 ID'],
        ['turn_id', 'VARCHAR(32)', '轮次 ID（关联 agent_turns.id）'],
        ['span_name', 'VARCHAR(64)', 'Span 名称（INTENT_CLASSIFICATION / RETRIEVAL / LLM_CALL 等）'],
        ['span_order', 'INT', 'Span 顺序（同一 trace 内从 1 递增）'],
        ['parent_span_id', 'BIGINT', '父 Span ID（支持嵌套，如 TOOL_CALL 在 LLM_CALL 内）'],
        ['start_time', 'DATETIME(3)', '开始时间（毫秒精度）'],
        ['end_time', 'DATETIME(3)', '结束时间（毫秒精度）'],
        ['duration_ms', 'INT', '耗时（毫秒）'],
        ['status', 'VARCHAR(16)', 'SUCCESS / ERROR / SKIPPED'],
        ['error_message', 'TEXT', '错误信息（status=ERROR 时）'],
        ['attributes', 'JSON', 'Span 属性（意图/模型名/召回数/Token数等，按 span_name 不同）'],
        ['created_at', 'DATETIME', '记录创建时间'],
    ],
    col_widths=[3.5, 3.5, 9.0])

add_heading_styled(doc, '3.3.3 Span 记录流程', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '在 AgentChatService 中，每个阶段开始和结束都记录 Span：', indent=True)
add_code_block(doc, '''// 阶段开始
SpanRecorder.Span span = spanRecorder.start("INTENT_CLASSIFICATION", turnId, traceId);
try {
    IntentResult intent = intentClassifier.classify(query);
    span.attribute("intent", intent.intent().name());
    span.attribute("confidence", intent.confidence());
    span.attribute("layer", intent.layer());
    span.success();
} catch (Exception e) {
    span.error(e.getMessage());
} finally {
    span.end();  // 计算耗时 + 异步写入 agent_trace_spans 表
}

// 阶段开始（RAG）
span = spanRecorder.start("RETRIEVAL", turnId, traceId);
try {
    List<RetrievalResult> results = retrievalService.retrieve(query, intent);
    span.attribute("recallCount", results.size());
    span.attribute("rerankCount", rerankCount);
    span.success();
} finally {
    span.end();
}''')

add_callout_box(doc, '异步写入：',
    'Span 记录采用异步写入（Mono.fromRunnable + subscribeOn(boundedElastic)），不阻塞主流程。'
    '写入失败不影响请求结果，仅 log.warn。',
    color_hex='E8F0FE', border_color='4285F4')

add_adr_box(doc,
    'ADR-OBS-02',
    'Span 级链路追踪表',
    'AgentTurn 表只记录整轮汇总耗时，无法下钻到单请求的多阶段（意图/RAG/LLM/工具）耗时。'
    '用户反馈"慢"时，无法定位是哪个阶段慢。',
    '新建 agent_trace_spans 表，每条 Span 记录一个阶段的开始/结束时间、耗时、状态、属性。'
    'SpanRecorder 封装 Span 生命周期管理，异步写入 DB。'
    '不引入 Jaeger/Zipkin/Tempo 等追踪后端，用 DB 表 + Grafana MySQL 数据源查询。',
    '1. DB 表方案零额外组件，部署简单，适合 V1.0；'
    '2. Grafana 支持 MySQL 数据源，可直接查询 agent_trace_spans 做链路面板；'
    '3. Jaeger/Zipkin 需要额外部署和维护，V1.0 投入产出比不高。',
    '正面：支持单请求多阶段下钻，定位慢请求精准到阶段。'
    '负面：DB 表查询性能不如专用追踪后端；后续 QPS 高时可升级到 Tempo。')

add_page_break(doc)

# ---- 3.4 统一 Metrics 体系 ----
add_heading_styled(doc, '3.4 统一 Metrics 体系（6 大维度）', level=2)

add_body_paragraph(doc, '设计目标：在现有 12 项指标基础上，补齐到 6 大维度全覆盖，让 Prometheus 抓取后能在 Grafana 展示完整看板。', indent=True)

add_heading_styled(doc, '3.4.1 6 大维度指标体系', level=3, color_hex='2E75B6')

add_table_styled(doc,
    headers=['维度', '指标名', '类型', '标签', '说明'],
    rows=[
        ['1. 延迟',
         'agent.chat.duration', 'Timer',
         'intent, route(short_circuit/rag)',
         '整轮对话耗时'],
        ['',
         'agent.chat.ttft', 'Timer',
         'model',
         'Time To First Token（首 Token 延迟）'],
        ['',
         'agent.span.duration', 'Timer',
         'span_name',
         '各阶段 Span 耗时'],
        ['2. Token',
         'agent.llm.tokens', 'Counter',
         'type(input/output), model, intent',
         'Token 消耗计数'],
        ['',
         'agent.llm.tokens.total', 'Gauge',
         'model',
         '当前累计 Token（每小时重置）'],
        ['3. 成本',
         'agent.cost.cents', 'Counter',
         'user_id, model, intent',
         '成本（分），按用户/模型/意图归因'],
        ['',
         'agent.cost.daily', 'Gauge',
         'model',
         '当日累计成本（分）'],
        ['4. 错误率',
         'agent.chat.errors', 'Counter',
         'error_type(llm_timeout/llm_429/rag_error/tool_error), intent',
         '错误计数'],
        ['',
         'agent.chat.error.rate', 'Gauge',
         'intent',
         '错误率（最近 5 分钟）'],
        ['5. 工具调用',
         'agent.tool.calls', 'Counter',
         'tool_name, result(success/error)',
         '工具调用计数'],
        ['',
         'agent.tool.duration', 'Timer',
         'tool_name',
         '工具调用耗时'],
        ['6. 缓存',
         'agent.cache.hits', 'Counter',
         'cache_type(intent/retrieval), result(hit/miss)',
         '缓存命中计数'],
        ['',
         'agent.cache.hit.rate', 'Gauge',
         'cache_type',
         '缓存命中率（最近 5 分钟）'],
    ],
    col_widths=[2.0, 4.0, 1.5, 4.5, 4.0])

add_heading_styled(doc, '3.4.2 指标采集点', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '各指标在代码中的采集位置：', indent=True)

add_table_styled(doc,
    headers=['指标', '采集位置', '采集方式'],
    rows=[
        ['agent.chat.duration', 'AgentChatService.chat() 流式完成时', 'System.nanoTime() 计算'],
        ['agent.chat.ttft', 'AgentChatService 流式第一个 ChatEvent 时', '记录首 Token 到达时间'],
        ['agent.span.duration', 'SpanRecorder.end()', 'start_time / end_time 差值'],
        ['agent.llm.tokens', 'AgentChatService.completeTurn()', 'DeepSeekResponse.Usage'],
        ['agent.cost.cents', 'AgentChatService.completeTurn()', 'tokens * 单价（配置文件）'],
        ['agent.chat.errors', 'AgentChatService.errorTurn()', '按异常类型分类'],
        ['agent.tool.calls', 'ToolExecutionEngine.execute()', '工具执行前后'],
        ['agent.cache.hits', 'IntentCache / RetrievalCache', 'get 时记录 hit/miss'],
    ],
    col_widths=[4.0, 5.5, 6.5])

add_heading_styled(doc, '3.4.3 成本单价配置', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '不同模型的 Token 单价不同，在 application.yml 中配置：', indent=True)
add_code_block(doc, '''app:
  observability:
    cost:
      # 模型 Token 单价（元/百万 Token）
      models:
        deepseek-v4-flash:
          input: 1.0
          output: 2.0
        qwen-turbo:
          input: 2.0
          output: 6.0
        deepseek-chat:
          input: 2.0
          output: 8.0
      # 默认单价（未配置的模型）
      default:
        input: 2.0
        output: 8.0''')

add_adr_box(doc,
    'ADR-OBS-03',
    '统一 Metrics 体系（6 大维度）',
    '现有 Metrics 只覆盖知识库（6项）+ 意图（4项）+ Prompt安全（2项）= 12 项，'
    '缺少 LLM 调用耗时/Token/成本、RAG 检索耗时、工具调用、缓存命中率、错误率等核心维度。'
    '无法回答"今天花了多少钱""错误率多少""缓存效果如何"。',
    '在 ObservabilityMetricsConfig 中集中定义 6 大维度 14 项指标（延迟3/Token2/成本2/错误2/工具2/缓存3），'
    '统一命名规范（agent.{module}.{metric}），标签低基数设计。'
    '成本指标通过 application.yml 配置模型单价，Counter 累加后用 Grafana 聚合展示日/周/月成本。',
    '1. 集中配置便于管理和发现；2. 低基数标签避免 Prometheus 高基数爆炸；'
    '3. 成本用 Counter 累加而非 Gauge，避免重启丢失。',
    '正面：6 大维度全覆盖，Grafana 可展示完整看板。'
    '负面：14 项指标需要各采集点改造，工作量中等。')

add_page_break(doc)

# ---- 3.5 结构化 JSON 日志 ----
add_heading_styled(doc, '3.5 结构化 JSON 日志', level=2)

add_body_paragraph(doc, '设计目标：日志从纯文本改为 JSON 格式，每条日志是一个 JSON 对象，字段可查询，便于 ELK/Loki 采集。', indent=True)

add_heading_styled(doc, '3.5.1 纯文本日志的问题', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '当前日志格式：', indent=True)
add_code_block(doc, '''2026-07-11 14:23:01 [boundedElastic-1] [a1b2c3d4] INFO  c.c.a.s.AgentChatService - Turn completed: sessionId=s123, turn=1, inputTokens=500, outputTokens=200, totalTokens=700, elapsedMs=1500, violation=null''')

add_body_paragraph(doc, '问题：', indent=True)
add_bullet(doc, '不可查询：想查"sessionId=s123 的所有日志"需要正则匹配，效率低')
add_bullet(doc, '不可聚合：想统计"ERROR 级别日志数"需要解析每行')
add_bullet(doc, '不可采集：Filebeat/Fluentd 需要配置复杂的 grok/pattern 解析规则')

add_heading_styled(doc, '3.5.2 JSON 结构化日志格式', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '改造后每条日志是一个 JSON 对象：', indent=True)
add_code_block(doc, '''{
  "@timestamp": "2026-07-11T14:23:01.123+08:00",
  "level": "INFO",
  "logger": "c.c.a.s.AgentChatService",
  "thread": "boundedElastic-1",
  "traceId": "a1b2c3d4",
  "sessionId": "s123",
  "turnId": "t456",
  "userId": "u789",
  "intent": "SEARCH",
  "model": "deepseek-v4-flash",
  "message": "Turn completed",
  "inputTokens": 500,
  "outputTokens": 200,
  "totalTokens": 700,
  "elapsedMs": 1500,
  "violation": null
}''')

add_heading_styled(doc, '3.5.3 logback-spring.xml 配置', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '使用 logstash-logback-encoder 输出 JSON 格式日志：', indent=True)
add_bullet(doc, '引入 logstash-logback-encoder 依赖')
add_bullet(doc, 'ConsoleAppender 输出 JSON 到控制台（开发环境）')
add_bullet(doc, 'RollingFileAppender 输出 JSON 到文件（生产环境），按天滚动')
add_bullet(doc, 'MDC 字段自动包含到 JSON 中（traceId/sessionId/turnId 等）')

add_heading_styled(doc, '3.5.4 MDC 结构化字段', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '在关键方法中向 MDC 写入结构化字段，日志自动包含：', indent=True)

add_table_styled(doc,
    headers=['MDC Key', '写入位置', '说明'],
    rows=[
        ['traceId', 'TraceIdFilter', '链路 ID'],
        ['sessionId', 'AgentChatService.prepareContext', '会话 ID'],
        ['turnId', 'AgentChatService.prepareContext', '轮次 ID'],
        ['userId', 'AgentChatService.prepareContext', '用户 ID'],
        ['intent', 'AgentChatService（意图识别后）', 'L1 意图'],
        ['model', 'AgentChatService（LLM 调用前）', '模型名'],
        ['spanName', 'SpanRecorder.start', '当前 Span 名称'],
    ],
    col_widths=[3.0, 5.5, 7.5])

add_adr_box(doc,
    'ADR-OBS-04',
    '结构化 JSON 日志',
    '当前日志是纯文本格式，无法被 ELK/Loki 结构化检索，排查问题需要 grep 正则匹配，效率低。'
    'MDC 中已有 traceId，但其他字段（sessionId/turnId/userId/intent）未写入 MDC。',
    '引入 logstash-logback-encoder，logback-spring.xml 配置 LogstashEncoder 输出 JSON 格式日志。'
    '在关键方法中向 MDC 写入结构化字段（sessionId/turnId/userId/intent/model），日志自动包含。'
    '生产环境用 RollingFileAppender 按天滚动，开发环境用 ConsoleAppender。',
    '1. JSON 日志是 ELK/Loki 的事实标准；2. MDC 字段自动包含到 JSON，无需手动拼接；'
    '3. logstash-logback-encoder 是 logback 生态最成熟的 JSON 编码器。',
    '正面：日志可结构化检索，排查效率大幅提升。'
    '负面：JSON 格式比纯文本多约 30% 体积；开发环境可读性略差（可配置双输出）。')

add_page_break(doc)

# ---- 3.6 告警规则体系 ----
add_heading_styled(doc, '3.6 告警规则体系（5 类告警）', level=2)

add_body_paragraph(doc, '设计目标：5 类告警覆盖 Agent 系统的核心风险，异常发生时主动通知，不等用户投诉。', indent=True)

add_heading_styled(doc, '3.6.1 5 类告警', level=3, color_hex='2E75B6')

add_table_styled(doc,
    headers=['告警类型', '告警名称', '触发条件', '持续时间', '严重级别'],
    rows=[
        ['1. 延迟',
         'AgentChatP99High',
         'agent_chat_duration_seconds_quantile{quantile="0.99"} > 5',
         '5 分钟',
         'WARNING'],
        ['',
         'AgentTTFTP99High',
         'agent_chat_ttft_seconds_quantile{quantile="0.99"} > 2',
         '5 分钟',
         'WARNING'],
        ['2. 错误率',
         'AgentErrorRateHigh',
         'rate(agent_chat_errors_total[5m]) / rate(agent_chat_total[5m]) > 0.05',
         '2 分钟',
         'CRITICAL'],
        ['',
         'AgentLLM429RateHigh',
         'rate(agent_chat_errors_total{error_type="llm_429"}[5m]) > 0.1',
         '1 分钟',
         'CRITICAL'],
        ['3. 成本',
         'AgentCostDailyHigh',
         'increase(agent_cost_cents_total[24h]) > 1000000',
         '即时',
         'WARNING'],
        ['',
         'AgentCostSpike',
         'rate(agent_cost_cents_total[10m]) > 200000',
         '10 分钟',
         'WARNING'],
        ['4. 可用率',
         'AgentAvailabilityLow',
         '1 - (rate(agent_chat_errors_total[5m]) / rate(agent_chat_total[5m])) < 0.99',
         '5 分钟',
         'CRITICAL'],
        ['5. 熔断',
         'AgentCircuitBreakerOpen',
         'resilience4j_circuitbreaker_state{name="deepseek"} == 1',
         '即时',
         'CRITICAL'],
    ],
    col_widths=[2.0, 3.5, 6.5, 2.0, 2.0])

add_heading_styled(doc, '3.6.2 告警通知渠道', level=3, color_hex='2E75B6')
add_bullet(doc, 'CRITICAL 级别：钉钉/飞书机器人 + 邮件 + 短信（值oncall）', bold_prefix='通知方式：')
add_bullet(doc, 'WARNING 级别：钉钉/飞书机器人 + 邮件', bold_prefix='通知方式：')
add_bullet(doc, 'CRITICAL 告警 1 分钟内通知；WARNING 告警 5 分钟内通知', bold_prefix='通知延迟：')
add_bullet(doc, '同一告警 10 分钟内不重复发送（告警抑制）', bold_prefix='抑制规则：')

add_heading_styled(doc, '3.6.3 多窗口多燃烧率', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '为避免告警风暴，采用多窗口多燃烧率算法（Multi-Window Multi-Burn-Rate）：', indent=True)
add_bullet(doc, '短窗口（5 分钟）+ 长窗口（1 小时）同时超阈值才告警 → 避免瞬时抖动误报')
add_bullet(doc, '快速燃烧（错误率 > 10%）：5m + 1h 双窗口 → 立即告警')
add_bullet(doc, '慢速燃烧（错误率 > 5%）：30m + 6h 双窗口 → 延迟告警')

add_adr_box(doc,
    'ADR-OBS-05',
    '告警规则体系（5 类告警）',
    '当前无任何告警规则，异常只能靠用户投诉发现。'
    'Agent 系统的常见异常包括延迟突增、错误率飙升、成本突增、可用率下降、熔断打开。',
    '定义 5 类告警（延迟P99/错误率/成本/可用率/熔断），8 条 Prometheus AlertRule。'
    'CRITICAL 级别钉钉+邮件+短信，WARNING 级别钉钉+邮件。'
    '采用多窗口多燃烧率算法避免告警风暴，同一告警 10 分钟内抑制。',
    '1. 多窗口多燃烧率是 SRE 最佳实践，平衡灵敏度和误报率；'
    '2. 分级通知避免值oncall 被低级别告警轰炸；'
    '3. Prometheus AlertRule 是 CNCF 标准，与 Prometheus 配套。',
    '正面：异常主动发现，不等用户投诉。'
    '负面：告警阈值需要根据实际流量调优，初期可能误报。')

add_page_break(doc)

# ---- 3.7 Grafana 仪表盘 ----
add_heading_styled(doc, '3.7 Grafana 仪表盘（4 大面板）', level=2)

add_body_paragraph(doc, '设计目标：4 大面板覆盖概览/链路/成本/错误，JSON 即开即用，导入即可看到完整看板。', indent=True)

add_heading_styled(doc, '3.7.1 4 大面板设计', level=3, color_hex='2E75B6')

add_table_styled(doc,
    headers=['面板', '数据源', '核心图表', '回答的问题'],
    rows=[
        ['概览面板',
         'Prometheus',
         'QPS / 错误率 / P99延迟 / 可用率 / 活跃会话数',
         '系统整体健康吗？'],
        ['链路面板',
         'MySQL (agent_trace_spans)',
         'Span 耗时分布 / 各阶段平均耗时 / 慢请求 Top10 / 链路详情下钻',
         '哪个请求慢？慢在哪一步？'],
        ['成本面板',
         'Prometheus',
         '日/周/月成本趋势 / 按模型/意图/用户分摊 / Token 消耗趋势',
         '花了多少钱？花在哪了？'],
        ['错误面板',
         'Prometheus + MySQL',
         '错误率趋势 / 错误类型分布 / 最近错误列表 / BadCase 列表',
         '出了什么错？为什么？'],
    ],
    col_widths=[2.5, 3.0, 6.5, 4.0])

add_heading_styled(doc, '3.7.2 概览面板详解', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '概览面板是日常巡检入口，一眼看到系统健康状态：', indent=True)

add_table_styled(doc,
    headers=['图表', '类型', 'PromQL / 查询', '说明'],
    rows=[
        ['QPS', 'Stat', 'rate(agent_chat_total[5m])', '每秒请求数'],
        ['错误率', 'Stat', 'rate(agent_chat_errors_total[5m]) / rate(agent_chat_total[5m])', '错误率百分比'],
        ['P99 延迟', 'Stat', 'histogram_quantile(0.99, agent_chat_duration_seconds_bucket)', 'P99 延迟（秒）'],
        ['可用率', 'Stat', '1 - (rate(agent_chat_errors_total[5m]) / rate(agent_chat_total[5m]))', '可用率百分比'],
        ['请求趋势', 'Time series', 'rate(agent_chat_total[5m]) by (intent)', '按意图分组的 QPS 趋势'],
        ['延迟分布', 'Heatmap', 'agent_chat_duration_seconds_bucket', '延迟分布热力图'],
        ['路由占比', 'Pie', 'agent_intent_route_total by (path)', '快路径 vs RAG 路径占比'],
    ],
    col_widths=[2.5, 2.0, 7.0, 4.5])

add_heading_styled(doc, '3.7.3 链路面板详解', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '链路面板使用 MySQL 数据源查询 agent_trace_spans 表，支持单请求下钻：', indent=True)

add_table_styled(doc,
    headers=['图表', '类型', 'SQL 查询', '说明'],
    rows=[
        ['各阶段平均耗时', 'Bar gauge',
         'SELECT span_name, AVG(duration_ms) FROM agent_trace_spans WHERE created_at > NOW() - INTERVAL 1 HOUR GROUP BY span_name',
         '各 Span 平均耗时'],
        ['Span 耗时分布', 'Box plot',
         'SELECT span_name, duration_ms FROM agent_trace_spans WHERE created_at > NOW() - INTERVAL 1 HOUR',
         '各 Span 耗时箱线图'],
        ['慢请求 Top10', 'Table',
         'SELECT trace_id, turn_id, SUM(duration_ms) as total FROM agent_trace_spans WHERE created_at > NOW() - INTERVAL 1 HOUR GROUP BY trace_id, turn_id ORDER BY total DESC LIMIT 10',
         '最慢的 10 个请求'],
        ['链路详情', 'Table',
         'SELECT span_name, span_order, duration_ms, status, attributes FROM agent_trace_spans WHERE trace_id = "$trace_id" ORDER BY span_order',
         '按 traceId 查看完整链路'],
    ],
    col_widths=[2.5, 2.5, 8.5, 2.5])

add_heading_styled(doc, '3.7.4 成本面板详解', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '成本面板回答"花了多少钱，花在哪了"：', indent=True)

add_table_styled(doc,
    headers=['图表', '类型', 'PromQL', '说明'],
    rows=[
        ['日成本趋势', 'Time series', 'increase(agent_cost_cents_total[1d]) / 100', '每日成本（元）'],
        ['按模型分摊', 'Pie', 'sum by (model) (increase(agent_cost_cents_total[1d]))', '各模型成本占比'],
        ['按意图分摊', 'Pie', 'sum by (intent) (increase(agent_cost_cents_total[1d]))', '各意图成本占比'],
        ['Token 消耗趋势', 'Time series', 'rate(agent_llm_tokens_total[5m]) by (type)', '输入/输出 Token 趋势'],
        ['成本 Top10 用户', 'Table', 'topk(10, sum by (user_id) (increase(agent_cost_cents_total[1d])))', '消耗最多的 10 个用户'],
    ],
    col_widths=[3.0, 2.5, 6.5, 4.0])

add_adr_box(doc,
    'ADR-OBS-06',
    'Grafana 仪表盘（4 大面板）',
    '当前无任何可视化仪表盘，运维和运营无法直观看到系统状态和成本。',
    '设计 4 大面板（概览/链路/成本/错误），使用 Grafana + Prometheus + MySQL 三数据源。'
    '概览/成本/错误面板用 Prometheus 数据源（聚合指标），链路面板用 MySQL 数据源（查询 agent_trace_spans 表）。'
    'Dashboard JSON 随项目提交，导入即用。',
    '1. Grafana 是开源标准，支持多数据源；2. Prometheus 适合聚合指标，MySQL 适合单请求下钻；'
    '3. JSON 随项目提交便于团队共享。',
    '正面：可视化运维，巡检效率大幅提升。'
    '负面：Dashboard 需要随指标变化持续维护。')

add_page_break(doc)

# ---- 3.8 BadCase 自动采集 ----
add_heading_styled(doc, '3.8 BadCase 自动采集', level=2)

add_body_paragraph(doc, '设计目标：从线上自动采集失败案例（BadCase），供评估体系（EVAL）分析改进，形成数据飞轮。', indent=True)

add_heading_styled(doc, '3.8.1 BadCase 来源', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'BadCase 有两类来源：', indent=True)

add_table_styled(doc,
    headers=['来源', '触发条件', '采集字段', '优先级'],
    rows=[
        ['ERROR 状态',
         'AgentTurn.status = "ERROR"（LLM 超时/熔断/异常）',
         'turnId / sessionId / userId / userMessage / errorMessage / traceId / span 链路',
         '高（系统错误，必须修复）'],
        ['DISLIKE 反馈',
         'AgentTurn.feedback = "DISLIKE"（用户点踩）',
         'turnId / sessionId / userId / userMessage / assistantMessage / intent / retrievalContext / traceId',
         '高（用户不满意，需分析原因）'],
    ],
    col_widths=[2.5, 4.0, 7.0, 2.5])

add_heading_styled(doc, '3.8.2 采集流程', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'BadCase 采集采用定时任务 + 异步写入：', indent=True)

add_table_styled(doc,
    headers=['步骤', '动作', '频率', '产出'],
    rows=[
        ['1. 扫描', 'BadCaseCollector 定时扫描 agent_turns 表，查找 status=ERROR 或 feedback=DISLIKE 且未采集的记录',
         '每 10 分钟', '候选 BadCase 列表'],
        ['2. 去重', '按 userMessage 的 MD5 去重，避免重复采集',
         '实时', '去重后 BadCase'],
        ['3. 关联', '关联 agent_trace_spans 表，获取完整链路信息',
         '实时', '带链路的 BadCase'],
        ['4. 写入', '写入 agent_badcases 表，状态为 UN_TRIAGED（未分诊）',
         '实时', '待分诊 BadCase'],
        ['5. 分诊', '人工标注：问题类型（意图错误/检索无关/LLM 幻觉/超时/其他）+ 严重级别',
         '人工', '已分诊 BadCase'],
        ['6. 导出', '分诊后的 BadCase 导出为评估体系的测试用例',
         '按需', '黄金集测试用例'],
    ],
    col_widths=[1.5, 6.0, 2.0, 6.5])

add_heading_styled(doc, '3.8.3 agent_badcases 表设计', level=3, color_hex='2E75B6')

add_table_styled(doc,
    headers=['字段', '类型', '说明'],
    rows=[
        ['id', 'BIGINT AUTO_INCREMENT', '主键'],
        ['turn_id', 'VARCHAR(32)', '关联 agent_turns.id'],
        ['trace_id', 'VARCHAR(32)', '链路 ID'],
        ['session_id', 'VARCHAR(32)', '会话 ID'],
        ['user_id', 'VARCHAR(32)', '用户 ID'],
        ['source', 'VARCHAR(16)', '来源（ERROR / DISLIKE）'],
        ['user_message', 'TEXT', '用户原始消息'],
        ['assistant_message', 'TEXT', 'AI 回复内容'],
        ['error_message', 'TEXT', '错误信息（source=ERROR 时）'],
        ['intent', 'VARCHAR(32)', 'L1 意图'],
        ['retrieval_context', 'TEXT', '检索上下文 JSON'],
        ['span_summary', 'JSON', 'Span 链路摘要 JSON'],
        ['status', 'VARCHAR(16)', 'UN_TRIAGED / TRIAGED / RESOLVED / IGNORED'],
        ['issue_type', 'VARCHAR(32)', '分诊类型（INTENT_ERROR / RETRIEVAL_IRRELEVANT / LLM_HALLUCINATION / TIMEOUT / OTHER）'],
        ['severity', 'VARCHAR(16)', '分诊严重级别（LOW / MEDIUM / HIGH / CRITICAL）'],
        ['triage_note', 'TEXT', '分诊备注'],
        ['triaged_by', 'VARCHAR(32)', '分诊人'],
        ['triaged_at', 'DATETIME', '分诊时间'],
        ['message_md5', 'VARCHAR(32)', '用户消息 MD5（去重用）'],
        ['created_at', 'DATETIME', '采集时间'],
    ],
    col_widths=[3.0, 3.5, 9.5])

add_adr_box(doc,
    'ADR-OBS-07',
    'BadCase 自动采集',
    'AgentTurn 表中已有 status=ERROR 和 feedback=DISLIKE 的记录，但散落在表中无人分析和利用。'
    '线上失败案例是最有价值的改进数据，但缺少自动采集和分诊流程。',
    'BadCaseCollector 定时任务每 10 分钟扫描 agent_turns 表，采集 status=ERROR 或 feedback=DISLIKE 的记录，'
    '去重后写入 agent_badcases 表（状态 UN_TRIAGED）。'
    '人工分诊标注问题类型和严重级别，分诊后导出为评估体系的测试用例，形成"采集→分诊→改进→上线→再采集"的数据飞轮。',
    '1. 自动采集替代人工翻表，效率高；2. 去重避免重复采集；3. 分诊标注让 BadCase 可分类分析；'
    '4. 导出到评估体系形成闭环，是数据飞轮的关键环节。',
    '正面：线上失败案例系统化归集，为评估改进提供数据。'
    '负面：分诊需要人工参与，初期工作量较大；可通过 LLM 辅助分诊降低人工成本。')

add_page_break(doc)

# ---- 3.9 ADR 决策表 ----
add_heading_styled(doc, '3.9 ADR 决策表', level=2)

add_body_paragraph(doc, '本章 7 条 ADR 汇总如下：', indent=True)

add_table_styled(doc,
    headers=['ADR', '决策', '核心选择', '关键理由'],
    rows=[
        ['ADR-OBS-01', '全链路 TraceId 传递',
         '网关生成 → X-Trace-Id Header → agent-service MDC',
         '跨服务追踪，日志可串联'],
        ['ADR-OBS-02', 'Span 级链路追踪表',
         'agent_trace_spans 表，不引入 Jaeger/Zipkin',
         'DB 表零额外组件，V1.0 够用'],
        ['ADR-OBS-03', '统一 Metrics 体系',
         '6 大维度 14 项指标，Micrometer + Prometheus',
         '全覆盖，回答成本/错误/延迟/缓存问题'],
        ['ADR-OBS-04', '结构化 JSON 日志',
         'logstash-logback-encoder 输出 JSON',
         '可结构化检索，ELK/Loki 标准格式'],
        ['ADR-OBS-05', '告警规则体系',
         '5 类告警 8 条规则，多窗口多燃烧率',
         '主动发现异常，不等用户投诉'],
        ['ADR-OBS-06', 'Grafana 仪表盘',
         '4 大面板，Prometheus + MySQL 双数据源',
         '可视化运维，JSON 即开即用'],
        ['ADR-OBS-07', 'BadCase 自动采集',
         '定时扫描 + 去重 + 分诊 + 导出评估体系',
         '数据飞轮，持续改进'],
    ],
    col_widths=[2.0, 3.0, 5.0, 6.0])

add_page_break(doc)


# ==================== 第四章 核心代码 ====================

add_heading_styled(doc, '四、核心代码', level=1)

add_body_paragraph(doc, '本章给出可观测性体系的核心代码实现，基于 agent-service 现有代码架构改造。'
    '所有代码遵循 CampusShare 现有约定：WebFlux 响应式、MyBatis-Plus、Micrometer Metrics。', indent=True)

# ---- 4.1 文件架构 ----
add_heading_styled(doc, '4.1 文件架构', level=2)

add_body_paragraph(doc, '新增和改造的文件清单：', indent=True)

add_table_styled(doc,
    headers=['文件', '类型', '说明'],
    rows=[
        ['gateway-service/TraceIdFilter.java', '新增', '网关生成 traceId + 传递 X-Trace-Id Header'],
        ['agent-service/config/TraceIdFilter.java', '新增', 'agent-service 接收 X-Trace-Id Header → MDC'],
        ['agent-service/config/TraceContext.java', '新增', 'Reactor Context 传递 traceId（WebFlux 兼容）'],
        ['agent-service/entity/AgentTraceSpan.java', '新增', 'agent_trace_spans 表实体'],
        ['agent-service/mapper/AgentTraceSpanMapper.java', '新增', 'Span 持久化 Mapper'],
        ['agent-service/observability/SpanRecorder.java', '新增', 'Span 生命周期管理 + 异步写入'],
        ['agent-service/config/ObservabilityMetricsConfig.java', '新增', '6 大维度 14 项指标集中配置'],
        ['agent-service/entity/AgentBadCase.java', '新增', 'agent_badcases 表实体'],
        ['agent-service/mapper/AgentBadCaseMapper.java', '新增', 'BadCase 持久化 Mapper'],
        ['agent-service/observability/BadCaseCollector.java', '新增', '定时采集 BadCase'],
        ['agent-service/resources/logback-spring.xml', '新增', 'JSON 结构化日志配置'],
        ['agent-service/resources/alert_rules.yml', '新增', 'Prometheus 告警规则'],
        ['agent-service/resources/grafana/agent-observability.json', '新增', 'Grafana 仪表盘 JSON'],
        ['agent-service/service/AgentChatService.java', '改造', '集成 SpanRecorder + Metrics + MDC'],
        ['agent-service/resources/application.yml', '改造', '新增 observability 配置'],
        ['agent-service/pom.xml', '改造', '新增 logstash-logback-encoder 依赖'],
    ],
    col_widths=[6.5, 1.5, 8.0])

add_heading_styled(doc, '4.1.1 目录结构', level=3, color_hex='2E75B6')
add_code_block(doc, '''campushare-agent/
├── src/main/java/com/campushare/agent/
│   ├── config/
│   │   ├── TraceIdFilter.java          # 新增：接收 X-Trace-Id
│   │   ├── TraceContext.java           # 新增：Reactor Context 传递
│   │   ├── ObservabilityMetricsConfig.java  # 新增：6 大维度指标
│   │   ├── KnowledgeMetricsConfig.java      # 现有
│   │   └── IntentMetricsConfig.java         # 现有
│   ├── observability/                  # 新增包
│   │   ├── SpanRecorder.java           # Span 生命周期管理
│   │   └── BadCaseCollector.java       # BadCase 定时采集
│   ├── entity/
│   │   ├── AgentTurn.java              # 现有
│   │   ├── AgentTraceSpan.java         # 新增
│   │   └── AgentBadCase.java           # 新增
│   ├── mapper/
│   │   ├── AgentTurnMapper.java        # 现有
│   │   ├── AgentTraceSpanMapper.java   # 新增
│   │   └── AgentBadCaseMapper.java     # 新增
│   └── service/
│       └── AgentChatService.java       # 改造
├── src/main/resources/
│   ├── application.yml                 # 改造
│   ├── logback-spring.xml              # 新增
│   ├── alert_rules.yml                 # 新增
│   └── grafana/
│       └── agent-observability.json    # 新增''')

add_page_break(doc)

# ---- 4.2 TraceIdFilter ----
add_heading_styled(doc, '4.2 TraceIdFilter — 网关生成 + 传递', level=2)

add_body_paragraph(doc, 'gateway-service 中的 TraceIdFilter，负责在请求入口生成 traceId 并传递到下游：', indent=True)

add_code_block(doc, '''package com.campushare.gateway.config;

import org.slf4j.MDC;
import org.springframework.core.annotation.Order;
import org.springframework.http.server.reactive.ServerHttpRequest;
import org.springframework.stereotype.Component;
import org.springframework.web.server.ServerWebExchange;
import org.springframework.web.server.WebFilter;
import org.springframework.web.server.WebFilterChain;
import reactor.core.publisher.Mono;

import java.util.UUID;

/**
 * 网关 TraceId 过滤器。
 *
 * 职责：
 *  1. 检查请求是否携带 X-Trace-Id Header
 *  2. 无则生成 32 位 UUID 作为 traceId
 *  3. 放入 MDC（当前线程日志）+ Reactor Context（异步传递）
 *  4. 响应时回传 X-Trace-Id Header
 *
 * ADR-OBS-01：全链路 TraceId 传递
 */
@Component
@Order(-100)  // 最高优先级，确保所有后续过滤器都能看到 traceId
public class TraceIdFilter implements WebFilter {

    public static final String TRACE_ID_HEADER = "X-Trace-Id";
    public static final String TRACE_ID_KEY = "traceId";

    @Override
    public Mono<Void> filter(ServerWebExchange exchange, WebFilterChain chain) {
        ServerHttpRequest request = exchange.getRequest();

        // 1. 从 Header 读取或生成 traceId
        String traceId = request.getHeaders().getFirst(TRACE_ID_HEADER);
        if (traceId == null || traceId.isBlank()) {
            traceId = UUID.randomUUID().toString().replace("-", "");
        }

        // 2. 放入 MDC（当前线程）
        MDC.put(TRACE_ID_KEY, traceId);

        // 3. 回传响应 Header
        exchange.getResponse().getHeaders().add(TRACE_ID_HEADER, traceId);

        // 4. 转发到下游时携带 X-Trace-Id
        ServerHttpRequest mutatedRequest = request.mutate()
                .header(TRACE_ID_HEADER, traceId)
                .build();
        ServerWebExchange mutatedExchange = exchange.mutate().request(mutatedRequest).build();

        // 5. 放入 Reactor Context（异步线程间传递）
        final String finalTraceId = traceId;
        return chain.filter(mutatedExchange)
                .contextWrite(ctx -> ctx.put(TRACE_ID_KEY, finalTraceId))
                .doFinally(signal -> MDC.remove(TRACE_ID_KEY));
    }
}''')

add_body_paragraph(doc, '关键设计点：', indent=True)
add_bullet(doc, '@Order(-100) 确保最先执行，后续过滤器和路由都能看到 traceId', bold_prefix='优先级：')
add_bullet(doc, '32 位 UUID 去掉横线，简洁且唯一性足够', bold_prefix='traceId 格式：')
add_bullet(doc, '响应 Header 回传 X-Trace-Id，前端可用于问题反馈关联', bold_prefix='响应回传：')
add_bullet(doc, 'doFinally 清理 MDC，防止线程池复用时 traceId 串扰', bold_prefix='清理：')

add_page_break(doc)

# ---- 4.3 TraceContext ----
add_heading_styled(doc, '4.3 TraceContext — agent-service 接收', level=2)

add_body_paragraph(doc, 'agent-service 中的 TraceIdFilter 接收网关传来的 traceId，并使用 Reactor Context 在异步线程间传递：', indent=True)

add_code_block(doc, '''package com.campushare.agent.config;

import org.slf4j.MDC;
import org.springframework.core.annotation.Order;
import org.springframework.http.server.reactive.ServerHttpRequest;
import org.springframework.stereotype.Component;
import org.springframework.web.server.ServerWebExchange;
import org.springframework.web.server.WebFilter;
import org.springframework.web.server.WebFilterChain;
import reactor.core.publisher.Mono;

import java.util.UUID;

/**
 * agent-service TraceId 过滤器。
 *
 * 职责：
 *  1. 从 X-Trace-Id Header 读取 traceId（网关生成）
 *  2. 无则兜底生成（防御性设计）
 *  3. 放入 MDC + Reactor Context
 *  4. 响应回传 X-Trace-Id
 *
 * ADR-OBS-01：全链路 TraceId 传递
 */
@Component
@Order(-100)
public class TraceIdFilter implements WebFilter {

    public static final String TRACE_ID_HEADER = "X-Trace-Id";
    public static final String TRACE_ID_KEY = "traceId";

    @Override
    public Mono<Void> filter(ServerWebExchange exchange, WebFilterChain chain) {
        ServerHttpRequest request = exchange.getRequest();

        String traceId = request.getHeaders().getFirst(TRACE_ID_HEADER);
        if (traceId == null || traceId.isBlank()) {
            // 兜底：网关未传则自己生成
            traceId = UUID.randomUUID().toString().replace("-", "");
        }

        MDC.put(TRACE_ID_KEY, traceId);
        exchange.getResponse().getHeaders().add(TRACE_ID_HEADER, traceId);

        final String finalTraceId = traceId;
        return chain.filter(exchange)
                .contextWrite(ctx -> ctx.put(TRACE_ID_KEY, finalTraceId))
                .doFinally(signal -> MDC.remove(TRACE_ID_KEY));
    }
}''')

add_heading_styled(doc, '4.3.1 Reactor Context 异步传递', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'WebFlux 中 MDC 基于 ThreadLocal，但 Reactor 异步切换线程时 ThreadLocal 会丢失。'
    '解决方式是用 Reactor Context 传递 traceId，在需要的地方从 Context 读取：', indent=True)

add_code_block(doc, '''// 从 Reactor Context 读取 traceId
public static Mono<String> getTraceId() {
    return Mono.deferContextual(ctx -> {
        String traceId = ctx.getOrDefault("traceId", "unknown");
        MDC.put("traceId", traceId);  // 同步到当前线程 MDC
        return Mono.just(traceId);
    });
}

// 在 AgentChatService 中使用
public Flux<ChatEvent> chat(String userId, ChatRequest request) {
    return getTraceId()
        .flatMap(traceId -> {
            MDC.put("traceId", traceId);
            MDC.put("userId", userId);
            return Mono.fromCallable(() -> prepareContext(userId, request, traceId));
        })
        .flatMapMany(ctx -> {
            // ... 流式对话 ...
        });
}''')

add_callout_box(doc, 'WebFlux MDC 陷阱：',
    'Reactor 的 publishOn/subscribeOn 会切换线程，MDC（ThreadLocal）不会自动传递。'
    '必须用 Reactor Context 携带 traceId，在每个操作符中 deferContextual 读取并同步到 MDC。'
    '这是 WebFlux 可观测性的最大坑点。',
    color_hex='FCE4E4', border_color='E06666')

add_page_break(doc)

# ---- 4.4 AgentTraceSpan + SpanRecorder ----
add_heading_styled(doc, '4.4 AgentTraceSpan 实体 + SpanRecorder', level=2)

add_heading_styled(doc, '4.4.1 AgentTraceSpan 实体', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '对应 agent_trace_spans 表的实体类：', indent=True)

add_code_block(doc, '''package com.campushare.agent.entity;

import com.baomidou.mybatisplus.annotation.*;
import lombok.AllArgsConstructor;
import lombok.Builder;
import lombok.Data;
import lombok.NoArgsConstructor;

import java.io.Serializable;
import java.time.LocalDateTime;

@Data
@Builder
@NoArgsConstructor
@AllArgsConstructor
@TableName("agent_trace_spans")
public class AgentTraceSpan implements Serializable {

    private static final long serialVersionUID = 1L;

    @TableId(value = "id", type = IdType.ASSIGN_ID)
    private Long id;

    private String traceId;

    private String sessionId;

    private String turnId;

    private String spanName;

    private Integer spanOrder;

    private Long parentSpanId;

    private LocalDateTime startTime;

    private LocalDateTime endTime;

    private Integer durationMs;

    private String status;

    private String errorMessage;

    /** Span 属性 JSON（意图/模型名/召回数/Token 数等） */
    private String attributes;

    @TableField(fill = FieldFill.INSERT)
    private LocalDateTime createdAt;
}''')

add_heading_styled(doc, '4.4.2 AgentTraceSpanMapper', level=3, color_hex='2E75B6')
add_code_block(doc, '''package com.campushare.agent.mapper;

import com.baomidou.mybatisplus.core.mapper.BaseMapper;
import com.campushare.agent.entity.AgentTraceSpan;
import org.apache.ibatis.annotations.Mapper;
import org.apache.ibatis.annotations.Select;

import java.util.List;
import java.util.Map;

@Mapper
public interface AgentTraceSpanMapper extends BaseMapper<AgentTraceSpan> {

    /** 查询某 traceId 的完整链路（按 span_order 排序） */
    @Select("SELECT * FROM agent_trace_spans WHERE trace_id = #{traceId} ORDER BY span_order")
    List<AgentTraceSpan> selectByTraceId(String traceId);

    /** 查询最近 N 分钟各阶段平均耗时（Grafana 链路面板用） */
    @Select("SELECT span_name, AVG(duration_ms) as avg_duration, COUNT(*) as cnt " +
            "FROM agent_trace_spans " +
            "WHERE created_at > DATE_SUB(NOW(), INTERVAL #{minutes} MINUTE) " +
            "GROUP BY span_name ORDER BY avg_duration DESC")
    List<Map<String, Object>> selectAvgDurationBySpan(int minutes);

    /** 查询慢请求 Top N（Grafana 链路面板用） */
    @Select("SELECT trace_id, turn_id, SUM(duration_ms) as total_ms " +
            "FROM agent_trace_spans " +
            "WHERE created_at > DATE_SUB(NOW(), INTERVAL #{minutes} MINUTE) " +
            "GROUP BY trace_id, turn_id " +
            "ORDER BY total_ms DESC LIMIT #{limit}")
    List<Map<String, Object>> selectSlowRequests(int minutes, int limit);
}''')

add_heading_styled(doc, '4.4.3 SpanRecorder — Span 生命周期管理', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'SpanRecorder 封装 Span 的 start/end/attribute 逻辑，异步写入 DB：', indent=True)

add_code_block(doc, '''package com.campushare.agent.observability;

import com.campushare.agent.entity.AgentTraceSpan;
import com.campushare.agent.mapper.AgentTraceSpanMapper;
import com.fasterxml.jackson.databind.ObjectMapper;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.slf4j.MDC;
import org.springframework.stereotype.Component;
import reactor.core.publisher.Mono;
import reactor.core.scheduler.Schedulers;

import java.time.Duration;
import java.time.Instant;
import java.time.LocalDateTime;
import java.time.ZoneId;
import java.util.HashMap;
import java.util.Map;
import java.util.concurrent.atomic.AtomicInteger;

/**
 * Span 记录器。
 *
 * 使用方式：
 *   Span span = spanRecorder.start("INTENT_CLASSIFICATION", turnId, traceId);
 *   try {
 *       span.attribute("intent", "SEARCH");
 *       span.success();
 *   } catch (Exception e) {
 *       span.error(e.getMessage());
 *   } finally {
 *       span.end();  // 异步写入 DB
 *   }
 *
 * ADR-OBS-02：Span 级链路追踪表
 */
@Slf4j
@Component
@RequiredArgsConstructor
public class SpanRecorder {

    private final AgentTraceSpanMapper spanMapper;
    private final ObjectMapper objectMapper = new ObjectMapper();

    /**
     * 开始一个 Span。
     */
    public Span start(String spanName, String turnId, String traceId) {
        Span span = new Span();
        span.spanName = spanName;
        span.turnId = turnId;
        span.traceId = traceId;
        span.startTime = System.nanoTime();
        span.startInstant = Instant.now();
        span.attributes = new HashMap<>();
        span.status = "SUCCESS";  // 默认成功，异常时改为 ERROR

        // 从 MDC 获取 sessionId
        span.sessionId = MDC.get("sessionId");

        // Span 顺序（同一 trace 内递增，简化实现用原子计数器）
        span.spanOrder = spanCounter.incrementAndGet();

        MDC.put("spanName", spanName);
        return span;
    }

    private static final AtomicInteger spanCounter = new AtomicInteger(0);

    /**
     * Span 生命周期对象。
     */
    public class Span {
        private String spanName;
        private String turnId;
        private String traceId;
        private String sessionId;
        private long startTime;
        private Instant startInstant;
        private Map<String, Object> attributes;
        private String status;
        private String errorMessage;
        private int spanOrder;

        public Span attribute(String key, Object value) {
            attributes.put(key, value);
            return this;
        }

        public Span success() {
            this.status = "SUCCESS";
            return this;
        }

        public Span error(String message) {
            this.status = "ERROR";
            this.errorMessage = message;
            return this;
        }

        public void end() {
            long durationMs = (System.nanoTime() - startTime) / 1_000_000;
            Instant endInstant = Instant.now();

            AgentTraceSpan entity = AgentTraceSpan.builder()
                    .traceId(traceId)
                    .sessionId(sessionId)
                    .turnId(turnId)
                    .spanName(spanName)
                    .spanOrder(spanOrder)
                    .startTime(LocalDateTime.ofInstant(startInstant, ZoneId.systemDefault()))
                    .endTime(LocalDateTime.ofInstant(endInstant, ZoneId.systemDefault()))
                    .durationMs((int) durationMs)
                    .status(status)
                    .errorMessage(errorMessage)
                    .attributes(toJson(attributes))
                    .build();

            // 异步写入 DB，不阻塞主流程
            Mono.fromCallable(() -> {
                spanMapper.insert(entity);
                return null;
            })
            .subscribeOn(Schedulers.boundedElastic())
            .doOnError(e -> log.warn("Failed to record span: traceId={}, spanName={}", traceId, spanName, e))
            .subscribe();

            MDC.remove("spanName");
        }

        private String toJson(Map<String, Object> map) {
            try {
                return objectMapper.writeValueAsString(map);
            } catch (Exception e) {
                return "{}";
            }
        }
    }
}''')

add_callout_box(doc, '异步写入设计：',
    'Span 记录采用 Mono.fromRunnable + subscribeOn(boundedElastic) 异步写入，不阻塞主请求流程。'
    '写入失败仅 log.warn，不影响用户请求结果。'
    '这是可观测性的黄金准则——观测系统不能拖慢被观测系统。',
    color_hex='E8F0FE', border_color='4285F4')

add_page_break(doc)

# ---- 4.5 ObservabilityMetricsConfig ----
add_heading_styled(doc, '4.5 ObservabilityMetricsConfig — 6 大维度指标', level=2)

add_body_paragraph(doc, '集中定义 6 大维度 14 项指标，与现有 KnowledgeMetricsConfig / IntentMetricsConfig 并列：', indent=True)

add_code_block(doc, '''package com.campushare.agent.config;

import io.micrometer.core.instrument.Counter;
import io.micrometer.core.instrument.Gauge;
import io.micrometer.core.instrument.MeterRegistry;
import io.micrometer.core.instrument.Timer;
import lombok.RequiredArgsConstructor;
import org.springframework.beans.factory.annotation.Value;
import org.springframework.stereotype.Component;

import java.util.Map;
import java.util.concurrent.ConcurrentHashMap;
import java.util.concurrent.atomic.AtomicLong;

/**
 * 可观测性监控指标（6 大维度 14 项）。
 *
 * 维度 1. 延迟：agent.chat.duration / agent.chat.ttft / agent.span.duration
 * 维度 2. Token：agent.llm.tokens / agent.llm.tokens.total
 * 维度 3. 成本：agent.cost.cents / agent.cost.daily
 * 维度 4. 错误率：agent.chat.errors / agent.chat.error.rate
 * 维度 5. 工具调用：agent.tool.calls / agent.tool.duration
 * 维度 6. 缓存：agent.cache.hits / agent.cache.hit.rate
 *
 * ADR-OBS-03：统一 Metrics 体系
 */
@Component
@RequiredArgsConstructor
public class ObservabilityMetricsConfig {

    private final MeterRegistry registry;

    @Value("#{${app.observability.cost.models}}")
    private Map<String, Map<String, Double>> modelCostConfig;

    @Value("${app.observability.cost.default.input:2.0}")
    private double defaultInputCost;

    @Value("${app.observability.cost.default.output:8.0}")
    private double defaultOutputCost;

    // ==================== 维度 1：延迟 ====================

    public Timer.Sample startChatTimer() {
        return Timer.start(registry);
    }

    public void recordChatDuration(Timer.Sample sample, String intent, String route) {
        if (sample == null) return;
        sample.stop(Timer.builder("agent.chat.duration")
                .tag("intent", intent != null ? intent : "unknown")
                .tag("route", route != null ? route : "unknown")
                .register(registry));
    }

    public void recordTTFT(long ttftMs, String model) {
        Timer.builder("agent.chat.ttft")
                .tag("model", model != null ? model : "unknown")
                .register(registry)
                .record(java.time.Duration.ofMillis(ttftMs));
    }

    public void recordSpanDuration(String spanName, long durationMs) {
        Timer.builder("agent.span.duration")
                .tag("span_name", spanName != null ? spanName : "unknown")
                .register(registry)
                .record(java.time.Duration.ofMillis(durationMs));
    }

    // ==================== 维度 2：Token ====================

    public void recordTokens(int inputTokens, int outputTokens, String model, String intent) {
        Counter.builder("agent.llm.tokens")
                .tag("type", "input")
                .tag("model", model != null ? model : "unknown")
                .tag("intent", intent != null ? intent : "unknown")
                .register(registry)
                .increment(inputTokens);

        Counter.builder("agent.llm.tokens")
                .tag("type", "output")
                .tag("model", model != null ? model : "unknown")
                .tag("intent", intent != null ? intent : "unknown")
                .register(registry)
                .increment(outputTokens);
    }

    // ==================== 维度 3：成本 ====================

    private final Map<String, AtomicLong> dailyCostByModel = new ConcurrentHashMap<>();

    public void recordCost(int inputTokens, int outputTokens, String model, String userId, String intent) {
        double inputCost = getModelCost(model, "input");
        double outputCost = getModelCost(model, "output");

        // 成本（分）= inputTokens * inputCost / 1_000_000 * 100 + outputTokens * outputCost / 1_000_000 * 100
        long costCents = (long) ((inputTokens * inputCost + outputTokens * outputCost) / 10000.0);

        Counter.builder("agent.cost.cents")
                .tag("user_id", userId != null ? userId : "unknown")
                .tag("model", model != null ? model : "unknown")
                .tag("intent", intent != null ? intent : "unknown")
                .register(registry)
                .increment(costCents);

        // 按模型累加日成本
        AtomicLong daily = dailyCostByModel.computeIfAbsent(model, k -> new AtomicLong(0));
        daily.addAndGet(costCents);
        Gauge.builder("agent.cost.daily", daily, AtomicLong::doubleValue)
                .tag("model", model)
                .register(registry);
    }

    private double getModelCost(String model, String type) {
        if (model == null) return "input".equals(type) ? defaultInputCost : defaultOutputCost;
        Map<String, Double> cost = modelCostConfig.get(model);
        if (cost == null) return "input".equals(type) ? defaultInputCost : defaultOutputCost;
        return cost.getOrDefault(type, "input".equals(type) ? defaultInputCost : defaultOutputCost);
    }

    // ==================== 维度 4：错误率 ====================

    public void recordError(String errorType, String intent) {
        Counter.builder("agent.chat.errors")
                .tag("error_type", errorType != null ? errorType : "unknown")
                .tag("intent", intent != null ? intent : "unknown")
                .register(registry)
                .increment();
    }

    // ==================== 维度 5：工具调用 ====================

    public void recordToolCall(String toolName, boolean success, long durationMs) {
        Counter.builder("agent.tool.calls")
                .tag("tool_name", toolName != null ? toolName : "unknown")
                .tag("result", success ? "success" : "error")
                .register(registry)
                .increment();

        Timer.builder("agent.tool.duration")
                .tag("tool_name", toolName != null ? toolName : "unknown")
                .register(registry)
                .record(java.time.Duration.ofMillis(durationMs));
    }

    // ==================== 维度 6：缓存 ====================

    public void recordCacheHit(String cacheType, boolean hit) {
        Counter.builder("agent.cache.hits")
                .tag("cache_type", cacheType != null ? cacheType : "unknown")
                .tag("result", hit ? "hit" : "miss")
                .register(registry)
                .increment();
    }
}''')

add_body_paragraph(doc, '指标命名规范：', indent=True)
add_bullet(doc, '统一前缀 agent.，避免与其他服务冲突', bold_prefix='命名前缀：')
add_bullet(doc, 'agent.{module}.{metric}，如 agent.chat.duration / agent.llm.tokens / agent.cost.cents', bold_prefix='层级结构：')
add_bullet(doc, 'Counter 用复数（tokens/errors/calls），Timer/Gauge 用单数（duration/ttft）', bold_prefix='单复数：')
add_bullet(doc, '标签值低基数（intent 5 种、model 3 种、error_type 5 种），避免 user_id 等高基数标签', bold_prefix='标签基数：')

add_callout_box(doc, '高基数陷阱：',
    'Prometheus 标签如果用 user_id（数万种值），会导致时间序列爆炸（series cardinality explosion），'
    'Prometheus 内存飙升甚至 OOM。'
    '成本指标虽然有 user_id 标签，但仅用于 Grafana Top10 展示，需配合指标过期策略。'
    '生产环境建议用 group by 聚合后再存储，或用 Loki/ELK 处理高基数维度。',
    color_hex='FCE4E4', border_color='E06666')

add_page_break(doc)

# ---- 4.6 logback-spring.xml ----
add_heading_styled(doc, '4.6 logback-spring.xml — 结构化 JSON 日志', level=2)

add_body_paragraph(doc, '使用 logstash-logback-encoder 输出 JSON 格式日志：', indent=True)

add_heading_styled(doc, '4.6.1 pom.xml 依赖', level=3, color_hex='2E75B6')
add_code_block(doc, '''<!-- logstash-logback-encoder：JSON 结构化日志 -->
<dependency>
    <groupId>net.logstash.logback</groupId>
    <artifactId>logstash-logback-encoder</artifactId>
    <version>7.4</version>
</dependency>''')

add_heading_styled(doc, '4.6.2 logback-spring.xml', level=3, color_hex='2E75B6')
add_code_block(doc, '''<?xml version="1.0" encoding="UTF-8"?>
<configuration>
    <!-- 引入 Spring Boot 默认配置 -->
    <include resource="org/springframework/boot/logging/logback/defaults.xml"/>

    <!-- ==================== 开发环境：控制台 JSON 输出 ==================== -->
    <springProfile name="dev">
        <appender name="CONSOLE_JSON" class="ch.qos.logback.core.ConsoleAppender">
            <encoder class="net.logstash.logback.encoder.LogstashEncoder">
                <!-- 包含 MDC 所有字段 -->
                <includeMdcKeyName>traceId</includeMdcKeyName>
                <includeMdcKeyName>sessionId</includeMdcKeyName>
                <includeMdcKeyName>turnId</includeMdcKeyName>
                <includeMdcKeyName>userId</includeMdcKeyName>
                <includeMdcKeyName>intent</includeMdcKeyName>
                <includeMdcKeyName>model</includeMdcKeyName>
                <includeMdcKeyName>spanName</includeMdcKeyName>
                <!-- 时间戳格式 -->
                <customFields>{"app":"campushare-agent"}</customFields>
            </encoder>
        </appender>
        <root level="INFO">
            <appender-ref ref="CONSOLE_JSON"/>
        </root>
        <logger name="com.campushare" level="DEBUG"/>
    </springProfile>

    <!-- ==================== 生产环境：文件 JSON 输出 + 滚动 ==================== -->
    <springProfile name="prod">
        <appender name="FILE_JSON" class="ch.qos.logback.core.rolling.RollingFileAppender">
            <file>logs/agent-service.json</file>
            <rollingPolicy class="ch.qos.logback.core.rolling.SizeAndTimeBasedRollingPolicy">
                <fileNamePattern>logs/agent-service.%d{yyyy-MM-dd}.%i.json</fileNamePattern>
                <maxFileSize>100MB</maxFileSize>
                <maxHistory>30</maxHistory>
                <totalSizeCap>10GB</totalSizeCap>
            </rollingPolicy>
            <encoder class="net.logstash.logback.encoder.LogstashEncoder">
                <includeMdcKeyName>traceId</includeMdcKeyName>
                <includeMdcKeyName>sessionId</includeMdcKeyName>
                <includeMdcKeyName>turnId</includeMdcKeyName>
                <includeMdcKeyName>userId</includeMdcKeyName>
                <includeMdcKeyName>intent</includeMdcKeyName>
                <includeMdcKeyName>model</includeMdcKeyName>
                <includeMdcKeyName>spanName</includeMdcKeyName>
                <customFields>{"app":"campushare-agent"}</customFields>
            </encoder>
        </appender>
        <root level="INFO">
            <appender-ref ref="FILE_JSON"/>
        </root>
        <logger name="com.campushare" level="INFO"/>
    </springProfile>

    <!-- ==================== 测试环境：纯文本（可读性优先） ==================== -->
    <springProfile name="test">
        <appender name="CONSOLE" class="ch.qos.logback.core.ConsoleAppender">
            <encoder>
                <pattern>%d{yyyy-MM-dd HH:mm:ss} [%thread] [%X{traceId:-}] %-5level %logger{36} - %msg%n</pattern>
            </encoder>
        </appender>
        <root level="INFO">
            <appender-ref ref="CONSOLE"/>
        </root>
    </springProfile>
</configuration>''')

add_body_paragraph(doc, '三套 Profile 设计：', indent=True)
add_bullet(doc, 'dev 环境：JSON 输出到控制台，便于开发时用 jq 工具分析', bold_prefix='dev：')
add_bullet(doc, 'prod 环境：JSON 输出到文件，按天 + 100MB 滚动，保留 30 天 10GB', bold_prefix='prod：')
add_bullet(doc, 'test 环境：纯文本输出，可读性好，便于测试时快速查看', bold_prefix='test：')

add_page_break(doc)

# ---- 4.7 告警规则 + Grafana 仪表盘 ----
add_heading_styled(doc, '4.7 告警规则 + Grafana 仪表盘 JSON', level=2)

add_heading_styled(doc, '4.7.1 Prometheus 告警规则（alert_rules.yml）', level=3, color_hex='2E75B6')
add_code_block(doc, '''# Prometheus AlertRule — Agent 可观测性告警
groups:
  - name: agent_alerts
    rules:
      # ==================== 1. 延迟告警 ====================
      - alert: AgentChatP99High
        expr: histogram_quantile(0.99, agent_chat_duration_seconds_bucket) > 5
        for: 5m
        labels:
          severity: warning
        annotations:
          summary: "Agent 对话 P99 延迟过高"
          description: "P99 延迟 {{ $value }}s 超过阈值 5s，持续 5 分钟"

      - alert: AgentTTFTP99High
        expr: histogram_quantile(0.99, agent_chat_ttft_seconds_bucket) > 2
        for: 5m
        labels:
          severity: warning
        annotations:
          summary: "Agent TTFT P99 过高"
          description: "首 Token 延迟 P99 {{ $value }}s 超过阈值 2s"

      # ==================== 2. 错误率告警 ====================
      - alert: AgentErrorRateHigh
        expr: |
          rate(agent_chat_errors_total[5m]) / rate(agent_chat_total[5m]) > 0.05
        for: 2m
        labels:
          severity: critical
        annotations:
          summary: "Agent 错误率过高"
          description: "错误率 {{ $value | humanizePercentage }} 超过 5%"

      - alert: AgentLLM429RateHigh
        expr: rate(agent_chat_errors_total{error_type="llm_429"}[5m]) > 0.1
        for: 1m
        labels:
          severity: critical
        annotations:
          summary: "LLM API 限流（429）频率过高"
          description: "429 错误率 {{ $value }}/s，可能需要降级到备用模型"

      # ==================== 3. 成本告警 ====================
      - alert: AgentCostDailyHigh
        expr: increase(agent_cost_cents_total[24h]) > 1000000
        labels:
          severity: warning
        annotations:
          summary: "Agent 日成本过高"
          description: "24 小时成本 {{ $value }} 分（{{ $value | humanize }} 元）超过阈值"

      - alert: AgentCostSpike
        expr: rate(agent_cost_cents_total[10m]) > 200000
        for: 10m
        labels:
          severity: warning
        annotations:
          summary: "Agent 成本突增"
          description: "10 分钟内成本速率 {{ $value }} 分/s，可能有异常消耗"

      # ==================== 4. 可用率告警 ====================
      - alert: AgentAvailabilityLow
        expr: |
          1 - (rate(agent_chat_errors_total[5m]) / rate(agent_chat_total[5m])) < 0.99
        for: 5m
        labels:
          severity: critical
        annotations:
          summary: "Agent 可用率低于 99%"
          description: "可用率 {{ $value | humanizePercentage }} 低于 99%"

      # ==================== 5. 熔断告警 ====================
      - alert: AgentCircuitBreakerOpen
        expr: resilience4j_circuitbreaker_state{name="deepseek"} == 1
        labels:
          severity: critical
        annotations:
          summary: "DeepSeek 熔断器已打开"
          description: "熔断器处于 OPEN 状态，LLM 调用被拒绝"''')

add_heading_styled(doc, '4.7.2 Grafana 仪表盘 JSON（概览面板片段）', level=3, color_hex='2E75B6')
add_body_paragraph(doc, 'Grafana Dashboard JSON 完整文件随项目提交，此处展示概览面板核心片段：', indent=True)

add_code_block(doc, '''{
  "title": "CampusShare Agent 可观测性 - 概览",
  "datasource": "Prometheus",
  "panels": [
    {
      "title": "QPS",
      "type": "stat",
      "targets": [{
        "expr": "rate(agent_chat_total[5m])",
        "legendFormat": "QPS"
      }],
      "fieldConfig": {
        "defaults": { "unit": "reqps" }
      }
    },
    {
      "title": "错误率",
      "type": "stat",
      "targets": [{
        "expr": "rate(agent_chat_errors_total[5m]) / rate(agent_chat_total[5m])",
        "legendFormat": "错误率"
      }],
      "fieldConfig": {
        "defaults": {
          "unit": "percentunit",
          "thresholds": {
            "steps": [
              {"color": "green", "value": 0},
              {"color": "yellow", "value": 0.01},
              {"color": "red", "value": 0.05}
            ]
          }
        }
      }
    },
    {
      "title": "P99 延迟",
      "type": "stat",
      "targets": [{
        "expr": "histogram_quantile(0.99, agent_chat_duration_seconds_bucket)",
        "legendFormat": "P99"
      }],
      "fieldConfig": {
        "defaults": {
          "unit": "s",
          "thresholds": {
            "steps": [
              {"color": "green", "value": 0},
              {"color": "yellow", "value": 2},
              {"color": "red", "value": 5}
            ]
          }
        }
      }
    },
    {
      "title": "可用率",
      "type": "stat",
      "targets": [{
        "expr": "1 - (rate(agent_chat_errors_total[5m]) / rate(agent_chat_total[5m]))",
        "legendFormat": "可用率"
      }],
      "fieldConfig": {
        "defaults": {
          "unit": "percentunit",
          "thresholds": {
            "steps": [
              {"color": "red", "value": 0},
              {"color": "yellow", "value": 0.95},
              {"color": "green", "value": 0.99}
            ]
          }
        }
      }
    },
    {
      "title": "请求趋势（按意图）",
      "type": "timeseries",
      "targets": [{
        "expr": "rate(agent_chat_total[5m]) by (intent)",
        "legendFormat": "{{intent}}"
      }]
    },
    {
      "title": "路由占比",
      "type": "piechart",
      "targets": [{
        "expr": "agent_intent_route_total by (path)",
        "legendFormat": "{{path}}"
      }]
    }
  ]
}''')

add_heading_styled(doc, '4.7.3 链路面板（MySQL 数据源）', level=3, color_hex='2E75B6')
add_body_paragraph(doc, '链路面板使用 MySQL 数据源查询 agent_trace_spans 表：', indent=True)

add_code_block(doc, '''{
  "title": "各阶段平均耗时",
  "type": "bargauge",
  "datasource": "MySQL",
  "targets": [{
    "query": "SELECT span_name as metric, AVG(duration_ms) as value " +
             "FROM agent_trace_spans " +
             "WHERE created_at > DATE_SUB(NOW(), INTERVAL 1 HOUR) " +
             "GROUP BY span_name ORDER BY value DESC",
    "format": "table"
  }]
}

{
  "title": "慢请求 Top10",
  "type": "table",
  "datasource": "MySQL",
  "targets": [{
    "query": "SELECT trace_id, turn_id, SUM(duration_ms) as total_ms, " +
             "COUNT(*) as span_count " +
             "FROM agent_trace_spans " +
             "WHERE created_at > DATE_SUB(NOW(), INTERVAL 1 HOUR) " +
             "GROUP BY trace_id, turn_id " +
             "ORDER BY total_ms DESC LIMIT 10",
    "format": "table"
  }]
}

{
  "title": "链路详情（按 traceId 下钻）",
  "type": "table",
  "datasource": "MySQL",
  "targets": [{
    "query": "SELECT span_order, span_name, duration_ms, status, attributes " +
             "FROM agent_trace_spans " +
             "WHERE trace_id = '${trace_id}' ORDER BY span_order",
    "format": "table"
  }]
}''')

add_page_break(doc)

# ---- 4.8 AgentChatService 集成改造 ----
add_heading_styled(doc, '4.8 AgentChatService 集成改造', level=2)

add_body_paragraph(doc, 'AgentChatService 改造前后对比，展示如何集成 SpanRecorder 和 ObservabilityMetrics：', indent=True)

add_heading_styled(doc, '4.8.1 改造前（现有代码）', level=3, color_hex='2E75B6')
add_code_block(doc, '''// 现有：prepareContext 中生成 traceId + MDC
private ChatContext prepareContext(String userId, ChatRequest request) {
    MDC.put("traceId", UUID.randomUUID().toString().substring(0, 8));
    try {
        AgentSession session = getOrCreateSession(userId, request);
        IntentResult intent = classifyIntent(request.getMessage());
        List<RetrievalResult> results = retrievalService.retrieve(...);
        // ... 无 Span 记录、无 Metrics 采集 ...
        return new ChatContext(...);
    } finally {
        MDC.remove("traceId");
    }
}''')

add_heading_styled(doc, '4.8.2 改造后', level=3, color_hex='2E75B6')
add_code_block(doc, '''// 改造后：traceId 从网关传入 + Span 记录 + Metrics 采集
private ChatContext prepareContext(String userId, ChatRequest request, String traceId) {
    // traceId 从 Reactor Context 传入（网关生成），不再自己生成
    MDC.put("traceId", traceId);
    MDC.put("userId", userId);
    try {
        AgentSession session = getOrCreateSession(userId, request);
        MDC.put("sessionId", session.getId());

        // === Span 1：意图识别 ===
        SpanRecorder.Span intentSpan = spanRecorder.start("INTENT_CLASSIFICATION", turnId, traceId);
        IntentResult intent;
        try {
            intent = classifyIntent(request.getMessage());
            intentSpan.attribute("intent", intent.intent().name())
                     .attribute("confidence", intent.confidence())
                     .attribute("layer", intent.layer())
                     .success();
        } catch (Exception e) {
            intentSpan.error(e.getMessage());
            observabilityMetrics.recordError("intent_error", null);
            throw e;
        } finally {
            intentSpan.end();
            MDC.put("intent", intent != null ? intent.intent().name() : "unknown");
        }

        // === Span 2：RAG 检索 ===
        SpanRecorder.Span retrievalSpan = spanRecorder.start("RETRIEVAL", turnId, traceId);
        List<RetrievalResult> results;
        try {
            results = retrievalService.retrieve(...);
            retrievalSpan.attribute("recallCount", results.size())
                        .success();
        } catch (Exception e) {
            retrievalSpan.error(e.getMessage());
            observabilityMetrics.recordError("rag_error", intent.intent().name());
            throw e;
        } finally {
            retrievalSpan.end();
        }

        // === Span 3：上下文装配 ===
        SpanRecorder.Span contextSpan = spanRecorder.start("CONTEXT_ASSEMBLY", turnId, traceId);
        AssembledContext assembled;
        try {
            assembled = contextAssembler.assemble(...);
            contextSpan.attribute("totalTokens", assembled.totalTokens())
                       .success();
        } finally {
            contextSpan.end();
        }

        return new ChatContext(..., traceId);
    } finally {
        MDC.remove("traceId");
        MDC.remove("userId");
        MDC.remove("sessionId");
        MDC.remove("intent");
    }
}''')

add_heading_styled(doc, '4.8.3 completeTurn 改造（增加 Metrics）', level=3, color_hex='2E75B6')
add_code_block(doc, '''// 改造后：completeTurn 中记录 Token + 成本 + TTFT
private void completeTurn(AgentTurn turn, AgentSession session, String content,
        long elapsedMs, DeepSeekResponse.Usage usage, int inputTokens,
        String retrievalContextJson, IntentResult intentResult, String promptVersion,
        long ttftMs) {  // 新增 TTFT 参数
    try {
        int completionTokens;
        int totalTokens;
        if (usage != null && usage.getTotalTokens() != null) {
            totalTokens = usage.getTotalTokens();
            completionTokens = usage.getCompletionTokens() != null
                    ? usage.getCompletionTokens()
                    : TokenCounter.countTokens(content);
        } else {
            completionTokens = TokenCounter.countTokens(content);
            totalTokens = inputTokens + completionTokens;
        }

        // === 新增：Metrics 采集 ===
        // 维度 1：延迟（整轮 + TTFT）
        String intentName = intentResult != null ? intentResult.intent().name() : "unknown";
        String route = intentResult != null && intentResult.isShortCircuit() ? "short_circuit" : "rag";
        // chatTimer 在 chat() 方法开始时启动，这里 stop
        // observabilityMetrics.recordChatDuration(chatTimerSample, intentName, route);
        observabilityMetrics.recordTTFT(ttftMs, modelName);

        // 维度 2：Token
        observabilityMetrics.recordTokens(inputTokens, completionTokens, modelName, intentName);

        // 维度 3：成本
        observabilityMetrics.recordCost(inputTokens, completionTokens, modelName,
                session.getUserId(), intentName);

        // === 原有逻辑保持不变 ===
        turn.setAssistantMessage(content);
        turn.setStatus("COMPLETED");
        turn.setResponseTimeMs((int) elapsedMs);
        turn.setTokensUsed(totalTokens);
        turn.setInputTokens(inputTokens);
        turn.setOutputTokens(completionTokens);
        turnMapper.updateById(turn);
        // ... session 更新、Redis 记忆写入 ...

        log.info("Turn completed: sessionId={}, turn={}, intent={}, inputTokens={}, " +
                "outputTokens={}, totalTokens={}, elapsedMs={}, ttftMs={}, costCents={}",
                session.getId(), turn.getTurnNumber(), intentName, inputTokens,
                completionTokens, totalTokens, elapsedMs, ttftMs,
                calculateCostCents(inputTokens, completionTokens, modelName));
    } catch (Exception e) {
        log.error("Failed to complete turn {}", turn.getId(), e);
    }
}''')

add_heading_styled(doc, '4.8.4 errorTurn 改造（增加错误 Metrics）', level=3, color_hex='2E75B6')
add_code_block(doc, '''// 改造后：errorTurn 中按异常类型记录错误 Metrics
private void errorTurn(AgentTurn turn, String errorMessage) {
    try {
        turn.setStatus("ERROR");
        turn.setErrorMessage(errorMessage);
        turnMapper.updateById(turn);

        // === 新增：按异常类型分类记录错误 ===
        String errorType = classifyError(errorMessage);
        String intent = turn.getIntent() != null ? turn.getIntent() : "unknown";
        observabilityMetrics.recordError(errorType, intent);

        log.error("Turn failed: turnId={}, errorType={}, intent={}, error={}",
                turn.getId(), errorType, intent, errorMessage);
    } catch (Exception e) {
        log.error("Failed to update error turn {}", turn.getId(), e);
    }
}

/** 异常类型分类 */
private String classifyError(String errorMessage) {
    if (errorMessage == null) return "unknown";
    String lower = errorMessage.toLowerCase();
    if (lower.contains("timeout") || lower.contains("timed out")) return "llm_timeout";
    if (lower.contains("429") || lower.contains("rate limit")) return "llm_429";
    if (lower.contains("circuit breaker")) return "circuit_breaker";
    if (lower.contains("retrieval") || lower.contains("rag")) return "rag_error";
    if (lower.contains("tool")) return "tool_error";
    return "other";
}''')

add_heading_styled(doc, '4.8.5 BadCaseCollector 定时采集', level=3, color_hex='2E75B6')
add_code_block(doc, '''package com.campushare.agent.observability;

import com.baomidou.mybatisplus.core.conditions.query.LambdaQueryWrapper;
import com.campushare.agent.entity.AgentBadCase;
import com.campushare.agent.entity.AgentTurn;
import com.campushare.agent.mapper.AgentBadCaseMapper;
import com.campushare.agent.mapper.AgentTurnMapper;
import lombok.RequiredArgsConstructor;
import lombok.extern.slf4j.Slf4j;
import org.springframework.scheduling.annotation.Scheduled;
import org.springframework.stereotype.Component;

import java.security.MessageDigest;
import java.time.LocalDateTime;
import java.util.List;

/**
 * BadCase 定时采集器。
 *
 * 每 10 分钟扫描 agent_turns 表，采集 status=ERROR 或 feedback=DISLIKE 的记录，
 * 去重后写入 agent_badcases 表。
 *
 * ADR-OBS-07：BadCase 自动采集
 */
@Slf4j
@Component
@RequiredArgsConstructor
public class BadCaseCollector {

    private final AgentTurnMapper turnMapper;
    private final AgentBadCaseMapper badCaseMapper;

    @Scheduled(fixedDelay = 600_000)  // 每 10 分钟
    public void collect() {
        LocalDateTime since = LocalDateTime.now().minusHours(1);

        // 1. 扫描 ERROR 状态
        LambdaQueryWrapper<AgentTurn> errorQuery = new LambdaQueryWrapper<>()
                .eq(AgentTurn::getStatus, "ERROR")
                .gt(AgentTurn::getCreatedAt, since);
        List<AgentTurn> errorTurns = turnMapper.selectList(errorQuery);

        // 2. 扫描 DISLIKE 反馈
        LambdaQueryWrapper<AgentTurn> dislikeQuery = new LambdaQueryWrapper<>()
                .eq(AgentTurn::getFeedback, "DISLIKE")
                .gt(AgentTurn::getCreatedAt, since);
        List<AgentTurn> dislikeTurns = turnMapper.selectList(dislikeQuery);

        // 3. 去重 + 写入
        for (AgentTurn turn : errorTurns) {
            collectOne(turn, "ERROR");
        }
        for (AgentTurn turn : dislikeTurns) {
            collectOne(turn, "DISLIKE");
        }

        log.info("BadCase collected: errors={}, dislikes={}", errorTurns.size(), dislikeTurns.size());
    }

    private void collectOne(AgentTurn turn, String source) {
        String md5 = md5(turn.getUserMessage());

        // 去重：同一 message_md5 + source 已存在则跳过
        Long existing = badCaseMapper.selectCount(new LambdaQueryWrapper<AgentBadCase>()
                .eq(AgentBadCase::getMessageMd5, md5)
                .eq(AgentBadCase::getSource, source));
        if (existing > 0) return;

        AgentBadCase badCase = AgentBadCase.builder()
                .turnId(turn.getId())
                .traceId(null)  // TODO: 从 agent_trace_spans 关联
                .sessionId(turn.getSessionId())
                .userId(null)   // 从 session 关联
                .source(source)
                .userMessage(turn.getUserMessage())
                .assistantMessage(turn.getAssistantMessage())
                .errorMessage(turn.getErrorMessage())
                .intent(turn.getIntent())
                .retrievalContext(turn.getRetrievalContext())
                .status("UN_TRIAGED")
                .messageMd5(md5)
                .createdAt(LocalDateTime.now())
                .build();
        badCaseMapper.insert(badCase);
    }

    private String md5(String text) {
        try {
            MessageDigest md = MessageDigest.getInstance("MD5");
            byte[] digest = md.digest(text.getBytes());
            StringBuilder sb = new StringBuilder();
            for (byte b : digest) sb.append(String.format("%02x", b));
            return sb.toString();
        } catch (Exception e) {
            return String.valueOf(text.hashCode());
        }
    }
}''')

add_page_break(doc)

# ---- 4.9 application.yml 配置 ----
add_heading_styled(doc, '4.9 application.yml 配置', level=2)

add_body_paragraph(doc, '新增 observability 配置段，包含成本单价和 BadCase 采集开关：', indent=True)

add_code_block(doc, '''app:
  # === 可观测性配置（ADR-OBS-03 / ADR-OBS-07）===
  observability:
    # Span 追踪开关（false 时 SpanRecorder 不写入 DB）
    span:
      enabled: true
    # 成本追踪
    cost:
      # 模型 Token 单价（元/百万 Token）
      models:
        deepseek-v4-flash:
          input: 1.0
          output: 2.0
        qwen-turbo:
          input: 2.0
          output: 6.0
        deepseek-chat:
          input: 2.0
          output: 8.0
      # 默认单价（未配置的模型）
      default:
        input: 2.0
        output: 8.0
    # BadCase 采集
    badcase:
      enabled: true
      # 采集间隔（毫秒），默认 10 分钟
      collect-interval: 600000
      # 扫描时间窗口（小时），默认 1 小时
      scan-window-hours: 1

# 日志配置（logback-spring.xml 中的 springProfile 对应）
# 开发环境用 dev profile（JSON 控制台），生产用 prod profile（JSON 文件）
spring:
  profiles:
    active: dev

# Prometheus 指标导出（现有配置增强）
management:
  endpoints:
    web:
      exposure:
        include: health,info,prometheus,metrics,loggers
  prometheus:
    metrics:
      export:
        enabled: true
        step: 30s  # 抓取间隔 30 秒
  metrics:
    tags:
      application: ${spring.application.name}
    distribution:
      percentiles-histogram:
        '[http.server.requests]': true
        '[agent.chat.duration]': true
        '[agent.chat.ttft]': true
        '[agent.span.duration]': true
      percentiles:
        '[http.server.requests]': [0.5, 0.75, 0.9, 0.95, 0.99]
        '[agent.chat.duration]': [0.5, 0.75, 0.9, 0.95, 0.99]
        '[agent.chat.ttft]': [0.5, 0.75, 0.9, 0.95, 0.99]
        '[agent.span.duration]': [0.5, 0.75, 0.9, 0.95, 0.99]''')

add_callout_box(doc, 'percentiles-histogram 配置：',
    '为 agent.chat.duration / agent.chat.ttft / agent.span.duration 开启 percentile histogram，'
    'Prometheus 才能在 Grafana 中用 histogram_quantile(0.99, ...) 计算 P99 延迟。'
    '否则只能看到平均值，看不到长尾延迟。',
    color_hex='E8F0FE', border_color='4285F4')

add_page_break(doc)


# ==================== 第五章 目标 ====================

add_heading_styled(doc, '五、目标：实现效果', level=1)

add_body_paragraph(doc, '可观测性体系建成后，Agent 系统从"黑盒"变为"玻璃盒"，实现以下目标：', indent=True)

add_heading_styled(doc, '5.1 功能目标', level=2)

add_table_styled(doc,
    headers=['目标', '描述', '验收标准'],
    rows=[
        ['全链路可追踪',
         '一个请求从网关到 LLM 的每一步都可见',
         '任意请求的 traceId 可在日志/DB/Grafana 中查到完整链路'],
        ['指标实时可查',
         '延迟/Token/成本/错误率等指标 30 秒内更新到 Grafana',
         'Prometheus 抓取间隔 30s，Grafana 面板实时刷新'],
        ['异常自动告警',
         '异常发生时 1-5 分钟内通知',
         '8 条告警规则覆盖 5 类异常，CRITICAL 1 分钟、WARNING 5 分钟'],
        ['BadCase 可采集',
         '线上失败案例自动归集',
         '每 10 分钟自动扫描，去重后写入 agent_badcases 表'],
        ['日志可检索',
         '结构化 JSON 日志可按字段查询',
         'logback 输出 JSON，traceId/sessionId/intent 等字段可查询'],
    ],
    col_widths=[2.5, 5.0, 8.5])

add_heading_styled(doc, '5.2 性能目标', level=2)

add_table_styled(doc,
    headers=['指标', '目标值', '说明'],
    rows=[
        ['Span 记录延迟', '< 5ms', '异步写入 DB，不阻塞主流程'],
        ['Metrics 采集延迟', '< 1ms', 'Counter/Timer.increment 内存操作'],
        ['日志写入开销', '< 5%', 'JSON 编码 + 文件写入开销 < 5% CPU'],
        ['DB 写入压力', '< 100 span/s', '单请求 5-7 个 Span，100 QPS = 500-700 span/s，需批量写入优化'],
        ['Prometheus 抓取间隔', '30s', 'step=30s，平衡实时性和性能'],
    ],
    col_widths=[3.5, 3.0, 9.5])

add_heading_styled(doc, '5.3 质量目标', level=2)

add_table_styled(doc,
    headers=['指标', '目标值', '说明'],
    rows=[
        ['日志覆盖率', '100%', '所有关键方法都有 log.info/log.error，且包含 traceId'],
        ['Span 覆盖率', '7 个阶段全覆盖', '意图/RAG/上下文/LLM/工具/验证/记忆'],
        ['Metrics 覆盖率', '6 大维度 14 项', '延迟/Token/成本/错误/工具/缓存'],
        ['告警覆盖率', '5 类异常', '延迟/错误率/成本/可用率/熔断'],
        ['BadCase 采集率', '> 90%', 'ERROR + DISLIKE 记录 90% 以上被采集'],
    ],
    col_widths=[3.5, 3.0, 9.5])

add_heading_styled(doc, '5.4 成本目标', level=2)

add_table_styled(doc,
    headers=['指标', '目标值', '说明'],
    rows=[
        ['可观测性基础设施成本', '< 1 核 2G', 'Prometheus + Grafana + Alertmanager 总资源'],
        ['DB 存储（Span 表）', '< 1GB/天', '单 Span ~200B，700 span/s * 86400s ≈ 12M 条 ≈ 2.4GB/天，需定期清理'],
        ['日志存储', '< 10GB/天', 'JSON 日志按天滚动，保留 30 天 10GB'],
        ['DB 清理策略', '保留 7 天', 'agent_trace_spans 表保留 7 天，超期自动清理'],
    ],
    col_widths=[4.0, 2.5, 9.5])

add_callout_box(doc, '关键认知：',
    '可观测性系统自身的成本不能超过被观测系统。'
    '如果可观测性占用了 Agent 系统一半的资源，那就是过度设计。'
    'V1.0 用最轻量的方案（DB 表 + Prometheus + Grafana），后续 QPS 高时再升级。',
    color_hex='E8F0FE', border_color='4285F4')

add_page_break(doc)


# ==================== 第六章 测试评估与验收 ====================

add_heading_styled(doc, '六、测试评估与验收', level=1)

add_heading_styled(doc, '6.1 评估指标', level=2)

add_table_styled(doc,
    headers=['评估维度', '评估指标', '目标值', '评估方式'],
    rows=[
        ['链路完整性', 'traceId 贯穿率', '100%',
         '抽样 100 个请求，检查网关日志和 agent-service 日志 traceId 一致'],
        ['Span 覆盖率', '7 阶段 Span 记录率', '100%',
         '检查每个请求的 agent_trace_spans 记录数 >= 3（意图+RAG+LLM）'],
        ['Metrics 准确性', 'Token 指标 vs 实际 Token', '误差 < 1%',
         '对比 agent.llm.tokens 指标和 AgentTurn 表 tokensUsed 字段'],
        ['成本计算准确性', '成本指标 vs 实际账单', '误差 < 5%',
         '对比 agent.cost.cents 累计值和 DeepSeek 账单'],
        ['告警触发率', '告警准确率', '> 90%',
         '统计告警触发是否对应真实异常（误报率 < 10%）'],
        ['日志可检索性', 'JSON 格式正确率', '100%',
         '用 jq 工具解析日志文件，100% 可解析'],
        ['BadCase 采集率', 'ERROR/DISLIKE 采集率', '> 90%',
         '对比 agent_badcases 表和 agent_turns 表的 ERROR/DISLIKE 记录数'],
    ],
    col_widths=[2.5, 3.5, 2.5, 7.5])

add_heading_styled(doc, '6.2 测试用例', level=2)

add_heading_styled(doc, '6.2.1 链路追踪测试', level=3, color_hex='2E75B6')
add_table_styled(doc,
    headers=['用例', '操作', '预期结果', '验证方式'],
    rows=[
        ['TC-OBS-01：traceId 跨服务传递',
         '发起 Agent 对话请求，检查网关和 agent-service 日志',
         '网关日志 traceId == agent-service 日志 traceId',
         'grep traceId 网关日志 + agent-service 日志，比对一致'],
        ['TC-OBS-02：Span 记录完整性',
         '发起一个 RAG 请求，查询 agent_trace_spans 表',
         '至少 3 条 Span（INTENT_CLASSIFICATION + RETRIEVAL + LLM_CALL）',
         'SELECT * FROM agent_trace_spans WHERE trace_id = ? ORDER BY span_order'],
        ['TC-OBS-03：Span 耗时准确性',
         '发起请求，对比 Span durationMs 和实际耗时',
         'Span 耗时与日志记录的 elapsedMs 误差 < 10ms',
         '对比 agent_trace_spans.duration_ms 和日志时间戳'],
        ['TC-OBS-04：traceId 响应回传',
         '发起请求，检查响应 Header',
         '响应 Header 包含 X-Trace-Id',
         'curl -v 检查响应 Header'],
    ],
    col_widths=[3.0, 3.5, 4.5, 5.0])

add_heading_styled(doc, '6.2.2 Metrics 指标测试', level=3, color_hex='2E75B6')
add_table_styled(doc,
    headers=['用例', '操作', '预期结果', '验证方式'],
    rows=[
        ['TC-OBS-05：延迟指标',
         '发起 10 个请求，检查 Prometheus 指标',
         'agent_chat_duration_seconds_count 增加 10',
         'curl /actuator/prometheus | grep agent_chat_duration'],
        ['TC-OBS-06：Token 指标',
         '发起 1 个请求，检查 Token 指标',
         'agent_llm_tokens_total{type="input/output"} 增加对应值',
         'curl /actuator/prometheus | grep agent_llm_tokens'],
        ['TC-OBS-07：成本指标',
         '发起 1 个请求，检查成本指标',
         'agent_cost_cents_total 增加对应分值',
         'curl /actuator/prometheus | grep agent_cost_cents'],
        ['TC-OBS-08：错误率指标',
         '模拟 LLM 超时，检查错误指标',
         'agent_chat_errors_total{error_type="llm_timeout"} 增加',
         'curl /actuator/prometheus | grep agent_chat_errors'],
        ['TC-OBS-09：缓存命中率指标',
         '连续发起 2 个相同意图的请求，检查缓存指标',
         '第 2 次请求 agent_cache_hits{result="hit"} 增加',
         'curl /actuator/prometheus | grep agent_cache_hits'],
    ],
    col_widths=[2.5, 3.5, 5.0, 5.0])

add_heading_styled(doc, '6.2.3 日志测试', level=3, color_hex='2E75B6')
add_table_styled(doc,
    headers=['用例', '操作', '预期结果', '验证方式'],
    rows=[
        ['TC-OBS-10：JSON 格式正确',
         '启动 dev profile，发起请求，查看控制台日志',
         '每行日志是合法 JSON，可用 jq 解析',
         'jq . logs/agent-service.json'],
        ['TC-OBS-11：MDC 字段包含',
         '发起请求，检查日志 JSON',
         '包含 traceId/sessionId/turnId/userId/intent/model 字段',
         'jq .traceId logs/agent-service.json'],
        ['TC-OBS-12：文件滚动',
         '启动 prod profile，写入 > 100MB 日志',
         '日志文件按天 + 100MB 滚动',
         'ls -la logs/ 检查文件数和大小'],
    ],
    col_widths=[2.5, 4.0, 4.5, 5.0])

add_heading_styled(doc, '6.2.4 告警测试', level=3, color_hex='2E75B6')
add_table_styled(doc,
    headers=['用例', '操作', '预期结果', '验证方式'],
    rows=[
        ['TC-OBS-13：延迟告警',
         '模拟 P99 > 5s 持续 5 分钟',
         'AgentChatP99High 告警触发',
         'Alertmanager UI 查看告警'],
        ['TC-OBS-14：错误率告警',
         '模拟错误率 > 5% 持续 2 分钟',
         'AgentErrorRateHigh 告警触发',
         'Alertmanager UI 查看告警'],
        ['TC-OBS-15：熔断告警',
         '触发 DeepSeek 熔断器打开',
         'AgentCircuitBreakerOpen 告警立即触发',
         'Alertmanager UI 查看告警'],
        ['TC-OBS-16：告警抑制',
         '同一告警 10 分钟内多次触发',
         '只发送 1 次通知',
         '检查通知渠道重复次数'],
    ],
    col_widths=[2.5, 4.0, 4.5, 5.0])

add_heading_styled(doc, '6.2.5 BadCase 采集测试', level=3, color_hex='2E75B6')
add_table_styled(doc,
    headers=['用例', '操作', '预期结果', '验证方式'],
    rows=[
        ['TC-OBS-17：ERROR 采集',
         '触发 1 个 ERROR 请求，等待 10 分钟',
         'agent_badcases 表新增 1 条 source=ERROR 记录',
         'SELECT * FROM agent_badcases WHERE source="ERROR"'],
        ['TC-OBS-18：DISLIKE 采集',
         '对 1 个回复点踩，等待 10 分钟',
         'agent_badcases 表新增 1 条 source=DISLIKE 记录',
         'SELECT * FROM agent_badcases WHERE source="DISLIKE"'],
        ['TC-OBS-19：去重',
         '同一消息触发 2 次 ERROR',
         'agent_badcases 表只新增 1 条（按 message_md5 去重）',
         'SELECT COUNT(*) FROM agent_badcases WHERE message_md5=?'],
    ],
    col_widths=[2.5, 4.0, 4.5, 5.0])

add_heading_styled(doc, '6.3 验收标准', level=2)

add_table_styled(doc,
    headers=['验收项', '验收标准', '状态'],
    rows=[
        ['traceId 跨服务传递', 'TC-OBS-01 通过', '待验收'],
        ['Span 记录完整', 'TC-OBS-02/03 通过', '待验收'],
        ['6 大维度 Metrics 可查', 'TC-OBS-05~09 通过', '待验收'],
        ['JSON 日志可检索', 'TC-OBS-10~12 通过', '待验收'],
        ['5 类告警可触发', 'TC-OBS-13~16 通过', '待验收'],
        ['BadCase 自动采集', 'TC-OBS-17~19 通过', '待验收'],
        ['Grafana 仪表盘可视化', '4 大面板导入可显示数据', '待验收'],
    ],
    col_widths=[4.5, 7.5, 4.0])

add_heading_styled(doc, '6.4 监控与持续改进', level=2)

add_body_paragraph(doc, '可观测性体系本身也需要监控和持续改进：', indent=True)
add_bullet(doc, '监控 Prometheus 自身健康（内存/CPU/抓取成功率）', bold_prefix='监控可观测性系统：')
add_bullet(doc, '每月 review 告警准确率，调优阈值，清理无效告警', bold_prefix='告警调优：')
add_bullet(doc, '每月 review Metrics 使用情况，删除无用指标，新增缺失指标', bold_prefix='指标调优：')
add_bullet(doc, 'Span 表按天分区或定期清理，避免无限增长', bold_prefix='DB 清理：')
add_bullet(doc, '随着 QPS 增长，升级到 Tempo/Jaeger 等专用追踪后端', bold_prefix='演进：')

add_page_break(doc)


# ==================== 第七章 总结与边界声明 ====================

add_heading_styled(doc, '七、总结与边界声明', level=1)

add_heading_styled(doc, '7.1 核心总结', level=2)

add_body_paragraph(doc, '本文档设计了 CampusShare Agent 的可观测性体系，核心内容回顾：', indent=True)

add_table_styled(doc,
    headers=['维度', '方案', 'ADR'],
    rows=[
        ['TraceId 传递',
         '网关生成 32 位 UUID → X-Trace-Id Header → agent-service MDC + Reactor Context',
         'ADR-OBS-01'],
        ['Span 追踪',
         'agent_trace_spans 表记录 7 个阶段（意图/RAG/上下文/LLM/工具/验证/记忆）耗时和状态',
         'ADR-OBS-02'],
        ['Metrics 体系',
         '6 大维度 14 项指标（延迟3/Token2/成本2/错误2/工具2/缓存3），Micrometer + Prometheus',
         'ADR-OBS-03'],
        ['结构化日志',
         'logstash-logback-encoder 输出 JSON，MDC 7 个字段（traceId/sessionId/turnId/userId/intent/model/spanName）',
         'ADR-OBS-04'],
        ['告警规则',
         '5 类告警 8 条规则（延迟P99/错误率/成本/可用率/熔断），多窗口多燃烧率',
         'ADR-OBS-05'],
        ['Grafana 仪表盘',
         '4 大面板（概览/链路/成本/错误），Prometheus + MySQL 双数据源',
         'ADR-OBS-06'],
        ['BadCase 采集',
         '定时扫描 ERROR + DISLIKE 记录，去重写入 agent_badcases 表，分诊后导出评估体系',
         'ADR-OBS-07'],
    ],
    col_widths=[2.5, 9.5, 4.0])

add_callout_box(doc, '一句话总结：',
    '可观测性体系让 Agent 从"黑盒"变为"玻璃盒"——'
    'Metrics 告诉你"出问题了"，Trace 告诉你"问题在哪"，Logs 告诉你"为什么"，'
    '告警让你"主动发现"，BadCase 采集让"持续改进"形成闭环。',
    color_hex='E8F0FE', border_color='4285F4')

add_heading_styled(doc, '7.2 与其他文档的关系', level=2)

add_table_styled(doc,
    headers=['文档', '关系', '交互点'],
    rows=[
        ['安全护栏（SEC）',
         'SEC 依赖 OBS',
         'SEC 的安全审计日志通过 OBS 的结构化日志输出；SEC 的安全事件通过 OBS 的告警规则通知'],
        ['LLM 网关（GW）',
         'GW 依赖 OBS',
         'GW 的成本追踪通过 OBS 的 Metrics 体系采集；GW 的熔断状态通过 OBS 的告警规则监控'],
        ['性能 SLO 工程（SLO）',
         'SLO 依赖 OBS',
         'SLO 的 SLO 目标基于 OBS 的 Metrics 度量；SLO 的燃烧率告警基于 OBS 的告警规则'],
        ['评估体系（EVAL）',
         'EVAL 依赖 OBS',
         'EVAL 的黄金测试集扩充基于 OBS 的 BadCase 采集；EVAL 的线上质量监控基于 OBS 的 Metrics'],
        ['分层部署（DEPLOY）',
         '并行',
         'DEPLOY 的在线/异步/离线分层后，各层的可观测性策略不同（在线全链路追踪，离线采样追踪）'],
    ],
    col_widths=[3.5, 2.5, 10.0])

add_heading_styled(doc, '7.3 演进路线', level=2)

add_table_styled(doc,
    headers=['阶段', '时间', '演进内容', '触发条件'],
    rows=[
        ['V1.0（当前）',
         '2026 Q3',
         'DB 表 Span 追踪 + Prometheus Metrics + JSON 日志 + 8 条告警 + 4 大面板 + BadCase 采集',
         '初始建设'],
        ['V1.5',
         '2026 Q4',
         '引入 Grafana Loki 采集 JSON 日志，替代文件日志；LLM 辅助 BadCase 分诊',
         '日志量 > 10GB/天；BadCase 分诊人工成本高'],
        ['V2.0',
         '2027 Q1',
         '升级到 OpenTelemetry + Tempo 分布式追踪，替代 DB 表 Span',
         'QPS > 500；DB Span 表写入压力大'],
        ['V2.5',
         '2027 Q2',
         '引入采样策略（高 QPS 时按比例采样），降低 Span 表写入压力',
         'QPS > 1000；Span 表 > 10GB/天'],
        ['V3.0',
         '2027 Q3',
         '接入 APM（SkyWalking 或自研），全服务全链路追踪',
         '多服务协作场景多；需要跨服务端到端追踪'],
    ],
    col_widths=[2.5, 2.0, 7.5, 4.0])

add_heading_styled(doc, '7.4 边界声明', level=2)

add_body_paragraph(doc, '本文档不覆盖以下内容（明确边界，避免混淆）：', indent=True)
add_bullet(doc, '不覆盖 SLO 目标定义——SLO 定义"目标是什么"（如 P99 < 3s），OBS 提供"度量手段"（Metrics 指标）', bold_prefix='边界一：')
add_bullet(doc, '不覆盖评估指标体系——EVAL 定义"评估什么"（如意图准确率），OBS 提供"数据来源"（BadCase 采集）', bold_prefix='边界二：')
add_bullet(doc, '不覆盖日志平台部署——本文档设计日志格式（JSON），不覆盖 ELK/Loki 的部署和运维', bold_prefix='边界三：')
add_bullet(doc, '不覆盖分布式追踪后端——V1.0 用 DB 表，V2.0 升级到 Tempo/Jaeger', bold_prefix='边界四：')
add_bullet(doc, '不覆盖业务监控——本文档聚焦 Agent 系统可观测性，不覆盖业务指标（如 DAU/留存率）', bold_prefix='边界五：')

add_page_break(doc)


# ==================== 附录：ADR 摘要 ====================

add_heading_styled(doc, '附录：ADR 摘要', level=1)

add_body_paragraph(doc, '本文档共 7 条 ADR，完整摘要如下：', indent=True)

# ADR-OBS-01
add_adr_box(doc,
    'ADR-OBS-01',
    '全链路 TraceId 传递',
    'Agent 系统涉及网关 → agent-service → LLM API 多跳调用，当前 traceId 仅在 agent-service 内部生成，无法跨服务追踪。',
    'gateway-service 的 TraceIdFilter 生成 32 位 UUID traceId，放入 MDC + X-Trace-Id 响应 Header；'
    '转发到 agent-service 时携带 X-Trace-Id 请求 Header；agent-service 的 TraceIdFilter 读取该 Header 放入 MDC。'
    'WebFlux 场景使用 Reactor Context 传递，避免 ThreadLocal 丢失。',
    '1. 自定义 Header 比引入 W3C Trace Context 更简单；2. 网关生成保证唯一入口；3. Reactor Context 是 WebFlux 异步传递标准方案。',
    '正面：traceId 贯穿全链路，日志可串联。负面：自定义 Header 不兼容 OTel 生态。')

# ADR-OBS-02
add_adr_box(doc,
    'ADR-OBS-02',
    'Span 级链路追踪表',
    'AgentTurn 表只记录整轮汇总耗时，无法下钻到单请求的多阶段（意图/RAG/LLM/工具）耗时。',
    '新建 agent_trace_spans 表，每条 Span 记录一个阶段的开始/结束时间、耗时、状态、属性。'
    'SpanRecorder 封装 Span 生命周期管理，异步写入 DB。不引入 Jaeger/Zipkin/Tempo。',
    '1. DB 表方案零额外组件，适合 V1.0；2. Grafana 支持 MySQL 数据源；3. Jaeger/Zipkin 需要额外部署。',
    '正面：支持单请求多阶段下钻。负面：DB 表查询性能不如专用追踪后端。')

# ADR-OBS-03
add_adr_box(doc,
    'ADR-OBS-03',
    '统一 Metrics 体系（6 大维度）',
    '现有 Metrics 只覆盖 12 项，缺少 LLM 调用耗时/Token/成本、RAG 检索耗时、工具调用、缓存命中率、错误率等核心维度。',
    '在 ObservabilityMetricsConfig 中集中定义 6 大维度 14 项指标，统一命名规范（agent.{module}.{metric}），标签低基数设计。'
    '成本指标通过 application.yml 配置模型单价。',
    '1. 集中配置便于管理；2. 低基数标签避免 Prometheus 高基数爆炸；3. 成本用 Counter 累加避免重启丢失。',
    '正面：6 大维度全覆盖。负面：14 项指标需要各采集点改造。')

# ADR-OBS-04
add_adr_box(doc,
    'ADR-OBS-04',
    '结构化 JSON 日志',
    '当前日志是纯文本格式，无法被 ELK/Loki 结构化检索，排查问题需要 grep 正则匹配。',
    '引入 logstash-logback-encoder，logback-spring.xml 配置 LogstashEncoder 输出 JSON 格式日志。'
    '在关键方法中向 MDC 写入结构化字段（sessionId/turnId/userId/intent/model）。',
    '1. JSON 日志是 ELK/Loki 事实标准；2. MDC 字段自动包含到 JSON；3. logstash-logback-encoder 最成熟。',
    '正面：日志可结构化检索。负面：JSON 格式比纯文本多约 30% 体积。')

# ADR-OBS-05
add_adr_box(doc,
    'ADR-OBS-05',
    '告警规则体系（5 类告警）',
    '当前无任何告警规则，异常只能靠用户投诉发现。',
    '定义 5 类告警（延迟P99/错误率/成本/可用率/熔断），8 条 Prometheus AlertRule。'
    'CRITICAL 级别钉钉+邮件+短信，WARNING 级别钉钉+邮件。采用多窗口多燃烧率算法避免告警风暴。',
    '1. 多窗口多燃烧率是 SRE 最佳实践；2. 分级通知避免值oncall 被轰炸；3. Prometheus AlertRule 是 CNCF 标准。',
    '正面：异常主动发现。负面：告警阈值需根据实际流量调优。')

# ADR-OBS-06
add_adr_box(doc,
    'ADR-OBS-06',
    'Grafana 仪表盘（4 大面板）',
    '当前无任何可视化仪表盘，运维和运营无法直观看到系统状态和成本。',
    '设计 4 大面板（概览/链路/成本/错误），使用 Grafana + Prometheus + MySQL 三数据源。'
    'Dashboard JSON 随项目提交，导入即用。',
    '1. Grafana 是开源标准，支持多数据源；2. Prometheus 适合聚合指标，MySQL 适合单请求下钻。',
    '正面：可视化运维，巡检效率提升。负面：Dashboard 需持续维护。')

# ADR-OBS-07
add_adr_box(doc,
    'ADR-OBS-07',
    'BadCase 自动采集',
    'AgentTurn 表中已有 ERROR 和 DISLIKE 记录，但散落在表中无人分析。线上失败案例是最有价值的改进数据。',
    'BadCaseCollector 定时任务每 10 分钟扫描 agent_turns 表，采集 ERROR 或 DISLIKE 记录，'
    '去重后写入 agent_badcases 表。人工分诊后导出为评估体系的测试用例。',
    '1. 自动采集替代人工翻表；2. 去重避免重复；3. 分诊标注可分类分析；4. 导出评估体系形成闭环。',
    '正面：线上失败案例系统化归集。负面：分诊需人工参与。')


# ==================== 保存文档 ====================

output_path = r'e:\workspace_work\CampusShare\docs\agent-design\可观测性模块设计方案.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
